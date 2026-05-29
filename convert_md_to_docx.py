#!/usr/bin/env python
"""Convert the Framework Comparison Implementation Plan markdown to a Word document."""

import re
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_code_block_lines(doc, code_lines):
    """Render fenced blocks line-by-line (COMPLIANCE_AUDIT_SELF_HEAL.docx style)."""
    for line in code_lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()


def parse_markdown(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Framework Comparison — Implementation Plan', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Gap Analysis & Implementation Roadmap')
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    
    doc.add_paragraph()  # Spacer
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_table = False
    in_code_block = False
    table_lines = []
    code_lines = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block — one paragraph per line (preserves ASCII workflows)
                if code_lines:
                    add_code_block_lines(doc, code_lines)
                    code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line.rstrip('\n'))
            i += 1
            continue
        
        # Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        else:
            if in_table:
                # Process table
                if len(table_lines) >= 2:
                    process_table(doc, table_lines)
                table_lines = []
                in_table = False
        
        # Headings
        if stripped.startswith('# '):
            text = stripped[2:]
            doc.add_heading(text, level=1)
            i += 1
            continue
        if stripped.startswith('## '):
            text = stripped[3:]
            doc.add_heading(text, level=2)
            i += 1
            continue
        if stripped.startswith('### '):
            text = stripped[4:]
            doc.add_heading(text, level=3)
            i += 1
            continue
        if stripped.startswith('#### '):
            text = stripped[5:]
            doc.add_heading(text, level=4)
            i += 1
            continue
        
        # Bold list items like "| Function | Purpose | ..." header separator handling already done above
        if stripped.startswith('**') and stripped.endswith('**') and not stripped.startswith('|'):
            p = doc.add_paragraph()
            run = p.add_run(stripped.replace('**', ''))
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(text, style='List Bullet')
            # Handle inline bold/italic
            i += 1
            continue
        
        # Numbered lists
        match = re.match(r'^\d+\.\s+(.*)', stripped)
        if match:
            text = match.group(1)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Horizontal rule / separator
        if stripped == '---' or stripped == '***':
            doc.add_paragraph()
            i += 1
            continue
        
        # Regular paragraph
        if stripped:
            p = doc.add_paragraph(stripped)
        else:
            doc.add_paragraph()  # Empty line spacer
        
        i += 1
    
    # Save
    doc.save(docx_path)
    print(f"Document saved to: {docx_path}")

def process_table(doc, table_lines):
    """Convert markdown table lines to a Word table."""
    if not table_lines:
        return
    
    # Remove separator line (line with just dashes and pipes)
    clean_lines = []
    for line in table_lines:
        stripped = line.strip()
        if re.match(r'^\|[-\s|]+\|$', stripped):
            continue
        clean_lines.append(line)
    
    if not clean_lines:
        return
    
    # Parse cells
    rows = []
    for line in clean_lines:
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    
    if not rows:
        return
    
    num_cols = max(len(r) for r in rows)
    num_rows = len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx in range(num_cols):
            cell = row.cells[col_idx]
            if col_idx < len(row_data):
                text = row_data[col_idx]
                # Clean up markdown bold/italic in cells
                text = text.replace('**', '')
                text = text.replace('*', '')
                text = text.replace('`', '')
                cell.text = text
            else:
                cell.text = ''
            
            # Header row formatting
            if row_idx == 0:
                set_cell_shading(cell, '4472C4')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.size = Pt(10)
            else:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacer after table

if __name__ == '__main__':
    md_file = r'e:\KPMG-WS2-Final-Testing\Framework_Comparison_Implementation_Plan.md'
    docx_file = r'e:\KPMG-WS2-Final-Testing\Framework_Comparison_Implementation_Plan.docx'
    
    if not os.path.exists(md_file):
        print(f"Error: Markdown file not found at {md_file}")
        exit(1)
    
    parse_markdown(md_file, docx_file)
    print("Conversion complete!")
