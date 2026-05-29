"""Generate POLICY_OUTDATED_SELF_HEAL.docx — text workflow diagrams (no images), functions + usage."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

OUT_PATH = Path(__file__).resolve().parent / "POLICY_OUTDATED_SELF_HEAL.docx"

# Text-only workflow blocks (editable in Word, not images)
FLOW_FEATURE_1 = """
┌─────────────────────────────────────────────────────────────────┐
│  PART A — AUTOMATED JOB (runs on schedule, e.g. once per day)   │
└─────────────────────────────────────────────────────────────────┘

    START: Scheduled job runs
              │
              ▼
    Load all approved, active policies where TODAY is between
    Start Date and End Date (framework must also be approved/active)
              │
              ▼
    ┌─────────────────────────┐
    │ Due for review TODAY?   │──── NO ────► SKIP this policy
    │  • every N days from    │
    │    start (e.g. 30), OR  │
    │  • daily in last 7 days │
    │    before End Date      │
    └─────────────────────────┘
              │ YES
              ▼
    ┌─────────────────────────┐
    │ Reminder already sent   │──── YES ───► SKIP (one per day max)
    │ for this policy TODAY?  │
    └─────────────────────────┘
              │ NO
              ▼
    ┌─────────────────────────┐
    │ Is CREATOR active?    │──── YES ───► Send reminder to CREATOR
    └─────────────────────────┘              (email + in-app)
              │ NO
              ▼
    ┌─────────────────────────┐
    │ NEITHER available       │────► Escalate to Policy Managers
    │ (creator missing or     │      (see Feature 2)
    │  inactive)            │
    └─────────────────────────┘
              │
              ▼
    Record: reminder sent for today
              │
              ▼

  Note: If a manager already assigned someone earlier, that person is
  treated as the responsible user instead of the creator (custodian).
  Then the job sends the reminder to them — not to managers again.

┌─────────────────────────────────────────────────────────────────┐
│  PART B — USER (creator, or assigned person after Feature 2)    │
└─────────────────────────────────────────────────────────────────┘

    User receives reminder → opens Renewal Review screen
              │
              ▼
    ┌─────────────────┐
    │  User chooses   │
    └─────────────────┘
         │           │
         ▼           ▼
     APPROVE      UPDATE
         │           │
         ▼           ▼
  Policy stays   User edits policy in
  as-is; dates   tailoring and submits
  unchanged      new version for approval
         │           │
         └─────┬─────┘
               ▼
            END
"""

FLOW_FEATURE_2 = """
┌─────────────────────────────────────────────────────────────────┐
│  PART A — SYSTEM (same job as Feature 1, escalation branch)     │
└─────────────────────────────────────────────────────────────────┘

    Policy is DUE for review today
    BUT creator is NOT active (missing or inactive) — NEITHER case
              │
              ▼
    Find all active Policy Managers (for this tenant)
              │
              ▼
    Create / keep PENDING escalation record for this policy
              │
              ▼
    Notify every Policy Manager (in-app + email)
              │
              ▼
    Record: reminder sent for today (same daily limit)
              │
              ▼

┌─────────────────────────────────────────────────────────────────┐
│  PART B — POLICY MANAGER                                        │
└─────────────────────────────────────────────────────────────────┘

    Manager opens dashboard → sees policy in escalation list
              │
              ▼
    Manager picks an active user to take over renewal
              │
              ▼
    System saves assignment (policy owner updated, escalation = assigned)
              │
              ▼
    That user gets renewal reminder immediately
              │
              ▼

┌─────────────────────────────────────────────────────────────────┐
│  PART C — ASSIGNED USER (same renewal steps as creator)         │
└─────────────────────────────────────────────────────────────────┘

    Assigned user opens Renewal Review → Approve OR Update
              │
              ▼
    On NEXT scheduled runs: job checks creator path first;
    if assignment exists, reminders go to that user only
    (managers are NOT notified again for this policy)
              │
              ▼
            END
"""


def add_text_flow(doc: Document, flow_text: str) -> None:
    """Insert workflow as monospace paragraphs so it stays text in Word."""
    for line in flow_text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()


def add_functions_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Function"
    h[1].text = "Usage"
    for p in h[0].paragraphs + h[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for i, (fn, usage) in enumerate(rows):
        c = table.rows[i + 1].cells
        c[0].text = fn
        c[1].text = usage
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.add_heading("Policy outdated detection & creator escalation", 0)

    # ── Feature 1 ──
    doc.add_heading("Feature 1 — Outdated policy detection and renewal reminders", 1)

    doc.add_heading("What normally happens", 2)
    doc.add_paragraph(
        "A scheduled job checks which policies need a renewal review. When a policy is due, "
        "there are only two outcomes at first: the creator is active and gets the reminder, or "
        "neither is available and managers are involved (Feature 2). "
        "After a manager assigns someone, that person receives future reminders instead of the creator."
    )
    doc.add_paragraph(
        "Creator vs assigned person: The creator is the original policy owner. "
        "An assigned person is only used after Feature 2 when the creator could not be reached."
    )

    doc.add_heading("Workflow diagram", 2)
    doc.add_paragraph(
        "Read top to bottom. Branches show decisions (YES/NO or which path). "
        "This is plain text in the document — not a picture — so you can edit it in Word."
    )
    add_text_flow(doc, FLOW_FEATURE_1)

    doc.add_heading("Functions and usage", 2)
    add_functions_table(
        doc,
        [
            (
                "execute_policy_self_heal_reminders",
                "Main job: loops policies, applies detection, sends reminders or triggers escalation, returns sent/skipped counts.",
            ),
            (
                "policy_should_remind_today",
                "Decides if today is a reminder day (frequency from start date OR last 7 days before end).",
            ),
            (
                "_resolve_reminder_target",
                "Two main branches: active creator → remind creator; neither → escalate. "
                "If a manager already assigned someone, that user is used instead of creator (see note in diagram).",
            ),
            (
                "resolve_policy_creator_user_id",
                "Finds the creator — used for the “creator active?” check.",
            ),
            (
                "_assigned_custodian_user_id",
                "Only after Feature 2: returns the user id the manager assigned, so reminders skip escalation.",
            ),
            (
                "_user_is_active",
                "Checks the user account is still active before notifying.",
            ),
            (
                "_notify_renewal_user",
                "Sends reminder to the creator or to the manager-assigned user (notification + optional email).",
            ),
            (
                "policy_self_heal_decision",
                "Handles Approve or Update from the renewal review screen.",
            ),
            (
                "_can_perform_self_heal_decision",
                "Only the creator or the manager-assigned user may submit a renewal decision.",
            ),
            (
                "PolicyReviewReminderSent",
                "Stores that a reminder was sent today so the same policy is not reminded twice in one day.",
            ),
        ],
    )

    # ── Feature 2 ──
    doc.add_heading("Feature 2 — Creator unavailable → Policy Manager", 1)

    doc.add_heading("What normally happens", 2)
    doc.add_paragraph(
        "This runs only in the NEITHER case: creator is missing or inactive. "
        "Policy Managers are alerted. A manager assigns an active user to own renewal. "
        "That person then receives reminders and completes the renewal review like the creator would."
    )

    doc.add_heading("Workflow diagram", 2)
    doc.add_paragraph(
        "Part A is the NEITHER branch from Feature 1. Parts B and C are what managers and the assigned user do next."
    )
    add_text_flow(doc, FLOW_FEATURE_2)

    doc.add_heading("Functions and usage", 2)
    add_functions_table(
        doc,
        [
            (
                "_resolve_reminder_target",
                "Returns escalate in the NEITHER case (no active creator, and no prior manager assignment).",
            ),
            (
                "_policy_manager_user_ids",
                "Builds list of active Policy Manager user ids to notify.",
            ),
            (
                "_ensure_pending_escalation",
                "Creates pending escalation row if none exists.",
            ),
            (
                "_notify_policy_managers_escalation",
                "Notifies all Policy Managers for the policy.",
            ),
            (
                "list_pending_self_heal_escalations",
                "Loads policies waiting for custodian assignment on the dashboard.",
            ),
            (
                "assign_self_heal_custodian",
                "Manager assigns user, updates policy owner, marks escalation assigned, sends first reminder to them.",
            ),
            (
                "_is_policy_manager",
                "Ensures only Policy Managers can list escalations and assign custodians.",
            ),
            (
                "PolicySelfHealEscalation",
                "Tracks pending vs assigned state and who was assigned.",
            ),
        ],
    )

    doc.add_heading("How both features connect", 1)
    doc.add_paragraph(
        "Every run starts with Feature 1: creator active → remind creator; neither → Feature 2. "
        "After a manager assigns someone, later runs remind that user (creator path with an existing assignment) "
        "and managers are not notified again."
    )

    doc.save(OUT_PATH)
    print(f"Written: {OUT_PATH}")


if __name__ == "__main__":
    build()
