"""Generate POLICY_CONFLICT_RESOLUTION.docx — Feature 1 and Feature 2 separate; implemented + planned each."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

OUT_PATH = Path(__file__).resolve().parent / "POLICY_CONFLICT_RESOLUTION.docx"

# --- Feature 1 diagrams ---
FLOW_F1_IMPLEMENTED = """
    User on CREATE form (Framework / Policy / Compliance — Suggest button)
              │
              ▼
    User clicks SUGGEST with current form data
              │
              ▼
    Steps 1–6: clean → domain → embed → search → rerank → LLM
              │
              ▼
    Similar records listed (same type, same tenant; may be other framework)
              │
              ▼
    If no match / LOW risk → form may continue without modal
    If matches / MEDIUM or HIGH risk → user goes to Feature 2 modal
"""

FLOW_F1_PLANNED = """
    User on CREATE or UPDATE (all 4: Framework, Policy, SubPolicy, Compliance)
              │
              ▼
    Check runs automatically on Save / Submit (not only Suggest button)
              │
              ▼
    Same Steps 1–6 pipeline
              │
              ▼
    On UPDATE: message “This is already similar — keep your changes?”
              │
              ▼
    SubPolicy create/edit screens wired with same Suggest / auto-check
"""

# --- Feature 2 diagrams ---
FLOW_F2_IMPLEMENTED = """
    Modal opens after Feature 1 finds similar records
              │
              ▼
    User sees risk, advice, list of similar items (scores, reasons)
              │
              ▼
    ┌─────────────────────────────────────┐
    │ User choice (logged-in user today)  │
    └─────────────────────────────────────┘
              │
      ┌───────┼───────────────┬──────────────┐
      ▼       ▼               ▼
  CREATE    USE SELECTED    CANCEL
  ANYWAY    existing
      │       │
      ▼       ▼
  New       No duplicate;
  record    use existing id
  created
      │
      ▼
    END
"""

FLOW_F2_PLANNED = """
    After Feature 1 finds conflict (including across frameworks)
              │
              ▼
    Notify CREATOR to resolve (keep new / use existing / cancel)
              │
              ▼
    ┌─────────────────────────┐
    │ Is CREATOR active?      │──── YES ──► Creator acts in UI
    └─────────────────────────┘
              │ NO
              ▼
    ┌─────────────────────────┐
    │ NEITHER                 │────► Policy Managers notified
    └─────────────────────────┘
              │
              ▼
    Manager assigns who will resolve the conflict
              │
              ▼
    Assigned user completes same choices as creator
              │
              ▼
    END
"""


def add_text_flow(doc: Document, flow_text: str) -> None:
    for line in flow_text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()


def add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_functions_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Function"
    h[1].text = "Usage"
    for p in h[0].paragraphs + h[1].paragraphs:
        for r in p.runs:
            r.bold = True
    for i, (fn, usage) in enumerate(rows):
        c = table.rows[i + 1].cells
        c[0].text = fn
        c[1].text = usage
    doc.add_paragraph()


def section_implemented(doc: Document) -> None:
    doc.add_heading("Implemented", 3)


def section_planned(doc: Document) -> None:
    doc.add_heading("Planned", 3)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.add_heading("Policy conflict resolution (similarity detection)", 0)
    doc.add_paragraph(
        "Two features: (1) find similar records before save, (2) let the user resolve the conflict. "
        "Each feature below has its own implemented and planned sections."
    )

    # ===================== FEATURE 1 =====================
    doc.add_heading("Feature 1 — Similarity detection and suggestions", 1)

    section_implemented(doc)
    doc.add_paragraph(
        "When a user is creating a Framework, Policy, or Compliance record, they can click Suggest. "
        "The system compares the draft to existing records of the same type in the tenant using a six-step pipeline. "
        "Matches can appear from the same framework or from another framework. "
        "If nothing similar is found (low risk), the form can continue; otherwise the user is sent to Feature 2."
    )
    doc.add_heading("Workflow diagram (implemented)", 4)
    add_text_flow(doc, FLOW_F1_IMPLEMENTED)
    doc.add_heading("Functions and usage (implemented)", 4)
    add_functions_table(
        doc,
        [
            ("check_similarity", "API when user clicks Suggest; starts the pipeline and returns candidates."),
            ("initiate_similarity_check", "Runs Steps 1–6 and saves audit with risk and candidates."),
            ("step1_clean_text", "Normalizes name and description for the entity type."),
            ("step2_classify_domain", "Domain and industry for search context."),
            ("step3_generate_embedding", "Semantic vector for the draft record."),
            ("step4_search_similar", "ChromaDB search — same entity_type and tenant; optional parent filters."),
            ("step5_rerank_candidates", "Cross-encoder rescore of top matches."),
            ("step6_llm_decision", "LLM risk and labels (e.g. highly similar, related but different)."),
            ("get_similarity_result", "Reload check results by id for the modal."),
            ("index_record", "Indexes saved records so future checks can find them."),
        ],
    )
    doc.add_paragraph("Entity types in implemented UI:")
    add_bullet_list(
        doc,
        [
            "Framework — Create Framework form (Suggest).",
            "Policy — Create Policy form (Suggest).",
            "Compliance — Create Compliance form (Suggest).",
            "SubPolicy — pipeline supports SubPolicy in API; Suggest on SubPolicy create form not wired yet.",
        ],
    )

    section_planned(doc)
    doc.add_paragraph(
        "Detection should run on both create and update for all four entity types, without relying only on the Suggest button."
    )
    doc.add_heading("Workflow diagram (planned)", 4)
    add_text_flow(doc, FLOW_F1_PLANNED)
    doc.add_heading("Planned behaviour", 4)
    add_bullet_list(
        doc,
        [
            "Auto-run similarity check on Save / Submit for create and update.",
            "All four types: Framework, Policy, SubPolicy, Compliance — including SubPolicy on tailoring/create screens.",
            "On update: prompt that the record is already similar and ask whether to keep changes.",
            "Re-index embeddings when an existing record is updated so search stays accurate.",
        ],
    )
    doc.add_heading("Functions and usage (planned)", 4)
    add_functions_table(
        doc,
        [
            (
                "trigger_similarity_before_save",
                "Called from create/update submit handlers for all four entity types; runs initiate_similarity_check before persisting to database.",
            ),
            (
                "run_similarity_on_update",
                "Same pipeline as create but passes existing record id so Step 4 excludes self-match and compares revised text to index.",
            ),
            (
                "prompt_update_similarity_decision",
                "After check on update, shows UI: keep changes, revert to previous, or use existing similar record — maps to user decision enum.",
            ),
            (
                "wire_similarity_on_subpolicy_forms",
                "Frontend integration on SubPolicy create/tailoring screens; same Suggest or auto-submit hook as Policy and Compliance.",
            ),
            (
                "reindex_record_on_update",
                "After approved update, refresh embedding in ChromaDB via index_record so future checks see latest content.",
            ),
            (
                "should_block_save_without_check",
                "Optional guard: if form dirty and no check_id for this session, block Save until similarity pipeline completes.",
            ),
        ],
    )

    # ===================== FEATURE 2 =====================
    doc.add_heading("Feature 2 — Conflict resolution (user decision)", 1)

    section_implemented(doc)
    doc.add_paragraph(
        "After Feature 1 finds similar records, a modal shows risk level, advice, and the list of matches. "
        "The logged-in user chooses: create the new record anyway, use an existing selected record, or cancel. "
        "The choice is stored and the backend creates the record, links to existing, or stops."
    )
    doc.add_heading("Workflow diagram (implemented)", 4)
    add_text_flow(doc, FLOW_F2_IMPLEMENTED)
    doc.add_heading("Functions and usage (implemented)", 4)
    add_functions_table(
        doc,
        [
            ("user_similarity_decision", "Records Create Anyway, Use Selected, or Cancel on the audit."),
            ("finalize_similarity_check", "Runs Step 9 from the user decision."),
            ("_handle_create_anyway", "Creates new Framework/Policy/SubPolicy/Compliance and activates index."),
            ("_handle_use_existing", "Does not create; returns existing candidate for navigation."),
            ("_handle_cancel", "Aborts; cleans up pending draft embedding."),
            ("checkSimilarity", "Frontend — calls check API and opens modal when needed."),
            ("handleCreateAnyway", "Frontend — posts CREATE_ANYWAY and closes modal."),
            ("handleUseExisting", "Frontend — posts USE_EXISTING with selected candidate id."),
            ("handleCancel", "Frontend — posts CANCEL and closes modal."),
        ],
    )
    doc.add_paragraph("User choices (implemented):")
    add_bullet_list(
        doc,
        [
            "Create Anyway — keep the new record despite similarity.",
            "Use Selected — do not duplicate; use the existing record chosen from the list.",
            "Cancel — stop; no create.",
        ],
    )

    section_planned(doc)
    doc.add_paragraph(
        "Resolution should not depend only on whoever clicked Suggest. "
        "When there is a cross-framework conflict, the creator should be notified; "
        "if the creator is not available, Policy Managers assign someone to resolve — same pattern as outdated-policy escalation."
    )
    doc.add_heading("Workflow diagram (planned)", 4)
    add_text_flow(doc, FLOW_F2_PLANNED)
    doc.add_heading("Planned behaviour", 4)
    add_bullet_list(
        doc,
        [
            "Notify policy/compliance creator when a conflict is detected (including different framework).",
            "Creator active — creator resolves in UI (keep / use existing / cancel).",
            "Creator not available (neither) — notify Policy Managers; manager assigns resolver.",
            "Assigned resolver has the same decision options as the creator.",
            "Applies to policy, subpolicy, and compliance conflicts as well as framework-level duplicates where relevant.",
        ],
    )
    doc.add_heading("Functions and usage (planned)", 4)
    add_functions_table(
        doc,
        [
            (
                "notify_creator_on_similarity_conflict",
                "After Feature 1 finds medium/high risk, send in-app and email to resolved creator with link to conflict review screen.",
            ),
            (
                "resolve_conflict_target",
                "Returns creator user id if active; else escalate — same precedence as renewal self-heal (custodian assignment if already set, else creator, else neither).",
            ),
            (
                "resolve_policy_creator_user_id",
                "Reuse from self-heal: find creator from PolicyApproval or CreatedByName for policy/subpolicy/compliance parent policy.",
            ),
            (
                "_notify_conflict_escalation_managers",
                "On neither branch, notify all active Policy Managers for tenant with dashboard link to conflict assignment list.",
            ),
            (
                "_ensure_pending_conflict_escalation",
                "Create or return pending row on conflict_escalation table for the entity being created or updated.",
            ),
            (
                "list_pending_conflict_escalations",
                "API for Policy Manager dashboard: policies/subpolicies/compliances waiting for assigned resolver.",
            ),
            (
                "assign_conflict_resolver",
                "Manager assigns active user; store assigned_user_id; notify assignee; future decisions use assignee not creator.",
            ),
            (
                "_can_perform_conflict_decision",
                "Allow submit of keep / use existing / cancel only if request user is creator or assigned conflict resolver.",
            ),
            (
                "record_conflict_resolution_decision",
                "Persist final choice on SimilarityCheckAudit plus escalation status when manager path was used.",
            ),
        ],
    )

    doc.save(OUT_PATH)
    print(f"Written: {OUT_PATH}")


if __name__ == "__main__":
    build()
