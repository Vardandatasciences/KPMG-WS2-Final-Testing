"""Generate COMPLIANCE_AUDIT_SELF_HEAL.docx — single full text workflow only."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

OUT_PATH = Path(__file__).resolve().parent / "COMPLIANCE_AUDIT_SELF_HEAL.docx"

# One document = one workflow (rules, branches, functions, APIs, UI — all inline)
FLOW_COMPLETE = """
═══════════════════════════════════════════════════════════════════════════════
  MISSED COMPLIANCE & AUDIT TASK — FULL WORKFLOW (implementation)
  Manual audits (I/E/S) + AI audits (A). Same scheduler as policy_self_healing.py.
═══════════════════════════════════════════════════════════════════════════════

RULES (product) — BOTH manual and AI
  • Reminders to auditor / reviewer       → AUDITOR / REVIEWER on audit row (NOT Assignee)
  • Due soon                              → 4 days before DueDate or before review due
  • Review due                            → ReviewStartDate + 5 days (ReviewDate = when review ENDED)
  • Escalation                            → AFTER 3+ overdue reminder days → Compliance Manager
  • SubPolicy EndDate only                → NO EndDate on Compliance; STOP all after subpolicy end
  • Example: SubPolicy end 01/04/2026, DueDate 01/09/2026 → reminders STOP after 01/04/2026

MANUAL ONLY (TaskView /audit/{id}/tasks)
  • Missing compliance evidence           → audit_findings.Evidence empty
  • Frequency                             → AssignedDate + audit.Frequency (days on assign)
  • Missing audit evidence                → audit.Evidence empty

AI ONLY (AIAuditDocumentUpload /ai-audit/{id}/upload)
  • Missing evidence                      → no uploaded docs / ai_audit_data; use EvidenceReminderDays from assign
  • Frequency                             → AIAuditSchedule on upload page (cron/next_run), NOT audit.Frequency
  • Run / report                          → remind upload, run check, or download report if results exist
  • Merge check_ai_audit_evidence command → avoid duplicate “no evidence” alerts

SCHEDULING (how the job runs — use same as policy self-heal)
  Triggers (any one):
    (1) apps.py → _run_compliance_audit_self_heal_loop()
        when ENABLE_COMPLIANCE_AUDIT_SELF_HEAL_SCHEDULER=true
    (2) python manage.py run_compliance_audit_self_heal_reminders
    (3) POST /api/compliance-audit/self-healing/reminders/run/
        header X-Compliance-Audit-Self-Heal-Secret

  Main entry:
    execute_compliance_audit_self_heal_reminders(for_date=today)
      → returns { date, sent, skipped, escalations }

  New module: grc/routes/Compliance/compliance_audit_self_healing.py
  Reuse: NotificationService.send_multi_channel_notification
         notifications.py (_insert_db_notification, _append_memory_notification)
         get_user_email_from_id — same as policy_self_healing.py

───────────────────────────────────────────────────────────────────────────────
PART A — START OF EACH DAILY RUN
───────────────────────────────────────────────────────────────────────────────

    START
      │
      ▼
    Load ALL audits from DB (tenant-aware) — includes AuditType A (AI)
      │
      ▼
    For EACH audit ──────────────────────────────────────────────────────────┐
      │                                                                       │
      ▼                                                                       │
    self_heal_active_for_audit(audit, today)                                  │
      │ SubPolicy.EndDate cap; skip if fully complete                         │
      │                                                                       │
      ├── today > SubPolicy.EndDate (if set) ──────────► SKIP audit           │
      └── else ACTIVE ────────────────────────────────────────────────────────┘
      │
      ▼
    AuditType == 'A' ?  ──YES──► Run PARTS AI-B → AI-F (below)
      │                └──NO───► Run PARTS B → F (manual TaskView)
      │
      ▼
    PART G escalation (manual + AI) — max 1 notify per TYPE per day
      │
      ▼
    ComplianceAuditReminderSent.get_or_create(audit, type, date=today)  ◄── dedup
      │
      ▼
    END RUN

  Data (migration):
    SubPolicy.EndDate (NEW field on subpolicy only)
    Table compliance_audit_reminder_sent (dedup)
    Table compliance_audit_self_heal_escalation (pending / assigned)

───────────────────────────────────────────────────────────────────────────────
PART B — MANUAL ONLY — MISSING COMPLIANCE EVIDENCE (per audit_finding)
───────────────────────────────────────────────────────────────────────────────

    Audit Status in: Yet to Start, Work In Progress  (NOT Under review / Completed)
      │
      ▼
    For EACH row in audit_findings for this AuditId
      │
      ▼
    audit_findings.Evidence empty or whitespace?
      │
      NO ──────────────────────────────────────────────► next finding
      │
      YES
      ▼
    should_remind_missing_compliance_evidence(audit, finding, today)
      │
      ├── already sent type=missing_compliance_evidence today? ──► SKIP
      │
      NO
      ▼
    _notify_auditor_missing_evidence(audit, finding, audit.Auditor_id)
      • DB notification + in-app (Notifications.vue)
      • optional email via NotificationService
      • link: /audit/{AuditId}/tasks  (TaskView — Upload Compliance Evidence)
      │
      ▼
    Record ComplianceAuditReminderSent

  Why Auditor: assign_audit sets audit_findings.UserId = audit.Auditor_id;
               get_my_audits filters a.auditor = current user.

───────────────────────────────────────────────────────────────────────────────
PART C — MANUAL ONLY — FREQUENCY (audit.Frequency from Assign Audit)
───────────────────────────────────────────────────────────────────────────────

    audit.Frequency set and Frequency > 0  (skip “Only Once” / 0)
    audit.AssignedDate set
      │
      ▼
    days_since = (today - AssignedDate.date()).days
    days_since > 0 AND (days_since % Frequency == 0)   ◄── same idea as policy_should_remind_today
      │
      NO ──────────────────────────────────────────────► SKIP
      │
      YES
      ▼
    Status NOT Completed; default skip if Under review (auditor handed off)
      │
      ▼
    should_remind_frequency_today → dedup type=audit_frequency
      │
      ▼
    _notify_auditor_frequency(audit, audit.Auditor_id)
      • message: perform audit per schedule
      • link: /audit/{AuditId}/tasks

───────────────────────────────────────────────────────────────────────────────
PART D — MANUAL + AI — AUDIT DUE DATE (audit.DueDate set on assign for both)
───────────────────────────────────────────────────────────────────────────────

    audit.DueDate set
      │
      ▼
    work_complete_for_due(audit) ?   ◄── manual: audit_work_complete; AI: ai_audit_work_complete
      │
      YES ─────────────────────────────────────────────► SKIP due/overdue (Part D done)
      │
      NO
      ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │ DueDate - today == 4                                            │
    │   → should_remind type=audit_due_soon → _notify_auditor_due     │
    │ today == DueDate                                                │
    │   → type=audit_due_today → _notify_auditor_due                  │
    │ today > DueDate                                                 │
    │   → type=audit_overdue → _notify_auditor_overdue                 │
    └─────────────────────────────────────────────────────────────────┘
      │
      ▼
    overdue_reminder_days(audit)  ◄── count days type=audit_overdue sent since DueDate
      │
      ├── count < 3 ───────────────────────────────────► stay in Part D only
      │
      └── count >= 3 ──────────────────────────────────► go to PART G (auditor escalation)

───────────────────────────────────────────────────────────────────────────────
PART E — MANUAL + AI — REVIEW (Reviewer)
───────────────────────────────────────────────────────────────────────────────

  MANUAL: auditor Send for Review → send_audit_for_review (set ReviewStartDate)
  AI: all documents processed → ai_audit_api sets Status Under review + ReviewStartDate
      │
      ▼
    PATCH send_audit_for_review AND ai_audit_api status updates:
      Status = 'Under review'
      ReviewStartDate = COALESCE(ReviewStartDate, NOW())
      │
      ▼
    review_due_date(audit) = ReviewStartDate.date() + 5   ◄── COMPLIANCE_AUDIT_REVIEW_SLA_DAYS

  DAILY JOB (only if Status = 'Under review'):
      │
      ▼
    review_complete(audit) ?
      • ReviewStatus Accept(2) or Reject(3), OR ReviewDate set (review ENDED)
      │
      YES ─────────────────────────────────────────────► SKIP Part E
      │
      NO
      ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │ review_due_date - today == 4                                    │
    │   → type=review_due_soon → _notify_reviewer(audit.reviewer)     │
    │ today == review_due_date                                        │
    │   → type=review_due_today                                       │
    │ today > review_due_date                                         │
    │   → type=review_overdue                                         │
    │ link: /reviewer/task/{AuditId}  (ReviewTaskView — manual + AI)    │
    └─────────────────────────────────────────────────────────────────┘
      │
      ▼
    overdue review reminder days >= 3 ?
      │
      YES ─────────────────────────────────────────────► PART G (reviewer escalation)

  Fields meaning:
    ReviewStartDate = review START (use for SLA)
    ReviewDate      = review FINISHED — do NOT use as due date
    ReviewStatus    = 0 Yet to Start, 1 In Review, 2 Accept, 3 Reject

───────────────────────────────────────────────────────────────────────────────
PART F — MANUAL ONLY — AUDIT EVIDENCE (audit.Evidence on TaskView)
───────────────────────────────────────────────────────────────────────────────

    audit.Evidence empty at audit level (Upload Audit Evidence on TaskView)
      │
      ▼
    Same gates as Part B (active audit, auditor, dedup type=missing_audit_evidence)
      │
      ▼
    _notify_auditor_missing_audit_evidence → /audit/{AuditId}/tasks

───────────────────────────────────────────────────────────────────────────────
PART AI-B — AI ONLY — MISSING EVIDENCE (AIAuditDocumentUpload)
───────────────────────────────────────────────────────────────────────────────

    Page: upload documents, Schedule AI Audit, Download Report
    Assign Audit (AI): EvidenceReminderDays (default 7) stored on audit
      │
      ▼
    ai_audit_has_documents(audit_id) == false ?
      │
      YES → days since AssignedDate >= EvidenceReminderDays (or due-within-15 legacy)
      │
      ▼
    should_remind_ai_missing_evidence → dedup type=ai_missing_evidence
      │
      ▼
    _notify_auditor_ai_evidence(audit)
      • link: /ai-audit/{AuditId}/upload  (NOT TaskView)
      • replaces duplicate check_ai_audit_evidence when job is live

───────────────────────────────────────────────────────────────────────────────
PART AI-C — AI ONLY — SCHEDULE FREQUENCY (ai_audit_schedule table)
───────────────────────────────────────────────────────────────────────────────

    User sets Frequency on upload page: Daily / Weekly / Monthly / Quarterly / Yearly
    Stored in AIAuditSchedule (cron, next_run_at) — audit.Frequency is NULL for AI
      │
      ▼
    should_remind_ai_schedule_today(audit)
      • e.g. next_run_at - today == 4 → ai_schedule_due_soon
      • or no schedule but AssignedDate + N days → remind to upload/run once
      │
      ▼
    _notify_auditor_ai_schedule → /ai-audit/{AuditId}/upload

───────────────────────────────────────────────────────────────────────────────
PART AI-D — AI ONLY — RUN CHECK / REPORT (optional product)
───────────────────────────────────────────────────────────────────────────────

    Documents uploaded but no compliance results / no successful schedule run?
      → type=ai_run_pending → auditor: run check on upload page
    Results exist, report not generated/downloaded within X days? (optional)
      → type=ai_report_pending → auditor: Download Report on upload page

───────────────────────────────────────────────────────────────────────────────
PART G — ESCALATION → COMPLIANCE MANAGER (manual + AI, after overdue reminders)
───────────────────────────────────────────────────────────────────────────────

    should_escalate(audit, role='auditor' OR 'reviewer', today)
      │
      ▼
    _compliance_manager_user_ids(tenant_id)   ◄── RBAC role "Compliance Manager"
      │
      ▼
    _ensure_pending_escalation(audit, escalation_type)
      → ComplianceAuditSelfHealEscalation status=pending_assignment
      │
      ▼
    _notify_compliance_managers_escalation(audit, type)
      • in-app metadata → dashboard ?auditEscalations=1
      • dedup managers same day

  MANAGER (UI — mirror Policy Dashboard renewal escalations):
    GET  /api/compliance-audit/self-healing/escalations/pending/
         → list_pending_compliance_audit_escalations()
         → _is_compliance_manager(user)
      │
      ▼
    Manager picks new user
      │
      ▼
    POST /api/compliance-audit/{audit_id}/self-healing/assign/
         body: { role: "auditor"|"reviewer", assigned_user_id }
         → assign_compliance_audit_escalation()
      │
      ├── role=auditor  → UPDATE audit SET auditor = assigned_user_id
      └── role=reviewer → UPDATE audit SET reviewer = assigned_user_id
      │
      ▼
    escalation status = assigned
    _notify_escalation_assignee()  ◄── new person gets first reminder
      │
      ▼
    NEXT daily runs: reminders go to NEW auditor/reviewer (not manager again)

───────────────────────────────────────────────────────────────────────────────
FRONTEND / UI TOUCH
───────────────────────────────────────────────────────────────────────────────

  SubPolicy create/edit     → capture SubPolicy.EndDate (no field on Compliance)
  Compliance Manager panel  → list pending escalations, assign auditor/reviewer
  Notifications.vue         → TaskView, ReviewTaskView, /ai-audit/{id}/upload, escalation panel
  api.js                    → COMPLIANCE_AUDIT_SELF_HEAL_* endpoints
  Deprecate                 → check_ai_audit_evidence merges into AI-B when live

═══════════════════════════════════════════════════════════════════════════════
                              END OF WORKFLOW
     (Functions and why — see tables in document section below)
═══════════════════════════════════════════════════════════════════════════════
"""

# (Function, Why we use it)
FUNCTION_ROWS: list[tuple[str, str, str]] = [
    # category, function, why
    ("Scheduling", "execute_compliance_audit_self_heal_reminders", "Main daily job: loops audits, runs Parts B–G, returns sent/skipped/escalation counts. Same role as execute_policy_self_heal_reminders."),
    ("Scheduling", "run_compliance_audit_self_heal_reminders", "Webhook/cron entry for Scheduler microservice; validates COMPLIANCE_AUDIT_SELF_HEAL_CRON_SECRET."),
    ("Scheduling", "run_compliance_audit_self_heal_reminders (management command)", "Manual/ops run: python manage.py run_compliance_audit_self_heal_reminders."),
    ("Scheduling", "_run_compliance_audit_self_heal_loop (apps.py)", "Optional inline thread on runserver when ENABLE_COMPLIANCE_AUDIT_SELF_HEAL_SCHEDULER=true."),
    ("SubPolicy end", "SubPolicy.EndDate (new field)", "Only place for lifecycle end; Compliance has no EndDate. UI on subpolicy create/edit."),
    ("SubPolicy end", "effective_subpolicy_end_date(audit)", "Reads end date from audit.SubPolicyId; caps reminders even if audit.DueDate is later."),
    ("SubPolicy end", "self_heal_active_for_audit(audit, today)", "Master gate: subpolicy end + not fully complete — manual and AI."),
    ("Branch", "is_manual_audit(audit) / is_ai_audit(audit)", "AuditType != 'A' vs == 'A' — routes to Part B–F or AI-B–F."),
    ("Links", "audit_self_heal_action_url(audit)", "Manual → /audit/{id}/tasks; AI → /ai-audit/{id}/upload."),
    ("Evidence", "should_remind_missing_compliance_evidence", "True when audit_findings.Evidence empty, audit in progress, dedup not sent today."),
    ("Evidence", "_notify_auditor_missing_evidence", "Sends reminder to audit.Auditor; link /audit/{id}/tasks (TaskView upload compliance evidence)."),
    ("Evidence", "_notify_auditor_missing_audit_evidence", "Same for audit.Evidence empty (Upload Audit Evidence section)."),
    ("Frequency (manual)", "should_remind_frequency_today", "(today - AssignedDate) % audit.Frequency == 0; manual only."),
    ("Frequency (manual)", "_notify_auditor_frequency", "Perform audit per assign frequency; TaskView link."),
    ("AI evidence", "ai_audit_has_documents(audit_id)", "True if uploaded docs / ai_audit_data exist for audit."),
    ("AI evidence", "should_remind_ai_missing_evidence", "Uses audit.EvidenceReminderDays from Assign Audit (AI)."),
    ("AI evidence", "_notify_auditor_ai_evidence", "Upload page link; merge check_ai_audit_evidence logic here."),
    ("AI schedule", "should_remind_ai_schedule_today", "Uses AIAuditSchedule next_run / cron; not audit.Frequency."),
    ("AI schedule", "_notify_auditor_ai_schedule", "Remind before scheduled AI run or to configure schedule."),
    ("AI work", "ai_audit_work_complete", "Docs processed + results; stops due reminders for AI."),
    ("AI optional", "should_remind_ai_run_or_report", "Run check or download report pending on upload page."),
    ("Audit due", "audit_work_complete", "Stops due/overdue when all findings done + evidence present or Status Completed."),
    ("Audit due", "should_remind_audit_due_today", "Picks type audit_due_soon (4d before), audit_due_today, or audit_overdue."),
    ("Audit due", "_notify_auditor_due / _notify_auditor_overdue", "Notifies auditor around audit.DueDate."),
    ("Audit due", "overdue_reminder_days", "Counts calendar days audit_overdue sent after DueDate; triggers Part G when >= 3."),
    ("Review", "send_audit_for_review (PATCH)", "Must set ReviewStartDate=NOW when Status→Under review; today only sets Status."),
    ("Review", "review_due_date", "ReviewStartDate.date() + 5 days (COMPLIANCE_AUDIT_REVIEW_SLA_DAYS); not ReviewDate."),
    ("Review", "review_complete", "Stops review reminders when ReviewStatus Accept/Reject or ReviewDate set."),
    ("Review", "should_remind_review_today", "review_due_soon / review_due_today / review_overdue for reviewer."),
    ("Review", "_notify_reviewer_due / _notify_reviewer_overdue", "Notifies audit.reviewer; link /reviewer/task/{id}."),
    ("Escalation", "should_escalate", "True after 3+ overdue reminder days (auditor or reviewer path) per product rule."),
    ("Escalation", "_compliance_manager_user_ids", "Lists active RBAC Compliance Manager users (tenant); mirrors _policy_manager_user_ids."),
    ("Escalation", "_is_compliance_manager", "Guards manager-only list/assign APIs."),
    ("Escalation", "_ensure_pending_escalation", "Creates ComplianceAuditSelfHealEscalation pending_assignment if missing."),
    ("Escalation", "_notify_compliance_managers_escalation", "Email + in-app to managers; dashboard link ?auditEscalations=1."),
    ("Escalation", "list_pending_compliance_audit_escalations", "GET API: manager dashboard list."),
    ("Escalation", "assign_compliance_audit_escalation", "POST API: updates audit.auditor or audit.reviewer; notifies assignee."),
    ("Escalation", "_notify_escalation_assignee", "First reminder to newly assigned auditor/reviewer after manager assign."),
    ("Dedup", "ComplianceAuditReminderSent", "One row per audit + reminder_type + calendar day; prevents duplicate spam."),
    ("Dedup", "ComplianceAuditSelfHealEscalation", "Tracks pending vs assigned escalation and who manager picked."),
    ("Reuse", "NotificationService.send_multi_channel_notification", "Same email/channel pipeline as policy self-heal and rest of GRC."),
    ("Reuse", "_insert_db_notification / _append_memory_notification", "Copy pattern from policy_self_healing.py for DB + in-app list."),
    ("Reuse", "get_user_email_from_id", "Resolve recipient email for auditor/reviewer/manager."),
    ("Reuse", "RBACUtils.get_user_id_from_request", "JWT user on manager assign APIs."),
    ("Reuse", "get_tenant_id_from_request / tenant_filter", "Multi-tenancy on all queries."),
    ("Reuse", "policy_should_remind_today (pattern only)", "Frequency uses same modulo-day logic from policy self-heal."),
    ("AI merge", "check_ai_audit_evidence (deprecate)", "Logic moves into should_remind_ai_missing_evidence — one alert path."),
    ("AI patch", "ai_audit_api.py ReviewStartDate", "When AI → Under review after all docs processed, set ReviewStartDate."),
    ("API", "POST /api/compliance-audit/self-healing/reminders/run/", "External scheduler trigger."),
    ("API", "GET /api/compliance-audit/self-healing/escalations/pending/", "Manager escalation queue."),
    ("API", "POST /api/compliance-audit/{audit_id}/self-healing/assign/", "Manager reassign auditor or reviewer."),
    ("Frontend", "SubPolicy UI (VV.vue / framework flows)", "Capture SubPolicy.EndDate only."),
    ("Frontend", "Compliance Manager escalation panel", "List pending + assign user (mirror PolicyDashboard renewal)."),
    ("Frontend", "Notifications.vue", "Deep links to TaskView, ReviewTaskView, escalation dashboard."),
    ("Frontend", "api.js COMPLIANCE_AUDIT_SELF_HEAL_*", "Frontend endpoints for panel and actions."),
]


def add_text_flow(doc: Document, flow_text: str) -> None:
    for line in flow_text.strip().split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
    doc.add_paragraph()


def add_functions_tables(doc: Document) -> None:
    doc.add_heading("Functions and why we use them", 1)
    doc.add_paragraph(
        "Each row: what to build or reuse, and why (aligned with policy_self_healing.py where noted)."
    )
    # Group by category
    categories: list[str] = []
    for cat, _, _ in FUNCTION_ROWS:
        if cat not in categories:
            categories.append(cat)
    for cat in categories:
        rows = [(fn, why) for c, fn, why in FUNCTION_ROWS if c == cat]
        doc.add_heading(cat, 2)
        table = doc.add_table(rows=1 + len(rows), cols=2)
        table.style = "Table Grid"
        h = table.rows[0].cells
        h[0].text = "Function"
        h[1].text = "Why we use it"
        for p in h[0].paragraphs + h[1].paragraphs:
            for r in p.runs:
                r.bold = True
        for i, (fn, why) in enumerate(rows):
            table.rows[i + 1].cells[0].text = fn
            table.rows[i + 1].cells[1].text = why
        doc.add_paragraph()


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.add_heading("Missed compliance & audit task self-healing", 0)
    doc.add_paragraph(
        "Section 1: full workflow (top to bottom). Section 2: functions in tables with why."
    )
    doc.add_heading("Workflow", 1)
    add_text_flow(doc, FLOW_COMPLETE)
    add_functions_tables(doc)

    doc.save(OUT_PATH)
    print(f"Written: {OUT_PATH}")


if __name__ == "__main__":
    build()
