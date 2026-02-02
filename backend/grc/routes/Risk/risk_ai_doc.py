"""
AI-Powered Risk Document Ingestion (Complete)
- Reads PDF/DOCX/XLSX/TXT
- Extracts risk data using OpenAI (GPT models)
- Fills missing fields with focused prompts
- Returns normalized, DB-ready JSON for `risk` table
"""

import os
import re
import json
import math
import time
import tempfile
from datetime import date, datetime
from typing import Any, Optional

import requests
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser

# RBAC imports
from ...rbac.decorators import rbac_required

# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import Risk  # , Users  # (Users not needed here but you can import if required)


# =========================
# OPENAI CONFIG
# =========================
# OpenAI API Configuration - Use Django settings
from django.conf import settings
OPENAI_API_KEY = getattr(settings, 'OPENAI_API_KEY', None)
OPENAI_API_URL = "https://api.openai.com/v1/chat/completions"
OPENAI_MODEL = getattr(settings, 'OPENAI_MODEL', 'gpt-3.5-turbo')

if not OPENAI_API_KEY:
    print("⚠️  WARNING: OPENAI_API_KEY not found in Django settings!")
    print("   Please set OPENAI_API_KEY in your .env file.")
else:
    print(f"🌐 OpenAI Configuration:")
    print(f"   API URL: {OPENAI_API_URL}")
    print(f"   Model: {OPENAI_MODEL}")
    print(f"   API Key: {'*' * (len(OPENAI_API_KEY) - 4) + OPENAI_API_KEY[-4:]}")

# Only the columns you want to fill (you said: do NOT fetch RiskId, ComplianceId, FrameworkId)
RISK_DB_FIELDS = [
    "RiskTitle",
    "Criticality",
    "PossibleDamage",
    "Category",
    "RiskType",
    "BusinessImpact",
    "RiskDescription",
    "RiskLikelihood",
    "RiskImpact",
    "RiskExposureRating",
    "RiskPriority",
    "RiskMitigation",
    "CreatedAt",
    "RiskMultiplierX",
    "RiskMultiplierY",
]

# Canonical choices / constraints to stabilize LLM outputs
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES    = ["Low", "Medium", "High", "Critical"]
CATEGORY_HINTS      = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]
RISKTYPE_HINTS      = ["Current", "Residual", "Inherent", "Emerging", "Accepted"]
DATE_FORMAT_HINT    = "YYYY-MM-DD (ISO)"

# Field-specific micro-prompts (used when a single field is missing/invalid)
# NOTE: RiskTitle is NEVER inferred by AI - it must always come from the document
# Enhanced prompts with examples and better guidance for improved accuracy
FIELD_PROMPTS = {
    "Criticality": f"""Return one of: {CRITICALITY_CHOICES}.
Examples:
- "Low": Minor operational issues, minimal business impact
- "Medium": Moderate impact, some business disruption
- "High": Significant impact, major business disruption, compliance issues
- "Critical": Severe impact, business-critical failure, regulatory violations
Base your decision on the severity described in the document.""",
    
    "PossibleDamage": """Describe concrete, specific damages that could occur.
Examples:
- "Data breach could result in loss of 10,000+ customer records, regulatory fines up to $2M, and reputation damage."
- "System downtime of 24+ hours could cause $500K revenue loss and breach SLA commitments."
- "Compliance violation could result in regulatory penalties, loss of certifications, and legal action."
Be specific about: data loss, downtime duration, financial impact, penalties, reputation damage.""",
    
    "Category": f"""Return one category from this list (best fit): {CATEGORY_HINTS}.
Examples:
- "Operational": Process failures, system outages, resource constraints
- "Financial": Revenue loss, cost overruns, budget issues
- "Compliance": Regulatory violations, audit failures, policy breaches
- "Information Security": Data breaches, cyber attacks, unauthorized access
- "Third-Party": Vendor failures, supplier issues, partner risks
- "Technical": Software bugs, infrastructure failures, technology obsolescence
Pick the category that best matches the risk described.""",
    
    "RiskType": f"""Return one of: {RISKTYPE_HINTS}.
- "Current": Risk that exists now in current state
- "Residual": Risk remaining after controls/mitigations applied
- "Inherent": Risk before any controls applied
- "Emerging": New risk that is developing or emerging
- "Accepted": Risk that has been acknowledged and accepted
Default to "Current" if unclear.""",
    
    "BusinessImpact": """Explain business impact in concrete business terms.
Examples:
- "Could result in 48-hour service outage affecting 50,000 users, violating SLA commitments and causing $1M revenue loss."
- "Regulatory non-compliance could lead to $5M fines, loss of operating license, and exclusion from government contracts."
- "Data breach could expose 100K customer records, resulting in $10M in fines, class-action lawsuits, and 30% customer churn."
Focus on: financial impact, operational disruption, compliance consequences, customer impact.""",
    
    "RiskDescription": """Write a precise, technical description (1–3 sentences) of how/why the risk arises.
Examples:
- "Insufficient access controls on customer database allow unauthorized users to view sensitive PII due to missing role-based permissions."
- "Legacy payment system lacks encryption, exposing credit card data during transmission to potential interception."
- "Third-party vendor has weak security practices, creating supply chain vulnerability for data exfiltration."
Be specific about: root cause, vulnerability, threat vector, affected systems.""",
    
    "RiskLikelihood": """Return an integer 1–10.
Scale:
- 1-2: Rare/Unlikely (less than 5% chance in next year)
- 3-4: Possible (5-25% chance)
- 5-6: Likely (25-50% chance)
- 7-8: Very Likely (50-75% chance)
- 9-10: Almost Certain (75%+ chance)
Consider: historical frequency, current controls, threat landscape, industry trends.""",
    
    "RiskImpact": """Return an integer 1–10.
Scale:
- 1-2: Negligible (minor inconvenience, easily recoverable)
- 3-4: Low (some disruption, recoverable within days)
- 5-6: Moderate (significant disruption, weeks to recover)
- 7-8: High (major disruption, months to recover, significant financial loss)
- 9-10: Catastrophic (business-critical failure, permanent damage, bankruptcy risk)
Consider: financial loss, operational impact, reputation damage, compliance consequences.""",
    
    "RiskExposureRating": """Return a float (0–100). 
Formula: RiskExposureRating = RiskLikelihood × RiskImpact × 10
Examples:
- Likelihood=3, Impact=4 → Exposure = 12.0
- Likelihood=7, Impact=8 → Exposure = 56.0
- Likelihood=9, Impact=9 → Exposure = 81.0
If Likelihood and Impact are provided, calculate using the formula. Otherwise, estimate based on document context.""",
    
    "RiskPriority": f"""Return one of: {PRIORITY_CHOICES}.
Priority should align with Exposure Rating and Criticality:
- Exposure 0-20 OR Criticality=Low → "Low"
- Exposure 21-40 OR Criticality=Medium → "Medium"  
- Exposure 41-60 OR Criticality=High → "High"
- Exposure 61-100 OR Criticality=Critical → "Critical"
Use the HIGHER priority if Exposure and Criticality suggest different levels.""",
    
    "RiskMitigation": """Return 2–4 actionable, specific mitigation steps.
Examples:
- "Implement multi-factor authentication for all admin accounts. Deploy intrusion detection system. Conduct quarterly security audits. Establish incident response plan."
- "Encrypt all data at rest and in transit. Implement role-based access controls. Regular penetration testing. Employee security training program."
Format as: numbered list or comma-separated steps. Be specific and actionable.""",
    
    "CreatedAt": f"""Return a plausible assessment date in {DATE_FORMAT_HINT}.
Look for dates in the document (assessment date, review date, document date).
If no date found, use today's date: {date.today().isoformat()}""",
    
    "RiskMultiplierX": """Return a float in 0.1–1.5 reflecting organizational impact weighting factor X.
- 0.1-0.3: Low organizational impact (affects single department)
- 0.4-0.7: Medium organizational impact (affects multiple departments)
- 0.8-1.0: High organizational impact (affects entire organization)
- 1.1-1.5: Critical organizational impact (affects multiple business units, strategic)
Default to 0.5 if unknown.""",
    
    "RiskMultiplierY": """Return a float in 0.1–1.5 reflecting organizational likelihood weighting factor Y.
- 0.1-0.3: Low likelihood in this organization (strong controls, rare occurrence)
- 0.4-0.7: Medium likelihood (moderate controls, occasional occurrence)
- 0.8-1.0: High likelihood (weak controls, frequent occurrence)
- 1.1-1.5: Very high likelihood (no controls, constant occurrence)
Default to 0.5 if unknown.""",
}

# Strict JSON schema block the LLM must follow
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON array. No markdown, no code blocks, no explanations.
Start with [ and end with ]. Use proper JSON syntax with double quotes.

Example structure (return array of risks like this):
[
  {{
    "RiskTitle": "Brief risk title here",
    "Criticality": "Medium",
    "PossibleDamage": "Description of damage",
    "Category": "Operational",
    "RiskType": "Current",
    "BusinessImpact": "Business impact description",
    "RiskDescription": "Risk description",
    "RiskLikelihood": 5,
    "RiskImpact": 6,
    "RiskExposureRating": 30.0,
    "RiskPriority": "Medium",
    "RiskMitigation": "Mitigation steps",
    "CreatedAt": "{date.today().isoformat()}",
    "RiskMultiplierX": 0.5,
    "RiskMultiplierY": 0.5,
    "_meta": {{
      "per_field": {{
        "RiskTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found in document"}},
        "Criticality": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from context"}}
      }}
    }}
  }}
]

Rules:
- Criticality must be: Low, Medium, High, or Critical
- RiskPriority must be: Low, Medium, High, or Critical
- RiskLikelihood and RiskImpact must be integers 1-10
- RiskExposureRating must be float 0-100
- Dates must be YYYY-MM-DD format
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON array, nothing else
"""

# =========================
# UTILITIES / VALIDATORS
# =========================
def _json_from_llm_text(text: str) -> Any:
    """Extract the first valid JSON array/object from the LLM response."""
    # Remove markdown code blocks if present
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    
    # Try to find JSON array or object
    m = re.search(r"(\[.*\]|\{.*\})", text, flags=re.S)
    block = m.group(1) if m else text.strip()
    
    # Clean up common JSON issues
    # Remove trailing commas before closing braces/brackets
    block = re.sub(r',(\s*[}\]])', r'\1', block)
    
    return json.loads(block)

def call_openai_json(prompt: str, retries: int = 3, timeout: int = 120) -> Any:
    """Call OpenAI API expecting JSON response."""
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY environment variable is not set")
    
    # OpenAI API format
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    }
    
    payload = {
        "model": OPENAI_MODEL,
        "messages": [
            {"role": "system", "content": "You are a GRC (Governance, Risk, and Compliance) analyst expert. Always return valid JSON responses as requested."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"} if "gpt-4" in OPENAI_MODEL or "gpt-3.5-turbo-1106" in OPENAI_MODEL else None
    }
    
    # Remove response_format if None (for older models)
    if payload["response_format"] is None:
        del payload["response_format"]
    
    print(f"🤖 Calling OpenAI API at {OPENAI_API_URL}")
    print(f"🤖 Model: {OPENAI_MODEL}")
    print(f"🤖 Prompt length: {len(prompt)} chars")
    
    for attempt in range(retries):
        print(f"🤖 Attempt {attempt + 1}/{retries}...")
        try:
            resp = requests.post(OPENAI_API_URL, json=payload, headers=headers, timeout=timeout)
            resp.raise_for_status()
            print(f"✅ OpenAI API responded with status {resp.status_code}")
            
            # OpenAI response format
            response_data = resp.json()
            raw = response_data.get("choices", [{}])[0].get("message", {}).get("content", "")
            print(f"📝 Response length: {len(raw)} chars")
            print(f"📝 First 200 chars: {raw[:200]}...")
            
            result = _json_from_llm_text(raw)
            print(f"✅ Successfully parsed JSON from OpenAI response")
            return result
            
        except json.JSONDecodeError as je:
            print(f"❌ JSON parsing error on attempt {attempt + 1}: {je}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            print(f"❌ All retries exhausted. Raw response: {raw[:500] if 'raw' in locals() else 'N/A'}...")
            raise RuntimeError(f"Failed to parse JSON from OpenAI response after {retries} attempts")
            
        except requests.exceptions.HTTPError as he:
            print(f"❌ HTTP error on attempt {attempt + 1}: {he}")
            if resp.status_code == 401:
                raise RuntimeError("OpenAI API authentication failed. Please check your OPENAI_API_KEY.")
            elif resp.status_code == 429:
                print(f"⚠️  Rate limit exceeded. Waiting 5 seconds...")
                if attempt < retries - 1:
                    time.sleep(5)
                    continue
            elif resp.status_code >= 500:
                print(f"⚠️  OpenAI server error. Retrying...")
                if attempt < retries - 1:
                    time.sleep(2)
                    continue
            raise RuntimeError(f"OpenAI API HTTP error: {he}")
            
        except requests.exceptions.ConnectionError as ce:
            print(f"❌ Connection error on attempt {attempt + 1}: {ce}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"Failed to connect to OpenAI API: {ce}")
            
        except requests.exceptions.Timeout as te:
            print(f"❌ Timeout error on attempt {attempt + 1}: {te}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 2 seconds...")
                time.sleep(2)
                continue
            raise RuntimeError(f"OpenAI API request timed out: {te}")
            
        except Exception as e:
            print(f"❌ Unexpected error on attempt {attempt + 1}: {type(e).__name__}: {e}")
            if attempt < retries - 1:
                print(f"⏳ Retrying in 1 second...")
                time.sleep(1)
                continue
            raise RuntimeError(f"Unexpected error calling OpenAI API: {e}")
    
    raise RuntimeError(f"Failed to get response from OpenAI API after {retries} attempts")

def clamp_int(v, lo, hi) -> Optional[int]:
    if v is None: return None
    try:
        return max(lo, min(hi, int(v)))
    except Exception:
        # try to find a number in string
        m = re.search(r"\b(\d{1,2})\b", str(v))
        if m:
            return max(lo, min(hi, int(m.group(1))))
        return None

def as_float_or_none(v) -> Optional[float]:
    if v is None: return None
    try:
        return float(v)
    except Exception:
        try:
            # extract first float from string
            m = re.search(r"[-+]?\d*\.?\d+", str(v))
            return float(m.group(0)) if m else None
        except Exception:
            return None

def as_date_or_none(s: str) -> Optional[str]:
    if not s: return None
    s = str(s).strip()
    return s if re.match(r"^\d{4}-\d{2}-\d{2}$", s) else None

def compute_exposure(lk, im) -> Optional[float]:
    if lk is None or im is None: return None
    try:
        return round(float(lk) * float(im), 2)
    except Exception:
        return None

def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None


# =========================
# FILE EXTRACTORS
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF (pdfplumber preferred; PyPDF2 fallback)."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        print(f"[PDF] Extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        print(f"[DOCX] Extraction error: {e}")
        return ""

def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        print(f"[XLSX] Extraction error: {e}")
        return ""

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""


# =========================
# DOCUMENT PARSING & DETECTION
# =========================
def detect_and_parse_risk_blocks(text: str) -> list[dict]:
    """
    Step 1: Detect how many risks are in the document by finding 'Risk X:' patterns.
    Step 2: For each risk, extract whatever fields are explicitly present.
    Returns a list of partially-filled risk dictionaries with only extracted fields.
    """
    print(f"🔍 Detecting risk blocks in document...")
    
    # Find all risk blocks using pattern "Risk X: Title"
    risk_pattern = r'Risk\s+(\d+):\s*([^\n]+)'
    risk_matches = list(re.finditer(risk_pattern, text, re.IGNORECASE))
    
    if not risk_matches:
        print(f"⚠️  No risks found with 'Risk X:' pattern")
        return []
    
    print(f"✅ Found {len(risk_matches)} risk(s) in document")
    
    risks = []
    for i, match in enumerate(risk_matches):
        risk_num = match.group(1)
        risk_title = match.group(2).strip()
        
        # Extract the text block for this risk (from this match to next risk or end of document)
        start_pos = match.end()
        if i + 1 < len(risk_matches):
            end_pos = risk_matches[i + 1].start()
        else:
            end_pos = len(text)
        
        risk_block = text[start_pos:end_pos]
        
        print(f"  📋 Risk {risk_num}: {risk_title[:60]}...")
        
        # Extract fields that are explicitly present in the block
        risk_data = {
            "RiskTitle": risk_title,  # Always extracted from document
            "_extracted_fields": ["RiskTitle"],  # Track what was extracted
            "_risk_block": risk_block  # Keep for AI context
        }
        
        # Field mapping: document field name -> DB field name
        field_mappings = {
            r'Description:\s*([^\n]+(?:\n(?!(?:Risk|Possible Damage|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'RiskDescription',
            r'Possible Damage:\s*([^\n]+(?:\n(?!(?:Risk|Description|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'PossibleDamage',
            r'Risk Priority:\s*([^\n]+)': 'RiskPriority',
            r'Status:\s*([^\n]+)': 'RiskExposureRating',
            r'Mitigation:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Status|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'RiskMitigation',
            r'Created At:\s*([^\n]+)': 'CreatedAt',
            r'Risk Type:\s*([^\n]+)': 'RiskType',
            r'Risk Likelihood:\s*([^\n]+)': 'RiskLikelihood',
            r'Risk Impact:\s*([^\n]+)': 'RiskImpact',
            r'Category:\s*([^\n]+)': 'Category',
            r'Criticality:\s*([^\n]+)': 'Criticality',
            r'Business Impact:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Status|Mitigation|Created At|Risk Type|Risk Likelihood|Risk Impact|—))[^\n]+)*)': 'BusinessImpact',
        }
        
        # Extract each field if present
        for pattern, field_name in field_mappings.items():
            match = re.search(pattern, risk_block, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if value:  # Only add if not empty
                    risk_data[field_name] = value
                    risk_data["_extracted_fields"].append(field_name)
                    print(f"    ✓ {field_name}: {value[:50]}...")
        
        risks.append(risk_data)
    
    return risks


# =========================
# AI EXTRACTION CORE
# =========================
def extract_smart_context(text: str, risk_block: str, field_name: str, current_record: dict, max_chars: int = 5000) -> str:
    """
    Extract the most relevant context for a specific field.
    Prioritizes: risk block > relevant keywords > surrounding context > full document.
    """
    # Start with the risk block (most relevant)
    context_parts = [risk_block] if risk_block else []
    
    # Add field-specific keywords to search for
    field_keywords = {
        "Criticality": ["critical", "severity", "serious", "urgent", "important", "high", "low", "medium"],
        "PossibleDamage": ["damage", "loss", "impact", "consequence", "harm", "penalty", "fine", "breach"],
        "Category": ["category", "type", "classification", "operational", "financial", "compliance", "security"],
        "RiskType": ["current", "residual", "inherent", "emerging", "accepted", "existing", "remaining"],
        "BusinessImpact": ["business", "revenue", "sla", "customer", "service", "operation", "financial"],
        "RiskDescription": ["description", "risk", "threat", "vulnerability", "issue", "problem"],
        "RiskLikelihood": ["likelihood", "probability", "chance", "frequency", "occurrence", "rate"],
        "RiskImpact": ["impact", "severity", "consequence", "effect", "damage", "loss"],
        "RiskExposureRating": ["exposure", "rating", "score", "risk level", "risk rating"],
        "RiskPriority": ["priority", "urgent", "important", "critical", "high priority"],
        "RiskMitigation": ["mitigation", "control", "prevent", "reduce", "address", "remedy", "solution"],
        "CreatedAt": ["date", "created", "assessed", "reviewed", "evaluated", "timestamp"],
    }
    
    keywords = field_keywords.get(field_name, [])
    
    # If we have keywords, find relevant sentences
    if keywords and len(context_parts[0]) < max_chars if context_parts else True:
        sentences = text.split('.')
        relevant_sentences = []
        for sent in sentences:
            sent_lower = sent.lower()
            if any(kw in sent_lower for kw in keywords):
                relevant_sentences.append(sent.strip())
        
        if relevant_sentences:
            keyword_context = '. '.join(relevant_sentences[:10])  # Top 10 relevant sentences
            if keyword_context not in context_parts[0] if context_parts else True:
                context_parts.append(keyword_context)
    
    # Combine and limit size
    combined = ' '.join(context_parts) if context_parts else text
    if len(combined) > max_chars:
        # Prioritize risk block, then truncate
        if risk_block and len(risk_block) < max_chars:
            remaining = max_chars - len(risk_block)
            combined = risk_block + ' ' + text[:remaining]
        else:
            combined = combined[:max_chars]
    
    return combined


def validate_cross_field_consistency(risk_record: dict) -> dict:
    """
    Validate and correct field relationships for consistency.
    Returns a dict of corrections/suggestions.
    """
    corrections = {}
    
    likelihood = risk_record.get("RiskLikelihood")
    impact = risk_record.get("RiskImpact")
    exposure = risk_record.get("RiskExposureRating")
    criticality = risk_record.get("Criticality")
    priority = risk_record.get("RiskPriority")
    
    # Validate RiskExposureRating matches Likelihood × Impact
    if likelihood and impact and exposure:
        expected_exposure = compute_exposure(likelihood, impact)
        if expected_exposure and abs(exposure - expected_exposure) > 5.0:
            corrections["RiskExposureRating"] = expected_exposure
            print(f"  ⚠️  Corrected RiskExposureRating: {exposure} → {expected_exposure} (Likelihood×Impact)")
    
    # Validate RiskPriority aligns with Exposure and Criticality
    if exposure is not None and priority:
        exposure_priority_map = {
            (0, 20): "Low",
            (21, 40): "Medium",
            (41, 60): "High",
            (61, 100): "Critical"
        }
        
        suggested_priority = None
        for (low, high), pri in exposure_priority_map.items():
            if low <= exposure <= high:
                suggested_priority = pri
                break
        
        # Also consider criticality
        criticality_priority_map = {
            "Low": "Low",
            "Medium": "Medium",
            "High": "High",
            "Critical": "Critical"
        }
        criticality_suggested = criticality_priority_map.get(criticality, "Medium")
        
        # Use the higher priority
        if suggested_priority and criticality_suggested:
            priority_order = ["Low", "Medium", "High", "Critical"]
            if priority_order.index(criticality_suggested) > priority_order.index(suggested_priority):
                suggested_priority = criticality_suggested
        
        if suggested_priority and suggested_priority != priority:
            corrections["RiskPriority"] = suggested_priority
            print(f"  ⚠️  Suggested RiskPriority correction: {priority} → {suggested_priority} (based on Exposure={exposure:.1f}, Criticality={criticality})")
    
    # Validate Criticality aligns with Impact
    if impact and criticality:
        impact_criticality_map = {
            (1, 3): "Low",
            (4, 6): "Medium",
            (7, 8): "High",
            (9, 10): "Critical"
        }
        suggested_criticality = None
        for (low, high), crit in impact_criticality_map.items():
            if low <= impact <= high:
                suggested_criticality = crit
                break
        
        if suggested_criticality and suggested_criticality != criticality:
            # Only suggest if significantly different
            criticality_order = ["Low", "Medium", "High", "Critical"]
            current_idx = criticality_order.index(criticality) if criticality in criticality_order else 1
            suggested_idx = criticality_order.index(suggested_criticality) if suggested_criticality in criticality_order else 1
            if abs(current_idx - suggested_idx) >= 2:  # At least 2 levels different
                corrections["Criticality"] = suggested_criticality
                print(f"  ⚠️  Suggested Criticality correction: {criticality} → {suggested_criticality} (based on Impact={impact})")
    
    return corrections


def infer_single_field(field_name: str, current_record: dict, document_context: str, risk_block: str = "") -> tuple[Any, dict]:
    """
    Enhanced focused prompt for ONE field using AI with better context extraction.
    Returns: (value, metadata_dict)
    """
    print(f"🤖 AI PREDICTING FIELD: {field_name}")
    
    # Extract smart context (prioritizes relevant parts)
    smart_context = extract_smart_context(document_context, risk_block, field_name, current_record, max_chars=5000)
    
    guidance = FIELD_PROMPTS.get(field_name, "Return a concise, professional value.")
    
    # Build context about related fields that might help
    related_fields_info = {}
    if field_name == "RiskPriority":
        if current_record.get("RiskLikelihood") and current_record.get("RiskImpact"):
            related_fields_info["RiskLikelihood"] = current_record.get("RiskLikelihood")
            related_fields_info["RiskImpact"] = current_record.get("RiskImpact")
        if current_record.get("RiskExposureRating"):
            related_fields_info["RiskExposureRating"] = current_record.get("RiskExposureRating")
        if current_record.get("Criticality"):
            related_fields_info["Criticality"] = current_record.get("Criticality")
    elif field_name == "RiskExposureRating":
        if current_record.get("RiskLikelihood") and current_record.get("RiskImpact"):
            related_fields_info["RiskLikelihood"] = current_record.get("RiskLikelihood")
            related_fields_info["RiskImpact"] = current_record.get("RiskImpact")
    
    # Enhanced prompt with better structure
    mini = f"""
You are an expert GRC (Governance, Risk, and Compliance) analyst with deep domain knowledge.
Your task: Infer ONLY the field "{field_name}" for this risk record.

DOCUMENT CONTEXT (most relevant excerpts):
\"\"\"{smart_context}\"\"\"

CURRENT RISK RECORD (already extracted fields):
{json.dumps({k: current_record.get(k) for k in RISK_DB_FIELDS if current_record.get(k)}, indent=2)}

RELATED FIELDS (to help with inference):
{json.dumps(related_fields_info, indent=2) if related_fields_info else "None"}

FIELD-SPECIFIC GUIDANCE:
{guidance}

IMPORTANT RULES:
1. Base your inference STRICTLY on the document context provided above
2. Use related fields to ensure consistency (e.g., Priority should align with Likelihood×Impact)
3. If the document provides clear evidence, confidence should be 0.8-1.0
4. If you must infer from limited context, confidence should be 0.5-0.7
5. If you cannot infer reliably, return {{"value": null, "confidence": 0.0, "rationale": "Not enough information"}}
6. Always provide a clear rationale explaining your reasoning

Return ONLY valid JSON (no markdown, no code blocks):
{{"value": <scalar or string or null>, "confidence": 0.0-1.0, "rationale": "brief explanation of your reasoning"}}
"""
    try:
        print(f"   📤 Sending enhanced prompt to OpenAI for {field_name}...")
        print(f"   📊 Context length: {len(smart_context)} chars")
        out = call_openai_json(mini)
        v = out.get("value") if isinstance(out, dict) else None
        confidence = out.get("confidence", 0.7) if isinstance(out, dict) else 0.7
        rationale = out.get("rationale", "AI predicted based on document context") if isinstance(out, dict) else "AI predicted based on document context"
        print(f"   ✅ AI PREDICTED {field_name}: '{v}' (confidence: {confidence:.2f})")
        print(f"   💭 Rationale: {rationale[:100]}...")
    except Exception as e:
        print(f"   ❌ AI FAILED to predict {field_name}: {str(e)}")
        v = None
        confidence = 0.0
        rationale = f"AI prediction failed: {str(e)}"

    # Create metadata for this field
    metadata = {
        "source": "AI_GENERATED",
        "confidence": confidence,
        "rationale": rationale
    }

    # Normalize after inference
    if field_name in ("RiskLikelihood", "RiskImpact"):
        v = clamp_int(v, 1, 10) or 5
    elif field_name == "RiskExposureRating":
        lk = clamp_int(current_record.get("RiskLikelihood"), 1, 10)
        im = clamp_int(current_record.get("RiskImpact"), 1, 10)
        computed = compute_exposure(lk, im)
        v = computed if computed is not None else (as_float_or_none(v) or 0.0)
    elif field_name in ("RiskMultiplierX", "RiskMultiplierY"):
        vf = as_float_or_none(v)
        v = max(0.1, min(1.5, vf)) if vf is not None else 0.5
    elif field_name == "CreatedAt":
        v = as_date_or_none(v) or date.today().isoformat()
    elif field_name == "Criticality":
        v = normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
    elif field_name == "RiskPriority":
        v = normalize_choice(v, PRIORITY_CHOICES) or "Medium"
    elif field_name == "RiskType":
        v = normalize_choice(v, RISKTYPE_HINTS) or "Current"
    elif isinstance(v, str):
        v = v.strip() or None
    
    return v, metadata


def refine_low_confidence_field(field_name: str, current_record: dict, document_context: str, risk_block: str = "", initial_value: Any = None, initial_confidence: float = 0.0) -> tuple[Any, dict]:
    """
    Refine a field that had low confidence (< 0.6) using a different approach.
    Uses the full risk record context and asks for validation.
    """
    print(f"   🔄 REFINING low-confidence field: {field_name} (initial confidence: {initial_confidence:.2f})")
    
    smart_context = extract_smart_context(document_context, risk_block, field_name, current_record, max_chars=6000)
    
    guidance = FIELD_PROMPTS.get(field_name, "Return a concise, professional value.")
    
    refine_prompt = f"""
You are an expert GRC analyst. A previous AI attempt to infer "{field_name}" had low confidence ({initial_confidence:.2f}).
Please re-analyze with more careful attention to the document context.

DOCUMENT CONTEXT:
\"\"\"{smart_context}\"\"\"

COMPLETE RISK RECORD (all fields so far):
{json.dumps({k: current_record.get(k) for k in RISK_DB_FIELDS if current_record.get(k)}, indent=2)}

PREVIOUS ATTEMPT:
- Value: {initial_value}
- Confidence: {initial_confidence:.2f}

FIELD GUIDANCE:
{guidance}

TASK:
1. Carefully re-read the document context
2. Look for explicit mentions or strong implicit evidence
3. Consider consistency with other fields in the risk record
4. Provide a more confident prediction if evidence exists
5. If still uncertain, explain why in the rationale

Return JSON:
{{"value": <scalar or string or null>, "confidence": 0.0-1.0, "rationale": "detailed explanation"}}
"""
    
    try:
        out = call_openai_json(refine_prompt)
        v = out.get("value") if isinstance(out, dict) else initial_value
        confidence = out.get("confidence", initial_confidence) if isinstance(out, dict) else initial_confidence
        rationale = out.get("rationale", "Refined prediction") if isinstance(out, dict) else "Refined prediction"
        
        if confidence > initial_confidence:
            print(f"   ✅ REFINED {field_name}: '{v}' (confidence improved: {initial_confidence:.2f} → {confidence:.2f})")
        else:
            print(f"   ⚠️  Refinement did not improve confidence for {field_name}")
        
        metadata = {
            "source": "AI_GENERATED_REFINED",
            "confidence": confidence,
            "rationale": rationale,
            "initial_confidence": initial_confidence
        }
        
        # Normalize the refined value
        if field_name in ("RiskLikelihood", "RiskImpact"):
            v = clamp_int(v, 1, 10) or 5
        elif field_name == "RiskExposureRating":
            lk = clamp_int(current_record.get("RiskLikelihood"), 1, 10)
            im = clamp_int(current_record.get("RiskImpact"), 1, 10)
            computed = compute_exposure(lk, im)
            v = computed if computed is not None else (as_float_or_none(v) or 0.0)
        elif field_name in ("RiskMultiplierX", "RiskMultiplierY"):
            vf = as_float_or_none(v)
            v = max(0.1, min(1.5, vf)) if vf is not None else 0.5
        elif field_name == "CreatedAt":
            v = as_date_or_none(v) or date.today().isoformat()
        elif field_name == "Criticality":
            v = normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
        elif field_name == "RiskPriority":
            v = normalize_choice(v, PRIORITY_CHOICES) or "Medium"
        elif field_name == "RiskType":
            v = normalize_choice(v, RISKTYPE_HINTS) or "Current"
        elif isinstance(v, str):
            v = v.strip() or None
        
        return v, metadata
    except Exception as e:
        print(f"   ❌ Refinement failed for {field_name}: {str(e)}")
        return initial_value, {
            "source": "AI_GENERATED",
            "confidence": initial_confidence,
            "rationale": f"Refinement failed: {str(e)}"
        }


def fallback_risk_extraction(text: str) -> list[dict]:
    """
    Minimal pattern-based fallback when AI fails completely.
    Produces at least one record using generic defaults.
    """
    risks = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    risk_keywords = ["risk", "threat", "vulnerability", "hazard", "danger", "exposure"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in risk_keywords):
            if current:
                risks.append(current)
            count += 1
            current = {
                "RiskTitle": f"Risk {count}: {ln[:100]}",
                "RiskDescription": ln,
                "Category": "Operational",
                "Criticality": "Medium",
                "RiskPriority": "Medium",
                "RiskType": "Current",
                "BusinessImpact": "Potential business impact",
                "PossibleDamage": "Potential damages vary (reputation, compliance, downtime).",
                "RiskMitigation": "Mitigation steps to be defined after assessment.",
                "RiskLikelihood": 5,
                "RiskImpact": 5,
                "RiskExposureRating": 25.0,
                "CreatedAt": date.today().isoformat(),
                "RiskMultiplierX": 0.5,
                "RiskMultiplierY": 0.5,
            }

    if current:
        risks.append(current)

    if not risks:
        risks.append({
            "RiskTitle": "Document Risk Analysis",
            "RiskDescription": f"Automated risk derived from document (length={len(text)} chars).",
            "Category": "Operational",
            "Criticality": "Medium",
            "RiskPriority": "Medium",
            "RiskType": "Current",
            "BusinessImpact": "Potential business impact in operations/compliance.",
            "PossibleDamage": "Data loss, downtime, penalties, reputation.",
            "RiskMitigation": "Review & implement standard mitigations for identified weaknesses.",
            "RiskLikelihood": 5,
            "RiskImpact": 5,
            "RiskExposureRating": 25.0,
            "CreatedAt": date.today().isoformat(),
            "RiskMultiplierX": 0.5,
            "RiskMultiplierY": 0.5,
        })

    return risks


def parse_risks_from_text(text: str) -> list[dict]:
    """
    NEW APPROACH:
    1. First detect how many risks are in the document by finding 'Risk X:' patterns
    2. Extract whatever fields are explicitly present in each risk block
    3. Use AI to fill ONLY the missing fields (not to create risk titles)
    """
    print(f"📊 parse_risks_from_text() called with {len(text)} chars of text")
    
    # Step 1: Detect and parse risk blocks from document
    detected_risks = detect_and_parse_risk_blocks(text)
    
    if not detected_risks:
        print(f"⚠️  No structured risks found. Falling back to old AI extraction...")
        # Fallback to old behavior if no "Risk X:" pattern found
        return fallback_risk_extraction(text)
    
    print(f"✅ Detected {len(detected_risks)} risk(s), now processing each...")
    
    # Step 2: For each detected risk, normalize extracted fields and fill missing ones
    completed_risks = []
    for idx, risk_data in enumerate(detected_risks, 1):
        print(f"\n🔧 Processing Risk {idx}: {risk_data.get('RiskTitle', 'Unknown')[:50]}...")
        
        # Start with empty item for all DB fields
        item = {k: None for k in RISK_DB_FIELDS}
        
        # Copy extracted fields
        for field in RISK_DB_FIELDS:
            if field in risk_data and risk_data[field]:
                item[field] = risk_data[field]
        
        # Get the risk block for AI context
        risk_block = risk_data.get("_risk_block", "")
        extracted_fields = risk_data.get("_extracted_fields", [])
        
        print(f"  📝 Extracted fields: {', '.join(extracted_fields)}")
        
        # Initialize metadata structure
        if "_meta" not in item:
            item["_meta"] = {}
        if "per_field" not in item["_meta"]:
            item["_meta"]["per_field"] = {}
        
        # Mark extracted fields in metadata
        for field in extracted_fields:
            if field in item and item[field]:
                item["_meta"]["per_field"][field] = {
                    "source": "EXTRACTED",
                    "confidence": 0.95,
                    "rationale": "Found explicitly in document"
                }
        
        # Normalize extracted fields
        if item.get("RiskLikelihood"):
            item["RiskLikelihood"] = clamp_int(item["RiskLikelihood"], 1, 10)
        if item.get("RiskImpact"):
            item["RiskImpact"] = clamp_int(item["RiskImpact"], 1, 10)
        if item.get("RiskExposureRating"):
            item["RiskExposureRating"] = as_float_or_none(item["RiskExposureRating"])
        if item.get("Criticality"):
            item["Criticality"] = normalize_choice(item["Criticality"], CRITICALITY_CHOICES)
        if item.get("RiskPriority"):
            item["RiskPriority"] = normalize_choice(item["RiskPriority"], PRIORITY_CHOICES)
        if item.get("RiskType"):
            item["RiskType"] = normalize_choice(item["RiskType"], RISKTYPE_HINTS)
        if item.get("CreatedAt"):
            item["CreatedAt"] = as_date_or_none(item["CreatedAt"])
        if item.get("RiskMultiplierX"):
            item["RiskMultiplierX"] = as_float_or_none(item["RiskMultiplierX"])
        if item.get("RiskMultiplierY"):
            item["RiskMultiplierY"] = as_float_or_none(item["RiskMultiplierY"])
        
        # Step 3: Use AI to fill ONLY missing fields (NEVER RiskTitle - must be in document)
        missing_fields = [f for f in RISK_DB_FIELDS if item.get(f) in (None, "", []) and f != "RiskTitle"]
        if missing_fields:
            print(f"  🤖 Missing fields: {', '.join(missing_fields)}")
            print(f"  🤖 Using AI to infer missing fields with enhanced accuracy...")
            
            # First pass: Infer all missing fields
            low_confidence_fields = []
            for field in missing_fields:
                print(f"    🔍 Inferring {field}...")
                value, metadata = infer_single_field(field, item, text, risk_block)
                item[field] = value
                item["_meta"]["per_field"][field] = metadata
                
                # Track low confidence fields for refinement
                if metadata.get("confidence", 0.7) < 0.6:
                    low_confidence_fields.append((field, value, metadata.get("confidence", 0.0)))
                print(f"    🏷️  Marked {field} as AI_GENERATED (confidence: {metadata.get('confidence', 0.7):.2f})")
            
            # Second pass: Refine low-confidence fields
            if low_confidence_fields:
                print(f"  🔄 Refining {len(low_confidence_fields)} low-confidence field(s)...")
                for field, initial_value, initial_conf in low_confidence_fields:
                    print(f"    🔄 Refining {field} (confidence: {initial_conf:.2f})...")
                    refined_value, refined_metadata = refine_low_confidence_field(
                        field, item, text, risk_block, initial_value, initial_conf
                    )
                    item[field] = refined_value
                    item["_meta"]["per_field"][field] = refined_metadata
        else:
            print(f"  ✅ All fields extracted from document!")
        
        # Step 4: Cross-field validation and consistency checks
        print(f"  🔍 Performing cross-field validation...")
        corrections = validate_cross_field_consistency(item)
        if corrections:
            print(f"  ⚠️  Applying {len(corrections)} field correction(s) for consistency...")
            for field, corrected_value in corrections.items():
                old_value = item.get(field)
                item[field] = corrected_value
                # Update metadata to reflect correction
                if field in item["_meta"]["per_field"]:
                    item["_meta"]["per_field"][field]["source"] = "AI_GENERATED_CORRECTED"
                    item["_meta"]["per_field"][field]["original_value"] = old_value
                    item["_meta"]["per_field"][field]["rationale"] += f" (Corrected for consistency: {old_value} → {corrected_value})"
                print(f"    ✓ Corrected {field}: {old_value} → {corrected_value}")
        
        # Final normalization and defaults
        item["RiskLikelihood"] = item["RiskLikelihood"] or 5
        item["RiskImpact"] = item["RiskImpact"] or 5
        
        # Recompute exposure rating based on final Likelihood and Impact (ensures consistency)
        computed_exposure = compute_exposure(item["RiskLikelihood"], item["RiskImpact"])
        if computed_exposure is not None:
            # Use computed value if it's significantly different from current (more than 10% difference)
            current_exposure = item.get("RiskExposureRating")
            if current_exposure is None or abs(current_exposure - computed_exposure) > max(5.0, computed_exposure * 0.1):
                if current_exposure is not None:
                    print(f"  ⚠️  Recalculating RiskExposureRating: {current_exposure} → {computed_exposure} (Likelihood×Impact)")
                item["RiskExposureRating"] = computed_exposure
                if "RiskExposureRating" in item["_meta"]["per_field"]:
                    item["_meta"]["per_field"]["RiskExposureRating"]["source"] = "AI_GENERATED_CALCULATED"
                    item["_meta"]["per_field"]["RiskExposureRating"]["rationale"] = f"Calculated from RiskLikelihood ({item['RiskLikelihood']}) × RiskImpact ({item['RiskImpact']})"
        else:
            item["RiskExposureRating"] = item.get("RiskExposureRating") or 25.0
        
        item["RiskExposureRating"] = float(max(0.0, min(100.0, item["RiskExposureRating"])))
        item["RiskMultiplierX"] = max(0.1, min(1.5, item["RiskMultiplierX"] or 0.5))
        item["RiskMultiplierY"] = max(0.1, min(1.5, item["RiskMultiplierY"] or 0.5))
        item["CreatedAt"] = item["CreatedAt"] or date.today().isoformat()
        item["Criticality"] = item["Criticality"] or "Medium"
        item["RiskPriority"] = item["RiskPriority"] or "Medium"
        item["RiskType"] = item["RiskType"] or "Current"
        
        # RiskTitle must ALWAYS come from document - never generate it
        if not item.get("RiskTitle"):
            raise ValueError(f"RiskTitle is missing for risk {idx}. All risk titles must be present in the document as 'Risk X: Title'.")
        
        # Debug: Print metadata summary
        ai_fields = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "AI_GENERATED"]
        extracted = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "EXTRACTED"]
        
        print(f"  📊 Metadata Summary:")
        print(f"     🤖 AI Generated: {len(ai_fields)} fields - {ai_fields}")
        print(f"     📄 Extracted: {len(extracted)} fields - {extracted}")
        
        completed_risks.append(item)
        print(f"  ✅ Risk {idx} completed!")
    
    return completed_risks


# =========================
# DJANGO API ENDPOINTS
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
def upload_and_process_risk_document(request):
    """
    Upload a document and process it to extract COMPLETE risk data
    (fills missing fields via Ollama).
    Saves the uploaded document to MEDIA_ROOT/ai_uploads/risk/
    """

    print(f"📤 Upload request for risk document")
    print(f"📤 Request data: {request.POST}")
    print(f"📤 Request files: {request.FILES}")
    print(f"📤 User ID: {request.POST.get('user_id', 'unknown')}")

    # CORS preflight support
    if request.method == 'OPTIONS':
        response = HttpResponse()
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Methods'] = 'POST, OPTIONS'
        response['Access-Control-Allow-Headers'] = 'Content-Type, Authorization'
        response['Access-Control-Max-Age'] = '86400'
        return response

    try:
        if 'file' not in request.FILES:
            resp = JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            resp = JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)
            resp['Access-Control-Allow-Origin'] = '*'
            return resp

        # Create the ai_uploads/risk directory if it doesn't exist
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'risk')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save the file with timestamp to avoid conflicts
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        # Write the uploaded file to disk
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        print(f"✅ File saved to: {file_path}")

        try:
            # Step 1: Extract text from the saved file
            print(f"🔍 STEP 1: Starting text extraction from {ext} file...")
            text = extract_text_from_file(file_path, ext)
            
            if not text or len(text.strip()) < 50:
                print(f"❌ ERROR: Could not extract meaningful text. Length: {len(text) if text else 0}")
                resp = JsonResponse({'status': 'error', 'message': 'Could not extract meaningful text from document'}, status=400)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp

            print(f"✅ STEP 1 COMPLETE: Extracted {len(text)} characters from document")
            print(f"📄 First 200 chars: {text[:200]}...")
            
            # Step 2: Check OpenAI API key
            print(f"🔍 STEP 2: Checking OpenAI API configuration...")
            if not OPENAI_API_KEY:
                print(f"❌ ERROR: OPENAI_API_KEY is not set")
                resp = JsonResponse({
                    'status': 'error', 
                    'message': 'OPENAI_API_KEY environment variable is not set. Please configure your OpenAI API key.'
                }, status=503)
                resp['Access-Control-Allow-Origin'] = '*'
                return resp
            
            print(f"✅ STEP 2 COMPLETE: OpenAI API key is configured")
            
            # Step 3: Process with AI
            print(f"🤖 STEP 3: Calling OpenAI model '{OPENAI_MODEL}' to extract risks...")
            risks = parse_risks_from_text(text)
            
            print(f"✅ STEP 3 COMPLETE: AI extracted {len(risks)} risk(s) from document")
            for idx, risk in enumerate(risks, 1):
                print(f"  Risk {idx}: {risk.get('RiskTitle', 'Untitled')[:50]}...")

            resp = JsonResponse({
                'status': 'success',
                'message': f'Successfully extracted {len(risks)} risk(s)',
                'document_name': file_name,
                'saved_path': safe_filename,
                'extracted_text_length': len(text),
                'risks': risks
            })
            resp['Access-Control-Allow-Origin'] = '*'
            return resp
        except Exception as process_error:
            # Clean up the file if processing fails
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        resp = JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)
        resp['Access-Control-Allow-Origin'] = '*'
        return resp


@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
def save_extracted_risks(request):
    """
    Save extracted/reviewed risks to the database.
    (Does not set RiskId, ComplianceId, FrameworkId here.)
    """
    try:
        data = json.loads(request.body or "{}")
        risks_data = data.get('risks', [])
        if not risks_data:
            return JsonResponse({'status': 'error', 'message': 'No risks provided'}, status=400)

        saved = []
        errors = []

        for idx, r in enumerate(risks_data):
            try:
                kwargs = {
                    # not setting: RiskId, ComplianceId, FrameworkId
                    'RiskTitle': r.get('RiskTitle', f'Risk {idx+1}'),
                    'Criticality': r.get('Criticality', 'Medium'),
                    'PossibleDamage': r.get('PossibleDamage', ''),
                    'Category': r.get('Category', ''),
                    'RiskType': r.get('RiskType', 'Current'),
                    'BusinessImpact': r.get('BusinessImpact', ''),
                    'RiskDescription': r.get('RiskDescription', ''),
                    'RiskLikelihood': int(r.get('RiskLikelihood') or 5),
                    'RiskImpact': int(r.get('RiskImpact') or 5),
                    'RiskExposureRating': float(r.get('RiskExposureRating') or 25.0),
                    'RiskPriority': r.get('RiskPriority', 'Medium'),
                    'RiskMitigation': r.get('RiskMitigation', ''),
                    'CreatedAt': r.get('CreatedAt', date.today().isoformat()),
                    'RiskMultiplierX': float(r.get('RiskMultiplierX') or 0.5),
                    'RiskMultiplierY': float(r.get('RiskMultiplierY') or 0.5),
                }
                risk = Risk.objects.create(**kwargs)
                saved.append({'risk_id': getattr(risk, "RiskId", None), 'risk_title': risk.RiskTitle})
            except Exception as ex:
                errors.append({'risk_index': idx, 'title': r.get('RiskTitle'), 'error': str(ex)})

        resp = {
            'status': 'success',
            'message': f'Saved {len(saved)} risk(s)' + (f' with {len(errors)} error(s)' if errors else ''),
            'saved': saved
        }
        if errors:
            resp['errors'] = errors
        return JsonResponse(resp)

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON payload'}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error saving risks: {str(e)}'}, status=500)


@api_view(['GET'])
@permission_classes([AllowAny])
@rbac_required(required_permission='view_all_risk')
def test_openai_connection(request):
    """Quick check that OpenAI API responds."""
    try:
        if not OPENAI_API_KEY:
            return JsonResponse({
                'status': 'error',
                'message': 'OPENAI_API_KEY is not set',
                'model': OPENAI_MODEL,
                'api_url': OPENAI_API_URL
            }, status=500)
        
        out = call_openai_json('Return JSON: {"ok": true}')
        return JsonResponse({
            'status': 'success', 
            'openai_reply': out, 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error', 
            'message': f'OpenAI error: {e}', 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        }, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
def test_file_upload(request):
    """Simple test endpoint for file uploads."""
    try:
        if 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file found in request'}, status=400)
        f = request.FILES['file']
        return JsonResponse({
            'status': 'success',
            'message': 'File upload test successful',
            'file_name': f.name,
            'file_size': f.size,
            'content_type': f.content_type
        })
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': f'Test upload error: {str(e)}'}, status=500)
