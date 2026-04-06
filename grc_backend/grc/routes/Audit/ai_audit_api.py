"""
AI Audit API Endpoints
Clean, single implementation for AI audit document upload and processing
"""

import base64
import logging
import json
import os
import uuid
import hashlib
import re
from datetime import datetime
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

from django.http import JsonResponse, FileResponse, HttpResponse, HttpResponseRedirect
from django.views.decorators.csrf import csrf_protect as csrf_exempt
from django.utils.decorators import method_decorator
from django.utils import timezone
from django.views import View
from django.core.files.storage import default_storage
from django.conf import settings

from django.db import connection, transaction
from rest_framework import status
from rest_framework.decorators import api_view, permission_classes, authentication_classes, parser_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.authentication import SessionAuthentication, BasicAuthentication
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.views import APIView
from django.core.files.base import ContentFile
from rest_framework.response import Response
from ...rbac.permissions import AuditConductPermission, AuditReviewPermission
from ...rbac.decorators import audit_conduct_required
from ...utils.file_compression import decompress_if_needed
from ...routes.Global.s3_fucntions import create_direct_mysql_client


def _compute_ai_audit_report_hash(report_data: Any) -> str:
    """
    Compute a stable SHA-256 hash for the AI audit report.
    Uses a canonical JSON representation so the same content
    always produces the same hash.
    """
    try:
        payload = json.dumps(report_data, sort_keys=True, default=str)
    except TypeError:
        def _fallback(o):
            try:
                return str(o)
            except Exception:
                return 'unserializable'

        payload = json.dumps(report_data, sort_keys=True, default=_fallback)
    return hashlib.sha256(payload.encode('utf-8')).hexdigest()
from ...authentication import verify_jwt_token
from .audit_views import create_audit_version
from ...debug_utils import debug_print

# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)

# Cached check: does ai_audit_data have compliance_id column? (avoids repeated try/except)
_ai_audit_data_has_compliance_id = None

def _check_ai_audit_data_has_compliance_id():
    global _ai_audit_data_has_compliance_id
    if _ai_audit_data_has_compliance_id is not None:
        return _ai_audit_data_has_compliance_id
    try:
        with connection.cursor() as cur:
            cur.execute("SELECT compliance_id FROM ai_audit_data LIMIT 0")
        _ai_audit_data_has_compliance_id = True
    except Exception:
        _ai_audit_data_has_compliance_id = False
    return _ai_audit_data_has_compliance_id


# DRF Session auth variant that skips CSRF enforcement for API clients
class CsrfExemptSessionAuthentication(SessionAuthentication):
    def enforce_csrf(self, request):
        return

logger = logging.getLogger(__name__)


def _sanitize_filename_part_for_audit(value: str) -> str:
    """
    Local copy of the filename sanitiser used in Document Handling.
    Converts arbitrary text (like an audit title) into a safe code that can
    be used for folder codes and filename prefixes.
    """
    if not value:
        return "na"
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9]+", "_", value)
    value = value.strip("_")
    return value or "na"


def _link_evidence_to_document_handling_folder(audit_id: int, file_operation_id: int) -> None:
    """
    Ensure that an AI Audit evidence file also appears in Document Handling:
    - Create a company folder named after the audit title (if missing).
    - Create a default subfolder 'AI Audit Evidence' (if missing).
    - Link the given file_operations row into that subfolder, reusing its URL/key.
    """
    try:
        logger.info(f"📁 Linking evidence to Document Handling: audit_id={audit_id}, file_operation_id={file_operation_id}")
        # Get audit title and tenant
        with connection.cursor() as cur:
            cur.execute(
                "SELECT Title, TenantId FROM audit WHERE AuditId = %s",
                [int(audit_id) if str(audit_id).isdigit() else audit_id],
            )
            row = cur.fetchone()
        if not row:
            logger.warning(f"⚠️ Could not link evidence to Document Handling: audit {audit_id} not found")
            return

        audit_title = row[0] or f"Audit {audit_id}"
        audit_tenant_id = row[1]
        # Unique folder per audit: name/code include audit_id so each audit gets its own folder
        # Truncate to 100 chars to match CompanyFolder.code max_length so lookup and create match
        folder_code = _sanitize_filename_part_for_audit(f"{audit_title}_{audit_id}")[:100]
        folder_name = f"{audit_title} (Audit {audit_id})"

        # Import models lazily to avoid heavy imports at module import time
        from ...models import CompanyFolder, CompanySubfolder, CompanySubfolderDocument, FileOperations, Tenant

        with transaction.atomic():
            # Find or create company folder. Look up by code only (code is globally unique) to avoid
            # duplicate-key on create when a folder exists with same code but different/null tenant.
            folder = CompanyFolder.objects.filter(code__iexact=folder_code).first()
            if not folder:
                tenant_obj = None
                if audit_tenant_id:
                    tenant_obj = Tenant.objects.filter(tenant_id=audit_tenant_id).first()
                folder = CompanyFolder.objects.create(
                    name=folder_name,
                    code=folder_code,
                    description=f"AI Audit evidence for audit {audit_id}",
                    is_active=True,
                    tenant=tenant_obj,
                )
                logger.info(f"📁 Created Document Handling company folder: {folder_name} (code={folder_code})")
            else:
                # Ensure folder is active so it shows in Document Handling UI (only when it was inactive)
                if not folder.is_active:
                    folder.is_active = True
                    folder.save(update_fields=["is_active"])
                    logger.info(f"📁 Reactivated company folder for audit {audit_id} (IsActive=1)")

            sub_code = "ai_audit_evidence"
            subfolder = (
                CompanySubfolder.objects.filter(
                    company_folder=folder, code__iexact=sub_code
                ).first()
            )
            if not subfolder:
                subfolder = CompanySubfolder.objects.create(
                    company_folder=folder,
                    name="AI Audit Evidence",
                    code=sub_code,
                    is_active=True,
                )
            else:
                if not subfolder.is_active:
                    subfolder.is_active = True
                    subfolder.save(update_fields=["is_active"])

            file_op = FileOperations.objects.filter(id=file_operation_id).first()
            if not file_op:
                # Row may exist in another connection (e.g. S3 client). Ensure it exists in Django's DB so we can link.
                try:
                    with connection.cursor() as cur:
                        cur.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        fw_row = cur.fetchone()
                        framework_id = fw_row[0] if fw_row and fw_row[0] else None
                    if framework_id:
                        with connection.cursor() as cur:
                            cur.execute("""
                                INSERT INTO file_operations
                                (id, operation_type, user_id, file_name, original_name, status, module, FrameworkId, created_at, updated_at)
                                VALUES (%s, 'upload', 'system', '', '', 'completed', 'document_handling', %s, NOW(), NOW())
                            """, [int(file_operation_id), framework_id])
                        file_op = FileOperations.objects.filter(id=file_operation_id).first()
                        if file_op:
                            logger.info(f"📁 Created stub file_operations row id={file_operation_id} in Django DB for linking")
                except Exception as insert_err:
                    if 'Duplicate' in str(insert_err) or 'duplicate' in str(insert_err).lower() or '1062' in str(insert_err):
                        file_op = FileOperations.objects.filter(id=file_operation_id).first()
                    else:
                        logger.warning(f"⚠️ Could not link evidence: FileOperations id={file_operation_id} not found and stub insert failed: {insert_err}")
                        return
            if not file_op:
                # ORM may not see row (e.g. different connection). Create link via raw SQL if row exists.
                with connection.cursor() as cur:
                    cur.execute("SELECT id FROM file_operations WHERE id = %s", [int(file_operation_id)])
                    if cur.fetchone():
                        cur.execute("""
                            INSERT INTO company_subfolder_documents (CompanySubfolderId, FileOperationId, DocumentLink, S3Key)
                            VALUES (%s, %s, %s, %s)
                            ON DUPLICATE KEY UPDATE DocumentLink = VALUES(DocumentLink), S3Key = VALUES(S3Key)
                        """, [subfolder.subfolder_id, int(file_operation_id), "", ""])
                        logger.info(f"📁 Linked via raw SQL (ORM did not see file_operations row): file_operation_id={file_operation_id}")
                        logger.info(
                            f"📁 Linked FileOperations {file_operation_id} to Document Handling folder "
                            f"'{folder.name}' / '{subfolder.name}' for audit {audit_id}"
                        )
                        return
                logger.warning(
                    f"⚠️ Could not link evidence to Document Handling: FileOperations id={file_operation_id} not found in DB"
                )
                return

            logger.info(f"📁 Found file_op id={file_operation_id}, folder_id={folder.folder_id}, subfolder_id={getattr(subfolder, 'subfolder_id', getattr(subfolder, 'pk', None))}")

            # Ensure file_name follows the company-folder naming convention so that
            # Document Handling can count and filter it by prefix. Use same sanitize as list_company_folders.
            from grc.routes.DocumentHandling.document import sanitize_filename_part as doc_sanitize
            prefix_for_count = f"{doc_sanitize(folder_code)}_"
            original_name = file_op.file_name or file_op.original_name or file_op.stored_name or "document"
            base, ext = os.path.splitext(original_name)
            safe_base = _sanitize_filename_part_for_audit(base)
            prefixed_name = f"{prefix_for_count}{sub_code}_{safe_base}_{file_operation_id}{ext}"
            if file_op.file_name != prefixed_name:
                file_op.file_name = prefixed_name
                file_op.save(update_fields=["file_name"])

            # Use latest s3_url from file_operations so DocumentLink is always populated
            file_op.refresh_from_db()
            document_link = file_op.s3_url or ""
            s3_key = (file_op.s3_key or file_op.stored_name or "")[:1000]

            try:
                csd, created = CompanySubfolderDocument.objects.update_or_create(
                    company_subfolder=subfolder,
                    file_operation=file_op,
                    defaults={
                        "document_link": document_link,
                        "s3_key": s3_key,
                    },
                )
                logger.info(f"📁 CompanySubfolderDocument {'created' if created else 'updated'} id={csd.id}")
            except Exception as orm_err:
                # Fallback: insert link via raw SQL so subfolder document is always saved
                logger.warning(f"⚠️ CompanySubfolderDocument update_or_create failed: {orm_err}, trying raw INSERT")
                try:
                    with connection.cursor() as cur:
                        cur.execute("""
                            INSERT INTO company_subfolder_documents (CompanySubfolderId, FileOperationId, DocumentLink, S3Key)
                            VALUES (%s, %s, %s, %s)
                            ON DUPLICATE KEY UPDATE DocumentLink = VALUES(DocumentLink), S3Key = VALUES(S3Key)
                        """, [subfolder.subfolder_id, int(file_operation_id), document_link[:65535] if document_link else "", s3_key])
                except Exception as raw_err:
                    logger.warning(f"⚠️ Raw insert company_subfolder_documents failed: {raw_err}")
                    raise

            logger.info(
                f"📁 Linked FileOperations {file_operation_id} to Document Handling folder "
                f"'{folder.name}' / '{subfolder.name}' for audit {audit_id}"
            )
    except Exception as link_err:
        import traceback
        logger.warning(f"⚠️ Failed to link evidence to Document Handling folder: {link_err}\n{traceback.format_exc()}")


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def trigger_database_analysis(request, audit_id):
    """
    Manually trigger database analysis for a specific audit.
    This will analyze all database records and create evidence entries in ai_audit_data.
    
    POST /api/ai-audit/{audit_id}/trigger-database-analysis/
    MULTI-TENANCY: Only triggers analysis for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        logger.info(f"🔍 Manual database analysis triggered for audit {audit_id}")
        
        # Get audit details to find framework_id
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT FrameworkId, Title, Status
                FROM audit
                WHERE AuditId = %s AND TenantId = %s
            """, [int(audit_id), tenant_id])
            audit_row = cursor.fetchone()
            
            if not audit_row:
                return JsonResponse({
                    'success': False,
                    'error': f'Audit {audit_id} not found'
                }, status=404)
            
            framework_id = audit_row[0]
            audit_title = audit_row[1] or 'Unknown'
            audit_status = audit_row[2] or 'Unknown'
            
            if not framework_id:
                return JsonResponse({
                    'success': False,
                    'error': f'Audit {audit_id} has no FrameworkId assigned'
                }, status=400)
        
        logger.info(f"📋 Audit {audit_id}: {audit_title}, Framework: {framework_id}, Status: {audit_status}")
        
        # Import S3 functions class
        from ..Global.s3_fucntions import S3Functions
        s3_func = S3Functions()
        
        # Trigger database analysis in background thread
        import threading
        
        def run_database_analysis():
            try:
                logger.info(f"🚀 Starting database analysis for audit {audit_id} (framework {framework_id})")
                
                # Get audit details
                audit_details = s3_func._get_audit_details(int(audit_id))
                
                # Get ALL database data for this framework
                all_database_data = s3_func._get_all_database_data(framework_id)
                db_count = sum(len(v) for v in all_database_data.values())
                logger.info(f"💾 Found {db_count} database records across all tables for framework {framework_id}")
                
                if db_count == 0:
                    logger.warning(f"⚠️ No database records found for framework {framework_id}")
                    return
                
                # Analyze database data relevance
                db_results = s3_func._analyze_database_data_relevance(
                    audit_id=int(audit_id),
                    audit_details=audit_details,
                    database_data=all_database_data
                )
                
                logger.info(f"✅ Database analysis completed. Found {len(db_results)} tables with relevant records")
                
                # Update JSON index with database analyses
                for table_name, analyses in db_results.items():
                    for analysis in analyses:
                        record_id = analysis.get("record_id", 0)
                        if record_id:
                            s3_func._update_json_index_database(
                                framework_id=framework_id,
                                audit_id=int(audit_id),
                                table_name=table_name,
                                record_id=record_id,
                                analysis_result=analysis
                            )
                
                # Create ai_audit_data evidence from database results
                s3_func._create_ai_evidence_from_database_results(
                    framework_id=framework_id,
                    audit_id=int(audit_id),
                    db_results=db_results
                )
                
                logger.info(f"✅ Database analysis and evidence creation completed for audit {audit_id}")
                
            except Exception as e:
                logger.error(f"❌ Error in database analysis for audit {audit_id}: {str(e)}", exc_info=True)
        
        # Start analysis in background thread
        analysis_thread = threading.Thread(
            target=run_database_analysis,
            daemon=True,
            name=f"DatabaseAnalysis-{audit_id}"
        )
        analysis_thread.start()
        
        return JsonResponse({
            'success': True,
            'message': f'Database analysis started for audit {audit_id}',
            'audit_id': audit_id,
            'framework_id': framework_id,
            'audit_title': audit_title,
            'status': 'analysis_started'
        })
        
    except Exception as e:
        logger.error(f"❌ Error triggering database analysis: {str(e)}", exc_info=True)
        return JsonResponse({
            'success': False,
            'error': f'Failed to trigger database analysis: {str(e)}'
        }, status=500)


def call_ai_api(prompt, audit_id=None, document_id=None, model_type='compliance'):
    """
    Unified AI API call. Uses OpenAI when OPENAI_API_KEY is set (better, more consistent
    output for DOCUMENT IS ABOUT, DATA REQUIRED, EXPECTED DOCUMENT TYPE, WHAT IS NEEDED);
    otherwise falls back to Ollama.
    
    Args:
        prompt: The prompt to send to the AI
        audit_id: Audit ID for logging / determinism
        document_id: Document ID for logging / determinism
        model_type: Type of model call ('compliance', 'analysis', 'recommendations')
    
    Returns:
        str: AI response text
    """
    api_key = getattr(settings, 'OPENAI_API_KEY', '') or ''
    api_key = (api_key.strip() if isinstance(api_key, str) else '') or ''
    if api_key:
        logger.info(f"🤖 [AI] Using OpenAI for {model_type} (audit_id={audit_id}, document_id={document_id})")
        return _call_openai_api(prompt, audit_id, document_id, model_type)
    logger.warning(
        "🤖 [AI] OPENAI_API_KEY is not set or empty. Using Ollama. "
        "Set OPENAI_API_KEY in .env (in grc_backend folder) and restart the server to use OpenAI."
    )
    return _call_ollama_api(prompt, audit_id, document_id, model_type)


def _call_ollama_api(prompt, audit_id=None, document_id=None, model_type='compliance'):
    """
    Call Ollama chat API for AI processing (replaces OpenAI usage).
    Expects an Ollama server running (either local or configured in settings).
    """
    from django.conf import settings
    import requests
    import json
    
    # Resolve Ollama base URL and model from settings or fall back to sensible defaults
    base_url = getattr(settings, 'OLLAMA_BASE_URL', 'http://127.0.0.1:11434').rstrip('/')
    # Safety: if someone set OLLAMA_BASE_URL without scheme (e.g. "13.205.15.232:11434"),
    # prepend "http://" so requests can use it.
    if not base_url.startswith("http://") and not base_url.startswith("https://"):
        base_url = f"http://{base_url}"
    model = getattr(settings, 'OLLAMA_MODEL', 'llama3.1:8b')
    temperature = getattr(settings, 'OLLAMA_TEMPERATURE', 0.1)
    timeout = getattr(settings, 'OLLAMA_TIMEOUT', 120)
    
    logger.info(f"🤖 Calling Ollama at {base_url} with model: {model}, temperature: {temperature}")
    
    system_message = (
        "You are an expert GRC (Governance, Risk & Compliance) auditor with deep expertise in "
        "regulatory frameworks, compliance standards, and audit methodologies. "
        "You excel at conducting comprehensive compliance assessments, identifying gaps, "
        "evaluating risks, and providing actionable recommendations. "
        "Always provide accurate, detailed, and consistent analysis in valid JSON format when requested."
    )
    
    options = {"temperature": float(temperature)}
    # Optional: only cap length if explicitly set in settings
    max_tokens = getattr(settings, 'OLLAMA_MAX_TOKENS', None)
    if max_tokens is not None and max_tokens > 0:
        options["num_predict"] = int(max_tokens)
    
    payload = {
        "model": str(model).strip(),
        "messages": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ],
        "stream": False,
        "format": "json",  # Force JSON response format
        "options": options,
    }
    
    try:
        response = requests.post(
            f"{base_url}/api/chat",
            json=payload,
            timeout=timeout,
        )
        response.raise_for_status()
        
        data = response.json()
        # Standard Ollama response shape: {"message": {"role": "...", "content": "..."}, ...}
        if isinstance(data, dict) and "message" in data and "content" in data["message"]:
            content = data["message"]["content"]
        else:
            # Fallback: try generic 'content' key
            content = data.get("content", "")
        
        content = content or ""
        logger.info(f"🤖 TinyLlama/Ollama response length: {len(content)} characters")
        return content
    
    except requests.exceptions.Timeout:
        raise Exception(f"Ollama API timeout after {timeout} seconds")
    except requests.exceptions.RequestException as e:
        raise Exception(f"Ollama API request failed: {str(e)}")
    except Exception as e:
        raise Exception(f"Ollama API error: {str(e)}")


def _call_openai_api(prompt, audit_id=None, document_id=None, model_type='compliance'):
    """
    Call OpenAI chat completions for compliance analysis. Used when OPENAI_API_KEY is set.
    Produces more consistent, specific output for DOCUMENT IS ABOUT, DATA REQUIRED,
    EXPECTED DOCUMENT TYPE, and WHAT IS NEEDED (fewer placeholders/generic phrases).
    """
    import requests as req_lib
    api_key = (getattr(settings, 'OPENAI_API_KEY', '') or '').strip()
    if not api_key:
        raise Exception("OPENAI_API_KEY is not set")
    model = getattr(settings, 'OPENAI_MODEL', 'gpt-4o-mini')
    if not model:
        model = 'gpt-4o-mini'
    model = str(model).strip().strip('"').strip("'")
    timeout = getattr(settings, 'OLLAMA_TIMEOUT', 120)
    # Use fully deterministic output for compliance assessments
    temperature = 0.0
    system_message = (
        "You are an expert GRC (Governance, Risk & Compliance) auditor with deep expertise in "
        "regulatory frameworks, compliance standards, and audit methodologies. "
        "You excel at conducting comprehensive compliance assessments, identifying gaps, "
        "evaluating risks, and providing actionable recommendations. "
        "Always provide accurate, detailed, and consistent analysis in valid JSON format when requested. "
        "For DOCUMENT IS ABOUT, DATA REQUIRED FOR THIS AUDIT, EXPECTED DOCUMENT TYPE, and WHAT IS NEEDED, "
        "give specific, concrete content—never placeholders or generic phrases like 'specific data points and metrics'."
    )
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_message},
            {"role": "user", "content": prompt},
        ],
        "temperature": temperature,
        "response_format": {"type": "json_object"},
    }
    # Optional: only cap length if explicitly set in settings
    max_tokens = getattr(settings, 'OLLAMA_MAX_TOKENS', None)
    if max_tokens is not None and max_tokens > 0:
        payload["max_tokens"] = int(max_tokens)
    logger.info(f"🤖 Calling OpenAI model={model} for {model_type}")
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    try:
        response = req_lib.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=timeout,
        )
        response.raise_for_status()
        data = response.json()
        content = (data.get("choices") or [{}])[0].get("message", {}).get("content", "")
        content = content or ""
        logger.info(f"🤖 OpenAI response length: {len(content)} characters")
        return content
    except req_lib.exceptions.Timeout:
        raise Exception(f"OpenAI API timeout after {timeout} seconds")
    except req_lib.exceptions.RequestException as e:
        raise Exception(f"OpenAI API request failed: {str(e)}")
    except Exception as e:
        raise Exception(f"OpenAI API error: {str(e)}")


def generate_deterministic_seed(document_id, audit_id, requirement_hash=None):
    """
    Generate a deterministic seed based on document and audit IDs for consistent results.
    
    Args:
        document_id: The document ID
        audit_id: The audit ID  
        requirement_hash: Optional hash of requirement content
        
    Returns:
        int: Deterministic seed value
    """
    import hashlib
    
    # Create a deterministic seed based on document and audit IDs
    seed_string = f"{document_id}_{audit_id}"
    if requirement_hash:
        seed_string += f"_{requirement_hash}"
    
    # Generate a hash and convert to integer seed
    hash_obj = hashlib.md5(seed_string.encode())
    seed_int = int(hash_obj.hexdigest()[:8], 16)
    
    return seed_int


def cleanup_ai_audit_temp_files(document_id, doc_id=None):
    """
    Clean up temporary files created during AI audit processing.
    
    Args:
        document_id: The document ID from the API request
        doc_id: The internal document ID from database (optional)
    """
    try:
        import tempfile
        import glob
        temp_dir = tempfile.gettempdir()
        
        # Look for temporary files created during AI audit processing
        temp_file_patterns = []
        
        if document_id:
            temp_file_patterns.extend([
                f"ai_audit_{document_id}_*",
                f"*ai_audit_{document_id}*",
                f"*compliance_check_{document_id}*",
                f"*temp_audit_{document_id}*"
            ])
            
        if doc_id and doc_id != document_id:
            temp_file_patterns.extend([
                f"ai_audit_{doc_id}_*",
                f"*ai_audit_{doc_id}*"
            ])
        
        cleaned_files = []
        for pattern in temp_file_patterns:
            temp_files = glob.glob(os.path.join(temp_dir, pattern))
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                        cleaned_files.append(temp_file)
                        logger.info(f"🗑️ Cleaned up temporary file: {temp_file}")
                except Exception as temp_e:
                    logger.warning(f"⚠️ Could not delete temporary file {temp_file}: {temp_e}")
        
        return cleaned_files
        
    except Exception as e:
        logger.warning(f"⚠️ Error during temporary file cleanup: {e}")
        return []

@method_decorator(csrf_exempt, name='dispatch')
class AIAuditDocumentUploadView(APIView):
    """AI Audit Document Upload API - Single clean implementation"""
    
    parser_classes = [MultiPartParser, FormParser]
    authentication_classes = []  # Handle auth manually via JWT
    permission_classes = []  # Handle permissions manually
    
    def post(self, request, audit_id):
        """Upload document for AI audit processing"""
        try:
            # Log content type early to debug parser issues
            logger.info(f"📤 Request Content-Type: {request.content_type}")
            logger.info(f"📤 Request META CONTENT_TYPE: {request.META.get('CONTENT_TYPE', 'N/A')}")
            debug_print("=" * 80)
            debug_print("AI AUDIT UPLOAD ENDPOINT CALLED - NEW CODE VERSION")
            debug_print(f"Upload request for audit {audit_id}")
            debug_print("=" * 80)
            logger.info("🚀🚀🚀 AI AUDIT UPLOAD ENDPOINT CALLED - NEW CODE VERSION 🚀🚀🚀")
            logger.info(f"📤 Upload request for audit {audit_id}")
            logger.info(f"📤 Upload audit_id type: {type(audit_id)}")
            logger.info(f"📤 Upload audit_id value: '{audit_id}'")
            
            # Check authentication using JWT (like other endpoints)
            from ...rbac.utils import RBACUtils
            user_id = RBACUtils.get_user_id_from_request(request)
            
            if not user_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Authentication required'
                }, status=401)
            
            logger.info(f"📤 Authenticated user ID: {user_id}")
            
            # MULTI-TENANCY: Extract tenant_id from request
            tenant_id = get_tenant_id_from_request(request)
            if not tenant_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Tenant ID is required'
                }, status=400)
            
            logger.info(f"📤 Tenant ID: {tenant_id}")
            
            # CRITICAL: Validate that the audit exists in the correct table (audit)
            try:
                from django.db import connection
                with connection.cursor() as cursor:
                    cursor.execute("SELECT AuditId, AuditType, Status FROM audit WHERE AuditId = %s AND TenantId = %s", [audit_id, tenant_id])
                    audit_row = cursor.fetchone()
                    
                if audit_row:
                    audit_type = audit_row[1] if len(audit_row) > 1 else 'Unknown'
                    audit_status = (audit_row[2] or '').strip() if len(audit_row) > 2 else ''
                    if audit_status == 'Completed':
                        return JsonResponse({
                            'success': False,
                            'error': 'This audit is closed. No documents or evidence can be uploaded.'
                        }, status=403)
                    logger.info(f"✅ Found audit {audit_id} in audit table with type: {audit_type}")
                else:
                    logger.error(f"❌ Audit {audit_id} not found in audit table")
                    return JsonResponse({
                        'success': False,
                        'error': f'Audit {audit_id} not found. Please ensure the audit exists before uploading documents.'
                    }, status=404)
            except Exception as e:
                logger.error(f"❌ Error validating audit {audit_id}: {e}")
                return JsonResponse({
                    'success': False,
                    'error': f'Error validating audit: {e}'
                }, status=500)
            
            # Get optional mapping fields - support single or multiple mappings
            policy_id = request.POST.get('policy_id') or None
            subpolicy_id = request.POST.get('subpolicy_id') or None
            
            # Check for multiple mappings (JSON array)
            mappings_json = request.POST.get('mappings', None)
            mappings = []
            logger.info(f"📤 Received mappings_json: {mappings_json}")
            if mappings_json:
                try:
                    import json
                    mappings = json.loads(mappings_json)
                    if not isinstance(mappings, list):
                        mappings = []
                    logger.info(f"📤 Parsed {len(mappings)} mapping(s) from JSON")
                except Exception as e:
                    logger.warning(f"⚠️ Failed to parse mappings JSON: {e}")
                    mappings = []
            
            # If single mapping provided, convert to list format
            if not mappings and policy_id:
                mappings = [{'policy_id': policy_id, 'subpolicy_id': subpolicy_id}]
                logger.info(f"📤 Using single mapping: policy_id={policy_id}, subpolicy_id={subpolicy_id}")
            elif not mappings:
                # Default: no mapping
                mappings = [{'policy_id': None, 'subpolicy_id': None}]
                logger.info("📤 No mappings provided, using default (no mapping)")
            
            logger.info(f"📤 Final mappings count: {len(mappings)}")

            # Check if this is a metadata-only request (for already uploaded files)
            already_uploaded = request.POST.get('already_uploaded') == 'true'
            
            # Check if this is a file_operation_id request (reusing existing file from Document Handling)
            file_operation_id = request.POST.get('file_operation_id')
            
            # Replica check: if uploading a new file, see if the same document already exists in any folder (same tenant).
            # If so, reuse that file_operation instead of uploading again.
            if not file_operation_id and request.FILES:
                file_for_check = None
                for key in ['file', 'document', 'upload', 'file_upload']:
                    if key in request.FILES:
                        file_for_check = request.FILES[key]
                        break
                if file_for_check and getattr(file_for_check, 'name', None) and getattr(file_for_check, 'size', None) is not None:
                    original_name_check = file_for_check.name
                    file_size_check = file_for_check.size
                    try:
                        with connection.cursor() as cursor:
                            # Use correct column names: company_folders.FolderId (PK), company_subfolders.SubfolderId (PK)
                            cursor.execute("""
                                SELECT fo.id FROM file_operations fo
                                INNER JOIN company_subfolder_documents csd ON csd.FileOperationId = fo.id
                                INNER JOIN company_subfolders cs ON cs.SubfolderId = csd.CompanySubfolderId
                                INNER JOIN company_folders cf ON cf.FolderId = cs.CompanyFolderId
                                WHERE (cf.TenantId = %s OR cf.TenantId IS NULL) AND fo.status = 'completed'
                                  AND fo.file_size = %s
                                  AND (fo.original_name = %s OR (fo.original_name IS NULL AND fo.file_name = %s))
                                ORDER BY fo.created_at DESC
                                LIMIT 1
                            """, [tenant_id, file_size_check, original_name_check, original_name_check])
                            row = cursor.fetchone()
                            if row:
                                file_operation_id = str(row[0])
                                logger.info(
                                    f"📋 REPLICA=YES: same document already exists (original_name={original_name_check}, size={file_size_check}), "
                                    f"reusing file_operation_id={file_operation_id} instead of uploading again"
                                )
                            else:
                                logger.info(
                                    f"📋 REPLICA=NO: no existing document found (original_name={original_name_check}, size={file_size_check}), "
                                    f"uploading new file"
                                )
                    except Exception as replica_err:
                        logger.warning(f"⚠️ Replica check failed (continuing with upload): {replica_err}")
            
            # Initialize variables that might be used later
            s3_key = None
            stored_name = None
            aws_file_link = None
            doc_handling_operation_id = None  # When set, file is also in Document Handling (file_operations)
            all_relevant_documents_to_process = []  # List of (operation_id, matched_compliances) tuples - will be populated if processing via file_operation_id
            
            if file_operation_id:
                # Handle reuse of existing file from file_operations table
                logger.info(f"📤 Reusing file from file_operations ID: {file_operation_id}")
                try:
                    with connection.cursor() as cursor:
                        cursor.execute("""
                            SELECT file_name, original_name, s3_url, s3_key, file_size, 
                                   file_type, content_type, stored_name
                            FROM file_operations 
                            WHERE id = %s AND status = 'completed'
                        """, [int(file_operation_id)])
                        file_op = cursor.fetchone()
                        
                        if not file_op:
                            return JsonResponse({
                                'success': False,
                                'error': f'File operation {file_operation_id} not found or not completed'
                            }, status=404)
                        
                        # Extract file information
                        # IMPORTANT: Use stored_name if available (most consistent), otherwise file_name, then original_name
                        # This ensures document_name is consistent for grouping across multiple uploads of the same file
                        stored_name = file_op[7]  # stored_name
                        file_name_from_db = file_op[0]  # file_name
                        original_name = file_op[1]  # original_name
                        file_name = stored_name or file_name_from_db or original_name or 'unknown_file'
                        aws_file_link = file_op[2]  # s3_url
                        s3_key = file_op[3]  # s3_key
                        file_size = file_op[4] or 0  # file_size
                        file_type = file_op[5]  # file_type
                        content_type = file_op[6]  # content_type
                        
                        logger.info(f"📤 Found file: file_name={file_name} (from stored_name={stored_name}, file_name_db={file_name_from_db}, original_name={original_name}), s3_url: {aws_file_link}, s3_key: {s3_key}, file_size: {file_size}")

                        # Ensure this reused evidence file also appears in Document Handling
                        try:
                            _link_evidence_to_document_handling_folder(audit_id, int(file_operation_id))
                        except Exception as link_err:
                            logger.warning(f"⚠️ Could not link reused evidence to Document Handling folder: {link_err}")
                        
                        # CRITICAL: When uploading ANY document, process ALL relevant documents from JSON
                        # The JSON contains all documents that have been analyzed for relevance
                        # Documents with matched_compliances are the relevant ones that should be processed
                        # (Note: all_relevant_documents_to_process is initialized at the top level)
                        
                        # STEP 1: Read documents_analysis.json to find ALL relevant documents (those with matched_compliances)
                        try:
                            # Get framework_id first
                            cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                            framework_row = cursor.fetchone()
                            if framework_row and framework_row[0]:
                                framework_id_for_json = framework_row[0]
                                
                                # Construct path to documents_analysis.json
                                json_path = os.path.join(
                                    settings.MEDIA_ROOT,
                                    'audit_indexes',
                                    f'framework_{framework_id_for_json}',
                                    f'audit_{audit_id}',
                                    'documents_analysis.json'
                                )
                                
                                logger.info(f"📋 Looking for documents_analysis.json at: {json_path}")
                                
                                if os.path.exists(json_path):
                                    try:
                                        with open(json_path, 'r', encoding='utf-8') as f:
                                            json_data = json.load(f)
                                        
                                        documents_in_json = json_data.get('documents', [])
                                        logger.info(f"📋 Found {len(documents_in_json)} document(s) in documents_analysis.json - will process ALL documents with matched compliances")
                                        
                                        for doc_entry in documents_in_json:
                                            doc_operation_id = doc_entry.get('operation_id')
                                            matched_compliances_for_doc = doc_entry.get('matched_compliances', [])
                                            
                                            if matched_compliances_for_doc and doc_operation_id:
                                                # This document is relevant - add it to the processing list
                                                all_relevant_documents_to_process.append((doc_operation_id, matched_compliances_for_doc))
                                                logger.info(f"📋 Document operation_id={doc_operation_id} has {len(matched_compliances_for_doc)} matched compliances: {matched_compliances_for_doc} - will be processed")
                                            elif doc_operation_id:
                                                logger.info(f"📋 Document operation_id={doc_operation_id} has NO matched compliances - skipping")
                                        
                                        logger.info(f"✅✅✅ WILL PROCESS {len(all_relevant_documents_to_process)} relevant document(s) from JSON")
                                        
                                    except Exception as json_err:
                                        logger.warning(f"⚠️ Could not read/parse documents_analysis.json: {json_err}")
                                        import traceback
                                        logger.warning(f"⚠️ Error details: {traceback.format_exc()}")
                                else:
                                    logger.info(f"ℹ️ documents_analysis.json not found at {json_path}, will check database instead")
                            else:
                                logger.warning(f"⚠️ Could not get framework_id for audit {audit_id}")
                        except Exception as json_read_err:
                            logger.warning(f"⚠️ Error reading documents_analysis.json: {json_read_err}")
                        
                        # STEP 2: If no documents found in JSON, add the uploaded document to processing list
                        if not all_relevant_documents_to_process:
                            logger.info(f"📋 No relevant documents found in JSON - will process only the uploaded document (file_operation_id={file_operation_id})")
                            # Try to get matched compliances from database for the uploaded document
                            # IMPORTANT: Check by s3_key/stored_name (not just file_operation_id) to handle same file uploaded with different modules
                            try:
                                # First, try by file_operation_id (direct match)
                                cursor.execute("""
                                    SELECT matched_compliances, matched_policies, matched_subpolicies
                                    FROM document_audit_relevance
                                    WHERE file_operation_id = %s AND audit_id = %s
                                    LIMIT 1
                                """, [int(file_operation_id), int(audit_id) if str(audit_id).isdigit() else audit_id])
                                relevance_row = cursor.fetchone()
                                
                                # If not found by file_operation_id, try by s3_key/stored_name (same physical file, different module)
                                if not relevance_row or not relevance_row[0]:
                                    logger.info(f"📋 No relevance found by file_operation_id={file_operation_id}, checking by s3_key/stored_name...")
                                    # Get s3_key/stored_name from current file_operation
                                    file_key = s3_key or stored_name
                                    if file_key:
                                        # Find other file_operations with the same s3_key/stored_name
                                        if s3_key:
                                            cursor.execute("""
                                                SELECT id FROM file_operations 
                                                WHERE (s3_key = %s OR stored_name = %s) AND id != %s AND status = 'completed'
                                                ORDER BY created_at DESC
                                                LIMIT 1
                                            """, [s3_key, s3_key, int(file_operation_id)])
                                        else:
                                            cursor.execute("""
                                                SELECT id FROM file_operations 
                                                WHERE stored_name = %s AND id != %s AND status = 'completed'
                                                ORDER BY created_at DESC
                                                LIMIT 1
                                            """, [stored_name, int(file_operation_id)])
                                        
                                        other_operation = cursor.fetchone()
                                        if other_operation:
                                            other_op_id = other_operation[0]
                                            logger.info(f"📋 Found other file_operation_id={other_op_id} with same s3_key/stored_name, checking relevance...")
                                            # Check if that other operation has relevance for this audit
                                            cursor.execute("""
                                                SELECT matched_compliances, matched_policies, matched_subpolicies
                                                FROM document_audit_relevance
                                                WHERE file_operation_id = %s AND audit_id = %s
                                                LIMIT 1
                                            """, [other_op_id, int(audit_id) if str(audit_id).isdigit() else audit_id])
                                            relevance_row = cursor.fetchone()
                                            if relevance_row:
                                                logger.info(f"✅ Found relevance record for same file (different module upload) - file_operation_id={other_op_id}")
                                
                                if relevance_row and relevance_row[0]:
                                    try:
                                        matched_compliances_for_uploaded = json.loads(relevance_row[0]) if isinstance(relevance_row[0], str) else relevance_row[0]
                                        if matched_compliances_for_uploaded:
                                            all_relevant_documents_to_process.append((file_operation_id, matched_compliances_for_uploaded))
                                            logger.info(f"✅✅✅ MATCHED COMPLIANCES FOUND in database for file_operation_id={file_operation_id}, audit_id={audit_id}")
                                            logger.info(f"📋 Found {len(matched_compliances_for_uploaded)} matched compliances from database: {matched_compliances_for_uploaded}")
                                    except Exception as e:
                                        logger.warning(f"⚠️ Could not parse matched_compliances from relevance analysis: {e}")
                                        import traceback
                                        logger.warning(f"⚠️ Error details: {traceback.format_exc()}")
                                else:
                                    logger.warning(f"❌❌❌ NO MATCHED COMPLIANCES FOUND in document_audit_relevance for file_operation_id={file_operation_id}, audit_id={audit_id}")
                                    logger.warning(f"❌ relevance_row: {relevance_row}")
                            except Exception as e:
                                logger.warning(f"❌❌❌ ERROR looking up matched compliances for file_operation_id {file_operation_id} and audit {audit_id}: {e}")
                                import traceback
                                logger.warning(f"❌ Error details: {traceback.format_exc()}")
                        
                        # STEP 3: Only add uploaded document if it has matched compliances (relevant documents only)
                        uploaded_doc_in_list = any(op_id == file_operation_id for op_id, _ in all_relevant_documents_to_process)
                        if not uploaded_doc_in_list:
                            # Check if uploaded document has matched compliances from database
                            try:
                                cursor.execute("""
                                    SELECT matched_compliances
                                    FROM document_audit_relevance
                                    WHERE file_operation_id = %s AND audit_id = %s
                                    LIMIT 1
                                """, [int(file_operation_id), int(audit_id) if str(audit_id).isdigit() else audit_id])
                                relevance_row = cursor.fetchone()
                                if relevance_row and relevance_row[0]:
                                    try:
                                        matched_compliances_for_uploaded = json.loads(relevance_row[0]) if isinstance(relevance_row[0], str) else relevance_row[0]
                                        if matched_compliances_for_uploaded and len(matched_compliances_for_uploaded) > 0:
                                            all_relevant_documents_to_process.append((file_operation_id, matched_compliances_for_uploaded))
                                            logger.info(f"✅ Added uploaded document (file_operation_id={file_operation_id}) with {len(matched_compliances_for_uploaded)} matched compliances")
                                        else:
                                            logger.info(f"⏭️ Skipping uploaded document (file_operation_id={file_operation_id}) - no matched compliances (not relevant)")
                                    except Exception as e:
                                        logger.warning(f"⚠️ Could not parse matched_compliances: {e}")
                                        logger.info(f"⏭️ Skipping uploaded document (file_operation_id={file_operation_id}) - invalid relevance data")
                                else:
                                    logger.info(f"⏭️ Skipping uploaded document (file_operation_id={file_operation_id}) - no relevance found (not relevant)")
                            except Exception as e:
                                logger.warning(f"⚠️ Error checking relevance for uploaded document: {e}")
                                logger.info(f"⏭️ Skipping uploaded document (file_operation_id={file_operation_id}) - cannot verify relevance")
                        
                        # STEP 4: Log what documents we found
                        logger.info(f"🔄 Found {len(all_relevant_documents_to_process)} relevant document(s) to process from JSON:")
                        for op_id, comps in all_relevant_documents_to_process:
                            logger.info(f"  - operation_id={op_id} with {len(comps)} matched compliances: {comps}")
                        
                        # Skip single-document processing setup - we'll process all documents after getting user_id and framework_id
                        # Set mappings to empty so we don't process the uploaded document separately
                        mappings = []  # Clear mappings - all documents will be processed in the multi-document loop
                        matched_compliances_from_analysis = []  # Not needed - we have all_relevant_documents_to_process
                        
                        # Use s3_key or stored_name as document path (for reference, but won't be used if we process all documents)
                        document_path = s3_key or stored_name or f"file_operations/{file_operation_id}/{file_name}"
                        upload_type = 'evidence'
                        already_uploaded = True  # Treat as already uploaded
                        
                except Exception as e:
                    logger.error(f"❌ Error fetching file_operation {file_operation_id}: {e}")
                    return JsonResponse({
                        'success': False,
                        'error': f'Error fetching file from file_operations: {str(e)}'
                    }, status=500)
            
            elif already_uploaded:
                # Handle metadata-only document creation for already uploaded files
                file_name = request.POST.get('file_name')
                file_size = int(request.POST.get('file_size', 0))
                aws_file_link = request.POST.get('aws_file_link')
                s3_key = request.POST.get('s3_key')
                stored_name = request.POST.get('stored_name')
                file_id = request.POST.get('file_id')
                upload_type = request.POST.get('upload_type', 'evidence')
                
                logger.info(f"📤 S3 file metadata: s3_key={s3_key}, stored_name={stored_name}, aws_file_link={aws_file_link}")
                
                if not file_name:
                    return JsonResponse({
                        'success': False,
                        'error': 'File name is required for metadata-only upload'
                    }, status=400)
                
                # Use the stored_name or file_id as the document path for already uploaded files
                document_path = stored_name or f"uploaded/{file_id}" if file_id else f"uploaded/{file_name}"
                
                logger.info(f"📤 Creating metadata record for already uploaded file: {file_name}")
            else:
                # Get uploaded file for new uploads
                # Debug: Log request details
                logger.info(f"📤 Request method: {request.method}")
                logger.info(f"📤 Content-Type: {request.content_type}")
                logger.info(f"📤 Has request.data: {hasattr(request, 'data')}")
                logger.info(f"📤 request.data keys: {list(request.data.keys()) if hasattr(request, 'data') else 'N/A'}")
                logger.info(f"📤 request.FILES keys: {list(request.FILES.keys())}")
                logger.info(f"📤 request.POST keys: {list(request.POST.keys())}")
                logger.info(f"📤 Has request.FILES: {hasattr(request, 'FILES')}")
                logger.info(f"📤 request.FILES: {dict(request.FILES) if hasattr(request, 'FILES') else 'N/A'}")
                
                # Try to access request.data to trigger parser if needed
                if hasattr(request, 'data'):
                    logger.info(f"📤 request.data type: {type(request.data)}")
                    # Check if file is in request.data (DRF sometimes puts files there)
                    if 'file' in request.data:
                        logger.info(f"📤 Found file in request.data!")
                
                # Check if file is in FILES
                if 'file' not in request.FILES:
                    # Additional check: maybe the file is being sent with a different key
                    available_files = list(request.FILES.keys())
                    logger.error(f"❌ File not found in request.FILES. Available keys: {available_files}")
                    logger.error(f"❌ Request content type: {request.content_type}")
                    logger.error(f"❌ Request META CONTENT_TYPE: {request.META.get('CONTENT_TYPE', 'N/A')}")
                    
                    # Try to get file with different possible keys from FILES
                    file = None
                    for possible_key in ['file', 'document', 'upload', 'file_upload']:
                        if possible_key in request.FILES:
                            file = request.FILES[possible_key]
                            logger.info(f"📤 Found file with key '{possible_key}': {file.name}")
                            break
                    
                    # Also check request.data (DRF parser might put files there)
                    if not file and hasattr(request, 'data'):
                        for possible_key in ['file', 'document', 'upload', 'file_upload']:
                            if possible_key in request.data:
                                file_obj = request.data[possible_key]
                                # Check if it's a file-like object
                                if hasattr(file_obj, 'read') or hasattr(file_obj, 'name'):
                                    file = file_obj
                                    logger.info(f"📤 Found file in request.data with key '{possible_key}': {getattr(file, 'name', 'unknown')}")
                                    break
                    
                    if not file:
                        return JsonResponse({
                            'success': False,
                            'error': 'No file provided. Please ensure the file is included in the upload request.',
                            'debug_info': {
                                'content_type': request.content_type,
                                'available_files': available_files,
                                'request_method': request.method
                            }
                        }, status=400)
                else:
                    file = request.FILES['file']
                
                logger.info(f"📤 Uploading file: {file.name} ({file.size} bytes)")
                
                # Resolve uploader user id and name (for filename and Document Handling "Uploaded By")
                upload_framework_id = None
                upload_user_id = None
                try:
                    with connection.cursor() as cur:
                        cur.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        fw_row = cur.fetchone()
                        if fw_row and fw_row[0]:
                            upload_framework_id = fw_row[0]
                    if hasattr(request, 'user') and request.user and getattr(request.user, 'is_authenticated', True):
                        upload_user_id = (
                            getattr(request.user, 'UserId', None)
                            or getattr(request.user, 'id', None)
                            or getattr(request.user, 'pk', None)
                        )
                    # Fallback: frontend may send user_id when request.user is not populated (e.g. token auth)
                    if upload_user_id is None:
                        post_user_id = (request.POST.get('user_id') or (request.data.get('user_id') if hasattr(request, 'data') and isinstance(getattr(request, 'data', None), dict) else None))
                        if post_user_id is not None and str(post_user_id).strip():
                            upload_user_id = int(post_user_id) if str(post_user_id).strip().isdigit() else str(post_user_id).strip()
                except Exception as e:
                    logger.warning(f"Could not get framework_id/user_id for upload: {e}")
                
                # Owner name for filename: prefer username (UserName) so filenames always show the login name
                uploader_name = "Unknown"
                if upload_user_id is not None:
                    try:
                        with connection.cursor() as cur:
                            if isinstance(upload_user_id, int) or (isinstance(upload_user_id, str) and upload_user_id.isdigit()):
                                cur.execute(
                                    "SELECT FirstName, LastName, UserName FROM users WHERE UserId = %s",
                                    [int(upload_user_id) if isinstance(upload_user_id, str) else upload_user_id],
                                )
                            else:
                                cur.execute(
                                    "SELECT FirstName, LastName, UserName FROM users WHERE UserName = %s",
                                    [upload_user_id],
                                )
                            row = cur.fetchone()
                        if row:
                            first_name = (row[0] or "").strip()
                            last_name = (row[1] or "").strip()
                            user_name = (row[2] or "").strip()
                            # Prefer username for filenames; fall back to full name if username missing
                            if user_name:
                                uploader_name = user_name
                            elif first_name or last_name:
                                uploader_name = f"{first_name} {last_name}".strip()
                    except Exception as e:
                        logger.warning(f"Could not get uploader name from DB: {e}")
                if uploader_name == "Unknown" and hasattr(request, 'user') and request.user and getattr(request.user, 'is_authenticated', True):
                    if callable(getattr(request.user, 'get_full_name', None)):
                        uploader_name = (request.user.get_full_name() or "").strip()
                    if not uploader_name:
                        uploader_name = (getattr(request.user, 'username', None) or getattr(request.user, 'UserName', None) or "Unknown")
                    if isinstance(uploader_name, bytes):
                        uploader_name = uploader_name.decode('utf-8', errors='replace')
                if uploader_name == "Unknown" and upload_user_id is not None:
                    uploader_name = str(upload_user_id)
                # Sanitize for filename: keep alphanumeric, underscore, hyphen; replace rest with underscore; limit length
                owner_safe = re.sub(r'[^\w\-]', '_', str(uploader_name))[:64].strip('_') or "Uploader"
                
                # Generate unique file ID and path (filename includes owner name)
                file_id = str(uuid.uuid4())
                file_extension = os.path.splitext(file.name)[1].lower()
                unique_filename = f"{owner_safe}_{file_id}{file_extension}"
                
                # Save file to MEDIA_ROOT under a per-audit folder so that
                # all evidence for the same audit is grouped together.
                audit_folder = f"audit_{audit_id}"
                file_path = os.path.join('ai_audit_documents', audit_folder, unique_filename)
                full_path = os.path.join(settings.MEDIA_ROOT, file_path)
                
                # Ensure directory exists
                os.makedirs(os.path.dirname(full_path), exist_ok=True)
                
                # Save file
                with open(full_path, 'wb+') as destination:
                    for chunk in file.chunks():
                        destination.write(chunk)
                
                # Decompress if needed (client-side compression)
                compression_metadata = None
                full_path, was_compressed, compression_stats = decompress_if_needed(full_path)
                if was_compressed:
                    compression_metadata = compression_stats
                    # Update file extension after decompression (remove .gz)
                    file_extension = os.path.splitext(full_path)[1].lower()
                    # Update file_path to reflect decompressed filename while
                    # keeping it inside the same per-audit folder
                    file_path = os.path.join('ai_audit_documents', audit_folder, os.path.basename(full_path))
                    logger.info(f"📦 Decompressed file: {compression_stats['ratio']}% reduction, saved {compression_stats['bandwidth_saved_kb']} KB")
                
                # Use size of file on disk (after decompression) so same file always gets same size for duplicate check
                file_size_on_disk = os.path.getsize(full_path) if os.path.exists(full_path) else file.size
                logger.info(f"📤 File saved to: {full_path} (size on disk: {file_size_on_disk} bytes)")
                
                # Upload to S3 via document_handling flow (upload_framework_id, upload_user_id set above) so file also appears in Document Handling section
                try:
                    logger.info(f"☁️ Uploading file to S3 (document_handling)...")
                    s3_client = create_direct_mysql_client()
                    connection_test = s3_client.test_connection()
                    if connection_test.get('overall_success', False):
                        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                        s3_filename = f"ai_audit_{timestamp}_{os.path.basename(full_path)}"
                        upload_result = s3_client.upload(
                            file_path=full_path,
                            user_id=str(upload_user_id) if upload_user_id else 'system',
                            custom_file_name=s3_filename,
                            module='document_handling',
                            framework_id=upload_framework_id
                        )
                        if upload_result.get('success'):
                            # Direct S3 client returns operation_id, bucket, key (no file_info)
                            file_info = upload_result.get('file_info') or {}
                            aws_file_link = file_info.get('url') or upload_result.get('url') or ''
                            s3_key = file_info.get('s3Key') or upload_result.get('key') or ''
                            stored_name = file_info.get('storedName') or upload_result.get('key') or s3_filename
                            doc_handling_operation_id = upload_result.get('operation_id')
                            logger.info(f"✅ File uploaded to S3 (Document Handling): operation_id={doc_handling_operation_id}")
                            # Store user's original file name and size (on-disk size after decompression) for replica/duplicate matching
                            if doc_handling_operation_id and file:
                                try:
                                    with connection.cursor() as cur:
                                        cur.execute(
                                            "UPDATE file_operations SET original_name = %s, file_size = %s WHERE id = %s",
                                            [file.name, file_size_on_disk, int(doc_handling_operation_id)],
                                        )
                                    logger.info(f"📋 Set file_operations.original_name={file.name}, file_size={file_size_on_disk} for replica matching")
                                except Exception as upd_err:
                                    logger.warning(f"⚠️ Could not set original_name on file_operations: {upd_err}")
                            # Link this newly uploaded evidence file into a Document Handling folder named after the audit
                            if doc_handling_operation_id:
                                try:
                                    _link_evidence_to_document_handling_folder(audit_id, int(doc_handling_operation_id))
                                    logger.info(f"📁 Linked evidence to Document Handling folder for audit {audit_id}")
                                except Exception as link_err:
                                    logger.warning(f"⚠️ Could not link new evidence to Document Handling folder: {link_err}")
                        else:
                            logger.warning(f"⚠️ S3 upload failed: {upload_result.get('error', 'Unknown error')}")
                    else:
                        logger.warning(f"⚠️ S3 service unavailable, continuing with local file")
                except Exception as s3_error:
                    logger.warning(f"⚠️ S3 upload error (continuing with local file): {str(s3_error)}")
                
                # If S3 worked but no operation_id (e.g. S3 client created without MySQL), create file_operations row in Django DB and link folder
                if not doc_handling_operation_id and file and upload_framework_id:
                    try:
                        with connection.cursor() as cur:
                            cur.execute("""
                                INSERT INTO file_operations
                                (operation_type, user_id, file_name, original_name, file_size, status, module, FrameworkId, stored_name, s3_key, s3_url)
                                VALUES (%s, %s, %s, %s, %s, 'completed', 'document_handling', %s, %s, %s, %s)
                            """, [
                                'upload',
                                str(upload_user_id) if upload_user_id else 'system',
                                file.name,
                                file.name,
                                file_size_on_disk,
                                upload_framework_id,
                                file_path,
                                file_path,
                                aws_file_link or '',
                            ])
                            doc_handling_operation_id = cur.lastrowid
                            if doc_handling_operation_id:
                                _link_evidence_to_document_handling_folder(audit_id, int(doc_handling_operation_id))
                                logger.info(f"📁 Created file_operations row and linked evidence to Document Handling folder (audit {audit_id}, operation_id={doc_handling_operation_id})")
                    except Exception as fallback_err:
                        logger.warning(f"⚠️ Fallback file_operations + link failed: {fallback_err}")
                
                logger.info(f"📤 File will be stored ONCE with path: {file_path}")
                logger.info(f"📤 This path will be reused for all {len(mappings)} mapping(s)")
                
                # Set variables for database insertion (use size on disk so same file = same size for duplicate check)
                file_name = file.name
                file_size = file_size_on_disk
                document_path = file_path  # Same path for ALL mappings
                upload_type = 'evidence'
            
            
            # Get user ID - handle case where user might not exist in database
            user_id = None
            if hasattr(request, 'user') and request.user.is_authenticated:
                try:
                    # Verify user exists in database
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT UserId FROM grc_users WHERE UserId = %s", [request.user.id])
                        db_user = cursor.fetchone()
                        if db_user:
                            user_id = request.user.id
                        else:
                            logger.warning(f"User {request.user.id} not found in grc_users table")
                except Exception as e:
                    logger.warning(f"Error checking user in database: {e}")
            
            # If no valid user found, use NULL (database now allows this)
            if user_id is None:
                logger.info("No valid user found, using NULL for UploadedBy")
            
            # STEP: Process ALL relevant documents from JSON (if we collected them via file_operation_id)
            all_created_document_ids = []  # Accumulate all document IDs created from processing all documents
            
            if all_relevant_documents_to_process and len(all_relevant_documents_to_process) > 0:
                logger.info(f"🔄 Processing ALL {len(all_relevant_documents_to_process)} relevant document(s) from JSON")
                
                # Get framework_id first (needed for all documents)
                with connection.cursor() as cursor:
                    try:
                        cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        framework_row = cursor.fetchone()
                        if framework_row and framework_row[0]:
                            framework_id_for_processing = framework_row[0]
                            logger.info(f"✅ Got FrameworkId {framework_id_for_processing} for processing all documents")
                        else:
                            logger.error(f"❌ No FrameworkId found for audit {audit_id}")
                            framework_id_for_processing = None
                    except Exception as framework_err:
                        logger.error(f"❌ Error querying FrameworkId: {framework_err}")
                        framework_id_for_processing = None
                
                if framework_id_for_processing:
                    # Track processed operation_ids to avoid duplicates
                    processed_operation_ids = set()
                    skipped_due_to_duplicate_count = 0
                    
                    # Process each document from the list
                    for doc_operation_id, doc_matched_compliances in all_relevant_documents_to_process:
                        try:
                            # Skip if we've already processed this operation_id in this run
                            if doc_operation_id in processed_operation_ids:
                                logger.info(f"⏭️ Skipping duplicate document operation_id={doc_operation_id} (already processed in this batch)")
                                continue
                            
                            logger.info(f"📄 Processing document operation_id={doc_operation_id} with {len(doc_matched_compliances)} matched compliances: {doc_matched_compliances}")
                            
                            # Fetch file info for this document
                            with connection.cursor() as cursor:
                                # Check if this document is already processed for this audit
                                doc_s3_key = None
                                doc_stored_name = None
                                cursor.execute("""
                                    SELECT file_name, original_name, s3_url, s3_key, file_size, 
                                           file_type, content_type, stored_name
                                    FROM file_operations 
                                    WHERE id = %s AND status = 'completed'
                                """, [int(doc_operation_id)])
                                doc_file_op = cursor.fetchone()
                                
                                if not doc_file_op:
                                    logger.warning(f"⚠️ File operation {doc_operation_id} not found or not completed - skipping")
                                    continue
                                
                                # Extract stored_name and s3_key for duplicate check
                                doc_stored_name = doc_file_op[7]
                                doc_s3_key = doc_file_op[3]
                                doc_compact_external_id = doc_s3_key or doc_stored_name or str(doc_operation_id)
                                if doc_compact_external_id and len(doc_compact_external_id) > 100:
                                    doc_compact_external_id = doc_compact_external_id[-100:]
                                
                                # Check if this document (by operation_id/external_id) is already in ai_audit_data for this audit
                                cursor.execute("""
                                    SELECT COUNT(*) FROM ai_audit_data 
                                    WHERE audit_id = %s 
                                      AND external_source = 'evidence_attachment'
                                      AND (external_id = %s OR external_id = %s OR external_id = %s)
                                      AND ai_processing_status != 'failed'
                                """, [
                                    int(audit_id) if str(audit_id).isdigit() else audit_id,
                                    str(doc_operation_id),
                                    doc_compact_external_id,
                                    doc_stored_name if doc_stored_name else ''
                                ])
                                count_result = cursor.fetchone()
                                existing_count = count_result[0] if count_result else 0
                                
                                if existing_count > 0:
                                    logger.info(f"⏭️ Document operation_id={doc_operation_id} already exists in ai_audit_data for audit {audit_id} ({existing_count} record(s)) - skipping to avoid duplicate")
                                    processed_operation_ids.add(doc_operation_id)
                                    skipped_due_to_duplicate_count += 1
                                    continue
                                
                                # Mark as processed to avoid duplicates within this batch
                                processed_operation_ids.add(doc_operation_id)
                                
                                # Extract remaining file information for this document (some already extracted above)
                                doc_file_name_from_db = doc_file_op[0]
                                doc_original_name = doc_file_op[1]
                                doc_file_name = doc_stored_name or doc_file_name_from_db or doc_original_name or 'unknown_file'
                                doc_aws_file_link = doc_file_op[2]
                                doc_file_size = doc_file_op[4] or 0
                                doc_file_type = doc_file_op[5]
                                doc_content_type = doc_file_op[6]
                                
                                # Create mappings from matched compliances for this document
                                doc_mappings = []
                                if doc_matched_compliances and len(doc_matched_compliances) > 0:
                                    try:
                                        cursor.execute(f"""
                                            SELECT DISTINCT c.ComplianceId, c.SubPolicyId, sp.PolicyId
                                            FROM compliance c
                                            JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                                            WHERE c.ComplianceId IN ({','.join(['%s'] * len(doc_matched_compliances))})
                                        """, doc_matched_compliances)
                                        doc_compliance_mappings = cursor.fetchall()
                                        
                                        for comp_id, subpol_id, pol_id in doc_compliance_mappings:
                                            doc_mappings.append({
                                                'policy_id': pol_id,
                                                'subpolicy_id': subpol_id,
                                                'compliance_id': comp_id
                                            })
                                        logger.info(f"📋 Created {len(doc_mappings)} mapping(s) for document operation_id={doc_operation_id}")
                                    except Exception as mapping_err:
                                        logger.error(f"❌ Error creating mappings for document {doc_operation_id}: {mapping_err}")
                                        continue
                                
                                # Insert records for this document (doc_compact_external_id already set above)
                                doc_document_path = doc_s3_key or doc_stored_name or f"file_operations/{doc_operation_id}/{doc_file_name}"
                                doc_normalized_file_name = (doc_file_name or '').strip()
                                doc_normalized_file_size = int(doc_file_size or 0)
                                
                                # Insert ai_audit_data records for each mapping
                                doc_created_count = 0  # Count records created for this specific document
                                for mapping in doc_mappings:
                                    try:
                                        map_policy_id = mapping.get('policy_id') or None
                                        map_subpolicy_id = mapping.get('subpolicy_id') or None
                                        map_compliance_id = mapping.get('compliance_id')
                                        
                                        if map_compliance_id:
                                            cursor.execute("""
                                                INSERT INTO ai_audit_data 
                                                (audit_id, document_id, document_name, document_path, document_type, 
                                                 file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                                 policy_id, subpolicy_id, compliance_id, upload_status, external_source, external_id,
                                                 FrameworkId, created_at, updated_at)
                                                VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                       %s, %s, %s, 'uploaded', %s, %s, %s, NOW(), NOW())
                                            """, [
                                                int(audit_id) if str(audit_id).isdigit() else audit_id,
                                                doc_normalized_file_name,
                                                doc_document_path,
                                                'evidence'[:50],
                                                doc_normalized_file_size,
                                                'application/octet-stream',
                                                user_id,
                                                map_policy_id,
                                                map_subpolicy_id,
                                                map_compliance_id,
                                                'evidence_attachment',
                                                doc_compact_external_id,
                                                framework_id_for_processing
                                            ])
                                        else:
                                            cursor.execute("""
                                                INSERT INTO ai_audit_data 
                                                (audit_id, document_id, document_name, document_path, document_type, 
                                                 file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                                 policy_id, subpolicy_id, upload_status, external_source, external_id,
                                                 FrameworkId, created_at, updated_at)
                                                VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                       %s, %s, 'uploaded', %s, %s, %s, NOW(), NOW())
                                            """, [
                                                int(audit_id) if str(audit_id).isdigit() else audit_id,
                                                doc_normalized_file_name,
                                                doc_document_path,
                                                'evidence'[:50],
                                                doc_normalized_file_size,
                                                'application/octet-stream',
                                                user_id,
                                                map_policy_id,
                                                map_subpolicy_id,
                                                'evidence_attachment',
                                                doc_compact_external_id,
                                                framework_id_for_processing
                                            ])
                                        
                                        cursor.execute("SELECT LAST_INSERT_ID()")
                                        result = cursor.fetchone()
                                        record_id = result[0] if result else None
                                        
                                        if record_id:
                                            cursor.execute("UPDATE ai_audit_data SET document_id = %s WHERE id = %s", [record_id, record_id])
                                            all_created_document_ids.append(record_id)
                                            doc_created_count += 1
                                            logger.info(f"✅ Created record ID {record_id} for document operation_id={doc_operation_id}, compliance_id={map_compliance_id}")
                                    except Exception as insert_err:
                                        logger.error(f"❌ Error inserting record for document {doc_operation_id}: {insert_err}")
                                        continue
                                
                                logger.info(f"✅ Finished processing document operation_id={doc_operation_id} - created {doc_created_count} record(s) for this document")
                                
                                # Commit the transaction (Django auto-commits, but explicit commit ensures it)
                                try:
                                    connection.commit()
                                    logger.info(f"✅ Committed database transaction for document operation_id={doc_operation_id}")
                                except Exception as commit_err:
                                    logger.error(f"❌ Error committing transaction for document {doc_operation_id}: {commit_err}")
                                    connection.rollback()
                                
                        except Exception as doc_process_err:
                            logger.error(f"❌ Error processing document {doc_operation_id}: {doc_process_err}")
                            import traceback
                            logger.error(f"❌ Error details: {traceback.format_exc()}")
                            try:
                                connection.rollback()
                            except:
                                pass
                            continue
                
                logger.info(f"✅✅✅ Processed ALL {len(all_relevant_documents_to_process)} relevant document(s) from JSON - created {len(all_created_document_ids)} total record(s)")
                
                # When user re-uploads same evidence for same policy/subpolicy/compliance, show message
                if len(all_relevant_documents_to_process) > 0 and len(all_created_document_ids) == 0 and skipped_due_to_duplicate_count > 0:
                    return JsonResponse({
                        'success': True,
                        'document_id': None,
                        'document_ids': [],
                        'mappings_count': 0,
                        'message': 'Evidence already uploaded for the selected policy/subpolicy/compliance.',
                        'already_uploaded': True
                    })
                
                # AUTOMATIC COMPLIANCE CHECK FOR DOCUMENT HANDLING UPLOADS (JSON batch processing)
                # When documents are processed from JSON (Document Handling), automatically trigger compliance checks
                if len(all_created_document_ids) > 0:
                    logger.info(f"🔄 Document Handling batch upload detected - checking if automatic compliance check needed for {len(all_created_document_ids)} mapping(s)")
                    
                    # Get the first document's info to check if already completed
                    primary_document_id = all_created_document_ids[0]
                    with connection.cursor() as cursor:
                        cursor.execute("""
                            SELECT document_name, file_size, external_source
                            FROM ai_audit_data
                            WHERE id = %s AND audit_id = %s
                            LIMIT 1
                        """, [primary_document_id, int(audit_id) if str(audit_id).isdigit() else audit_id])
                        doc_info = cursor.fetchone()
                        
                        if doc_info:
                            doc_name, doc_size, doc_external_source = doc_info
                            
                            # Check if this file (by name + size) is already completed for this audit
                            cursor.execute("""
                                SELECT COUNT(*) FROM ai_audit_data 
                                WHERE document_name = %s
                                  AND file_size = %s
                                  AND audit_id = %s
                                  AND ai_processing_status = 'completed'
                                  AND (external_source = 'evidence_attachment' OR external_source IS NULL)
                            """, [doc_name, doc_size, int(audit_id) if str(audit_id).isdigit() else audit_id])
                            count_result = cursor.fetchone()
                            completed_count = count_result[0] if count_result else 0
                            
                            if completed_count > 0:
                                logger.info(f"⏭️ Skipping automatic compliance check - File '{doc_name}' (size: {doc_size}) is already completed for audit {audit_id}")
                                # Update the newly created records to 'completed' status to match existing ones
                                try:
                                    placeholders = ",".join(["%s"] * len(all_created_document_ids))
                                    cursor.execute(f"""
                                        UPDATE ai_audit_data 
                                        SET ai_processing_status = 'completed',
                                            updated_at = NOW()
                                        WHERE id IN ({placeholders})
                                          AND audit_id = %s
                                    """, all_created_document_ids + [int(audit_id) if str(audit_id).isdigit() else audit_id])
                                    updated_count = cursor.rowcount
                                    logger.info(f"✅ Updated {updated_count} newly created record(s) to 'completed' status to match existing completed file")
                                except Exception as update_err:
                                    logger.warning(f"⚠️ Could not update new records to completed status: {update_err}")
                            else:
                                logger.info(f"🚀 Documents are new/not audited - will automatically trigger compliance checks for {len(all_created_document_ids)} mapping(s)")
                                
                                # Trigger compliance checks asynchronously (don't block the upload response)
                                import threading
                                def auto_check_compliance():
                                    try:
                                        logger.info(f"🚀 Starting automatic compliance checks for Document Handling batch upload")
                                        
                                        # Call compliance check for the primary document (which will check all mappings for that file)
                                        # The check_document_compliance function handles all mappings for a file group
                                        try:
                                            logger.info(f"🔍 Auto-checking compliance for primary document_id={primary_document_id}")
                                            
                                            # Create a minimal request object for the compliance check
                                            from django.test import RequestFactory
                                            factory = RequestFactory()
                                            check_request = factory.post(f'/api/ai-audit/{audit_id}/documents/{primary_document_id}/check/')
                                            
                                            # Copy authentication from original request
                                            if 'HTTP_AUTHORIZATION' in request.META:
                                                check_request.META['HTTP_AUTHORIZATION'] = request.META['HTTP_AUTHORIZATION']
                                            
                                            # Set request data (empty - will use all compliances from mappings)
                                            check_request.data = {}
                                            
                                            # Call the compliance check function directly
                                            from rest_framework.response import Response as DRFResponse
                                            result = check_document_compliance(check_request, audit_id, primary_document_id)
                                            
                                            if isinstance(result, DRFResponse):
                                                if result.status_code == 200:
                                                    logger.info(f"✅ Automatic compliance check completed successfully for document_id={primary_document_id}")
                                                else:
                                                    logger.warning(f"⚠️ Automatic compliance check returned status {result.status_code}")
                                            else:
                                                logger.info(f"ℹ️ Compliance check triggered for document_id={primary_document_id}")
                                                
                                        except Exception as check_err:
                                            logger.warning(f"⚠️ Could not auto-check document_id={primary_document_id}: {check_err}")
                                            import traceback
                                            logger.warning(traceback.format_exc())
                                        
                                        logger.info(f"✅ Automatic compliance checks completed for Document Handling batch upload")
                                    except Exception as auto_err:
                                        logger.error(f"❌ Error in automatic compliance check: {auto_err}")
                                        import traceback
                                        logger.error(traceback.format_exc())
                                
                                # Start background thread for automatic compliance checks
                                auto_check_thread = threading.Thread(target=auto_check_compliance, daemon=True)
                                auto_check_thread.start()
                                logger.info(f"🚀 Started background thread for automatic compliance checks")
                    
                    # Return early (skip single document processing)
                    return JsonResponse({
                        'success': True,
                        'document_id': primary_document_id,
                        'document_ids': all_created_document_ids,
                        'file_name': 'multiple_documents',
                        'file_size': 0,
                        'file_type': 'multiple',
                        'mappings_count': len(all_created_document_ids),
                        'message': f'Processed {len(all_relevant_documents_to_process)} relevant document(s) from JSON with {len(all_created_document_ids)} total mapping(s)' + 
                                  (' - Automatic compliance check started' if doc_external_source == 'evidence_attachment' else '')
                    })
            
            # Store ONLY in ai_audit_data table (it has everything we need)
            try:
                with connection.cursor() as cursor:
                    # First insert with a temporary document_id (will be updated after)
                    # Determine mime type and external source
                    if already_uploaded:
                        # Documents coming from Document Handling reuse (file_operation_id)
                        # should continue to be treated as 'evidence_attachment'
                        mime_type = 'application/octet-stream'  # Default for already uploaded files
                        external_source = 'evidence_attachment'
                    else:
                        # Pure manual uploads for AI audit:
                        # - We still upload to S3 / file_operations so they appear in Document Handling
                        # - But we keep external_source='manual' so compliance checks use the local path
                        #   and do NOT require S3 metadata / relevance tables.
                        mime_type = file.content_type
                        external_source = 'manual'
                    
                    # Get FrameworkId from the audit record
                    try:
                        logger.info(f"🔍 Querying FrameworkId for audit {audit_id}")
                        cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        framework_result = cursor.fetchone()
                        logger.info(f"🔍 FrameworkId query result: {framework_result}")
                        
                        if framework_result and framework_result[0] is not None:
                            framework_id = framework_result[0]
                            logger.info(f"✅ Found FrameworkId {framework_id} for audit {audit_id}")
                        else:
                            logger.error(f"❌ No FrameworkId found for audit {audit_id}. Audit record may not exist or FrameworkId is NULL.")
                            return JsonResponse({
                                'success': False,
                                'error': f'Audit {audit_id} not found or has no FrameworkId assigned. Please ensure the audit exists and has a framework assigned.'
                            }, status=400)
                    except Exception as framework_err:
                        logger.error(f"❌ Error querying FrameworkId for audit {audit_id}: {framework_err}")
                        return JsonResponse({
                            'success': False,
                            'error': f'Database error while retrieving audit framework: {framework_err}'
                        }, status=500)
                    
                    # external_id column is VARCHAR(100). Store a compact identifier
                    # When file was uploaded via Document Handling (s3_functions), use operation_id so it links to file_operations
                    compact_external_id = None
                    if doc_handling_operation_id:
                        compact_external_id = str(doc_handling_operation_id)[:100]
                    elif s3_key:
                        compact_external_id = s3_key
                    elif stored_name:
                        compact_external_id = stored_name
                    elif file_id:
                        compact_external_id = str(file_id)
                    elif aws_file_link and 'amazonaws.com/' in aws_file_link:
                        compact_external_id = aws_file_link.split('amazonaws.com/')[-1].split('?')[0]
                    
                    # Ensure it fits the varchar(100)
                    if compact_external_id and len(compact_external_id) > 100:
                        compact_external_id = compact_external_id[-100:]
                    
                    # Create records for each mapping (file stored once, multiple mapping records)
                    # When mappings is empty (e.g. replica path with no relevance yet), use one default so we still create a record
                    if not mappings:
                        mappings = [{'policy_id': None, 'subpolicy_id': None}]
                        logger.info("📋 No mappings provided - using default single mapping so at least one ai_audit_data record is created")
                    # IMPORTANT: Normalize document_name and file_size to ensure exact match for frontend grouping
                    normalized_file_name = (file_name or '').strip()
                    normalized_file_size = int(file_size or 0)
                    
                    created_document_ids = []
                    first_insert_error = None
                    skipped_duplicate_count = 0
                    
                    for mapping_idx, mapping in enumerate(mappings):
                        map_policy_id = mapping.get('policy_id') or None
                        map_subpolicy_id = mapping.get('subpolicy_id') or None
                        map_compliance_id = mapping.get('compliance_id')  # May be None for non-matched-compliance mappings
                        
                        # Check if same file already linked for this policy/subpolicy/compliance (avoid duplicate)
                        # Match by: same external_id, OR same document_name+file_size, OR same file_size + same original name,
                        # OR same original name with size within 2KB (catches re-upload when size differs slightly e.g. compression)
                        _size_tolerance = 2048  # bytes
                        if _check_ai_audit_data_has_compliance_id():
                            cursor.execute("""
                                SELECT COUNT(*) FROM ai_audit_data a
                                LEFT JOIN file_operations fo ON (fo.id = a.external_id OR CAST(fo.id AS CHAR) = a.external_id) AND fo.status = 'completed'
                                WHERE a.audit_id = %s
                                  AND (a.policy_id <=> %s)
                                  AND (a.subpolicy_id <=> %s)
                                  AND (a.compliance_id <=> %s)
                                  AND (a.ai_processing_status IS NULL OR a.ai_processing_status != 'failed')
                                  AND (
                                    (a.external_id IS NOT NULL AND a.external_id = %s)
                                    OR (a.document_name = %s AND a.file_size = %s)
                                    OR (a.file_size = %s AND (fo.original_name = %s OR a.document_name = %s))
                                    OR ((fo.original_name = %s OR a.document_name = %s) AND ABS(COALESCE(a.file_size, 0) - %s) <= %s)
                                  )
                            """, [
                                int(audit_id) if str(audit_id).isdigit() else audit_id,
                                map_policy_id,
                                map_subpolicy_id,
                                map_compliance_id,
                                compact_external_id or '',
                                normalized_file_name,
                                normalized_file_size,
                                normalized_file_size,
                                normalized_file_name,
                                normalized_file_name,
                                normalized_file_name,
                                normalized_file_name,
                                normalized_file_size,
                                _size_tolerance
                            ])
                        else:
                            cursor.execute("""
                                SELECT COUNT(*) FROM ai_audit_data a
                                LEFT JOIN file_operations fo ON (fo.id = a.external_id OR CAST(fo.id AS CHAR) = a.external_id) AND fo.status = 'completed'
                                WHERE a.audit_id = %s
                                  AND (a.policy_id <=> %s)
                                  AND (a.subpolicy_id <=> %s)
                                  AND (a.ai_processing_status IS NULL OR a.ai_processing_status != 'failed')
                                  AND (
                                    (a.external_id IS NOT NULL AND a.external_id = %s)
                                    OR (a.document_name = %s AND a.file_size = %s)
                                    OR (a.file_size = %s AND (fo.original_name = %s OR a.document_name = %s))
                                    OR ((fo.original_name = %s OR a.document_name = %s) AND ABS(COALESCE(a.file_size, 0) - %s) <= %s)
                                  )
                            """, [
                                int(audit_id) if str(audit_id).isdigit() else audit_id,
                                map_policy_id,
                                map_subpolicy_id,
                                compact_external_id or '',
                                normalized_file_name,
                                normalized_file_size,
                                normalized_file_size,
                                normalized_file_name,
                                normalized_file_name,
                                normalized_file_name,
                                normalized_file_name,
                                normalized_file_size,
                                _size_tolerance
                            ])
                        dup_row = cursor.fetchone()
                        if dup_row and (dup_row[0] or 0) > 0:
                            logger.info(f"⏭️ Evidence already uploaded for this policy/subpolicy/compliance (audit_id={audit_id}) - skipping duplicate")
                            skipped_duplicate_count += 1
                            continue
                        
                        try:
                            debug_print("=" * 80)
                            debug_print(f"ABOUT TO INSERT INTO ai_audit_data (mapping {mapping_idx + 1}/{len(mappings)})")
                            debug_print(f"audit_id: {audit_id}")
                            debug_print(f"framework_id: {framework_id}")
                            debug_print(f"policy_id: {map_policy_id}, subpolicy_id: {map_subpolicy_id}")
                            if map_compliance_id:
                                debug_print(f"compliance_id: {map_compliance_id}")
                            debug_print("=" * 80)
                            logger.info(f"🔍 About to insert mapping {mapping_idx + 1}/{len(mappings)} into ai_audit_data:")
                            logger.info(f"  - audit_id: {audit_id}")
                            logger.info(f"  - file_name: {file_name}")
                            logger.info(f"  - document_path: {document_path} (shared across all mappings)")
                            logger.info(f"  - upload_type: {upload_type}")
                            logger.info(f"  - file_size: {file_size}")
                            logger.info(f"  - mime_type: {mime_type}")
                            logger.info(f"  - user_id: {user_id}")
                            logger.info(f"  - policy_id: {map_policy_id}")
                            logger.info(f"  - subpolicy_id: {map_subpolicy_id}")
                            if map_compliance_id:
                                logger.info(f"  - compliance_id: {map_compliance_id} (from matched compliances)")
                            logger.info(f"  - external_source: {external_source}")
                            logger.info(f"  - compact_external_id: {compact_external_id}")
                            logger.info(f"  - framework_id: {framework_id}")
                            
                            try:
                                # Try INSERT with compliance_id first
                                if map_compliance_id:
                                    try:
                                        cursor.execute("""
                                            INSERT INTO ai_audit_data 
                                            (audit_id, document_id, document_name, document_path, document_type, 
                                             file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                             policy_id, subpolicy_id, compliance_id, upload_status, external_source, external_id,
                                             FrameworkId, created_at, updated_at)
                                            VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                   %s, %s, %s, 'uploaded', %s, %s, %s, NOW(), NOW())
                                        """, [
                                            int(audit_id) if str(audit_id).isdigit() else audit_id,
                                            normalized_file_name,  # Use normalized name to ensure grouping works
                                            document_path,  # Same file_path for all mappings
                                            upload_type[:50],  # Truncate to fit varchar(50)
                                            normalized_file_size,  # Use normalized size to ensure grouping works
                                            mime_type,
                                            user_id,
                                            map_policy_id,
                                            map_subpolicy_id,
                                            map_compliance_id,  # Store compliance_id
                                            external_source,
                                            compact_external_id,  # Compact identifier fits column limit
                                            framework_id  # Add FrameworkId
                                        ])
                                        logger.info(f"✅ Inserted with compliance_id={map_compliance_id}")
                                    except Exception as compliance_err:
                                        # If compliance_id column doesn't exist, fall back to INSERT without it
                                        if 'Unknown column' in str(compliance_err) and 'compliance_id' in str(compliance_err).lower():
                                            logger.warning("ai_audit_data lacks compliance_id column; inserting without it")
                                            raise compliance_err  # Will be caught by outer try-except
                                        else:
                                            raise
                                else:
                                    # No compliance_id, use standard INSERT
                                    cursor.execute("""
                                        INSERT INTO ai_audit_data 
                                        (audit_id, document_id, document_name, document_path, document_type, 
                                         file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                         policy_id, subpolicy_id, upload_status, external_source, external_id,
                                         FrameworkId, created_at, updated_at)
                                        VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                               %s, %s, 'uploaded', %s, %s, %s, NOW(), NOW())
                                    """, [
                                        int(audit_id) if str(audit_id).isdigit() else audit_id,
                                        normalized_file_name,  # Use normalized name to ensure grouping works
                                        document_path,  # Same file_path for all mappings
                                        upload_type[:50],  # Truncate to fit varchar(50)
                                        normalized_file_size,  # Use normalized size to ensure grouping works
                                        mime_type,
                                        user_id,
                                        map_policy_id,
                                        map_subpolicy_id,
                                        external_source,
                                        compact_external_id,  # Compact identifier fits column limit
                                        framework_id  # Add FrameworkId
                                    ])
                            except Exception as framework_err:
                                # Handle missing FrameworkId or compliance_id column
                                if 'Unknown column' in str(framework_err):
                                    if 'FrameworkId' in str(framework_err):
                                        logger.warning("ai_audit_data lacks FrameworkId column; inserting without it")
                                        # Try INSERT without FrameworkId and without compliance_id
                                        try:
                                            if map_compliance_id:
                                                cursor.execute("""
                                                    INSERT INTO ai_audit_data 
                                                    (audit_id, document_id, document_name, document_path, document_type, 
                                                     file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                                     policy_id, subpolicy_id, compliance_id, upload_status, external_source, external_id,
                                                     created_at, updated_at)
                                                    VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                           %s, %s, %s, 'uploaded', %s, %s, NOW(), NOW())
                                                """, [
                                                    int(audit_id) if str(audit_id).isdigit() else audit_id,
                                                    normalized_file_name,
                                                    document_path,
                                                    upload_type[:50],
                                                    normalized_file_size,
                                                    mime_type,
                                                    user_id,
                                                    map_policy_id,
                                                    map_subpolicy_id,
                                                    map_compliance_id,
                                                    external_source,
                                                    compact_external_id
                                                ])
                                            else:
                                                cursor.execute("""
                                                    INSERT INTO ai_audit_data 
                                                    (audit_id, document_id, document_name, document_path, document_type, 
                                                     file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                                     policy_id, subpolicy_id, upload_status, external_source, external_id,
                                                     created_at, updated_at)
                                                    VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                           %s, %s, 'uploaded', %s, %s, NOW(), NOW())
                                                """, [
                                                    int(audit_id) if str(audit_id).isdigit() else audit_id,
                                                    normalized_file_name,
                                                    document_path,
                                                    upload_type[:50],
                                                    normalized_file_size,
                                                    mime_type,
                                                    user_id,
                                                    map_policy_id,
                                                    map_subpolicy_id,
                                                    external_source,
                                                    compact_external_id
                                                ])
                                        except Exception as compliance_err2:
                                            # If compliance_id also doesn't exist, use basic INSERT
                                            if 'compliance_id' in str(compliance_err2).lower():
                                                logger.warning("ai_audit_data also lacks compliance_id column; inserting without both")
                                                cursor.execute("""
                                                    INSERT INTO ai_audit_data 
                                                    (audit_id, document_id, document_name, document_path, document_type, 
                                                     file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                                     policy_id, subpolicy_id, upload_status, external_source, external_id,
                                                     created_at, updated_at)
                                                    VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                           %s, %s, 'uploaded', %s, %s, NOW(), NOW())
                                                """, [
                                                    int(audit_id) if str(audit_id).isdigit() else audit_id,
                                                    normalized_file_name,
                                                    document_path,
                                                    upload_type[:50],
                                                    normalized_file_size,
                                                    mime_type,
                                                    user_id,
                                                    map_policy_id,
                                                    map_subpolicy_id,
                                                    external_source,
                                                    compact_external_id
                                                ])
                                            else:
                                                raise
                                    elif 'compliance_id' in str(framework_err).lower():
                                        logger.warning("ai_audit_data lacks compliance_id column; inserting without it")
                                        cursor.execute("""
                                            INSERT INTO ai_audit_data 
                                            (audit_id, document_id, document_name, document_path, document_type, 
                                             file_size, mime_type, uploaded_by, ai_processing_status, uploaded_date,
                                             policy_id, subpolicy_id, upload_status, external_source, external_id,
                                             FrameworkId, created_at, updated_at)
                                            VALUES (%s, 0, %s, %s, %s, %s, %s, %s, 'pending', NOW(),
                                                   %s, %s, 'uploaded', %s, %s, %s, NOW(), NOW())
                                        """, [
                                            int(audit_id) if str(audit_id).isdigit() else audit_id,
                                            normalized_file_name,
                                            document_path,
                                            upload_type[:50],
                                            normalized_file_size,
                                            mime_type,
                                            user_id,
                                            map_policy_id,
                                            map_subpolicy_id,
                                            external_source,
                                            compact_external_id,
                                            framework_id
                                        ])
                                    else:
                                        raise
                                else:
                                    raise
                            
                            logger.info(f"✅ Insert mapping {mapping_idx + 1} into ai_audit_data successful")
                            
                            # Get the auto-generated ID for this mapping
                            cursor.execute("SELECT LAST_INSERT_ID()")
                            result = cursor.fetchone()
                            record_id = result[0] if result else None
                            
                            # Update the document_id column
                            if record_id:
                                cursor.execute("""
                                    UPDATE ai_audit_data 
                                    SET document_id = %s 
                                    WHERE id = %s
                                """, [record_id, record_id])
                                created_document_ids.append(record_id)
                                logger.info(f"📤 Created mapping record with ID: {record_id}")
                        except Exception as insert_err:
                            if first_insert_error is None:
                                first_insert_error = insert_err
                            logger.error(f"❌ Insert error for mapping {mapping_idx + 1}: {insert_err}")
                            # Continue with other mappings even if one fails
                            continue
                    
                    if not created_document_ids:
                        if skipped_duplicate_count > 0:
                            return JsonResponse({
                                'success': True,
                                'document_id': None,
                                'document_ids': [],
                                'mappings_count': 0,
                                'message': 'Evidence already uploaded for the selected policy/subpolicy/compliance.',
                                'already_uploaded': True
                            })
                        err_msg = "Failed to create any mapping records."
                        if first_insert_error:
                            err_msg += f" First error: {first_insert_error}"
                        raise Exception(err_msg)
                    
                    logger.info(f"✅ Created {len(created_document_ids)} mapping record(s) for file {file_name}")
            except Exception as e:
                logger.error(f"❌ Could not store in ai_audit_data table: {e}")
                raise

            # All document data is now stored in ai_audit_data table with policy mappings
            # Return the first document_id as primary, and list of all created IDs
            primary_document_id = created_document_ids[0] if created_document_ids else None
            
            # AUTOMATIC COMPLIANCE CHECK FOR DOCUMENT HANDLING UPLOADS
            # When documents are uploaded from Document Handling (evidence_attachment),
            # automatically trigger compliance checks in the background
            # BUT ONLY if document is new and not already audited
            if external_source == 'evidence_attachment' and created_document_ids:
                logger.info(f"🔄 Document Handling upload detected - checking if automatic compliance check needed for {len(created_document_ids)} mapping(s)")
                
                # Check if document is already audited (status = 'completed') FOR THIS SPECIFIC AUDIT
                # Important: Check by file identity (document_name + file_size) not just new IDs
                # Same file uploaded again should not be re-audited if already completed
                with connection.cursor() as cursor:
                    # Check if this file (by name + size) is already completed for this audit
                    cursor.execute("""
                            SELECT COUNT(*) FROM ai_audit_data 
                        WHERE document_name = %s
                          AND file_size = %s
                              AND audit_id = %s
                              AND ai_processing_status = 'completed'
                          AND (external_source = 'evidence_attachment' OR external_source IS NULL)
                    """, [file_name, file_size, int(audit_id) if str(audit_id).isdigit() else audit_id])
                    count_result = cursor.fetchone()
                    completed_count = count_result[0] if count_result else 0
                
                if completed_count > 0:
                    logger.info(f"⏭️ Skipping automatic compliance check - File '{file_name}' (size: {file_size}) is already completed for audit {audit_id}")
                    # Update the newly created records to 'completed' status to match existing ones
                    try:
                        with connection.cursor() as cursor:
                            placeholders = ",".join(["%s"] * len(created_document_ids))
                            cursor.execute(f"""
                                UPDATE ai_audit_data 
                                SET ai_processing_status = 'completed',
                                    updated_at = NOW()
                                WHERE id IN ({placeholders})
                                  AND audit_id = %s
                            """, created_document_ids + [int(audit_id) if str(audit_id).isdigit() else audit_id])
                            updated_count = cursor.rowcount
                            logger.info(f"✅ Updated {updated_count} newly created record(s) to 'completed' status to match existing completed file")
                    except Exception as update_err:
                        logger.warning(f"⚠️ Could not update new records to completed status: {update_err}")
                else:
                    logger.info(f"🚀 Document is new/not audited - will automatically trigger compliance checks for {len(created_document_ids)} mapping(s)")
                    
                    # Trigger compliance checks asynchronously (don't block the upload response)
                    import threading
                    def auto_check_compliance():
                        try:
                            logger.info(f"🚀 Starting automatic compliance checks for Document Handling upload")
                            
                            # Call compliance check for the primary document (which will check all mappings for that file)
                            # The check_document_compliance function handles all mappings for a file group
                            try:
                                logger.info(f"🔍 Auto-checking compliance for primary document_id={primary_document_id}")
                                
                                # Create a minimal request object for the compliance check
                                from django.test import RequestFactory
                                factory = RequestFactory()
                                check_request = factory.post(f'/api/ai-audit/{audit_id}/documents/{primary_document_id}/check/')
                                
                                # Copy authentication from original request
                                if 'HTTP_AUTHORIZATION' in request.META:
                                    check_request.META['HTTP_AUTHORIZATION'] = request.META['HTTP_AUTHORIZATION']
                                
                                # Set request data (empty - will use all compliances from mappings)
                                check_request.data = {}
                                
                                # Call the compliance check function directly
                                from rest_framework.response import Response as DRFResponse
                                result = check_document_compliance(check_request, audit_id, primary_document_id)
                                
                                if isinstance(result, DRFResponse):
                                    if result.status_code == 200:
                                        logger.info(f"✅ Automatic compliance check completed successfully for document_id={primary_document_id}")
                                    else:
                                        logger.warning(f"⚠️ Automatic compliance check returned status {result.status_code}")
                                else:
                                    logger.info(f"ℹ️ Compliance check triggered for document_id={primary_document_id}")
                                    
                            except Exception as check_err:
                                logger.warning(f"⚠️ Could not auto-check document_id={primary_document_id}: {check_err}")
                                import traceback
                                logger.warning(traceback.format_exc())
                            
                            logger.info(f"✅ Automatic compliance checks completed for Document Handling upload")
                        except Exception as auto_err:
                            logger.error(f"❌ Error in automatic compliance check: {auto_err}")
                            import traceback
                            logger.error(traceback.format_exc())
                    
                    # Update audit status to "Work In Progress" when auto-check starts
                    try:
                        from ...models import Audit
                        audit_obj = Audit.objects.get(AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id=tenant_id)
                        if audit_obj.Status == 'Yet to Start':
                            audit_obj.Status = 'Work In Progress'
                            audit_obj.save()
                            logger.info(f"✅ Updated audit {audit_id} status to 'Work In Progress' (auto-check started)")
                    except Exception as status_err:
                        logger.warning(f"⚠️ Could not update audit status: {status_err}")
                    
                    # Start background thread for automatic compliance checks
                    auto_check_thread = threading.Thread(target=auto_check_compliance, daemon=True)
                    auto_check_thread.start()
                    logger.info(f"🚀 Started background thread for automatic compliance checks")
            
            # Update audit status to "Work In Progress" for manual uploads as well
            # (when documents are uploaded through AI audit upload page)
            if external_source == 'manual' and created_document_ids:
                try:
                    from ...models import Audit
                    audit_obj = Audit.objects.get(AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id)
                    if audit_obj.Status == 'Yet to Start':
                        audit_obj.Status = 'Work In Progress'
                        audit_obj.save()
                        logger.info(f"✅ Updated audit {audit_id} status to 'Work In Progress' (manual document upload)")
                except Exception as status_err:
                    logger.warning(f"⚠️ Could not update audit status for manual upload: {status_err}")
            
            return JsonResponse({
                'success': True,
                'document_id': primary_document_id,  # Primary document ID
                'document_ids': created_document_ids,  # All created document IDs
                'file_name': file_name,
                'file_size': file_size,
                'file_type': mime_type,
                'mappings_count': len(created_document_ids),
                'message': f'File uploaded successfully with {len(created_document_ids)} mapping(s)' + 
                          (' - Automatic compliance check started' if external_source == 'evidence_attachment' else '')
            })
            
        except Exception as e:
            logger.error(f"❌ Upload error: {e}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)
    
 
def save_ai_compliance_to_checklist(audit_id, document_id, analyses, user_id, framework_id, policy_id=None, subpolicy_id=None, tenant_id=None):
    """
    Save AI compliance results to lastchecklistitemverified table for standard compliance tracking.
   
    Args:
        audit_id: The audit ID
        document_id: The document ID that was analyzed
        analyses: List of compliance analysis results from AI (or dict with 'compliance_analyses' key)
        user_id: User ID who triggered the check (can be numeric UserId or username string)
        framework_id: Framework ID from the audit
        policy_id: Policy ID (optional, will be resolved from compliance if not provided)
        subpolicy_id: Sub-policy ID (optional, will be resolved from compliance if not provided)
        tenant_id: Tenant ID (optional, will be resolved from audit if not provided)
    """
    from django.db import connection
    from django.db.utils import IntegrityError
    from django.utils import timezone
    import json
   
    try:
        # MULTI-TENANCY: Get tenant_id from audit if not provided
        if not tenant_id:
            with connection.cursor() as cursor:
                cursor.execute("SELECT TenantId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                tenant_row = cursor.fetchone()
                if tenant_row:
                    tenant_id = tenant_row[0]
                else:
                    logger.warning(f"⚠️ Could not find tenant_id for audit {audit_id}")
                    return
        # Convert user_id to numeric UserId if it's a username string
        numeric_user_id = None
        if user_id is not None:
            try:
                # Try to convert to int directly
                numeric_user_id = int(user_id)
            except (ValueError, TypeError):
                # If conversion fails, try to resolve username to UserId
                try:
                    with connection.cursor() as cursor:
                        cursor.execute(
                            "SELECT UserId FROM users WHERE UserName = %s LIMIT 1",
                            (str(user_id),)
                        )
                        user_row = cursor.fetchone()
                        if user_row:
                            numeric_user_id = user_row[0]
                            logger.info(f"✅ Resolved username '{user_id}' to UserId {numeric_user_id}")
                        else:
                            logger.warning(f"⚠️ Could not resolve username '{user_id}' to UserId; using NULL")
                            numeric_user_id = None
                except Exception as user_lookup_err:
                    logger.warning(f"⚠️ Error resolving username '{user_id}' to UserId: {user_lookup_err}")
                    numeric_user_id = None
        
        current_datetime = timezone.now()
        current_date = current_datetime.date()
        current_time = current_datetime.time()
       
        # Handle different analysis formats
        if isinstance(analyses, dict):
            # If it's a dict, extract the compliance_analyses array
            analyses_list = analyses.get('compliance_analyses', [])
        elif isinstance(analyses, list):
            analyses_list = analyses
        else:
            analyses_list = []
       
        if not analyses_list or len(analyses_list) == 0:
            logger.warning(f"No compliance analyses found for document {document_id}")
            return
       
        logger.info(f"💾 Saving {len(analyses_list)} compliance results to lastchecklistitemverified for audit {audit_id}")
       
        with connection.cursor() as cursor:
            # Get audit's policy and subpolicy IDs if not provided
            if not policy_id or not subpolicy_id:
                cursor.execute("""
                    SELECT PolicyId, SubPolicyId
                    FROM audit
                    WHERE AuditId = %s
                """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
                audit_row = cursor.fetchone()
               
                if audit_row:
                    if not policy_id:
                        policy_id = audit_row[0]
                    if not subpolicy_id:
                        subpolicy_id = audit_row[1]
           
            saved_count = 0
            updated_count = 0
            
            # Check if this is an AI audit (only check once, not per compliance)
            cursor.execute("SELECT AuditType, Auditor FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
            audit_info = cursor.fetchone()
            is_ai_audit = audit_info and audit_info[0] == 'A' if audit_info else False
            auditor_id = audit_info[1] if audit_info and audit_info[1] else numeric_user_id
            
            # Get AssignedDate from audit (for creating new findings)
            cursor.execute("SELECT AssignedDate FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
            assigned_date_result = cursor.fetchone()
            assigned_date = assigned_date_result[0] if assigned_date_result and assigned_date_result[0] else timezone.now()
            
            # Process each compliance requirement result
            for analysis in analyses_list:
                if not isinstance(analysis, dict):
                    continue
                   
                compliance_id = analysis.get('compliance_id')
                if not compliance_id:
                    logger.warning(f"⚠️ Analysis item missing compliance_id: {analysis}")
                    continue
               
                # Map compliance_status to Complied field (0, 1, 2)
                # Check for compliance_status field first
                compliance_status = analysis.get('compliance_status', '').upper()
               
                # If compliance_status not found, check for 'status' field as fallback
                if not compliance_status:
                    status_value = analysis.get('status', '').upper()
                    if status_value in ['COMPLIANT', 'COMPLIED']:
                        compliance_status = 'COMPLIANT'
                    elif status_value in ['PARTIALLY_COMPLIANT', 'PARTIALLY_COMPLIED']:
                        compliance_status = 'PARTIALLY_COMPLIANT'
                    elif status_value in ['NON_COMPLIANT', 'NON_COMPLIED', 'NOT_COMPLIANT']:
                        compliance_status = 'NON_COMPLIANT'
               
                # Map to Complied value
                if compliance_status == 'COMPLIANT':
                    complied_value = '2'  # Fully Compliant
                elif compliance_status == 'PARTIALLY_COMPLIANT':
                    complied_value = '1'  # Partially Compliant
                elif compliance_status == 'NON_COMPLIANT':
                    complied_value = '0'  # Not Compliant
                else:
                    # Fallback: use compliance_score to determine status
                    compliance_score = analysis.get('compliance_score', 0.0)
                    if isinstance(compliance_score, str):
                        try:
                            compliance_score = float(compliance_score)
                        except:
                            compliance_score = 0.0
                   
                    if compliance_score >= 0.7:
                        complied_value = '2'  # Fully Compliant
                    elif compliance_score >= 0.4:
                        complied_value = '1'  # Partially Compliant
                    else:
                        complied_value = '0'  # Not Compliant
               
                # Get subpolicy_id and policy_id from compliance requirement if not available
                current_subpolicy_id = subpolicy_id
                current_policy_id = policy_id
               
                if not current_subpolicy_id or not current_policy_id:
                    cursor.execute("""
                        SELECT c.SubPolicyId, sp.PolicyId
                        FROM compliance c
                        JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                        WHERE c.ComplianceId = %s
                    """, [compliance_id])
                   
                    compliance_row = cursor.fetchone()
                    if compliance_row:
                        if not current_subpolicy_id:
                            current_subpolicy_id = compliance_row[0]
                        if not current_policy_id:
                            current_policy_id = compliance_row[1]
               
                if not current_subpolicy_id or not current_policy_id:
                    logger.warning(f"⚠️ Could not resolve SubPolicyId/PolicyId for compliance {compliance_id}")
                    continue
               
                # Build comments from AI analysis
                comments_parts = []
                if analysis.get('reason'):
                    comments_parts.append(str(analysis['reason']))
               
                # Add evidence found
                if analysis.get('evidence') and isinstance(analysis['evidence'], list) and len(analysis['evidence']) > 0:
                    evidence_text = ', '.join(str(e) for e in analysis['evidence'][:3])  # Limit to first 3
                    comments_parts.append(f"Evidence: {evidence_text}")
               
                # Add gaps/missing elements
                if analysis.get('missing') and isinstance(analysis['missing'], list) and len(analysis['missing']) > 0:
                    missing_text = ', '.join(str(m) for m in analysis['missing'][:3])  # Limit to first 3
                    comments_parts.append(f"Gaps: {missing_text}")
               
                # Add recommendations
                if analysis.get('recommendations') and isinstance(analysis['recommendations'], list) and len(analysis['recommendations']) > 0:
                    rec_text = ', '.join(str(r) for r in analysis['recommendations'][:2])  # Limit to first 2
                    comments_parts.append(f"Recommendations: {rec_text}")
                
                # Add compliance score and status
                compliance_score = analysis.get('compliance_score', 0.0)
                if isinstance(compliance_score, str):
                    try:
                        compliance_score = float(compliance_score)
                    except:
                        compliance_score = 0.0
                comments_parts.append(f"AI Score: {round(compliance_score * 100, 1)}% | Status: {compliance_status}")
               
                # Add document reference (or indicate combined/scheduled check when no single document)
                if document_id:
                    comments_parts.append(f"[AI Audit - Document ID: {document_id}]")
                else:
                    comments_parts.append("[AI Audit - Combined/Scheduled check]")
               
                comments = " | ".join(comments_parts)
                # Truncate if too long (some databases have limits)
                if len(comments) > 1000:
                    comments = comments[:997] + "..."
               
                # Try INSERT first; if a row with the same primary key already exists,
                # gracefully fall back to UPDATE instead of raising an IntegrityError.
                try:
                    cursor.execute(
                        """
                        INSERT INTO lastchecklistitemverified (
                            ComplianceId, SubPolicyId, PolicyId, FrameworkId,
                            Date, Time, User, Complied, Comments, Count
                        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 1)
                        """,
                        [
                            compliance_id,
                            current_subpolicy_id,
                            current_policy_id,
                            framework_id,
                            current_date,
                            current_time,
                            numeric_user_id,
                            complied_value,
                            comments,
                        ],
                    )
                    saved_count += 1
                    logger.info(
                        f"✅ Inserted into lastchecklistitemverified for compliance {compliance_id} (Complied={complied_value})"
                    )
                except IntegrityError:
                    # Row already exists with this (ComplianceId, SubPolicyId, PolicyId, FrameworkId);
                    # update it and increment Count instead of erroring out.
                    cursor.execute(
                        """
                        UPDATE lastchecklistitemverified
                        SET SubPolicyId = %s,
                            PolicyId = %s,
                            FrameworkId = %s,
                            Date = %s,
                            Time = %s,
                            User = %s,
                            Complied = %s,
                            Comments = %s,
                            Count = IFNULL(Count, 0) + 1
                        WHERE ComplianceId = %s
                          AND SubPolicyId = %s
                          AND PolicyId = %s
                          AND (FrameworkId = %s OR FrameworkId IS NULL)
                        """,
                        [
                        current_subpolicy_id,
                        current_policy_id,
                        framework_id,
                        current_date,
                        current_time,
                        numeric_user_id,
                        complied_value,
                        comments,
                        compliance_id,
                        current_subpolicy_id,
                        current_policy_id,
                        framework_id,
                        ],
                    )
                    updated_count += 1
                    logger.info(
                        f"✅ Updated existing lastchecklistitemverified for compliance {compliance_id} (Complied={complied_value})"
                    )
                
                # For AI audits, also create/update audit findings when compliance check completes
                try:
                    if is_ai_audit:
                        # Get compliance description from database
                        cursor.execute("""
                            SELECT ComplianceItemDescription, ComplianceTitle 
                            FROM compliance 
                            WHERE ComplianceId = %s
                        """, [compliance_id])
                        compliance_row = cursor.fetchone()
                        compliance_description = ''
                        if compliance_row:
                            compliance_description = compliance_row[0] or compliance_row[1] or f"Compliance ID {compliance_id}"
                        else:
                            compliance_description = f"Compliance ID {compliance_id}"
                        
                        # Get audit context for AI generation
                        audit_context_data = None
                        try:
                            cursor.execute("""
                                SELECT Title, Objective, Scope, BusinessUnit, AuditType, DueDate
                                FROM audit 
                                WHERE AuditId = %s
                            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
                            audit_row = cursor.fetchone()
                            if audit_row:
                                audit_context_data = {
                                    'title': audit_row[0] or '',
                                    'objective': audit_row[1] or '',
                                    'scope': audit_row[2] or '',
                                    'business_unit': audit_row[3] or '',
                                    'audit_type': audit_row[4] or '',
                                    'due_date': str(audit_row[5]) if audit_row[5] else ''
                                }
                        except Exception as ctx_err:
                            logger.warning(f"⚠️ Could not fetch audit context: {ctx_err}")
                        
                        # Generate comprehensive audit finding fields using AI
                        logger.info(f"🤖 Generating comprehensive audit finding fields for compliance {compliance_id} using AI...")
                        try:
                            comprehensive_fields = generate_comprehensive_audit_finding_fields(
                                analysis=analysis,
                                compliance_description=compliance_description,
                                audit_context=audit_context_data
                            )
                            logger.info(f"✅ AI generation completed for compliance {compliance_id}. Fields: {list(comprehensive_fields.keys())}")
                        except Exception as ai_gen_err:
                            logger.error(f"❌ AI generation failed for compliance {compliance_id}: {ai_gen_err}")
                            import traceback
                            logger.error(f"❌ AI generation traceback: {traceback.format_exc()}")
                            # Use fallback - extract from analysis directly
                            comprehensive_fields = {
                                "HowToVerify": ' | '.join(str(s) for s in analysis.get('strengths', [])[:3]) if isinstance(analysis.get('strengths'), list) else '',
                                "Impact": ' | '.join(str(w) for w in analysis.get('weaknesses', [])[:3]) if isinstance(analysis.get('weaknesses'), list) else '',
                                "Recommendation": ' | '.join(str(r) for r in analysis.get('recommendations', [])[:5]) if isinstance(analysis.get('recommendations'), list) else '',
                                "DetailsOfFinding": f"Status: {analysis.get('compliance_status', 'UNKNOWN')}, Score: {analysis.get('compliance_score', 0):.2f}",
                                "MajorMinor": "major" if analysis.get('risk_level', '').upper() == 'HIGH' or analysis.get('compliance_score', 0) < 0.4 else "minor",
                                "SeverityRating": int(max(0, min(10, (1.0 - analysis.get('compliance_score', 0.5)) * 10))),
                                "PredictiveRisks": json.dumps({"risk_level": analysis.get('risk_level', 'MEDIUM'), "compliance_score": analysis.get('compliance_score', 0)}),
                                "CorrectiveActions": json.dumps({"recommendations": analysis.get('recommendations', [])[:5]}),
                                "UnderlyingCause": ' | '.join(str(w) for w in analysis.get('weaknesses', [])[:3]) if isinstance(analysis.get('weaknesses'), list) else '',
                                "WhyToVerify": analysis.get('strengths', [])[0] if isinstance(analysis.get('strengths'), list) and len(analysis.get('strengths', [])) > 0 else '',
                                "WhatToVerify": ' | '.join(str(m) for m in analysis.get('missing', [])[:3]) if isinstance(analysis.get('missing'), list) else '',
                                "SuggestedActionPlan": ' | '.join(str(r) for r in analysis.get('recommendations', [])[:5]) if isinstance(analysis.get('recommendations'), list) else ''
                            }
                            logger.info(f"✅ Using fallback fields for compliance {compliance_id}")
                        
                        # Extract evidence from analysis
                        evidence_text = ''
                        if 'evidence' in analysis:
                            evidence = analysis['evidence']
                            if isinstance(evidence, str):
                                evidence_text = evidence
                            elif isinstance(evidence, list):
                                evidence_text = ' | '.join(str(e) for e in evidence[:10])
                            elif isinstance(evidence, dict):
                                evidence_text = str(evidence)
                        elif 'evidence_found' in analysis:
                            evidence_text = str(analysis['evidence_found'])
                        if len(evidence_text) > 2000:
                            evidence_text = evidence_text[:1997] + "..."
                        
                        # Use AI-generated fields - ensure all are strings, not dicts/lists
                        def safe_string(value, max_len=2000):
                            """Convert value to string safely, handling dicts/lists"""
                            if value is None:
                                return ''
                            if isinstance(value, (dict, list)):
                                # Convert to JSON string
                                json_str = json.dumps(value)
                                return json_str[:max_len] if len(json_str) > max_len else json_str
                            elif isinstance(value, str):
                                return value[:max_len] if len(value) > max_len else value
                            else:
                                return str(value)[:max_len] if len(str(value)) > max_len else str(value)
                        
                        how_to_verify = safe_string(comprehensive_fields.get('HowToVerify'))
                        impact = safe_string(comprehensive_fields.get('Impact'))
                        recommendation = safe_string(comprehensive_fields.get('Recommendation'))
                        details_of_finding = safe_string(comprehensive_fields.get('DetailsOfFinding'))
                        # MajorMinor must be 'major' or 'minor' (not '0', '1', '2')
                        major_minor_raw = comprehensive_fields.get('MajorMinor')
                        if major_minor_raw:
                            major_minor_str = str(major_minor_raw).lower().strip()
                            # Convert old numeric values to text
                            if major_minor_str in ['0', '1', '2']:
                                if major_minor_str == '1':
                                    major_minor = 'major'
                                elif major_minor_str == '0':
                                    major_minor = 'minor'
                                else:
                                    major_minor = 'minor'  # Default for '2' (Not Applicable)
                            elif major_minor_str in ['major', 'minor']:
                                major_minor = major_minor_str
                            else:
                                # Try to infer from risk level or compliance score
                                risk_level = analysis.get('risk_level', '').upper()
                                compliance_score = analysis.get('compliance_score', 0.5)
                                if risk_level == 'HIGH' or compliance_score < 0.4:
                                    major_minor = 'major'
                                else:
                                    major_minor = 'minor'
                        else:
                            # Infer from risk level or compliance score
                            risk_level = analysis.get('risk_level', '').upper()
                            compliance_score = analysis.get('compliance_score', 0.5)
                            if risk_level == 'HIGH' or compliance_score < 0.4:
                                major_minor = 'major'
                            else:
                                major_minor = 'minor'
                        
                        # SeverityRating must be INT (0-10) or None
                        severity_rating_raw = comprehensive_fields.get('SeverityRating')
                        if severity_rating_raw is not None:
                            try:
                                severity_rating = float(severity_rating_raw)
                                # Convert from 0-100 scale to 0-10 scale if needed
                                if severity_rating > 10:
                                    severity_rating = severity_rating / 10.0
                                # Ensure it's within valid range (0-10)
                                severity_rating = int(max(0, min(10, severity_rating)))
                            except (ValueError, TypeError):
                                # If conversion fails, calculate from compliance_score
                                compliance_score = analysis.get('compliance_score', 0.5)
                                severity_rating = int(max(0, min(10, (1.0 - compliance_score) * 10)))
                        else:
                            # Calculate from compliance_score if not provided
                            compliance_score = analysis.get('compliance_score', 0.5)
                            severity_rating = int(max(0, min(10, (1.0 - compliance_score) * 10)))
                        underlying_cause = safe_string(comprehensive_fields.get('UnderlyingCause'))
                        why_to_verify = safe_string(comprehensive_fields.get('WhyToVerify'))
                        what_to_verify = safe_string(comprehensive_fields.get('WhatToVerify'))
                        suggested_action_plan = safe_string(comprehensive_fields.get('SuggestedActionPlan'))
                        
                        # Ensure JSON fields are properly formatted as strings
                        # Convert dict/list to JSON string if needed
                        predictive_risks_raw = comprehensive_fields.get('PredictiveRisks')
                        if predictive_risks_raw is not None:
                            if isinstance(predictive_risks_raw, (dict, list)):
                                predictive_risks = json.dumps(predictive_risks_raw)
                            elif isinstance(predictive_risks_raw, str):
                                predictive_risks = predictive_risks_raw
                            else:
                                predictive_risks = json.dumps(str(predictive_risks_raw))
                        else:
                            predictive_risks = None
                        
                        corrective_actions_raw = comprehensive_fields.get('CorrectiveActions')
                        if corrective_actions_raw is not None:
                            if isinstance(corrective_actions_raw, (dict, list)):
                                corrective_actions = json.dumps(corrective_actions_raw)
                            elif isinstance(corrective_actions_raw, str):
                                corrective_actions = corrective_actions_raw
                            else:
                                corrective_actions = json.dumps(str(corrective_actions_raw))
                        else:
                            corrective_actions = None
                        
                        # Final safety check - ensure NO dicts/lists are passed to SQL
                        # Convert any remaining dicts/lists to JSON strings
                        def ensure_string_for_sql(value):
                            """Final check to ensure value is a string (or None) for SQL"""
                            if value is None:
                                return None
                            if isinstance(value, (dict, list)):
                                return json.dumps(value)
                            return str(value) if not isinstance(value, str) else value
                        
                        # Re-check all fields one more time
                        how_to_verify = ensure_string_for_sql(how_to_verify) or ''
                        impact = ensure_string_for_sql(impact) or ''
                        recommendation = ensure_string_for_sql(recommendation) or ''
                        details_of_finding = ensure_string_for_sql(details_of_finding) or ''
                        underlying_cause = ensure_string_for_sql(underlying_cause) or ''
                        why_to_verify = ensure_string_for_sql(why_to_verify) or ''
                        what_to_verify = ensure_string_for_sql(what_to_verify) or ''
                        suggested_action_plan = ensure_string_for_sql(suggested_action_plan) or ''
                        evidence_text = ensure_string_for_sql(evidence_text) or ''
                        comments = ensure_string_for_sql(comments) or ''
                        
                        logger.info(f"✅ JSON fields formatted - PredictiveRisks type: {type(predictive_risks)}, CorrectiveActions type: {type(corrective_actions)}")
                        logger.info(f"✅ All fields converted to strings - Recommendation type: {type(recommendation)}, Impact type: {type(impact)}")
                        
                        # Check if audit finding already exists
                        cursor.execute("""
                            SELECT af.AuditFindingsId FROM audit_findings af
                            JOIN audit a ON af.AuditId = a.AuditId
                            WHERE af.AuditId = %s AND a.TenantId = %s AND af.ComplianceId = %s
                        """, [int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id, compliance_id])
                        
                        existing_finding = cursor.fetchone()
                        
                        # Map complied_value to Check field
                        # 0 = Not Compliant, 1 = Partially Compliant, 2 = Fully Compliant
                        # Ensure check_value is always numeric string ('0', '1', or '2')
                        if isinstance(complied_value, str):
                            # If it's already a string, ensure it's a valid numeric value
                            if complied_value in ['0', '1', '2', '3']:
                                check_value = complied_value
                            elif complied_value.upper() in ['NON_COMPLIANT', 'NOT_COMPLIANT', 'NOT COMPLIANT']:
                                check_value = '0'
                            elif complied_value.upper() in ['PARTIALLY_COMPLIANT', 'PARTIALLY COMPLIANT']:
                                check_value = '1'
                            elif complied_value.upper() in ['COMPLIANT', 'FULLY_COMPLIANT', 'FULLY COMPLIANT']:
                                check_value = '2'
                            else:
                                # Fallback: try to convert to int and back to string
                                try:
                                    int_val = int(float(complied_value))
                                    check_value = str(max(0, min(3, int_val)))  # Clamp between 0-3
                                except (ValueError, TypeError):
                                    check_value = '0'  # Default to Not Compliant
                        elif isinstance(complied_value, (int, float)):
                            # If it's a number, convert to string and clamp
                            int_val = int(complied_value)
                            check_value = str(max(0, min(3, int_val)))  # Clamp between 0-3
                        else:
                            # Unknown type, default to Not Compliant
                            logger.warning(f"⚠️ Unknown complied_value type: {type(complied_value)}, value: {complied_value}. Defaulting to '0'")
                            check_value = '0'
                        
                        # Final validation: ensure check_value is exactly '0', '1', '2', or '3'
                        if check_value not in ['0', '1', '2', '3']:
                            logger.error(f"❌ Invalid check_value '{check_value}' for compliance {compliance_id}. Defaulting to '0'")
                            check_value = '0'
                        
                        if existing_finding:
                            # Update existing finding with ALL fields
                            cursor.execute("""
                                UPDATE audit_findings af
                                JOIN audit a ON af.AuditId = a.AuditId
                                SET af.`Check` = %s,
                                    af.Comments = %s,
                                    af.Evidence = %s,
                                    af.HowToVerify = %s,
                                    af.Impact = %s,
                                    af.Recommendation = %s,
                                    af.DetailsOfFinding = %s,
                                    af.MajorMinor = %s,
                                    af.SeverityRating = %s,
                                    af.PredictiveRisks = %s,
                                    af.CorrectiveActions = %s,
                                    af.UnderlyingCause = %s,
                                    af.WhyToVerify = %s,
                                    af.WhatToVerify = %s,
                                    af.SuggestedActionPlan = %s,
                                    af.CheckedDate = NOW()
                                WHERE af.AuditId = %s AND a.TenantId = %s AND af.ComplianceId = %s
                            """, [
                                check_value,
                                comments[:1000] if len(comments) > 1000 else comments,
                                evidence_text,
                                how_to_verify,
                                impact,
                                recommendation,
                                details_of_finding,
                                major_minor,
                                severity_rating,
                                predictive_risks,
                                corrective_actions,
                                underlying_cause,
                                why_to_verify,
                                what_to_verify,
                                suggested_action_plan,
                                int(audit_id) if str(audit_id).isdigit() else audit_id,
                                tenant_id,
                                compliance_id
                            ])
                            logger.info(f"✅ Updated audit finding (ID: {existing_finding[0]}) for compliance {compliance_id} with ALL AI-generated fields")
                        else:
                            # Create new finding with ALL fields
                            cursor.execute("""
                                INSERT INTO audit_findings (
                                    AuditId, ComplianceId, UserId, Evidence, 
                                    `Check`, Comments, MajorMinor, AssignedDate, FrameworkId, ReviewRejected,
                                    HowToVerify, Impact, Recommendation, DetailsOfFinding,
                                    SeverityRating, PredictiveRisks, CorrectiveActions,
                                    UnderlyingCause, WhyToVerify, WhatToVerify, SuggestedActionPlan,
                                    CheckedDate
                                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, NOW())
                            """, [
                                int(audit_id) if str(audit_id).isdigit() else audit_id,
                                compliance_id,
                                auditor_id,
                                evidence_text,
                                check_value,
                                comments[:1000] if len(comments) > 1000 else comments,
                                major_minor,
                                assigned_date,
                                framework_id,
                                0,  # ReviewRejected
                                how_to_verify,
                                impact,
                                recommendation,
                                details_of_finding,
                                severity_rating,
                                predictive_risks,
                                corrective_actions,
                                underlying_cause,
                                why_to_verify,
                                what_to_verify,
                                suggested_action_plan
                            ])
                            finding_id = cursor.lastrowid
                            logger.info(f"✅ Created audit finding (ID: {finding_id}) for compliance {compliance_id} with ALL AI-generated fields")
                            logger.info(f"✅ Inserted values - Check: {check_value}, MajorMinor: {major_minor}, SeverityRating: {severity_rating}")
                except Exception as finding_err:
                    import traceback
                    logger.error(f"❌ Error creating/updating audit finding for compliance {compliance_id} in audit {audit_id}: {finding_err}")
                    logger.error(f"❌ Traceback: {traceback.format_exc()}")
                    # Don't re-raise - continue with other compliances
        
        logger.info(f"✅ Saved AI compliance results: {saved_count} inserted, {updated_count} updated in lastchecklistitemverified")
       
    except Exception as e:
        logger.error(f"❌ Error saving to lastchecklistitemverified: {e}")
        import traceback
        logger.error(traceback.format_exc())
 
 
@method_decorator(csrf_exempt, name='dispatch')
class AIAuditDocumentsView(View):
    """Get uploaded documents for an audit"""
    
    def get(self, request, audit_id):
        """Get all documents for an audit"""
        try:
            logger.info(f"📋 Getting documents for audit {audit_id}")
            
            # Check authentication using JWT
            from ...rbac.utils import RBACUtils
            user_id = RBACUtils.get_user_id_from_request(request)
            
            if not user_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Authentication required'
                }, status=401)
            framework_id = None
            with connection.cursor() as cursor:
                # Debug the audit_id parameter
                converted_audit_id = int(audit_id) if str(audit_id).isdigit() else audit_id
                logger.info(f"📋 Querying documents for audit_id: {audit_id} (type: {type(audit_id)})")
                logger.info(f"📋 Converted audit_id: {converted_audit_id} (type: {type(converted_audit_id)})")
                
                cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [converted_audit_id])
                framework_row = cursor.fetchone()
                if framework_row:
                    framework_id = framework_row[0]
                # Select with or without compliance_id based on schema (cached check)
                cols = "id, document_id, document_name, document_type, file_size, created_at, upload_status, ai_processing_status, external_source, external_id, mime_type, document_path, compliance_status, confidence_score, compliance_analyses, processing_completed_at, policy_id, subpolicy_id"
                if _check_ai_audit_data_has_compliance_id():
                    cols += ", compliance_id"
                cursor.execute(f"""
                    SELECT {cols}
                    FROM ai_audit_data 
                    WHERE audit_id = %s 
                      AND (external_source != 'database_record' AND document_type != 'db_record')
                """, [converted_audit_id])
                
                rows = cursor.fetchall()
                columns = [col[0] for col in cursor.description]
                logger.info(f"📋 Raw SQL result rows: {len(rows)}")
                
                # Sort in Python to avoid MySQL sort memory issues
                # Sort by created_at DESC (newest first)
                created_at_idx = columns.index('created_at') if 'created_at' in columns else None
                if created_at_idx is not None:
                    rows = sorted(rows, key=lambda x: x[created_at_idx] if x[created_at_idx] else None, reverse=True)

                # Also load the audit's policy/subpolicy names to show mapping
                policy_name = None
                subpolicy_name = None
                policy_id_from_audit = None
                subpolicy_id_from_audit = None
                try:
                    cursor.execute(
                        """
                        SELECT p.PolicyName, sp.SubPolicyName, a.PolicyId, a.SubPolicyId
                        FROM audit a
                        LEFT JOIN policies p ON a.PolicyId = p.PolicyId
                        LEFT JOIN subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                        WHERE a.AuditId = %s
                        """,
                        [converted_audit_id]
                    )
                    row = cursor.fetchone()
                    if row:
                        from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value as _decrypt
                        policy_name = _decrypt(row[0]) if row[0] else row[0]
                        subpolicy_name = _decrypt(row[1]) if row[1] else row[1]
                        policy_id_from_audit, subpolicy_id_from_audit = row[2], row[3]
                except Exception as e:
                    logger.warning(f"ℹ️ Could not fetch policy/subpolicy names for audit {audit_id}: {e}")
                
            # OPTIMIZATION: Fetch all unique policy/subpolicy names in bulk to avoid N+1 queries
            unique_policy_ids = set()
            unique_subpolicy_ids = set()
            compliance_ids_need_resolution = set()
            for row in rows:
                doc_dict = dict(zip(columns, row))
                if doc_dict.get('policy_id'):
                    unique_policy_ids.add(doc_dict.get('policy_id'))
                if doc_dict.get('subpolicy_id'):
                    unique_subpolicy_ids.add(doc_dict.get('subpolicy_id'))
                # Documents with compliance_id but missing policy_id/subpolicy_id (e.g. from scheduled run) need names resolved from compliance
                cid = doc_dict.get('compliance_id') if 'compliance_id' in doc_dict else None
                if cid and (not doc_dict.get('policy_id') or not doc_dict.get('subpolicy_id')):
                    compliance_ids_need_resolution.add(cid)

            # Resolve policy_id/subpolicy_id from compliance for docs that have compliance_id but no policy/subpolicy
            compliance_to_policy_subpolicy = {}
            if compliance_ids_need_resolution:
                try:
                    with connection.cursor() as res_cursor:
                        placeholders = ','.join(['%s'] * len(compliance_ids_need_resolution))
                        res_cursor.execute(
                            f"""
                            SELECT c.ComplianceId, c.SubPolicyId, sp.PolicyId
                            FROM compliance c
                            JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                            WHERE c.ComplianceId IN ({placeholders})
                            """,
                            list(compliance_ids_need_resolution)
                        )
                        for res_row in res_cursor.fetchall():
                            cid, subpol_id, pol_id = res_row[0], res_row[1], res_row[2]
                            compliance_to_policy_subpolicy[cid] = (pol_id, subpol_id)
                            if pol_id:
                                unique_policy_ids.add(pol_id)
                            if subpol_id:
                                unique_subpolicy_ids.add(subpol_id)
                except Exception as e:
                    logger.warning(f"⚠️ Could not resolve policy/subpolicy from compliance: {e}")

            # Bulk fetch policy names (decrypt PolicyName - stored encrypted via EncryptedFieldsMixin)
            from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
            policy_names_map = {}
            if unique_policy_ids:
                try:
                    with connection.cursor() as bulk_cursor:
                        placeholders = ','.join(['%s'] * len(unique_policy_ids))
                        bulk_cursor.execute(
                            f"SELECT PolicyId, PolicyName FROM policies WHERE PolicyId IN ({placeholders})",
                            list(unique_policy_ids)
                        )
                        for policy_row in bulk_cursor.fetchall():
                            raw_name = policy_row[1]
                            policy_names_map[policy_row[0]] = decrypt_any_encrypted_value(raw_name) if raw_name else raw_name
                except Exception as e:
                    logger.warning(f"⚠️ Could not bulk fetch policy names: {e}")

            # Bulk fetch subpolicy names (decrypt SubPolicyName - stored encrypted via EncryptedFieldsMixin)
            subpolicy_names_map = {}
            if unique_subpolicy_ids:
                try:
                    with connection.cursor() as bulk_cursor:
                        placeholders = ','.join(['%s'] * len(unique_subpolicy_ids))
                        bulk_cursor.execute(
                            f"SELECT SubPolicyId, SubPolicyName FROM subpolicies WHERE SubPolicyId IN ({placeholders})",
                            list(unique_subpolicy_ids)
                        )
                        for subpolicy_row in bulk_cursor.fetchall():
                            raw_name = subpolicy_row[1]
                            subpolicy_names_map[subpolicy_row[0]] = decrypt_any_encrypted_value(raw_name) if raw_name else raw_name
                except Exception as e:
                    logger.warning(f"⚠️ Could not bulk fetch subpolicy names: {e}")
            
            documents = []
            for row in rows:
                doc_dict = dict(zip(columns, row))
                
                # Use bulk-fetched policy/sub-policy names (no additional queries per document)
                doc_policy_name = policy_names_map.get(doc_dict.get('policy_id')) if doc_dict.get('policy_id') else None
                doc_subpolicy_name = subpolicy_names_map.get(doc_dict.get('subpolicy_id')) if doc_dict.get('subpolicy_id') else None
                # When document has compliance_id but no policy/subpolicy names (e.g. scheduled run), use resolved IDs
                if (not doc_policy_name or not doc_subpolicy_name) and doc_dict.get('compliance_id'):
                    resolved = compliance_to_policy_subpolicy.get(doc_dict.get('compliance_id'))
                    if resolved:
                        pol_id, subpol_id = resolved
                        if not doc_policy_name and pol_id:
                            doc_policy_name = policy_names_map.get(pol_id)
                        if not doc_subpolicy_name and subpol_id:
                            doc_subpolicy_name = subpolicy_names_map.get(subpol_id)
                
                # Use document-specific names, fallback to audit-level names
                final_policy_name = doc_policy_name or policy_name
                final_subpolicy_name = doc_subpolicy_name or subpolicy_name
                
                # Parse compliance analyses if available
                compliance_analyses = None
                if doc_dict.get('compliance_analyses'):
                    try:
                        import json
                        from grc.utils.auto_decrypt_helper import decrypt_all_encrypted_in_dict
                        compliance_analyses = json.loads(doc_dict['compliance_analyses'])
                        # Decrypt any encrypted fields (requirement_title, compliance_title, etc.) whether list or dict
                        if compliance_analyses:
                            compliance_analyses = decrypt_all_encrypted_in_dict(compliance_analyses)
                    except (json.JSONDecodeError, TypeError):
                        compliance_analyses = None

                # NOTE: We used to auto-save compliance results to lastchecklistitemverified
                # here for EVERY document returned by this endpoint. That caused large audits
                # with many AI‑processed documents to become very slow on each refresh
                # (dozens of extra DB writes on every GET /documents/).
                # The compliance check endpoint already persists results, so we disable the
                # auto-save on read to keep the UI fast.

                # Decrypt document_name and document_path (may be stored encrypted)
                raw_doc_name = doc_dict.get('document_name')
                raw_doc_path = doc_dict.get('document_path')
                doc_name_decrypted = decrypt_any_encrypted_value(raw_doc_name) if raw_doc_name else raw_doc_name
                doc_path_decrypted = decrypt_any_encrypted_value(raw_doc_path) if raw_doc_path else raw_doc_path

                documents.append({
                        'id': doc_dict.get('id'),  # Primary key - use this for delete operations
                        'document_id': doc_dict.get('id') or doc_dict.get('document_id'),  # Use id as document_id for backward compatibility
                        'document_name': doc_name_decrypted,
                        'file_name': doc_name_decrypted,  # Frontend compatibility
                        'file_type': doc_dict.get('document_type'),
                        'file_size': doc_dict.get('file_size'),
                        'uploaded_date': doc_dict.get('created_at').isoformat() if doc_dict.get('created_at') else None,
                        'upload_status': doc_dict.get('upload_status'),
                        'processing_status': doc_dict.get('ai_processing_status'),
                        'ai_processing_status': doc_dict.get('ai_processing_status'),  # Frontend compatibility
                        'external_source': doc_dict.get('external_source'),
                        'external_id': doc_dict.get('external_id'),
                        'mime_type': doc_dict.get('mime_type'),
                        'document_path': doc_path_decrypted,
                        'compliance_status': doc_dict.get('compliance_status'),
                        'confidence_score': doc_dict.get('confidence_score'),
                        'compliance_analyses': compliance_analyses,
                        'processing_completed_at': doc_dict.get('processing_completed_at').isoformat() if doc_dict.get('processing_completed_at') else None,
                        'policy_id': doc_dict.get('policy_id'),
                        'subpolicy_id': doc_dict.get('subpolicy_id'),
                        'compliance_id': doc_dict.get('compliance_id') if 'compliance_id' in doc_dict else None,  # Add compliance_id if column exists
                        # Use document-specific policy/sub-policy names
                        'policy_name': final_policy_name,
                        'subpolicy_name': final_subpolicy_name,
                        'mapped_policy': final_policy_name,  # Frontend compatibility
                        'mapped_subpolicy': final_subpolicy_name,  # Frontend compatibility
                        'processing_results': compliance_analyses,  # Use compliance analyses as processing results
                        'compliance_mapping': compliance_analyses,  # Use compliance analyses as mapping
                        'extracted_text': None  # Not implemented in this table yet
                    })
            
            logger.info(f"✅ Returning {len(documents)} documents for audit {audit_id}")
            
            return JsonResponse({
                'success': True,
                'audit_id': audit_id,
                'documents': documents,
                'total_documents': len(documents)
            })
            
        except Exception as e:
            logger.error(f"❌ Error getting documents: {e}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)


# =====================
# Compliance Results API (for scheduled/combined-check results visibility)
# =====================
@method_decorator(csrf_exempt, name='dispatch')
class AIAuditComplianceResultsView(View):
    """Get compliance-level results for an audit (from audit_findings).
    Used to show results from scheduled runs and combined checks when document-level
    results may not be visible in the Documents grid."""
    
    def get(self, request, audit_id):
        try:
            from ...rbac.utils import RBACUtils
            user_id = RBACUtils.get_user_id_from_request(request)
            if not user_id:
                return JsonResponse({'success': False, 'error': 'Authentication required'}, status=401)
            
            converted_audit_id = int(audit_id) if str(audit_id).isdigit() else audit_id
            tenant_id = get_tenant_id_from_request(request)

            # Filter by schedule scope: when schedule_id is provided, show ONLY compliances checked by that schedule
            compliance_ids_param = request.GET.get('compliance_ids', '')
            schedule_id_param = request.GET.get('schedule_id', '').strip()
            compliance_ids = [int(x) for x in compliance_ids_param.split(',') if x.strip().isdigit()]

            # If schedule_id provided, use that schedule's scope (selected_compliance_ids or framework compliances)
            if schedule_id_param and schedule_id_param.isdigit():
                with connection.cursor() as cur:
                    cur.execute(
                        "SELECT selected_compliance_ids FROM ai_audit_schedule WHERE id = %s AND AuditId = %s LIMIT 1",
                        [int(schedule_id_param), converted_audit_id]
                    )
                    row = cur.fetchone()
                if row and row[0]:
                    sel = row[0]
                    if isinstance(sel, list):
                        compliance_ids = [int(x) for x in sel if x is not None and str(x).strip().isdigit()]
                    elif isinstance(sel, str):
                        try:
                            parsed = json.loads(sel)
                            compliance_ids = [int(x) for x in (parsed if isinstance(parsed, list) else []) if x is not None and str(x).strip().isdigit()]
                        except Exception:
                            pass
                elif row:
                    # Schedule has no selected_compliance_ids - use framework compliances for audit
                    with connection.cursor() as cur2:
                        cur2.execute("""
                            SELECT DISTINCT c.ComplianceId FROM audit a
                            JOIN compliance c ON c.FrameworkId = a.FrameworkId
                            WHERE a.AuditId = %s AND a.FrameworkId IS NOT NULL
                        """, [converted_audit_id])
                        compliance_ids = [r[0] for r in cur2.fetchall()]

            with connection.cursor() as cursor:
                params = [converted_audit_id]
                tenant_clause = ""
                if tenant_id:
                    tenant_clause = " AND a.TenantId = %s"
                    params.append(tenant_id)
                compliance_filter = ""
                if compliance_ids:
                    placeholders = ",".join(["%s"] * len(compliance_ids))
                    compliance_filter = f" AND af.ComplianceId IN ({placeholders})"
                    params.extend(compliance_ids)
                cursor.execute("""
                    SELECT af.AuditFindingsId, af.ComplianceId, af.Check, af.Evidence, af.DetailsOfFinding,
                           af.HowToVerify, af.Impact, af.Recommendation, af.MajorMinor, af.Comments,
                           af.CheckedDate, c.ComplianceTitle, c.ComplianceItemDescription,
                           sp.SubPolicyName, p.PolicyName
                    FROM audit_findings af
                    JOIN compliance c ON af.ComplianceId = c.ComplianceId
                    LEFT JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                    LEFT JOIN policies p ON sp.PolicyId = p.PolicyId
                    JOIN audit a ON af.AuditId = a.AuditId
                    WHERE af.AuditId = %s
                """ + tenant_clause + compliance_filter + """
                    ORDER BY p.PolicyName, sp.SubPolicyName, c.ComplianceId
                """, params)
                rows = cursor.fetchall()
                columns = [col[0] for col in cursor.description]
            
            from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value as _decrypt_result
            results = []
            for row in rows:
                d = dict(zip(columns, row))
                comp_title = d.get('ComplianceTitle') or d.get('ComplianceItemDescription') or f"Compliance {d.get('ComplianceId')}"
                policy_name_raw = d.get('PolicyName')
                subpolicy_name_raw = d.get('SubPolicyName')
                results.append({
                    'finding_id': d.get('AuditFindingsId'),
                    'compliance_id': d.get('ComplianceId'),
                    'compliance_title': _decrypt_result(comp_title) if comp_title else comp_title,
                    'check': d.get('Check'),
                    'evidence': d.get('Evidence'),
                    'details_of_finding': d.get('DetailsOfFinding'),
                    'how_to_verify': d.get('HowToVerify'),
                    'impact': d.get('Impact'),
                    'recommendation': d.get('Recommendation'),
                    'major_minor': d.get('MajorMinor'),
                    'comments': d.get('Comments'),
                    'checked_date': (lambda v: v.isoformat() if hasattr(v, 'isoformat') else str(v) if v else None)(d.get('CheckedDate')),
                    'policy_name': _decrypt_result(policy_name_raw) if policy_name_raw else policy_name_raw,
                    'subpolicy_name': _decrypt_result(subpolicy_name_raw) if subpolicy_name_raw else subpolicy_name_raw,
                })
            
            return JsonResponse({
                'success': True,
                'audit_id': audit_id,
                'results': results,
                'total': len(results)
            })
        except Exception as e:
            logger.error(f"❌ Error getting compliance results: {e}")
            return JsonResponse({'success': False, 'error': str(e)}, status=500)


# =====================
# Mapping Update API
# =====================
@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def map_audit_document_api(request, audit_id, document_id):
    """Update the policy/sub-policy mapping for a specific uploaded document.
    MULTI-TENANCY: Only maps documents for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        policy_id = request.data.get('policy_id')
        subpolicy_id = request.data.get('subpolicy_id')

        if not policy_id and not subpolicy_id:
            return Response({
                'success': False,
                'error': 'Provide at least policy_id or subpolicy_id'
            }, status=status.HTTP_400_BAD_REQUEST)

        with connection.cursor() as cursor:
            # Verify the document belongs to this audit
            cursor.execute(
                """
                SELECT ad.document_id FROM audit_document ad
                JOIN audit a ON ad.audit_id = a.AuditId
                WHERE ad.document_id = %s AND ad.audit_id = %s AND a.TenantId = %s
                """,
                [int(document_id), int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id]
            )
            if not cursor.fetchone():
                return Response({
                    'success': False,
                    'error': 'Document not found for this audit'
                }, status=status.HTTP_404_NOT_FOUND)

            # Optional: validate subpolicy belongs to policy if both are provided
            if policy_id and subpolicy_id:
                try:
                    cursor.execute(
                        """
                        SELECT 1 FROM subpolicies WHERE SubPolicyId = %s AND PolicyId = %s
                        """,
                        [int(subpolicy_id), int(policy_id)]
                    )
                    if cursor.fetchone() is None:
                        return Response({
                            'success': False,
                            'error': 'Sub-policy does not belong to the provided policy'
                        }, status=status.HTTP_400_BAD_REQUEST)
                except Exception:
                    # If schema differs, skip strict validation
                    pass

            # Update mapping columns when available
            updated = 0
            try:
                cursor.execute(
                    """
                    UPDATE audit_document ad
                    JOIN audit a ON ad.audit_id = a.AuditId
                    SET ad.policy_id = COALESCE(%s, ad.policy_id),
                        ad.subpolicy_id = COALESCE(%s, ad.subpolicy_id)
                    WHERE ad.document_id = %s AND a.TenantId = %s
                    """,
                    [policy_id, subpolicy_id, int(document_id), tenant_id]
                )
                updated = cursor.rowcount
            except Exception as e:
                logger.warning(f"ℹ️ Mapping columns may not exist on audit_document: {e}")
                return Response({
                    'success': False,
                    'error': 'Mapping columns not available on audit_document table'
                }, status=status.HTTP_400_BAD_REQUEST)

            # Fetch human-readable names (decrypt if stored encrypted)
            from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value as _decrypt_map
            policy_name = None
            subpolicy_name = None
            try:
                if policy_id:
                    cursor.execute("SELECT PolicyName FROM policies WHERE PolicyId = %s", [int(policy_id)])
                    row = cursor.fetchone()
                    if row and row[0]:
                        policy_name = _decrypt_map(row[0])
                if subpolicy_id:
                    cursor.execute("SELECT SubPolicyName FROM subpolicies WHERE SubPolicyId = %s", [int(subpolicy_id)])
                    row = cursor.fetchone()
                    if row and row[0]:
                        subpolicy_name = _decrypt_map(row[0])
            except Exception:
                pass

        return Response({
            'success': True,
            'document_id': int(document_id),
            'audit_id': audit_id,
            'policy_id': policy_id,
            'subpolicy_id': subpolicy_id,
            'policy_name': policy_name,
            'subpolicy_name': subpolicy_name,
            'updated': updated
        })
    except Exception as e:
        logger.error(f"❌ Error updating mapping: {e}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@method_decorator(csrf_exempt, name='dispatch')
class AIAuditStatusView(View):
    """Get AI audit processing status"""
    
    def get(self, request, audit_id):
        """Get processing status for an audit"""
        try:
            logger.info(f"📊 Getting AI audit status for audit {audit_id}")
            
            # Check authentication using JWT
            from ...rbac.utils import RBACUtils
            user_id = RBACUtils.get_user_id_from_request(request)
            
            if not user_id:
                return JsonResponse({
                    'success': False,
                    'error': 'Authentication required'
                }, status=401)
            
            with connection.cursor() as cursor:
                # Get all documents from ai_audit_data table (where actual documents are stored)
                cursor.execute("""
                    SELECT document_name, file_size, ai_processing_status
                    FROM ai_audit_data 
                    WHERE audit_id = %s 
                    ORDER BY document_name, file_size
                """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
                
                all_records = cursor.fetchall()
            
            # Group by physical file (document_name + file_size) to count unique files, not mapping records
            file_groups = {}  # Key: (document_name, file_size), Value: list of statuses
            for doc_name, file_size, status in all_records:
                file_key = (doc_name or '', file_size or 0)
                if file_key not in file_groups:
                    file_groups[file_key] = []
                file_groups[file_key].append(status)
            
            # Determine status for each unique physical file
            # A file is "completed" if at least one mapping is completed
            # A file is "failed" if all mappings failed and none are completed/processing
            # A file is "processing" if any mapping is processing and none are completed
            # Otherwise it's "pending"
            status_counts = {
                'pending': 0,
                'processing': 0,
                'completed': 0,
                'failed': 0
            }
            
            for file_key, statuses in file_groups.items():
                # Determine file-level status based on its mappings
                has_completed = any(s == 'completed' for s in statuses)
                has_processing = any(s == 'processing' for s in statuses)
                has_failed = any(s == 'failed' for s in statuses)
                
                if has_completed:
                    status_counts['completed'] += 1
                elif has_processing:
                    status_counts['processing'] += 1
                elif has_failed and not has_completed and not has_processing:
                    status_counts['failed'] += 1
                else:
                    status_counts['pending'] += 1
            
            # Calculate progress
            total_documents = len(file_groups)  # Count unique physical files
            completed_documents = status_counts['completed']
            progress_percentage = (completed_documents / total_documents * 100) if total_documents > 0 else 0
            
            return JsonResponse({
                'success': True,
                'audit_id': audit_id,
                'processing_status': {
                    'total_documents': total_documents,
                    'pending': status_counts['pending'],
                    'processing': status_counts['processing'],
                    'completed': status_counts['completed'],
                    'failed': status_counts['failed'],
                    'progress_percentage': round(progress_percentage, 2)
                },
                'is_complete': status_counts['pending'] == 0 and status_counts['processing'] == 0
            })
            
        except Exception as e:
            logger.error(f"❌ Error getting AI status: {e}")
            return JsonResponse({
                'success': False,
                'error': str(e)
            }, status=500)

def trigger_audit_reprocessing_on_json_update(audit_id, framework_id):
    """
    Trigger audit re-processing when JSON files are updated.
    This marks existing documents as 'pending' so they get re-processed with new matched compliances.
    
    Args:
        audit_id: The audit ID
        framework_id: The framework ID
    
    Returns:
        dict with status and count of documents marked for re-processing
    """
    try:
        logger.info(f"🔄 Triggering audit re-processing for audit {audit_id} due to JSON update")
        
        with connection.cursor() as cursor:
            # Mark all documents in ai_audit_data for this audit as 'pending' (except database records)
            # This will cause them to be re-processed with the new matched compliances from JSON
            cursor.execute("""
                UPDATE ai_audit_data
                SET ai_processing_status = 'pending',
                    updated_at = NOW()
                WHERE audit_id = %s
                  AND (external_source != 'database_record' AND document_type != 'db_record')
                  AND ai_processing_status != 'failed'
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
            
            updated_count = cursor.rowcount
            logger.info(f"✅ Marked {updated_count} document(s) as 'pending' for re-processing in audit {audit_id}")
            
            return {
                'success': True,
                'updated_count': updated_count,
                'message': f'Marked {updated_count} document(s) for re-processing'
            }
    except Exception as e:
        logger.error(f"❌ Error triggering audit re-processing: {e}")
        import traceback
        logger.error(f"❌ Error details: {traceback.format_exc()}")
        return {
            'success': False,
            'error': str(e),
            'updated_count': 0
        }


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def start_ai_audit_processing_api(request, audit_id):
    """Start AI processing for all pending documents with real AI/ML analysis
    MULTI-TENANCY: Only processes documents for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        logger.info(f"🚀 Starting AI processing for audit {audit_id}")
        
        # Check authentication using JWT
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        logger.info(f"🚀 Processing request from user: {user_id}")
        
        # Get pending documents from new audit_document table
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT document_id, document_name, document_path, document_type, file_size, mime_type
                FROM audit_document ad
                JOIN audit a ON ad.audit_id = a.AuditId
                WHERE ad.audit_id = %s AND a.TenantId = %s AND ad.ai_processing_status = 'pending'
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id])
            
            documents = cursor.fetchall()
        
        if not documents:
            return Response({
                'success': False,
                'error': 'No pending documents to process'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        logger.info(f"🚀 Found {len(documents)} documents to process")
        
        # Update status to processing in new table
        with connection.cursor() as cursor:
            cursor.execute("""
                UPDATE audit_document 
                SET ai_processing_status = 'processing' 
                WHERE audit_id = %s AND ai_processing_status = 'pending'
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
            
            updated_count = cursor.rowcount
        
        logger.info(f"🚀 Updated {updated_count} documents to processing status")
        
        # Start AI processing for each document
        processing_results = []
        for doc in documents:
            doc_id, doc_name, file_path, doc_type, file_size, mime_type = doc
            try:
                logger.info(f"🤖 Processing document: {doc_name}")
                
                # Use MIME type instead of DocumentType for better accuracy
                actual_doc_type = mime_type if mime_type else doc_type
                logger.info(f"🔍 Using document type: {actual_doc_type}")
                
                # Process the document with AI/ML
                result = process_document_with_ai(doc_id, doc_name, file_path, actual_doc_type, audit_id)
                processing_results.append(result)
                
                # Update document status to completed (store minimal fields available)
                with connection.cursor() as cursor:
                    cursor.execute("""
                        UPDATE audit_document 
                        SET ai_processing_status = 'completed'
                        WHERE document_id = %s
                    """, [doc_id])
                
                logger.info(f"✅ Document {doc_name} processed successfully")
                
            except Exception as e:
                logger.error(f"❌ Error processing document {doc_name}: {e}")
                logger.error(f"❌ Error type: {type(e).__name__}")
                import traceback
                logger.error(f"❌ Traceback: {traceback.format_exc()}")
                
                # Mark document as failed (minimal update)
                with connection.cursor() as cursor:
                    cursor.execute("""
                        UPDATE audit_document 
                        SET ai_processing_status = 'failed'
                        WHERE document_id = %s
                    """, [doc_id])
        
        return Response({
            'success': True,
            'message': f'AI processing completed for {len(processing_results)} documents',
            'audit_id': audit_id,
            'documents_processed': len(processing_results),
            'processing_results': processing_results
        })
        
    except Exception as e:
        logger.error(f"❌ Error starting AI processing: {e}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def start_ai_audit_selective_processing_api(request, audit_id):
    """
    Start AI processing for a specific set of compliance requirements.
    - Creates new Compliance records for any custom_compliance_labels.
    - Links them to the audit via audit_findings.
    - Runs combined-evidence AI check only for the selected + newly created compliances.
    """
    tenant_id = get_tenant_id_from_request(request)

    # Get user from JWT
    try:
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
    except Exception:
        user_id = None

    if not user_id:
        return Response({
            'success': False,
            'error': 'Authentication required'
        }, status=status.HTTP_401_UNAUTHORIZED)

    # Import models lazily to avoid cycles
    from ...models import Audit, Compliance, SubPolicy
    from django.utils import timezone as _tz

    try:
        audit = Audit.objects.get(
            AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id,
            tenant_id=tenant_id
        )
    except Audit.DoesNotExist:
        return Response({
            'success': False,
            'error': f'Audit {audit_id} not found for this tenant'
        }, status=status.HTTP_404_NOT_FOUND)

    if getattr(audit, 'Status', None) == 'Completed':
        return Response({
            'success': False,
            'error': 'This audit is closed. No further evidence or processing is allowed.'
        }, status=status.HTTP_403_FORBIDDEN)

    # --- Parse selected compliance IDs ---
    raw_selected = request.data.get('selected_compliance_ids') or request.data.get('compliance_ids') or []
    selected_ids: set[int] = set()
    if isinstance(raw_selected, str):
        try:
            import json as _json
            parsed = _json.loads(raw_selected)
            if isinstance(parsed, list):
                raw_selected = parsed
        except Exception:
            pass
    if isinstance(raw_selected, (list, tuple, set)):
        for x in raw_selected:
            try:
                if x is not None and str(x).strip() != '':
                    selected_ids.add(int(x))
            except (ValueError, TypeError):
                continue

    # --- Combine compliance IDs (custom ones are now created via a separate endpoint) ---
    all_ids = sorted(selected_ids)
    if not all_ids:
        return Response({
            'success': False,
            'error': 'No compliance IDs or custom labels provided for selective processing.'
        }, status=status.HTTP_400_BAD_REQUEST)

    logger.info(
        f"🚀 Starting selective AI processing for audit {audit_id} with "
        f"{len(all_ids)} compliance IDs."
    )

    try:
        # Run combined-evidence compliance check only for the selected IDs
        result = _check_compliance_with_combined_evidence_internal(
            audit_id=audit_id,
            compliance_ids=all_ids,
            user_id=user_id
        )

        if not isinstance(result, dict) or not result.get('success'):
            return Response(result or {
                'success': False,
                'error': 'Selective AI processing failed.'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Try to update AI audit status based on processing results
        try:
            status_update = check_and_update_ai_audit_status(audit_id)
            logger.info(f"🔁 Selective processing: status update result for audit {audit_id}: {status_update}")
        except Exception as status_err:
            logger.error(f"⚠️ Selective processing: could not update audit status: {status_err}")

        payload = {
            'success': True,
            'message': 'Selective AI processing completed successfully.',
            'audit_id': audit_id,
            'selected_compliance_ids': all_ids,
        }
        # Merge any extra details from internal result
        if isinstance(result, dict):
            payload.update({k: v for k, v in result.items() if k not in payload})

        return Response(payload, status=status.HTTP_200_OK)

    except Exception as e:
        logger.error(f"❌ Error during selective AI processing: {e}")
        import traceback as _tb
        logger.error(_tb.format_exc())
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def add_custom_compliance_to_ai_audit(request, audit_id):
    """
    Immediately create a new Compliance under a specific Policy/SubPolicy for an AI audit,
    link it via audit_findings, and return the created compliance details.
    Used by the AI Audit Document Upload page when the user clicks Add.
    """
    tenant_id = get_tenant_id_from_request(request)

    # Get user from JWT
    try:
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
    except Exception:
        user_id = None

    if not user_id:
        return Response({
            'success': False,
            'error': 'Authentication required'
        }, status=status.HTTP_401_UNAUTHORIZED)

    from ...models import Audit, Compliance, SubPolicy
    from django.utils import timezone as _tz

    try:
        audit = Audit.objects.get(
            AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id,
            tenant_id=tenant_id
        )
    except Audit.DoesNotExist:
        return Response({
            'success': False,
            'error': f'Audit {audit_id} not found for this tenant'
        }, status=status.HTTP_404_NOT_FOUND)

    if getattr(audit, 'Status', None) == 'Completed':
        return Response({
            'success': False,
            'error': 'This audit is closed. No documents or evidence can be added.'
        }, status=status.HTTP_403_FORBIDDEN)

    data = request.data or {}
    try:
        policy_id = int(data.get('policy_id')) if data.get('policy_id') is not None else None
    except (TypeError, ValueError):
        policy_id = None
    try:
        subpolicy_id = int(data.get('subpolicy_id')) if data.get('subpolicy_id') is not None else None
    except (TypeError, ValueError):
        subpolicy_id = None
    name = (data.get('name') or '').strip()

    if not policy_id or not subpolicy_id:
        return Response({
            'success': False,
            'error': 'Policy and Sub Policy required',
            'details': 'Please select a Policy and Sub‑policy for the new compliance.'
        }, status=status.HTTP_400_BAD_REQUEST)

    if not name:
        return Response({
            'success': False,
            'error': 'Compliance name required'
        }, status=status.HTTP_400_BAD_REQUEST)

    # Validate that subpolicy belongs to policy and tenant
    try:
        sp = SubPolicy.objects.get(SubPolicyId=subpolicy_id, PolicyId_id=policy_id, tenant_id=tenant_id)
    except SubPolicy.DoesNotExist:
        return Response({
            'success': False,
            'error': 'Invalid policy/sub‑policy combination'
        }, status=status.HTTP_400_BAD_REQUEST)

    # Create compliance
    new_compliance = Compliance.objects.create(
        tenant_id=tenant_id,
        SubPolicy_id=sp.SubPolicyId,
        ComplianceTitle=name[:1000],
        Identifier=name[:500],
        ComplianceVersion='1.0',
        Status='Under Review',
        ActiveInactive='Active',
        PermanentTemporary='Temporary',
    )

    # Ensure audit has AssignedDate
    assigned_dt = audit.AssignedDate or _tz.now()
    if audit.AssignedDate is None:
        audit.AssignedDate = assigned_dt
        audit.save(update_fields=['AssignedDate'])

    # Link via audit_findings
    with connection.cursor() as cursor:
        cursor.execute("""
            INSERT INTO audit_findings (
                `AuditId`, `ComplianceId`, `UserId`, `Evidence`,
                `Check`, `Comments`, `MajorMinor`, `AssignedDate`,
                `FrameworkId`, `ReviewRejected`, `TenantId`
            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """, [
            audit.AuditId,
            new_compliance.ComplianceId,
            audit.Auditor_id,
            '',
            '0',
            '',
            None,
            assigned_dt,
            audit.FrameworkId_id,
            0,
            tenant_id,
        ])

    payload = {
        'success': True,
        'compliance': {
            'ComplianceId': new_compliance.ComplianceId,
            'compliance_id': new_compliance.ComplianceId,
            'ComplianceTitle': new_compliance.ComplianceTitle,
            'compliance_title': new_compliance.ComplianceTitle,
            'policy_id': policy_id,
            'policy_name': sp.Policy.PolicyName if hasattr(sp, 'Policy') else None,
            'subpolicy_id': sp.SubPolicyId,
            'subpolicy_name': sp.SubPolicyName,
        }
    }
    return Response(payload, status=status.HTTP_201_CREATED)


def process_document_with_ai(doc_id, doc_name, file_path, doc_type, audit_id):
    """Process document with real AI/ML to extract metadata and determine compliance"""
    try:
        logger.info(f"🤖 Starting AI processing for document: {doc_name}")
        
        # Get full file path - handle both relative and absolute paths
        if os.path.isabs(file_path):
            full_path = file_path
        else:
            full_path = os.path.join(settings.MEDIA_ROOT, file_path)
        
        logger.info(f"🔍 Checking file path: {full_path}")
        if not os.path.exists(full_path):
            logger.error(f"❌ File not found: {full_path}")
            raise Exception(f"File not found: {full_path}")
        
        logger.info(f"✅ File found, size: {os.path.getsize(full_path)} bytes")
        
        # Extract text from document
        full_text = extract_text_from_document(full_path, doc_type)
        
        if not full_text or len(full_text.strip()) < 50:
            raise Exception(f"Document {doc_name} has insufficient text for AI analysis")
        
        # Use structured compliance checking
        # TODO: Implement structured_compliance_checker module
        # from .structured_compliance_checker import check_document_structured_compliance
        
        # Check compliance against specific policies/subpolicies
        # compliance_result = check_document_structured_compliance(audit_id, doc_id, full_text, doc_name)
        
        # if not compliance_result['success']:
        #     raise Exception(f"Compliance checking failed: {compliance_result['error']}")
        
        # Extract AI analysis and compliance matrix
        # TODO: Implement structured_compliance_checker module
        # ai_analysis = compliance_result['ai_analysis']
        # compliance_matrix = compliance_result['compliance_matrix']
        # compliance_summary = compliance_result['summary']
        
        # Default values until structured_compliance_checker is implemented
        ai_analysis = {}
        compliance_matrix = {}
        compliance_summary = {}
        
        # Generate AI-powered recommendations
        # TODO: Implement generate_ai_recommendations function
        # ai_recommendations = generate_ai_recommendations(full_text, ai_analysis)
        ai_recommendations = []
        
        # Extract metadata using AI
        # TODO: Implement extract_ai_metadata function
        # metadata = extract_ai_metadata(full_path, doc_type, full_text, ai_analysis)
        metadata = {}
        
        result = {
            'document_id': doc_id,
            'document_name': doc_name,
            'processing_results': {
                'text_length': len(full_text),
                'metadata': metadata,
                'processing_timestamp': datetime.now().isoformat(),
                'ai_analysis': ai_analysis,
                'data_quality_score': ai_analysis.get('data_quality_score', 0.0) if ai_analysis else 0.0
            },
            'compliance_mapping': ai_analysis.get('compliance_analysis', {}) if ai_analysis else {},
            'compliance_matrix': compliance_matrix,  # Detailed compliance matrix
            'compliance_summary': compliance_summary,  # Overall compliance summary
            'extracted_text': full_text[:1000] + '...' if len(full_text) > 1000 else full_text,
            'ai_recommendations': ai_recommendations,
            'compliance_status': compliance_summary.get('overall_status', 'unknown') if compliance_summary else 'unknown',
            'risk_level': ai_analysis.get('risk_level', 'medium') if ai_analysis else 'medium',
            'confidence_score': ai_analysis.get('confidence_score', 0.0) if ai_analysis else 0.0
        }
        
        logger.info(f"✅ AI processing completed for {doc_name}")
        return result
        
    except Exception as e:
        logger.error(f"❌ AI processing failed for {doc_name}: {e}")
        raise


def _infer_doc_type_from_path(file_path):
    """Infer MIME-like doc type from file extension when DB type is missing/wrong."""
    if not file_path:
        return None
    ext = (os.path.splitext(file_path)[1] or "").lower()
    if ext == '.pdf':
        return 'application/pdf'
    if ext in ('.xlsx', '.xlsm', '.xltx', '.xltm'):
        return 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    if ext in ('.xls', '.xlt'):
        return 'application/vnd.ms-excel'
    if ext in ('.docx',):
        return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    if ext in ('.doc',):
        return 'application/msword'
    if ext in ('.txt',):
        return 'text/plain'
    if ext in ('.xbrl', '.xml'):
        # Treat XBRL/XML as plain text so tags and values are visible to AI
        return 'text/plain'
    return None


def extract_text_from_document(file_path, doc_type):
    """Extract text from various document types using AI/ML.
    Falls back to extension-based detection if doc_type is missing or wrong (e.g. PDF stored as Excel).
    """
    try:
        # Prefer extension when stored type might be wrong (e.g. PDF sent to Excel extractor)
        inferred = _infer_doc_type_from_path(file_path)
        effective = (doc_type if doc_type in (
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'application/vnd.ms-excel',
            'text/plain',
            'application/xml',  # XBRL and other XML: extract as text so AI sees tags and values
        ) else None) or inferred
        # Treat application/xml (e.g. XBRL) as text for extraction
        if effective == 'application/xml':
            effective = 'text/plain'

        if effective == 'application/pdf':
            return extract_text_from_pdf(file_path)
        if effective in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            return extract_text_from_word(file_path)
        if effective in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', 'application/vnd.ms-excel']:
            return extract_text_from_excel(file_path)
        if effective == 'text/plain':
            return extract_text_from_txt(file_path)
        return f"Document type {doc_type} not supported for text extraction"
    except Exception as e:
        logger.error(f"Error extracting text: {e}")
        return f"Error extracting text: {str(e)}"


def extract_text_from_pdf(file_path):
    """Extract text from PDF using PyPDF2"""
    try:
        import PyPDF2
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return f"PDF extraction failed: {str(e)}"


def extract_text_from_word(file_path):
    """Extract text from Word document using python-docx"""
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Word extraction error: {e}")
        return f"Word extraction failed: {str(e)}"


def _json_serial(val):
    """Make Excel cell values (e.g. datetime, date) JSON-serializable."""
    from datetime import datetime, date
    if isinstance(val, (datetime, date)):
        return val.isoformat()
    return val


def extract_text_from_excel(file_path):
    """Extract text from Excel document using openpyxl"""
    try:
        import openpyxl
        workbook = openpyxl.load_workbook(file_path)
        text = ""
        inferred = { 'sheets': [] }
        for sheet_name in workbook.sheetnames[:3]:
            sheet = workbook[sheet_name]
            rows_iter = list(sheet.iter_rows(values_only=True))
            headers = []
            sample = []
            if rows_iter:
                headers = [str(h).strip() if h is not None else '' for h in (rows_iter[0] or [])]
                for r in rows_iter[1:11]:
                    row_dict = {}
                    for i in range(max(len(headers), len(r or []))):
                        key = headers[i] if i < len(headers) else f'col_{i+1}'
                        raw = r[i] if i < len(r) else None
                        row_dict[key] = _json_serial(raw)
                    sample.append(row_dict)
            inferred['sheets'].append({ 'name': sheet_name, 'headers': headers, 'sample_rows': sample })
            # Also flatten into text for generic analysis
            if headers:
                text += ' | '.join(headers) + '\n'
            for r in rows_iter[1:101]:
                line = ' | '.join([str(v) for v in (r or []) if v is not None])
                if line:
                    text += line + '\n'
        # Attach inferred schema JSON at the end (truncated in prompt if needed)
        text += f"\n__INFERRED_SCHEMA_START__\n{json.dumps(inferred, ensure_ascii=False)}\n__INFERRED_SCHEMA_END__\n"
        return text
    except Exception as e:
        logger.error(f"Excel extraction error: {e}")
        return f"Excel extraction failed: {str(e)}"


def extract_text_from_txt(file_path):
    """Extract text from plain text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return f"TXT extraction failed: {str(e)}"


def generate_comprehensive_audit_finding_fields(analysis, compliance_description, audit_context=None):
    """
    Generate comprehensive audit finding fields using AI based on compliance analysis results.
    
    Args:
        analysis: The compliance analysis result dict containing evidence, strengths, weaknesses, etc.
        compliance_description: The compliance requirement description
        audit_context: Optional audit context (objective, scope, etc.)
    
    Returns:
        dict: Comprehensive audit finding fields
    """
    try:
        # Extract key information from analysis
        evidence = analysis.get('evidence', [])
        strengths = analysis.get('strengths', [])
        weaknesses = analysis.get('weaknesses', [])
        missing = analysis.get('missing', [])
        recommendations = analysis.get('recommendations', [])
        compliance_status = analysis.get('compliance_status', 'UNKNOWN')
        compliance_score = analysis.get('compliance_score', 0.0)
        risk_level = analysis.get('risk_level', 'MEDIUM')
        
        # Build context for AI prompt
        evidence_text = ' | '.join(str(e) for e in evidence[:10]) if isinstance(evidence, list) else str(evidence)
        strengths_text = ' | '.join(str(s) for s in strengths[:5]) if isinstance(strengths, list) else str(strengths)
        weaknesses_text = ' | '.join(str(w) for w in weaknesses[:5]) if isinstance(weaknesses, list) else str(weaknesses)
        missing_text = ' | '.join(str(m) for m in missing[:7]) if isinstance(missing, list) else str(missing)
        recommendations_text = ' | '.join(str(r) for r in recommendations[:5]) if isinstance(recommendations, list) else str(recommendations)
        
        audit_objective = audit_context.get('objective', 'N/A') if audit_context else 'N/A'
        audit_scope = audit_context.get('scope', 'N/A') if audit_context else 'N/A'
        
        # Create comprehensive AI prompt
        prompt = f"""You are an expert GRC auditor. Based on the compliance analysis results below, generate comprehensive audit finding fields.

=== COMPLIANCE REQUIREMENT ===
{compliance_description}

=== AUDIT CONTEXT ===
Audit Objective: {audit_objective}
Audit Scope: {audit_scope}
Compliance Status: {compliance_status}
Compliance Score: {compliance_score:.2f}
Risk Level: {risk_level}

=== ANALYSIS RESULTS ===
Evidence Found: {evidence_text[:1000] if len(evidence_text) > 1000 else evidence_text}
Strengths: {strengths_text[:1000] if len(strengths_text) > 1000 else strengths_text}
Weaknesses: {weaknesses_text[:1000] if len(weaknesses_text) > 1000 else weaknesses_text}
Missing Elements: {missing_text[:1000] if len(missing_text) > 1000 else missing_text}
Recommendations: {recommendations_text[:1000] if len(recommendations_text) > 1000 else recommendations_text}

=== YOUR TASK ===
Generate comprehensive audit finding fields in JSON format. Return ONLY valid JSON (no markdown, no explanations).

Required fields:
1. **HowToVerify**: Step-by-step instructions on how to verify this compliance requirement. Include specific methods, tools, or procedures.
2. **Impact**: Detailed description of the business impact if this compliance requirement is not met. Include financial, operational, reputational, and regulatory impacts.
3. **Recommendation**: Actionable recommendations to improve compliance. Be specific and prioritized.
4. **DetailsOfFinding**: Comprehensive details of the audit finding. Combine strengths, weaknesses, and missing elements into a cohesive narrative.
5. **MajorMinor**: Determine if this is a "major" or "minor" finding based on risk level and compliance score.
6. **SeverityRating**: Numeric severity rating from 0-10 based on compliance score and risk level (0 = no risk, 10 = critical risk).
7. **PredictiveRisks**: JSON object containing identified risks, their likelihood, and potential impact.
8. **CorrectiveActions**: JSON object containing prioritized corrective actions with timelines and responsible parties.
9. **UnderlyingCause**: Root cause analysis explaining why the compliance gap exists.
10. **WhyToVerify**: Explanation of why this requirement needs to be verified and its importance.
11. **WhatToVerify**: Specific items, processes, or controls that need to be verified.
12. **SuggestedActionPlan**: Detailed action plan with steps, timelines, and milestones to achieve compliance.

Return JSON in this exact format:
{{
    "HowToVerify": "Step-by-step verification instructions...",
    "Impact": "Detailed business impact description...",
    "Recommendation": "Actionable recommendations...",
    "DetailsOfFinding": "Comprehensive finding details...",
    "MajorMinor": "major",
    "SeverityRating": 7,
    "PredictiveRisks": {{
        "identified_risks": ["Risk 1", "Risk 2"],
        "risk_level": "HIGH",
        "likelihood": 0.7,
        "impact": "High financial and regulatory impact"
    }},
    "CorrectiveActions": {{
        "actions": [
            {{"action": "Action 1", "priority": "High", "timeline": "30 days", "responsible": "IT Team"}},
            {{"action": "Action 2", "priority": "Medium", "timeline": "60 days", "responsible": "Compliance Team"}}
        ],
        "estimated_completion": "90 days"
    }},
    "UnderlyingCause": "Root cause analysis...",
    "WhyToVerify": "Explanation of why verification is needed...",
    "WhatToVerify": "Specific items to verify...",
    "SuggestedActionPlan": "Detailed action plan with steps and timelines..."
}}

CRITICAL RULES:
- All fields MUST be filled with actual, specific content - NO placeholders or generic text
- Base your analysis on the provided evidence, strengths, weaknesses, and missing elements
- MajorMinor: "major" for Major (HIGH risk, score < 0.4), "minor" for Minor (MEDIUM/LOW risk, score >= 0.4)
- SeverityRating: 0-10 scale (0 = no risk, 10 = critical risk). Calculate as: (1.0 - compliance_score) * 10
- PredictiveRisks and CorrectiveActions must be valid JSON objects
- Be specific and actionable in all recommendations and action plans
- Consider the audit objective and scope when generating fields

Return JSON now:"""
        
        # Call AI API
        logger.info(f"🤖 Generating comprehensive audit finding fields for compliance: {compliance_description[:50]}...")
        logger.info(f"🤖 Prompt length: {len(prompt)} characters")
        try:
            ai_response = call_ai_api(prompt, None, None, 'analysis')
            logger.info(f"✅ AI API call completed. Response length: {len(ai_response) if ai_response else 0}")
        except Exception as api_err:
            logger.error(f"❌ AI API call failed: {api_err}")
            import traceback
            logger.error(f"❌ AI API traceback: {traceback.format_exc()}")
            raise
        
        # Parse JSON response
        import re
        cleaned_response = ai_response.strip() if ai_response else ""
        
        # Remove markdown code blocks if present
        code_block_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', cleaned_response, re.DOTALL)
        if code_block_match:
            cleaned_response = code_block_match.group(1)
        
        # Find first JSON object
        json_start = cleaned_response.find('{')
        if json_start >= 0:
            cleaned_response = cleaned_response[json_start:]
            # Find matching closing brace
            brace_count = 0
            json_end = -1
            for i, char in enumerate(cleaned_response):
                if char == '{':
                    brace_count += 1
                elif char == '}':
                    brace_count -= 1
                    if brace_count == 0:
                        json_end = i + 1
                        break
            if json_end > 0:
                cleaned_response = cleaned_response[:json_end]
        
        # Parse JSON with better error handling and recovery
        try:
            finding_fields = json.loads(cleaned_response)
            logger.info(f"✅ Successfully generated comprehensive audit finding fields. Keys: {list(finding_fields.keys())}")
            return finding_fields
        except json.JSONDecodeError as json_err:
            logger.error(f"❌ JSON parsing failed: {json_err}")
            logger.error(f"❌ Response content (first 1000 chars): {cleaned_response[:1000]}")
            logger.error(f"❌ Response content (last 500 chars): {cleaned_response[-500:]}")
            
            # Try to fix truncated JSON by finding the last complete field
            try:
                # Look for the last complete key-value pair before the error
                # Try to extract valid JSON up to the error point
                error_pos = json_err.pos if hasattr(json_err, 'pos') else len(cleaned_response)
                # Try to find a valid JSON object before the error
                truncated_json = cleaned_response[:error_pos]
                # Try to close any open strings/objects
                if truncated_json.count('{') > truncated_json.count('}'):
                    # Add missing closing braces
                    missing_braces = truncated_json.count('{') - truncated_json.count('}')
                    truncated_json += '}' * missing_braces
                # Try to close any open strings
                if truncated_json.count('"') % 2 != 0:
                    truncated_json = truncated_json.rstrip().rstrip('"') + '"'
                    if truncated_json.count('{') > truncated_json.count('}'):
                        truncated_json += '}'
                
                finding_fields = json.loads(truncated_json)
                logger.info(f"✅ Fixed truncated JSON by closing structures. Keys: {list(finding_fields.keys())}")
                return finding_fields
            except:
                logger.error(f"❌ Could not fix JSON, using fallback fields")
                raise
        
    except Exception as e:
        logger.error(f"❌ Error generating comprehensive audit finding fields: {e}")
        import traceback
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        
        # Return fallback fields if AI generation fails
        return {
            "HowToVerify": "Review evidence and verify compliance against requirement",
            "Impact": "Non-compliance may result in regulatory penalties and operational risks",
            "Recommendation": "Address identified gaps and strengthen controls",
            "DetailsOfFinding": f"Compliance status: {compliance_status}. Score: {compliance_score:.2f}",
            "MajorMinor": "major" if risk_level == "HIGH" or compliance_score < 0.4 else "minor",
            "SeverityRating": int(max(0, min(10, (1.0 - compliance_score) * 10))),
            "PredictiveRisks": json.dumps({"risk_level": risk_level, "compliance_score": compliance_score}),
            "CorrectiveActions": json.dumps({"recommendations": recommendations[:5] if isinstance(recommendations, list) else []}),
            "UnderlyingCause": ' | '.join(str(w) for w in weaknesses[:3]) if isinstance(weaknesses, list) and weaknesses else "Root cause analysis needed",
            "WhyToVerify": "Verification ensures compliance with regulatory requirements",
            "WhatToVerify": compliance_description,
            "SuggestedActionPlan": ' | '.join(str(r) for r in recommendations[:5]) if isinstance(recommendations, list) and recommendations else "Action plan needed"
        }


# =============================
# Compliance Checking Endpoints
# =============================

def _get_policy_requirements(policy_id: int, subpolicy_id: int = None):
    """Fetch compliance requirements for a policy or subpolicy from DB."""
    try:
        with connection.cursor() as cursor:
            # If subpolicy_id is provided, only fetch requirements for that specific subpolicy
            if subpolicy_id:
                cursor.execute(
                    """
                    SELECT c.ComplianceId, c.ComplianceTitle, c.ComplianceItemDescription,
                           c.ComplianceType, c.Criticality, c.MandatoryOptional,
                           sp.SubPolicyId, sp.SubPolicyName, p.PolicyId, p.PolicyName
                    FROM compliance c
                    JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                    JOIN policies p ON sp.PolicyId = p.PolicyId
                    WHERE sp.SubPolicyId = %s
                    """,
                    [int(subpolicy_id)]
                )
            else:
                cursor.execute(
                """
                SELECT c.ComplianceId, c.ComplianceTitle, c.ComplianceItemDescription,
                       c.ComplianceType, c.Criticality, c.MandatoryOptional,
                       sp.SubPolicyId, sp.SubPolicyName, p.PolicyId, p.PolicyName
                FROM compliance c
                JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                JOIN policies p ON sp.PolicyId = p.PolicyId
                WHERE sp.PolicyId = %s
                """,
                [int(policy_id)]
            )
            rows = cursor.fetchall()
            reqs = []
            used_titles = set()  # Track used titles to prevent duplicates
            used_descriptions = set()  # Track used descriptions to prevent duplicates
            
            for r in rows:
                # Create a unique title by combining title with compliance ID or description
                base_title = r[1] or 'Compliance Requirement'
                description = r[2] or ''
                
                # Use the full description as the title if it's available and meaningful
                if description and len(description) > 10 and description != base_title:
                    # Use full description as the main title
                    unique_title = description
                else:
                    unique_title = f"{base_title} (ID: {r[0]})"
                
                # Check for duplicates based on description (most reliable)
                if description and description in used_descriptions:
                    logger.info(f"🔄 Skipping duplicate compliance {r[0]}: '{description[:50]}...'")
                    continue
                
                # Mark this description as used
                if description:
                    used_descriptions.add(description)
                
                # Ensure title is unique by adding ID if duplicate
                counter = 1
                original_title = unique_title
                while unique_title in used_titles:
                    unique_title = f"{original_title} (ID: {r[0]})"
                    counter += 1
                
                used_titles.add(unique_title)
                logger.info(f"🔍 Requirement {r[0]}: '{base_title}' -> '{unique_title}'")
                
                reqs.append({
                    'compliance_id': r[0],
                    'title': unique_title,
                    'description': description,
                    'type': r[3] or 'General',
                    'risk': r[4] or 'Medium',
                    'mandatory': (r[5] or '').lower() == 'mandatory',
                    'subpolicy_id': r[6],
                    'subpolicy_name': r[7],
                    'policy_id': r[8],
                    'policy_name': r[9]
                })
            return reqs
    except Exception as e:
        logger.error(f"❌ Error loading requirements for policy {policy_id}: {e}")
        return []


def _ai_score_requirements_with_openai(document_text: str, requirements: list, schema: dict = None, audit_id=None, document_id=None, audit_context=None):
    """Call OpenAI in batches to score requirements against text."""
    import requests, json, re
    
    results = []
    logger.info(f"🔍 Processing {len(requirements)} requirements sequentially with 600s timeout each")
    
    # Process requirements sequentially
    for i, req in enumerate(requirements):
        global_idx = i + 1
        logger.info(f"🔍 Processing requirement {global_idx}: {req.get('title', 'Requirement')}")
        batch_results = _process_single_requirement_batch(document_text, [req], global_idx, audit_id, document_id, audit_context=audit_context)
        results.extend(batch_results)
        logger.info(f"✅ Completed requirement {global_idx}")
    
    return results

def _process_single_requirement_batch(document_text: str, batch: list, global_idx: int, audit_id=None, document_id=None, audit_context=None):
    """Process a single requirement batch"""
    import json, re
    from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
    
    req = batch[0]  # Single requirement
    
    # CRITICAL: Decrypt all encrypted fields in the requirement BEFORE using them
    if req.get('title'):
        req['title'] = decrypt_any_encrypted_value(req['title'])
    if req.get('description'):
        req['description'] = decrypt_any_encrypted_value(req['description'])
    if req.get('compliance_title'):
        req['compliance_title'] = decrypt_any_encrypted_value(req['compliance_title'])
    if req.get('policy_name'):
        req['policy_name'] = decrypt_any_encrypted_value(req['policy_name'])
    if req.get('subpolicy_name'):
        req['subpolicy_name'] = decrypt_any_encrypted_value(req['subpolicy_name'])
    
    logger.info("🤖 Using unified AI API for compliance checking")
    logger.info(f"⏱️ Processing requirement {global_idx} with 600s timeout (10 minutes)...")
    
    # Build audit context section for prompt - Enhanced with clear audit purpose
    audit_context_section = ""
    if audit_context:
        # Decrypt audit context fields
        if audit_context.get('title'):
            audit_context['title'] = decrypt_any_encrypted_value(audit_context['title'])
        if audit_context.get('objective'):
            audit_context['objective'] = decrypt_any_encrypted_value(audit_context['objective'])
        if audit_context.get('scope'):
            audit_context['scope'] = decrypt_any_encrypted_value(audit_context['scope'])
        if audit_context.get('policy_name'):
            audit_context['policy_name'] = decrypt_any_encrypted_value(audit_context['policy_name'])
        if audit_context.get('subpolicy_name'):
            audit_context['subpolicy_name'] = decrypt_any_encrypted_value(audit_context['subpolicy_name'])
        
        audit_context_section = "\n=== AUDIT CONTEXT (How This Audit Works) ===\n"
        audit_context_section += "This is a compliance audit that verifies whether the organization meets specific regulatory and policy requirements.\n\n"
        
        if audit_context.get('title'):
            audit_context_section += f"AUDIT PURPOSE: {audit_context['title']}\n"
        if audit_context.get('objective'):
            audit_context_section += f"WHAT WE ARE CHECKING: {audit_context['objective']}\n"
        if audit_context.get('scope'):
            audit_context_section += f"AUDIT SCOPE: {audit_context['scope']}\n"
        if audit_context.get('business_unit'):
            audit_context_section += f"BUSINESS UNIT UNDER AUDIT: {audit_context['business_unit']}\n"
        if audit_context.get('audit_type'):
            audit_context_section += f"AUDIT TYPE: {audit_context['audit_type']}\n"
        if audit_context.get('due_date'):
            audit_context_section += f"DUE DATE: {audit_context['due_date']}\n"
        
        # CRITICAL: Include policy and subpolicy information for relevance checking
        if audit_context.get('policy_name'):
            audit_context_section += f"\nPOLICY BEING AUDITED: {audit_context['policy_name']}\n"
        if audit_context.get('subpolicy_name'):
            audit_context_section += f"SUB-POLICY BEING AUDITED: {audit_context['subpolicy_name']}\n"
        
        audit_context_section += "\nYOUR TASK: Analyze the uploaded document to determine if it provides evidence that the organization complies with the requirement below.\n"
        audit_context_section += "If the document is NOT about this policy/subpolicy or requirement topic, it is IRRELEVANT.\n"
        audit_context_section += "==========================================\n\n"
    
    # Detect if this is combined evidence (document + database)
    is_combined_evidence = "=== DOCUMENT EVIDENCE" in document_text and "=== DATABASE EVIDENCE" in document_text
    
    # Set evidence section label and instructions based on evidence type
    evidence_section_label = "COMBINED EVIDENCE (Documents + Database Records)" if is_combined_evidence else "DOCUMENT CONTENT"
    
    if is_combined_evidence:
        evidence_instruction = """CRITICAL ANALYSIS INSTRUCTIONS FOR COMBINED EVIDENCE:

You have TWO types of evidence to analyze:

1. **DOCUMENT EVIDENCE SECTION** (Intent/Design):
   - Contains policies, procedures, design documents, and written standards
   - Shows what the organization INTENDS to do or HAS DESIGNED
   - Look for: policy statements, procedure descriptions, control frameworks, documented processes
   - MANDATORY: You MUST extract and list specific evidence from this section

2. **DATABASE EVIDENCE SECTION** (Operational/Factual):
   - Contains actual operational records, data, and factual information
   - Shows what the organization ACTUALLY DOES in practice
   - Look for: actual data values, operational records, implementation details, factual evidence
   - MANDATORY: You MUST extract and list specific evidence from this section

YOUR ANALYSIS MUST:
- **STEP 1: Analyze DOCUMENT EVIDENCE** - Read the document section carefully and extract ALL evidence that relates to the compliance requirement
- **STEP 2: Analyze DATABASE EVIDENCE** - Read the database section carefully and extract ALL evidence that relates to the compliance requirement
- **STEP 3: Compare and Assess** - Compare what is documented vs what is actually implemented
- **STEP 4: Identify Missing Elements** - Determine what is missing based on BOTH sections

REQUIRED OUTPUT FOR "evidence" ARRAY:
- MUST include at least 2-3 evidence items from the DOCUMENT section (actual text excerpts, policy statements, procedure descriptions)
- MUST include at least 2-3 evidence items from the DATABASE section (actual data values, operational records, implementation details)
- Each evidence item should be a concrete, specific finding from the provided content
- Do NOT use placeholder text, examples, or generic statements
- Do NOT prefix with "From document:" or "From database:" - just provide the actual evidence content

REQUIRED OUTPUT FOR "missing" ARRAY:
- Identify specific gaps found by comparing the requirement against BOTH document and database evidence
- Consider: What is required but not documented? What is documented but not implemented? What is neither documented nor implemented?
- Provide concrete, specific missing elements - NOT generic examples like "Requirement X" or placeholder text
- If nothing is missing, return an empty array []

- Assess compliance holistically: Both design (documents) AND implementation (database) matter
- Your compliance_status, compliance_score, and all arrays must reflect your ACTUAL analysis of the real evidence provided
- If you find evidence in BOTH sections that demonstrates compliance, mark as COMPLIANT or PARTIALLY_COMPLIANT
- If you find evidence in only ONE section, mark as PARTIALLY_COMPLIANT and note the missing section in "missing" array
- If you find NO evidence in either section, mark as NON_COMPLIANT"""
    else:
        evidence_instruction = "Analyze the document content to assess compliance."
    
    # Note: Examples are shown in the prompt format section, but AI must provide ACTUAL analysis results
    
    # Build dynamic requirement description from actual database data
    req_title = req.get('title', 'Requirement')
    req_desc = req.get('description', req_title)
    req_text = f"{req_title}" + (f"\n{req_desc}" if req_desc and req_desc != req_title else "")
    
    # Get policy/subpolicy from requirement data (from database) or audit context
    policy_name = req.get('policy_name') or audit_context.get('policy_name', 'N/A') if audit_context else 'N/A'
    subpolicy_name = req.get('subpolicy_name') or audit_context.get('subpolicy_name', 'N/A') if audit_context else 'N/A'
    
    # Build missing elements context - AI will analyze what data is required for THIS AUDIT
    audit_objective_text = audit_context.get('objective', '') if audit_context else ''
    audit_scope_text = audit_context.get('scope', '') if audit_context else ''
    
    # Create enhanced prompt with clear audit process explanation
    prompt = f"""You are a GRC compliance auditor performing an audit. Your job is to analyze whether the uploaded document provides evidence that the organization meets the compliance requirement.

IMPORTANT: The uploaded content may be:
- A narrative document (policy, procedure, memo, minutes, email, etc.)
- A structured table, spreadsheet, or CSV-style data (rows and columns with headers)
- A mix of both

You MUST treat structured/tabular data (including CSV-style rows) as valid evidence. Each row in a table/CSV represents a real record or instance, and column headers with cell values contain factual data that can demonstrate compliance. Do not ignore structured data just because it is in tabular format - extract evidence from the actual data values in the rows and columns.

{audit_context_section if audit_context else ''}=== COMPLIANCE REQUIREMENT TO CHECK ===
{req_text}

=== UPLOADED DOCUMENT CONTENT ===
{evidence_section_label}:
{document_text}

=== HOW TO PERFORM THIS AUDIT ===

STEP 1: RELEVANCE CHECK (MUST DO FIRST):
   - CRITICAL: This audit is about: {audit_context.get('title', 'N/A') if audit_context else 'N/A'}
   - CRITICAL: This audit's objective is: {audit_objective_text if audit_context and audit_context.get('objective') else 'N/A'}
   - CRITICAL: This audit's scope is: {audit_scope_text if audit_context and audit_context.get('scope') else 'N/A'}
   - The requirement being checked: "{req_desc}"
   - The document content provided above (including any tables / CSV-style data)
   
   IMPORTANT FOR STRUCTURED DATA: If the document contains comma-separated values, tab-separated values, or table rows with headers:
   - The first line typically contains column headers that describe what data is in each column
   - Each subsequent line represents a data record with values separated by commas/tabs
   - This structured data format contains actual evidence - treat the header row and data rows as factual information
   - Check if the column headers relate to the requirement topic - if yes, the document is relevant
   
   RELEVANCE QUESTIONS (ALL must be YES for document to be relevant):
   1. Is the document about the SAME TOPIC as the requirement "{req_desc}"? (For structured data: do the column headers indicate data related to this requirement?)
   2. Is the document related to Policy "{policy_name}" / Sub-Policy "{subpolicy_name}"?
   3. Does the document relate to THIS AUDIT's objective: {audit_objective_text if audit_context and audit_context.get('objective') else 'N/A'}?
   4. Does the document fall within THIS AUDIT's scope: {audit_scope_text if audit_context and audit_context.get('scope') else 'N/A'}?
   5. Would this document help prove compliance for THIS SPECIFIC AUDIT (not just any audit)?
   
   IF ANY ANSWER IS NO → Document is IRRELEVANT:
   → Set evidence = []
   → Set compliance_status = 'NON_COMPLIANT'
   → Set compliance_score = 0.0-0.2
   → You STILL MUST generate a full "missing" array with 7 items, each containing your own analysis, NOT a generic template
   → STOP HERE - do not proceed to Step 2

STEP 2: EVIDENCE EXTRACTION (ONLY IF RELEVANT):
   - Remember: This audit is checking: {audit_objective_text if audit_context and audit_context.get('objective') else 'N/A'}
   - Requirement asks for: "{req_desc}"
   - Search document for text/data that directly addresses this requirement
   - Evidence MUST match requirement topic EXACTLY
   - Evidence MUST be relevant to the audit objective and scope
   - If document mentions requirement topic but doesn't provide evidence → evidence = []
   - Extract actual quotes, numbers, or facts from document
   
   CRITICAL FOR STRUCTURED/TABULAR DATA (CSV, tables, spreadsheets):
   - If content has comma-separated or tab-separated values with a header row:
       * Identify the header row (first line with column names)
       * Each subsequent data row represents a real instance/record
       * Extract evidence by combining column header names with values from each data row
       * Count how many data rows exist - this shows quantity/frequency
       * Extract ALL relevant data rows that match the requirement - do not skip any
       * Treat each data row as factual evidence that proves or demonstrates compliance
   - Do NOT ignore structured data - it contains real evidence in organized format
   - If you see multiple data rows, each one is separate evidence of compliance activity

STEP 3: COMPLIANCE VERIFICATION (ONLY IF EVIDENCE FOUND):
   - Remember: This audit objective is: {audit_objective_text if audit_context and audit_context.get('objective') else 'N/A'}
   - Does the evidence show the organization MEETS the requirement?
   - Check: Are ALL aspects of the requirement addressed?
   - Check: Is the evidence complete and sufficient?
   - Check: Does evidence demonstrate actual compliance (not just intent)?
   - Check: Does the evidence align with what the audit objective is checking?
   - ANALYZE: How does each piece of evidence relate to the requirement AND the audit objective?
   - ANALYZE: What does the evidence prove about compliance in the context of this audit?

STEP 4: AUDIT CONCLUSION & REASONING (REQUIRED):
   - Explain HOW you performed this audit (what you checked, what you found)
   - Explain WHY you determined the compliance status:
     * If COMPLIANT: Explain how the evidence demonstrates full compliance with all aspects of the requirement
     * If PARTIALLY_COMPLIANT: Explain what evidence was found and what specific gaps remain
     * If NON_COMPLIANT: Explain why the evidence is insufficient or missing
   - Provide clear reasoning that connects the evidence to the requirement

=== COMPLIANCE STATUS DECISION ===
- COMPLIANT: Document is relevant + Strong evidence found + All requirement aspects met + Evidence clearly demonstrates compliance
- PARTIALLY_COMPLIANT: Document is relevant + Some evidence found + But specific gaps exist + Not all requirement aspects met
- NON_COMPLIANT: Document is irrelevant OR No evidence found OR Evidence doesn't show compliance OR Evidence is insufficient

=== EVIDENCE RULES ===
- Only include evidence if it directly matches the requirement topic
- Extract actual text, numbers, or data from the document
- Do NOT include generic statements or unrelated content
- If no relevant evidence: evidence = [], and you MUST still generate a full "missing" array with 7 items, each containing your own, document-specific analysis

=== MISSING ELEMENTS RULES ===
If evidence = [] (no evidence found):
   - missing MUST include ALL 7 items from the template below - DO NOT skip any field
   - The missing array MUST have exactly these 7 items (in this order). Each item MUST start with the fixed heading shown below, followed by your own analysis text based on THIS document and THIS audit (do NOT copy any placeholder text from this prompt):
   1. "NO EVIDENCE FOUND: ..." → explain in your own words why no evidence was found for this requirement
   2. "REQUIREMENT NEEDS: ..." → explain in your own words what the requirement needs (based on: {req_desc})
   3. "DOCUMENT IS ABOUT: ..." → explain in your own words what the document is actually about
   4. "DATA REQUIRED FOR THIS AUDIT: ..." → explain in your own words what data is required for this audit
   5. "EXPECTED DOCUMENT TYPE: ..." → explain in your own words what type of document/data would provide evidence
   6. "WHAT IS NEEDED: ..." → explain in your own words what concrete document, data, or evidence is needed to prove compliance
   7. "POLICY CONTEXT: ..." → explain in your own words how this requirement fits Policy "{policy_name}" / Sub-Policy "{subpolicy_name}" and why the document must relate to this policy area
   
   - CRITICAL: When filling in the missing array, you MUST REPLACE ALL bracketed placeholders with ACTUAL ANALYSIS
   - DO NOT include square brackets [] or placeholder text in your output
   - DO NOT skip any of the 7 required fields
   - Each field must contain actual analysis, not instruction text
   
   - For "DOCUMENT IS ABOUT": 
    * Read the document content provided above
    * Analyze what the document actually contains (NOT what the requirement says)
    * If the content is a table/CSV, explicitly name key column headers and what the rows represent
    * Describe the topic, purpose, and content in a way that could only be written after actually reading THIS document (no generic sentences)
   
   - For "DATA REQUIRED FOR THIS AUDIT": 
    * Consider the Audit Objective: {audit_objective_text if audit_context and audit_context.get('objective') else 'N/A'}
    * Consider the Audit Scope: {audit_scope_text if audit_context and audit_context.get('scope') else 'N/A'}
    * Consider the Requirement: "{req_desc}"
    * Consider the Policy Context: Policy "{policy_name}" / Sub-Policy "{subpolicy_name}"
    * Analyze and specify: What concrete data points, metrics, logs, tables, or records would prove compliance for THIS specific audit. Do NOT use vague phrases like "specific data points, metrics, records, measurements, or evidence".
    
   - For "EXPECTED DOCUMENT TYPE": 
    * Name concrete document types that would provide evidence for this requirement, not generic descriptions like "a document that provides detailed information..."
   
   - For "WHAT IS NEEDED": 
    * List the specific additional documents and/or data sets that are still missing for this requirement in this audit
    * Name at least 2-3 concrete items needed for this audit
    * Do NOT use generic phrases like "Specific data points and metrics to demonstrate compliance" – always mention the actual type of data/log/report needed
   
   - Be specific and detailed in your analysis for ALL fields
   - Reference the policy/subpolicy context: Policy "{policy_name}" / Sub-Policy "{subpolicy_name}"

If evidence found but incomplete:
   - List specific gaps: What parts of requirement are missing?
   - What additional data points or evidence would be needed?
   - What specific actions, measurements, or data are missing?

=== OUTPUT FORMAT ===
Return ONLY valid JSON (no markdown, no explanations, no code blocks).

JSON Structure:
{{
  "analysis": [{{
    "index": {global_idx},
    "compliance_id": {req.get('compliance_id', global_idx)},
    "relevance": 0.0-1.0,
    "compliance_status": "COMPLIANT|PARTIALLY_COMPLIANT|NON_COMPLIANT",
    "compliance_score": 0.0-1.0,
    "risk_level": "LOW|MEDIUM|HIGH",
    "confidence": 0.0-1.0,
    "evidence": ["actual text quotes from document that prove compliance"],
    "missing": ["MUST include ALL 7 fields from template if no evidence: NO EVIDENCE FOUND, REQUIREMENT NEEDS, DOCUMENT IS ABOUT, DATA REQUIRED FOR THIS AUDIT, EXPECTED DOCUMENT TYPE, WHAT IS NEEDED, POLICY CONTEXT"],
    "strengths": ["HOW AUDIT WAS PERFORMED: Explain the audit process - what you checked, what evidence you found, and how it relates to the requirement. REASON FOR STATUS: Explain why this is COMPLIANT/PARTIALLY_COMPLIANT/NON_COMPLIANT - what evidence supports this conclusion"],
    "weaknesses": ["What specific aspects of the requirement are not met or are missing. Explain what gaps exist and why they prevent full compliance"],
    "recommendations": ["Actionable recommendations to improve compliance, based on the specific gaps found during this audit"]
  }}]
}}

CRITICAL OUTPUT RULES:
- If document is IRRELEVANT: 
  * evidence = []
  * missing = a full 7-item array that YOU generate with your own analysis (see MISSING ELEMENTS RULES)
  * compliance_status = 'NON_COMPLIANT'
  * strengths = ["AUDIT PROCESS: Document was analyzed for relevance to requirement '{req_desc}'. REASON: Document is about [topic] which is not related to requirement '{req_desc}' or Policy '{policy_name}'. Therefore, no evidence can be found and status is NON_COMPLIANT."]

- If document is RELEVANT but no evidence: 
  * evidence = []
  * missing = specific gaps
  * compliance_status = 'NON_COMPLIANT'
  * strengths = ["AUDIT PROCESS: Document was analyzed and found relevant to requirement '{req_desc}'. However, no specific evidence was found that demonstrates compliance. REASON: Document mentions the topic but does not provide the required data/evidence to prove compliance."]

- If evidence found: 
  * evidence = [actual quotes from document]
  * missing = specific gaps only (if any)
  * strengths = ["AUDIT PROCESS: [Explain how you performed the audit - what you checked, what evidence you found]. REASON FOR [COMPLIANT/PARTIALLY_COMPLIANT]: [Explain why the evidence supports this status - how each piece of evidence relates to the requirement and what it proves]"]
  * weaknesses = [specific gaps if PARTIALLY_COMPLIANT, or empty if COMPLIANT]

- relevance: 0.0-0.3 if irrelevant, 0.4-0.7 if partially relevant, 0.8-1.0 if highly relevant
- compliance_score: 0.0-0.3 if non-compliant, 0.4-0.7 if partially compliant, 0.8-1.0 if compliant

CRITICAL: When filling the "missing" array:
- You MUST include ALL 7 required fields - DO NOT skip any field
- You MUST REPLACE all bracketed placeholders with ACTUAL ANALYSIS
- The missing array MUST contain exactly 7 items:
  1. NO EVIDENCE FOUND message
  2. REQUIREMENT NEEDS with the requirement text
  3. DOCUMENT IS ABOUT with your analysis
  4. DATA REQUIRED FOR THIS AUDIT with your analysis
  5. EXPECTED DOCUMENT TYPE with your analysis
  6. WHAT IS NEEDED with your analysis
  7. POLICY CONTEXT message
- For "DOCUMENT IS ABOUT": Analyze the document content and describe what it actually contains
- For "DATA REQUIRED FOR THIS AUDIT": Analyze what specific data/evidence is needed based on the audit objective, scope, requirement, and policy context - provide actual analysis
- For "EXPECTED DOCUMENT TYPE": Determine and specify the actual document type needed
- For "WHAT IS NEEDED": Specify what actual document, data, or evidence is needed
- DO NOT include square brackets [], instruction text, or placeholder text
- DO NOT skip any of the 7 required fields - ALL must be present

REQUIRED: The "strengths" field MUST explain:
1. HOW the audit was performed (what was checked, what evidence was found)
2. WHY the compliance status was determined (reasoning connecting evidence to requirement)

Return JSON now:"""
    
    # Use unified AI API call
    data = call_ai_api(prompt, audit_id, document_id, 'compliance')
    
    logger.info(f"🤖 TinyLlama response length: {len(data)} characters")
    
    # Parse JSON response - Enhanced with markdown code block handling and truncation repair
    try:
        # Remove markdown code blocks if present
        cleaned_data = data.strip()
        
        # Remove any leading text before JSON (common issue: "Here is the analysis:" or similar)
        # Use regex to find the first valid JSON object start, handling markdown formatting
        import re
        
        # First, try to find JSON wrapped in markdown code blocks
        code_block_match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', cleaned_data, re.DOTALL)
        if code_block_match:
            cleaned_data = code_block_match.group(1)
            logger.info(f"🔧 Extracted JSON from markdown code block")
        else:
            # Look for first occurrence of { or [ that starts a valid JSON structure
            # Skip common prefixes like "Here is the analysis:", "**evidence**", etc.
            json_start = -1
            for i, char in enumerate(cleaned_data):
                if char in ['{', '[']:
                    # Check if this looks like the start of a JSON object/array
                    # (not part of markdown formatting or explanatory text)
                    if i > 0:
                        # Look back a bit to see if we're in markdown context
                        context = cleaned_data[max(0, i-50):i].lower()
                        # Skip if it's part of markdown formatting like ** or * 
                        if not (context.endswith('*') or context.endswith('**') or context.endswith('#')):
                            json_start = i
                            break
                    else:
                        json_start = i
                        break
            
            if json_start > 0:
                # Extract from json_start, but only take the first complete JSON object/array
                # This prevents "Extra data" errors when there's text after the JSON
                temp_data = cleaned_data[json_start:]
                
                # Find the matching closing brace/bracket for the root object/array
                if temp_data[0] == '{':
                    # Find matching closing brace
                    brace_count = 0
                    in_string = False
                    escape_next = False
                    end_pos = len(temp_data)
                    
                    for i, char in enumerate(temp_data):
                        if escape_next:
                            escape_next = False
                            continue
                        if char == '\\':
                            escape_next = True
                            continue
                        if char == '"' and not escape_next:
                            in_string = not in_string
                            continue
                        if in_string:
                            continue
                        if char == '{':
                            brace_count += 1
                        elif char == '}':
                            brace_count -= 1
                            if brace_count == 0:
                                end_pos = i + 1
                                break
                    
                    cleaned_data = temp_data[:end_pos]
                    logger.info(f"🔧 Removed {json_start} characters of leading text and extracted complete JSON object (length: {len(cleaned_data)})")
                elif temp_data[0] == '[':
                    # Find matching closing bracket
                    bracket_count = 0
                    in_string = False
                    escape_next = False
                    end_pos = len(temp_data)
                    
                    for i, char in enumerate(temp_data):
                        if escape_next:
                            escape_next = False
                            continue
                        if char == '\\':
                            escape_next = True
                            continue
                        if char == '"' and not escape_next:
                            in_string = not in_string
                            continue
                        if in_string:
                            continue
                        if char == '[':
                            bracket_count += 1
                        elif char == ']':
                            bracket_count -= 1
                            if bracket_count == 0:
                                end_pos = i + 1
                                break
                    
                    cleaned_data = temp_data[:end_pos]
                    logger.info(f"🔧 Removed {json_start} characters of leading text and extracted complete JSON array (length: {len(cleaned_data)})")
                else:
                    cleaned_data = temp_data
                    logger.info(f"🔧 Removed {json_start} characters of leading text before JSON")
            elif json_start == -1:
                # If we didn't find { or [, try to find JSON-like structure
                # Look for "analysis": pattern which indicates JSON structure
                analysis_match = re.search(r'["\']?analysis["\']?\s*:\s*\[', cleaned_data, re.IGNORECASE)
                if analysis_match:
                    # Find the opening brace before "analysis"
                    brace_pos = cleaned_data.rfind('{', 0, analysis_match.start())
                    if brace_pos >= 0:
                        cleaned_data = cleaned_data[brace_pos:]
                        logger.info(f"🔧 Found JSON structure starting with 'analysis' key")
        
        # Remove markdown code block markers if still present
        if cleaned_data.startswith('```json'):
            cleaned_data = cleaned_data[7:]  # Remove ```json
        if cleaned_data.startswith('```'):
            cleaned_data = cleaned_data[3:]   # Remove ```
        if cleaned_data.endswith('```'):
            cleaned_data = cleaned_data[:-3]  # Remove trailing ```
        cleaned_data = cleaned_data.strip()
        
        # Remove any remaining markdown formatting that might interfere
        # Remove lines that are just markdown headers or formatting
        lines = cleaned_data.split('\n')
        cleaned_lines = []
        for line in lines:
            line_stripped = line.strip()
            # Skip markdown headers (##, ###), bold markers (**), or bullet points that aren't in JSON strings
            # Keep the line if it's not pure markdown formatting
            is_markdown_header = line_stripped.startswith('##')
            is_bold_marker = line_stripped.startswith('**') and line_stripped.endswith('**') and '{' not in line_stripped
            is_bullet_point = line_stripped.startswith('* ') and not line_stripped.startswith('*"')
            
            if not (is_markdown_header or is_bold_marker or is_bullet_point):
                cleaned_lines.append(line)
        cleaned_data = '\n'.join(cleaned_lines).strip()
        
        # Additional check: If the response is still not JSON (starts with text like "Here is my analysis:"),
        # try to find JSON structure more aggressively by looking for common JSON patterns
        if not cleaned_data.startswith('{') and not cleaned_data.startswith('['):
            # Look for first JSON object or array and extract the complete structure
            json_obj_match = re.search(r'\{', cleaned_data)
            json_array_match = re.search(r'\[', cleaned_data)
            
            if json_obj_match and (not json_array_match or json_obj_match.start() < json_array_match.start()):
                # Extract complete JSON object starting from first {
                start_pos = json_obj_match.start()
                temp_data = cleaned_data[start_pos:]
                brace_count = 0
                in_string = False
                escape_next = False
                end_pos = len(temp_data)
                
                for i, char in enumerate(temp_data):
                    if escape_next:
                        escape_next = False
                        continue
                    if char == '\\':
                        escape_next = True
                        continue
                    if char == '"' and not escape_next:
                        in_string = not in_string
                        continue
                    if in_string:
                        continue
                    if char == '{':
                        brace_count += 1
                    elif char == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            end_pos = i + 1
                            break
                
                cleaned_data = temp_data[:end_pos]
                logger.info(f"🔧 Extracted complete JSON object from markdown text response (length: {len(cleaned_data)})")
            elif json_array_match:
                # Extract complete JSON array starting from first [
                start_pos = json_array_match.start()
                temp_data = cleaned_data[start_pos:]
                bracket_count = 0
                in_string = False
                escape_next = False
                end_pos = len(temp_data)
                
                for i, char in enumerate(temp_data):
                    if escape_next:
                        escape_next = False
                        continue
                    if char == '\\':
                        escape_next = True
                        continue
                    if char == '"' and not escape_next:
                        in_string = not in_string
                        continue
                    if in_string:
                        continue
                    if char == '[':
                        bracket_count += 1
                    elif char == ']':
                        bracket_count -= 1
                        if bracket_count == 0:
                            end_pos = i + 1
                            break
                
                cleaned_data = temp_data[:end_pos]
                logger.info(f"🔧 Extracted complete JSON array from markdown text response (length: {len(cleaned_data)})")
            else:
                # Last resort: Look for "analysis" key pattern and extract surrounding JSON
                analysis_key_match = re.search(r'["\']?analysis["\']?\s*:\s*\[', cleaned_data, re.IGNORECASE)
                if analysis_key_match:
                    # Find the opening brace before "analysis"
                    start_pos = cleaned_data.rfind('{', 0, analysis_key_match.start())
                    if start_pos >= 0:
                        # Try to find matching closing brace
                        brace_count = 0
                        end_pos = start_pos
                        for i in range(start_pos, len(cleaned_data)):
                            if cleaned_data[i] == '{':
                                brace_count += 1
                            elif cleaned_data[i] == '}':
                                brace_count -= 1
                                if brace_count == 0:
                                    end_pos = i + 1
                                    break
                        if end_pos > start_pos:
                            cleaned_data = cleaned_data[start_pos:end_pos]
                            logger.info(f"🔧 Extracted JSON structure around 'analysis' key from markdown response")
                else:
                    # No JSON structure found at all - this is a text-only response
                    logger.warning(f"⚠️ No JSON structure detected in response. Response appears to be plain text.")
                    logger.warning(f"⚠️ Response preview (first 200 chars): {cleaned_data[:200]}")
                    # Set cleaned_data to empty to trigger the fallback handler
                    cleaned_data = ""
        
        # Try to parse the JSON (if we have any)
        if cleaned_data and (cleaned_data.startswith('{') or cleaned_data.startswith('[')):
            parsed = json.loads(cleaned_data)
        else:
            # No valid JSON found - raise an error to trigger fallback
            raise ValueError("No JSON structure found in response")
    except (json.JSONDecodeError, ValueError) as e:
        # If JSON is truncated or malformed, try to repair it
        error_pos = getattr(e, 'pos', 0)
        error_msg = getattr(e, 'msg', str(e))
        logger.warning(f"⚠️ JSON parsing error at position {error_pos}: {error_msg}. Attempting to repair...")
        
        # If it's a ValueError (no JSON structure found), skip repair attempts
        skip_repair = isinstance(e, ValueError) and "No JSON structure found" in str(e)
        if skip_repair:
            logger.warning(f"⚠️ No JSON structure detected. Skipping repair attempts and going to fallback handler.")
        
        # First, try to fix unterminated strings by finding last complete key-value pair
        if not skip_repair and ("Unterminated string" in str(error_msg) or "Unterminated string" in str(e)):
            try:
                error_pos = error_pos if error_pos > 0 else len(cleaned_data)
                logger.info(f"🔧 Attempting to repair unterminated string at position {error_pos}")
                
                # Strategy 1: Find the last complete key-value pair before the error
                safe_cut = error_pos
                # Look backwards from error position to find a safe cut point
                # Check for complete array elements, object properties, etc.
                for i in range(min(error_pos - 1, len(cleaned_data) - 1), max(0, error_pos - 2000), -1):
                    char = cleaned_data[i]
                    # Look for complete array elements: ", " or ", " or ]
                    if char == ',' or char == ']' or char == '}':
                        # Check if this is outside a string by counting quotes
                        quotes_before = cleaned_data[:i+1].count('"')
                        # Also check for escaped quotes
                        escaped_quotes = cleaned_data[:i+1].count('\\"')
                        # Real quote count (excluding escaped)
                        real_quotes = quotes_before - escaped_quotes
                        if real_quotes % 2 == 0:  # Not inside a string
                            # Found a safe cut point
                            if char == ',':
                                safe_cut = i + 1
                            elif char == ']' or char == '}':
                                safe_cut = i + 1
                            break
                
                # Strategy 2: If we're in an array with an unterminated string, try to close it
                if safe_cut >= error_pos:
                    # Try to find where the array started and close it properly
                    # Look backwards for the opening of the array containing the error
                    bracket_count = 0
                    brace_count = 0
                    in_string = False
                    escape_next = False
                    array_start = -1
                    
                    for i in range(error_pos - 1, max(0, error_pos - 1000), -1):
                        char = cleaned_data[i]
                        if escape_next:
                            escape_next = False
                            continue
                        if char == '\\':
                            escape_next = True
                            continue
                        if char == '"' and not escape_next:
                            in_string = not in_string
                            continue
                        if in_string:
                            continue
                        if char == ']':
                            bracket_count += 1
                        elif char == '[':
                            bracket_count -= 1
                            if bracket_count == 0:
                                array_start = i
                                break
                        elif char == '}':
                            brace_count += 1
                        elif char == '{':
                            brace_count -= 1
                    
                    if array_start >= 0:
                        # Found the array start, try to close the string and array
                        # Find the last complete element before the error
                        last_comma = cleaned_data.rfind(',', array_start, error_pos)
                        if last_comma > array_start:
                            # Check if comma is outside string
                            quotes_before_comma = cleaned_data[array_start:last_comma+1].count('"')
                            if quotes_before_comma % 2 == 0:
                                safe_cut = last_comma + 1
                
                # If we found a safe cut point, truncate there
                if safe_cut < error_pos and safe_cut > 0:
                    repaired = cleaned_data[:safe_cut]
                    # Close any open structures
                    open_braces = repaired.count('{') - repaired.count('}')
                    open_brackets = repaired.count('[') - repaired.count(']')
                    # Ensure strings are closed (should be even number of quotes)
                    quote_count = repaired.count('"')
                    escaped_quotes = repaired.count('\\"')
                    real_quotes = quote_count - escaped_quotes
                    if real_quotes % 2 != 0:
                        # Remove the last incomplete string by finding the last unclosed quote
                        # Find the last quote that's not escaped
                        for i in range(len(repaired) - 1, -1, -1):
                            if repaired[i] == '"' and (i == 0 or repaired[i-1] != '\\'):
                                repaired = repaired[:i] + '"' + repaired[i+1:]
                                break
                    # Close arrays first, then objects
                    repaired += ']' * open_brackets
                    repaired += '}' * open_braces
                    
                    # Try to parse the repaired JSON
                    try:
                        parsed = json.loads(repaired)
                        logger.info(f"✅ Successfully repaired unterminated string JSON by truncating at position {safe_cut}")
                    except json.JSONDecodeError as parse_err:
                        logger.warning(f"⚠️ Repaired JSON still invalid: {parse_err}, trying alternative repair...")
                        cleaned_data = repaired
                else:
                    # Strategy 3: Try to extract valid fields using regex as fallback
                    logger.info(f"🔧 Could not find safe cut point, trying field extraction...")
                    # This will be handled in the analysis extraction section below
            except Exception as str_repair_err:
                logger.warning(f"⚠️ String repair failed: {str_repair_err}, trying analysis extraction...")
        
        # Try to extract just the analysis array if the JSON is truncated
        # Skip if we already know there's no JSON structure
        if skip_repair:
            # No JSON structure found - raise error immediately
            logger.error(f"❌ No JSON structure found in AI response. Response preview (first 500 chars): {data[:500]}")
            raise ValueError(f"No JSON structure found in AI response. The model returned plain text instead of structured JSON. Response preview: {data[:200]}...")
        
        try:
            import re
            # Look for the analysis array pattern: "analysis": [{...}]
            analysis_match = re.search(r'"analysis"\s*:\s*\[', cleaned_data)
            if analysis_match:
                start_pos = analysis_match.end() - 1  # Position of '['
                # Try to find the first complete object in the array
                brace_count = 0
                in_string = False
                escape_next = False
                obj_start = start_pos + 1
                obj_end = obj_start
                
                # Skip whitespace after '['
                while obj_start < len(cleaned_data) and cleaned_data[obj_start] in ' \n\t\r':
                    obj_start += 1
                
                # Find the end of the first complete object
                for i in range(obj_start, len(cleaned_data)):
                    char = cleaned_data[i]
                    if escape_next:
                        escape_next = False
                        continue
                    if char == '\\':
                        escape_next = True
                        continue
                    if char == '"' and not escape_next:
                        in_string = not in_string
                        continue
                    if in_string:
                        continue
                    if char == '{':
                        brace_count += 1
                    elif char == '}':
                        brace_count -= 1
                        if brace_count == 0:
                            obj_end = i + 1
                            break
                
                if obj_end > obj_start:
                    # Extract the first complete object
                    obj_json = cleaned_data[obj_start:obj_end]
                    # Try to parse it as a single analysis object
                    try:
                        obj_parsed = json.loads(obj_json)
                        # Wrap it in the expected structure
                        parsed = {"analysis": [obj_parsed]}
                        logger.info(f"✅ Successfully extracted first analysis object from truncated JSON")
                    except:
                        # If that fails, try to build a minimal valid structure
                        # Extract key fields using regex as fallback
                        compliance_id_match = re.search(r'"compliance_id"\s*:\s*(\d+)', obj_json)
                        status_match = re.search(r'"compliance_status"\s*:\s*"([^"]+)"', obj_json)
                        score_match = re.search(r'"(?:compliance_score|relevance)"\s*:\s*([0-9.]+)', obj_json)
                        
                        if compliance_id_match:
                            parsed = {
                                "analysis": [{
                                    "compliance_id": int(compliance_id_match.group(1)),
                                    "compliance_status": status_match.group(1) if status_match else "PARTIALLY_COMPLIANT",
                                    "compliance_score": float(score_match.group(1)) if score_match else 0.5,
                                    "relevance": float(score_match.group(1)) if score_match else 0.5
                                }]
                            }
                            logger.info(f"✅ Successfully extracted analysis from truncated JSON using regex fallback")
                        else:
                            raise
                else:
                    raise
            else:
                # No "analysis" key found, but might have fields directly
                # Try to extract fields using regex as a last resort
                logger.info(f"🔧 No 'analysis' key found, trying to extract fields directly from response...")
                import re
                
                # Try to extract key fields from the response
                compliance_status_match = re.search(r'"compliance_status"\s*:\s*"([^"]+)"', cleaned_data)
                relevance_match = re.search(r'"relevance"\s*:\s*([0-9.]+)', cleaned_data)
                compliance_score_match = re.search(r'"compliance_score"\s*:\s*([0-9.]+)', cleaned_data)
                
                # Extract evidence array (might be truncated)
                # Look for "evidence": [ and extract until ] or end of data
                evidence_start = cleaned_data.find('"evidence"')
                evidence_array_start = -1
                if evidence_start >= 0:
                    bracket_pos = cleaned_data.find('[', evidence_start)
                    if bracket_pos >= 0:
                        evidence_array_start = bracket_pos + 1
                
                # Extract missing array (might be truncated)
                missing_start = cleaned_data.find('"missing"')
                missing_array_start = -1
                if missing_start >= 0:
                    bracket_pos = cleaned_data.find('[', missing_start)
                    if bracket_pos >= 0:
                        missing_array_start = bracket_pos + 1
                
                # Extract strengths and weaknesses if present
                strengths_start = cleaned_data.find('"strengths"')
                strengths_array_start = -1
                if strengths_start >= 0:
                    bracket_pos = cleaned_data.find('[', strengths_start)
                    if bracket_pos >= 0:
                        strengths_array_start = bracket_pos + 1
                
                weaknesses_start = cleaned_data.find('"weaknesses"')
                weaknesses_array_start = -1
                if weaknesses_start >= 0:
                    bracket_pos = cleaned_data.find('[', weaknesses_start)
                    if bracket_pos >= 0:
                        weaknesses_array_start = bracket_pos + 1
                
                # Build a minimal valid structure if we found at least one key field
                # Check for any valid fields: status fields OR array fields
                has_status_fields = compliance_status_match or relevance_match or compliance_score_match
                has_array_fields = evidence_array_start >= 0 or missing_array_start >= 0 or strengths_array_start >= 0 or weaknesses_array_start >= 0
                
                if has_status_fields or has_array_fields:
                    extracted = {}
                    if compliance_status_match:
                        extracted['compliance_status'] = compliance_status_match.group(1)
                    if relevance_match:
                        extracted['relevance'] = float(relevance_match.group(1))
                    if compliance_score_match:
                        extracted['compliance_score'] = float(compliance_score_match.group(1))
                    elif relevance_match:
                        extracted['compliance_score'] = float(relevance_match.group(1))
                    
                    # Set default values if status fields are missing but we have array fields
                    if not has_status_fields:
                        # Default to PARTIALLY_COMPLIANT if we have missing items
                        extracted['compliance_status'] = 'PARTIALLY_COMPLIANT'
                        extracted['relevance'] = 0.5
                        extracted['compliance_score'] = 0.5
                    
                    # Try to extract evidence and missing (even if truncated)
                    if evidence_array_start >= 0:
                        # Find the end of the array (either ] or end of data)
                        evidence_end = cleaned_data.find(']', evidence_array_start)
                        if evidence_end < 0:
                            evidence_end = len(cleaned_data)
                        evidence_str = cleaned_data[evidence_array_start:evidence_end].strip()
                        if evidence_str:
                            # Try to parse as JSON array, fallback to empty if fails
                            try:
                                extracted['evidence'] = json.loads('[' + evidence_str + ']')
                            except:
                                # Extract complete strings from the array
                                string_matches = re.findall(r'"([^"]*)"', evidence_str)
                                extracted['evidence'] = string_matches[:10] if string_matches else []
                        else:
                            extracted['evidence'] = []
                    
                    if missing_array_start >= 0:
                        # Find the end of the array (either ] or end of data)
                        missing_end = cleaned_data.find(']', missing_array_start)
                        if missing_end < 0:
                            missing_end = len(cleaned_data)
                        missing_str = cleaned_data[missing_array_start:missing_end].strip()
                        if missing_str:
                            missing_items = []
                            
                            # Extract all complete quoted strings first
                            # This regex will find all "text" patterns where text doesn't contain quotes
                            string_matches = re.findall(r'"([^"]*)"', missing_str)
                            if string_matches:
                                missing_items = string_matches[:5]  # Limit to first 5 complete items
                            
                            # Check if there's an unterminated string after the last complete string
                            # Find the position after the last complete string
                            last_complete_pos = -1
                            if string_matches:
                                # Find the position of the last closing quote
                                last_quote_pos = missing_str.rfind('"')
                                if last_quote_pos >= 0:
                                    # Check if there's more content after the last quote
                                    remaining = missing_str[last_quote_pos+1:].strip()
                                    # If remaining starts with comma, there might be another item
                                    if remaining.startswith(','):
                                        remaining = remaining[1:].strip()
                                        # Check if there's an opening quote for an incomplete string
                                        next_quote = remaining.find('"')
                                        if next_quote >= 0:
                                            # Extract the incomplete string content
                                            incomplete_content = remaining[next_quote+1:].strip()
                                            if incomplete_content:
                                                # Limit length and add note
                                                max_len = 150
                                                if len(incomplete_content) > max_len:
                                                    incomplete_content = incomplete_content[:max_len] + "..."
                                                missing_items.append(f"REQUIREMENT NEEDS: {incomplete_content}")
                            
                            # If no strings found at all but there's content, try to extract something
                            if not missing_items and missing_str.strip():
                                # Look for any quoted content, even if incomplete
                                first_quote = missing_str.find('"')
                                if first_quote >= 0:
                                    # Extract content after first quote (might be incomplete)
                                    content_after_quote = missing_str[first_quote+1:].strip()
                                    if content_after_quote:
                                        max_len = 150
                                        if len(content_after_quote) > max_len:
                                            content_after_quote = content_after_quote[:max_len] + "..."
                                        missing_items.append(f"REQUIREMENT NEEDS: {content_after_quote}")
                            
                            extracted['missing'] = missing_items if missing_items else []
                        else:
                            extracted['missing'] = []
                    
                    if strengths_array_start >= 0:
                        strengths_end = cleaned_data.find(']', strengths_array_start)
                        if strengths_end < 0:
                            strengths_end = len(cleaned_data)
                        strengths_str = cleaned_data[strengths_array_start:strengths_end].strip()
                        if strengths_str:
                            try:
                                extracted['strengths'] = json.loads('[' + strengths_str + ']')
                            except:
                                string_matches = re.findall(r'"([^"]*)"', strengths_str)
                                extracted['strengths'] = string_matches[:3] if string_matches else []
                        else:
                            extracted['strengths'] = []
                    
                    if weaknesses_array_start >= 0:
                        weaknesses_end = cleaned_data.find(']', weaknesses_array_start)
                        if weaknesses_end < 0:
                            weaknesses_end = len(cleaned_data)
                        weaknesses_str = cleaned_data[weaknesses_array_start:weaknesses_end].strip()
                        if weaknesses_str:
                            try:
                                extracted['weaknesses'] = json.loads('[' + weaknesses_str + ']')
                            except:
                                string_matches = re.findall(r'"([^"]*)"', weaknesses_str)
                                extracted['weaknesses'] = string_matches[:3] if string_matches else []
                        else:
                            extracted['weaknesses'] = []
                    
                    # Wrap in analysis array
                    parsed = {"analysis": [extracted]}
                    logger.info(f"✅ Successfully extracted fields from truncated JSON using regex fallback")
                else:
                    raise
        except Exception as repair_err:
            # All JSON parsing and repair attempts failed
            logger.error(f"❌ All JSON parsing attempts failed. Error: {repair_err}")
            logger.error(f"❌ Response preview (first 500 chars): {data[:500]}")
            # Re-raise the error to be handled by the calling function
            raise ValueError(f"Failed to parse AI response as JSON. The AI model returned plain text instead of structured JSON. Response preview: {data[:200]}...") from repair_err
    
    # Check if the response has analysis fields directly (not wrapped in 'analysis' array)
    # This handles cases where AI returns: {"relevance": 0.5, "compliance_status": "...", ...}
    # Instead of: {"analysis": [{"relevance": 0.5, "compliance_status": "...", ...}]}
    if 'analysis' not in parsed or not isinstance(parsed.get('analysis'), list):
        # Check if parsed has analysis-like fields directly
        analysis_fields = ['relevance', 'compliance_status', 'compliance_score', 'strengths', 'weaknesses', 'evidence', 'missing']
        has_analysis_fields = any(field in parsed for field in analysis_fields)
        
        if has_analysis_fields and isinstance(parsed, dict):
            # Wrap the response in the expected format
            logger.info(f"🔧 Response has analysis fields directly, wrapping in 'analysis' array")
            # Add compliance_id from requirement if missing
            if 'compliance_id' not in parsed and req and req.get('compliance_id'):
                parsed['compliance_id'] = req.get('compliance_id')
            parsed = {"analysis": [parsed]}
    
    if 'analysis' in parsed and isinstance(parsed['analysis'], list):
        # Enhanced compliance analysis processing
        from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
        for a in parsed['analysis']:
            # Get requirement title - decrypt if encrypted
            req_title = req.get('title') or ''
            if req_title:
                req_title = decrypt_any_encrypted_value(req_title)
            
            existing_title = a.get('requirement_title') or ''
            if existing_title:
                existing_title = decrypt_any_encrypted_value(existing_title)
            
            a['requirement_title'] = req_title or existing_title or f"Requirement {global_idx}"
            # Ensure each analysis has compliance_id for frontend mapping filter
            if 'compliance_id' not in a or a.get('compliance_id') is None:
                a['compliance_id'] = req.get('compliance_id')
            
            # Add compliance_title (short name) if available from requirements
            if req.get('compliance_title'):
                compliance_title = decrypt_any_encrypted_value(req.get('compliance_title'))
                a['compliance_title'] = compliance_title
            elif req_title and len(req_title) > 100:
                # If title is very long (description), try to extract a shorter title
                # Use first sentence or first 80 chars
                first_sentence = req_title.split('.')[0] if '.' in req_title else req_title[:80]
                a['compliance_title'] = first_sentence.strip()
            else:
                a['compliance_title'] = req_title or f"Compliance {req.get('compliance_id', global_idx)}"
            
            # Handle enhanced response format with predefined compliance_status
            if 'compliance_status' in a:
                # Use the AI-determined compliance status directly
                compliance_status = a.get('compliance_status', '').upper()
                ai_score = a.get('compliance_score', None)
                ai_relevance = a.get('relevance', None)
                missing_elements = a.get('missing', [])
                has_gaps = missing_elements and len(missing_elements) > 0
                
                # Use compliance_score if available, otherwise use relevance
                if ai_score is None and ai_relevance is not None:
                    ai_score = ai_relevance
                
                if compliance_status == 'COMPLIANT':
                    a['status'] = 'compliant'
                    # If COMPLIANT but has gaps, reduce score slightly
                    if has_gaps:
                        adjusted_score = min(ai_score if ai_score is not None else 0.9, 0.85)
                        a['compliance_score'] = adjusted_score
                        a['relevance'] = adjusted_score  # Also update relevance for consistency
                    else:
                        final_score = ai_score if ai_score is not None else 0.9
                        a['compliance_score'] = final_score
                        a['relevance'] = final_score
                elif compliance_status == 'PARTIALLY_COMPLIANT':
                    a['status'] = 'partially_compliant'
                    # Adjust score for PARTIALLY_COMPLIANT: should be 0.4-0.8 range
                    if ai_score is not None:
                        # If AI gave high score but status is PARTIALLY_COMPLIANT, cap it appropriately
                        if ai_score >= 0.8:
                            # High score but partially compliant - likely has gaps, reduce to 0.6-0.75
                            adjusted_score = min(ai_score, 0.75) if has_gaps else min(ai_score, 0.8)
                        else:
                            adjusted_score = ai_score
                        a['compliance_score'] = adjusted_score
                        a['relevance'] = adjusted_score  # Also update relevance for consistency
                    else:
                        a['compliance_score'] = 0.6
                        a['relevance'] = 0.6
                elif compliance_status == 'NON_COMPLIANT':
                    a['status'] = 'non_compliant'
                    # NON_COMPLIANT should have low score (0.0-0.4)
                    if ai_score is not None:
                        adjusted_score = min(ai_score, 0.4)
                        a['compliance_score'] = adjusted_score
                        a['relevance'] = adjusted_score
                    else:
                        a['compliance_score'] = 0.2
                        a['relevance'] = 0.2
                else:
                    # Default to non-compliant if status is unclear
                    a['compliance_score'] = 0.2
                    a['relevance'] = 0.2
                    a['status'] = 'non_compliant'
            else:
                # If no compliance_status, check if we have gaps and adjust score accordingly
                ai_score = a.get('compliance_score') or a.get('relevance')
                missing_elements = a.get('missing', [])
                has_gaps = missing_elements and len(missing_elements) > 0
                
                if ai_score is not None:
                    # If high score but has gaps, reduce it
                    if has_gaps and ai_score >= 0.8:
                        adjusted_score = min(ai_score, 0.75)
                        a['compliance_score'] = adjusted_score
                        a['relevance'] = adjusted_score
                    else:
                        a['compliance_score'] = ai_score
                        if 'relevance' not in a:
                            a['relevance'] = ai_score
                else:
                    a['compliance_score'] = 0.2
                    a['relevance'] = 0.2
                    a['status'] = 'non_compliant'
            
            # Calculate compliance percentage
            a['compliance_percent'] = int(round(a['compliance_score'] * 100))
        
        # Final safety check: Decrypt ALL encrypted values in the analysis results before returning
        from grc.utils.auto_decrypt_helper import decrypt_all_encrypted_in_dict
        decrypted_analysis = decrypt_all_encrypted_in_dict(parsed['analysis'])
        return decrypted_analysis
    else:
        raise Exception(f"Unexpected response format: {data}")
        logger.info("🤖 Using OpenAI for compliance checking")
        logger.info(f"⏱️ Processing {len(batch)} requirement with 60s timeout...")
        req_lines = []
        for j, req in enumerate(batch):
            global_idx = i + j + 1  # Global index across all batches
            desc = (req.get('description') or '').strip()
            req_lines.append(f"{global_idx}. {req.get('title','Requirement')}")
        
        # Advanced multi-requirement compliance analysis prompt
        prompt = f"""You are an expert GRC compliance auditor conducting a comprehensive multi-requirement audit. Perform detailed compliance analysis for each requirement.

DOCUMENT CONTENT: {document_text}

COMPLIANCE REQUIREMENTS TO ANALYZE:
{chr(10).join(req_lines)}

ADVANCED MULTI-REQUIREMENT ANALYSIS TASK:
For EACH requirement (1 through {len(batch)}), perform:

1. **Comprehensive Evidence Detection**: Find ALL specific evidence demonstrating compliance
2. **Gap Analysis**: Identify ALL missing elements needed for full compliance  
3. **Compliance Assessment**: Determine exact compliance level (COMPLIANT/PARTIALLY_COMPLIANT/NON_COMPLIANT)
4. **Risk Evaluation**: Assess risk level of compliance gaps
5. **Quality Assessment**: Evaluate evidence quality and completeness

COMPLIANCE LEVEL DEFINITIONS:
- **COMPLIANT**: Strong evidence, all key elements present, well-documented, low risk
- **PARTIALLY_COMPLIANT**: Some evidence exists, but gaps/weaknesses identified, medium risk  
- **NON_COMPLIANT**: Little/no evidence, major gaps, high risk

REQUIRED JSON OUTPUT FORMAT:
{{
  "analysis": [
    {{
      "index": 1,
      "compliance_id": {requirements[i+0]['compliance_id']},
      "relevance": 0.0-1.0,
      "compliance_status": "COMPLIANT|PARTIALLY_COMPLIANT|NON_COMPLIANT",
      "compliance_score": 0.0-1.0,
      "risk_level": "LOW|MEDIUM|HIGH", 
      "confidence": 0.0-1.0,
      "evidence": ["Specific evidence found in document"],
      "missing": ["Specific gaps identified"],
      "strengths": ["What document does well"],
      "weaknesses": ["Areas needing improvement"],
      "recommendations": ["Specific actionable steps"]
    }}
  ]
}}

ANALYSIS STANDARDS:
- Look for policies, procedures, controls, metrics, documentation, training records
- Consider both explicit statements and implied compliance through practices
- Evaluate evidence completeness, accuracy, and implementation quality
- Assess regulatory context and industry standards
- Identify root causes of any compliance gaps

CRITICAL REQUIREMENTS:
- Analyze ALL {len(batch)} requirements listed above
- Return EXACTLY {len(batch)} analysis objects
- Use double quotes, square brackets, curly braces for JSON
- No explanations or additional text outside JSON
- Each requirement gets ONE object in the analysis array

JSON:"""
        
        try:
            # Use unified AI API call
            data = call_ai_api(prompt, audit_id, document_id, 'compliance')
            
            logger.info(f"🤖 TinyLlama response length: {len(data)} characters")
            logger.info(f"🤖 TinyLlama response preview: {data[:300]}...")
            logger.info(f"🤖 Full TinyLlama response: {data}")
            
            # Enhanced JSON parsing to handle markdown code blocks and malformed responses
            try:
                # Remove markdown code blocks if present
                cleaned_data = data.strip()
                if cleaned_data.startswith('```json'):
                    cleaned_data = cleaned_data[7:]  # Remove ```json
                if cleaned_data.startswith('```'):
                    cleaned_data = cleaned_data[3:]   # Remove ```
                if cleaned_data.endswith('```'):
                    cleaned_data = cleaned_data[:-3]  # Remove trailing ```
                
                parsed = json.loads(cleaned_data.strip())
            except json.JSONDecodeError as e:
                logger.error(f"JSON parsing failed: {e}")
                logger.error(f"Raw response: {data}")
                raise Exception(f"Failed to parse JSON response: {e}. Response: {data[:200]}...")
            
            # Enforce de-duplication by requirement title and index (across batches)
            for a in parsed.get('analysis', []):
                # Validate that analysis item is a dictionary
                if not isinstance(a, dict):
                    logger.error(f"❌ Analysis item is not a dict: {type(a)} - {a}")
                    raise Exception(f"OpenAI returned invalid analysis item: {a}")
                
                # Add the requirement title to the analysis result
                # Model returns index relative to the current batch (1..len(batch)).
                # Convert to GLOBAL requirement index using the batch offset 'i'.
                global_index = i + int(a.get('index', 1))
                req_index = global_index - 1  # 0-based
                if req_index < len(requirements):
                    a['requirement_title'] = requirements[req_index].get('title', f'Requirement {global_index}')
                    a['requirement_description'] = requirements[req_index].get('description', '')
                    a['compliance_id'] = requirements[req_index].get('compliance_id')
                    a['compliance_title'] = requirements[req_index].get('compliance_title', a.get('requirement_title', f'Compliance {global_index}'))
                else:
                    a['requirement_title'] = f'Requirement {global_index}'
                    a['requirement_description'] = ''
                    a['compliance_title'] = f'Compliance {global_index}'
                # Overwrite index with the global index so downstream consumers are consistent
                a['index'] = global_index
                # Skip duplicates by title or by index across all batches
                title_key = a['requirement_title'].strip().lower()
                if global_index in seen_indexes_global or title_key in seen_titles_global:
                    logger.info(f"🔄 Skipping duplicate analysis for global index {global_index} / title '{a['requirement_title']}'")
                    continue
                seen_indexes_global.add(global_index)
                seen_titles_global.add(title_key)
                results.append(a)
            
            logger.info(f"✅ Processed batch {i//BATCH + 1} successfully")
            
        except Exception as e:
            logger.error(f"❌ OpenAI batch {i//BATCH + 1} failed: {e}")
            raise Exception(f"OpenAI processing failed: {e}")

    logger.info(f"✅ OpenAI compliance check completed for {len(requirements)} requirements")
    return results


def _determine_status(requirements: list, analyses: list):
    """Overall status per user's rule:
    - If ALL requirements are COMPLIANT → overall 'compliant'
    - Else if ANY requirement is NON_COMPLIANT → overall 'non_compliant'
    - Else (mix of compliant/partially_compliant) → 'partially_compliant'
    Return also the average score for display.
    """
    if not analyses or len(analyses) == 0:
        return 'non_compliant', 0.0
    
    scores = []
    all_compliant = True
    has_non_compliant = False
    
    for a in analyses:
        # Get compliance status from analysis
        compliance_status = a.get('compliance_status', '').upper()
        status = a.get('status', '').upper()
        score = float(a.get('compliance_score', 0.0))
        scores.append(score)
        
        # Determine status from compliance_status or status field
        if not compliance_status:
            if status in ['COMPLIANT', 'COMPLIED']:
                compliance_status = 'COMPLIANT'
            elif status in ['PARTIALLY_COMPLIANT', 'PARTIALLY_COMPLIED']:
                compliance_status = 'PARTIALLY_COMPLIANT'
            elif status in ['NON_COMPLIANT', 'NON_COMPLIED', 'NOT_COMPLIANT']:
                compliance_status = 'NON_COMPLIANT'
            else:
                # Fallback to score-based determination
                if score >= 0.7:
                    compliance_status = 'COMPLIANT'
                elif score >= 0.4:
                    compliance_status = 'PARTIALLY_COMPLIANT'
                else:
                    compliance_status = 'NON_COMPLIANT'
        
        # Check if this requirement is compliant
        if compliance_status != 'COMPLIANT':
            all_compliant = False
        
        # Check if this requirement is non-compliant
        if compliance_status == 'NON_COMPLIANT':
            has_non_compliant = True
    
    avg = (sum(scores)/len(scores)) if scores else 0.0
    
    # Apply user's logic
    if all_compliant:
        return 'compliant', avg
    elif has_non_compliant:
        return 'non_compliant', avg
    else:
        return 'partially_compliant', avg


def _check_document_compliance_internal(audit_id, document_id, user_id=None, selected_compliance_ids=None):
    """
    Internal function to check document compliance without requiring an HttpRequest.
    This is used for automatic background processing.
    
    Returns: dict with 'success', 'error' (if failed), or result data (if successful)
    """
    try:
        logger.info(f"🔍 Internal compliance check: audit_id={audit_id}, document_id={document_id}")
        
        # MULTI-TENANCY: Extract tenant_id from audit
        tenant_id = None
        with connection.cursor() as cursor:
            cursor.execute("SELECT TenantId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
            tenant_row = cursor.fetchone()
            if tenant_row:
                tenant_id = tenant_row[0]
        
        # Lookup document path, mime, and policy mapping from ai_audit_data table
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT d.document_path, COALESCE(d.mime_type, d.document_type) AS doc_type,
                       COALESCE(d.policy_id, a.PolicyId) AS policy_id,
                       COALESCE(d.subpolicy_id, a.SubPolicyId) AS subpolicy_id,
                       d.external_source, d.external_id, d.document_name, d.file_size
                FROM ai_audit_data d
                JOIN audit a ON a.AuditId = d.audit_id
                WHERE d.document_id = %s AND d.audit_id = %s
                """,
                [int(document_id), int(audit_id) if str(audit_id).isdigit() else audit_id]
            )
            row = cursor.fetchone()
            if not row:
                return {'success': False, 'error': 'Document not found'}
            doc_path, doc_type, policy_id, subpolicy_id, external_source, external_id, document_name, file_size = row
            # Decrypt document_path if stored encrypted (so local fallback can match ai_audit_documents/...)
            try:
                from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
                doc_path = decrypt_any_encrypted_value(doc_path) if doc_path else doc_path
            except Exception:
                pass
            
            # IMPORTANT: If document is from document handling, ALWAYS check for matched_compliances
            # and use ONLY those - don't check all compliances in policy/subpolicy
            # This ensures documents from Document Handling only check the AI-selected relevant compliances
            if external_source == 'evidence_attachment' and external_id:
                try:
                    # Try to extract file_operation_id from external_id (could be numeric ID or S3 key)
                    file_operation_id_for_lookup = None
                    if external_id and str(external_id).isdigit():
                        file_operation_id_for_lookup = int(external_id)
                    else:
                        # Try to find file_operation_id by s3_key or stored_name
                        cursor.execute("""
                            SELECT id FROM file_operations 
                            WHERE (s3_key = %s OR stored_name = %s) AND status = 'completed'
                            LIMIT 1
                        """, [external_id, external_id])
                        fo_row = cursor.fetchone()
                        if fo_row:
                            file_operation_id_for_lookup = fo_row[0]
                    
                    if file_operation_id_for_lookup:
                        cursor.execute("""
                            SELECT matched_compliances
                            FROM document_audit_relevance
                            WHERE file_operation_id = %s AND audit_id = %s
                            LIMIT 1
                        """, [file_operation_id_for_lookup, int(audit_id) if str(audit_id).isdigit() else audit_id])
                        relevance_row = cursor.fetchone()
                        if relevance_row and relevance_row[0]:
                            import json
                            try:
                                matched_compliances_from_relevance = json.loads(relevance_row[0]) if isinstance(relevance_row[0], str) else relevance_row[0]
                                if matched_compliances_from_relevance and isinstance(matched_compliances_from_relevance, list) and len(matched_compliances_from_relevance) > 0:
                                    # Use ONLY the matched compliances - override any selected_compliance_ids
                                    selected_compliance_ids = matched_compliances_from_relevance
                                    logger.info(f"📋 Document from Document Handling: Using ONLY {len(selected_compliance_ids)} matched compliances from relevance analysis: {selected_compliance_ids}")
                                    logger.info(f"📋 Will NOT check all compliances in policy/subpolicy - only these specific ones")
                                else:
                                    logger.info(f"📋 Document from Document Handling but no matched_compliances found - will use policy/subpolicy defaults")
                            except Exception as e:
                                logger.warning(f"⚠️ Could not parse matched_compliances from relevance analysis: {e}")
                        else:
                            logger.info(f"📋 No relevance analysis found for file_operation_id {file_operation_id_for_lookup} - will use policy/subpolicy defaults")
                except Exception as e:
                    logger.info(f"ℹ️ Could not retrieve matched_compliances from document_audit_relevance: {e}")

        # Handle file path - check if it's S3 or local
        temp_file_created = False
        if external_source in ['s3', 'evidence_attachment'] and external_id:
            # Handle S3 file
            try:
                import json
                s3_metadata = None
                s3_key = None
                file_name_from_db = None
                s3_url_for_fallback = None
                stored_name_for_fallback = None
                
                # For documents from document handling (evidence_attachment), try to look up from file_operations table
                if external_source == 'evidence_attachment':
                    logger.info(f"🔍 Looking up file_operations for external_id: {external_id}")
                    try:
                        with connection.cursor() as cursor:
                            # First, try to find by operation_id (if external_id is numeric)
                            # Otherwise, try to match by s3_key, stored_name, etc.
                            if external_id and str(external_id).isdigit():
                                # external_id is an operation_id - look up upload row only (download rows have NULL s3_url)
                                cursor.execute("""
                                    SELECT s3_key, file_name, original_name, s3_url, stored_name
                                    FROM file_operations
                                    WHERE id = %s AND operation_type = 'upload' AND status = 'completed'
                                    LIMIT 1
                                """, [int(external_id)])
                            else:
                                # external_id might be an S3 key or stored_name - use upload rows only
                                cursor.execute("""
                                    SELECT s3_key, file_name, original_name, s3_url, stored_name
                                    FROM file_operations
                                    WHERE operation_type = 'upload'
                                      AND status = 'completed'
                                      AND (
                                        s3_key = %s
                                        OR stored_name = %s
                                        OR file_name = %s
                                        OR original_name = %s
                                        OR s3_key LIKE %s
                                        OR s3_key LIKE %s
                                      )
                                    ORDER BY 
                                        CASE 
                                            WHEN s3_key = %s THEN 1
                                            WHEN stored_name = %s THEN 2
                                            WHEN file_name = %s THEN 3
                                            WHEN original_name = %s THEN 4
                                            ELSE 5
                                        END
                                    LIMIT 1
                                """, [
                                    external_id, external_id, external_id, external_id,
                                    f'%/{external_id}', f'%{external_id}',
                                    external_id, external_id, external_id, external_id
                                ])
                            file_op = cursor.fetchone()
                            
                            if file_op:
                                s3_key_from_db = file_op[0]
                                file_name_from_db = file_op[1] or file_op[2]
                                s3_url_from_db = file_op[3]
                                stored_name_from_db = file_op[4] if len(file_op) > 4 else None
                                
                                if s3_url_from_db:
                                    if 'amazonaws.com/' in s3_url_from_db:
                                        extracted_key = s3_url_from_db.split('amazonaws.com/')[-1].split('?')[0]
                                        if extracted_key:
                                            s3_key = extracted_key
                                    elif 's3.' in s3_url_from_db or 's3-' in s3_url_from_db:
                                        try:
                                            from urllib.parse import urlparse
                                            parsed = urlparse(s3_url_from_db)
                                            if parsed.path:
                                                s3_key = parsed.path.lstrip('/')
                                        except Exception:
                                            pass
                                
                                if not s3_key and s3_key_from_db:
                                    s3_key = s3_key_from_db
                                
                                s3_url_for_fallback = s3_url_from_db if s3_url_from_db else None
                                stored_name_for_fallback = stored_name_from_db
                    except Exception as lookup_err:
                        logger.warning(f"⚠️ Error looking up file_operations: {lookup_err}")

                # If this is an evidence_attachment but we couldn't resolve anything from file_operations
                # and the document_path clearly points to a local ai_audit_documents file, fall back to
                # treating it as a local/manual upload instead of hard‑failing on missing S3 metadata.
                if external_source == 'evidence_attachment' and not s3_key and doc_path:
                    if doc_path.startswith('ai_audit_documents') or doc_path.startswith('/'):
                        logger.info("ℹ️ No S3 metadata found for evidence_attachment; falling back to local file path.")
                        external_source = 'manual'
                
                # If we didn't find it in file_operations, try to extract from external_id
                if not s3_key:
                    if isinstance(external_id, str):
                        raw_value = external_id.strip()
                        try:
                            s3_metadata = json.loads(raw_value)
                        except Exception:
                            if 'amazonaws.com/' in raw_value:
                                s3_key = raw_value.split('amazonaws.com/')[-1].split('?')[0]
                            else:
                                # Treat plain string as direct S3 key
                                s3_key = raw_value
                    elif isinstance(external_id, dict):
                        # Legacy structure where external_id is a JSON/dict payload
                        s3_metadata = external_id
                    else:
                        # external_id is neither string nor dict (e.g. numeric operation_id) –
                        # we've already tried file_operations lookup above, so nothing more to derive here
                        s3_metadata = None

                    if not s3_key and isinstance(s3_metadata, dict):
                        s3_key = s3_metadata.get('s3_key')
                        if not s3_key and s3_metadata.get('aws_file_link'):
                            aws_link = s3_metadata.get('aws_file_link')
                            if 'amazonaws.com/' in aws_link:
                                s3_key = aws_link.split('amazonaws.com/')[-1].split('?')[0]
                
                if not s3_key and doc_path:
                    if 'amazonaws.com/' in doc_path:
                        s3_key = doc_path.split('amazonaws.com/')[-1].split('?')[0]
                    elif not doc_path.startswith('/') and not doc_path.startswith('ai_audit_documents'):
                        s3_key = doc_path
                
                if not s3_key:
                    return {'success': False, 'error': 'S3 key not found in metadata'}
                
                # Download file from S3
                from ..Global.s3_fucntions import create_direct_mysql_client
                import tempfile
                import re as _re_module  # Use alias to avoid shadowing from later 'import re' in this function
                
                s3_client = create_direct_mysql_client()
                s3_basename = s3_key.split('/')[-1] if '/' in s3_key else s3_key
                if file_name_from_db:
                    file_name = file_name_from_db
                    # Document Handling: if document_name is an extracted page (document_XXX.pdf) but s3_key
                    # points to parent .docx, use s3_key basename so we download the actual S3 object
                    if (file_name and s3_basename and file_name != s3_basename
                            and _re_module.match(r'^document_\d+\.pdf$', (file_name or '').strip())
                            and s3_basename.lower().endswith(('.docx', '.doc', '.pdf'))):
                        file_name = s3_basename
                        logger.info(f"Using s3_key basename for download (extracted page pattern): {file_name}")
                else:
                    file_name = s3_basename
                
                # Default filename: use extension from document_name or s3_key so type is correct (e.g. .xlsx not .pdf)
                if not file_name or '.' not in file_name:
                    ext = '.pdf'
                    s3_basename = s3_key.split('/')[-1] if s3_key and '/' in s3_key else s3_key
                    for src in (file_name_from_db, document_name, (external_id if isinstance(external_id, str) else None), s3_basename):
                        if src and '.' in str(src):
                            candidate = ('.' + str(src).rsplit('.', 1)[-1].strip()).lower()
                            if candidate in ('.pdf', '.xlsx', '.xls', '.docx', '.doc', '.txt'):
                                ext = candidate
                                break
                    file_name = f"document_{document_id}{ext}"
                
                temp_dir = tempfile.gettempdir()
                temp_file_path = os.path.join(temp_dir, f"ai_audit_{document_id}_{file_name}")
                
                download_success = False
                last_error = None
                s3_keys_to_try = [s3_key]
                used_local_fallback = False
                
                if stored_name_for_fallback:
                    if stored_name_for_fallback not in s3_keys_to_try:
                        s3_keys_to_try.append(stored_name_for_fallback)
                
                for attempt_key in s3_keys_to_try:
                    try:
                        download_result = s3_client.download(attempt_key, file_name, temp_dir, str(user_id) if user_id else 'system')
                        if download_result.get('success'):
                            download_success = True
                            s3_key = attempt_key
                            break
                        else:
                            last_error = download_result.get('error', 'Unknown error')
                    except Exception as download_ex:
                        last_error = str(download_ex)
                
                if not download_success:
                    if s3_url_for_fallback:
                        try:
                            import requests as req_lib
                            url_response = req_lib.get(s3_url_for_fallback, timeout=60, stream=True)
                            if url_response.status_code == 200:
                                with open(temp_file_path, 'wb') as f:
                                    for chunk in url_response.iter_content(chunk_size=8192):
                                        f.write(chunk)
                                if os.path.exists(temp_file_path):
                                    download_success = True
                        except Exception:
                            pass
                    
                    # Fallback to local file when doc_path points to ai_audit_documents
                    if not download_success and doc_path:
                        local_full = doc_path if os.path.isabs(doc_path) else os.path.join(settings.MEDIA_ROOT, doc_path)
                        if (doc_path.startswith('ai_audit_documents') or doc_path.startswith('/')) and os.path.exists(local_full):
                            full_path = local_full
                            temp_file_created = False
                            download_success = True
                            used_local_fallback = True
                            logger.warning(f"⚠️ Using local file after S3 failure: {full_path}")
                    
                    if not download_success:
                        return {'success': False, 'error': f'Failed to download from S3: {last_error}'}
                
                if download_success and 'download_result' in locals() and not used_local_fallback:
                    temp_file_path = download_result.get('file_path', temp_file_path)
                    if not temp_file_path or not os.path.exists(temp_file_path):
                        return {'success': False, 'error': 'Failed to download file from S3'}
                    
                    full_path = temp_file_path
                    temp_file_created = True
            except Exception as e:
                logger.error(f"❌ Error handling S3 file: {e}")
                return {'success': False, 'error': f'S3 file handling error: {str(e)}'}
        else:
            # Handle local file
            full_path = doc_path if os.path.isabs(doc_path) else os.path.join(settings.MEDIA_ROOT, doc_path)
            if not os.path.exists(full_path):
                return {'success': False, 'error': 'File not found on server'}

        # =============================
        # Load requirements FIRST (before expensive text extraction)
        # =============================
        # Load requirements to analyze
        if selected_compliance_ids:
            # Use the selected compliance IDs (passed from auto-processing)
            try:
                with connection.cursor() as cursor:
                    placeholders = ",".join(["%s"] * len(selected_compliance_ids))
                    cursor.execute(
                        f"""
                        SELECT c.ComplianceId, c.ComplianceTitle, c.ComplianceItemDescription,
                               c.ComplianceType, c.Criticality, c.MandatoryOptional,
                               sp.SubPolicyId, sp.SubPolicyName, p.PolicyId, p.PolicyName
                        FROM compliance c
                        JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                        JOIN policies p ON sp.PolicyId = p.PolicyId
                        WHERE c.ComplianceId IN ({placeholders})
                        """,
                        list(selected_compliance_ids)
                    )
                    rows = cursor.fetchall()
                if not rows:
                    logger.warning(f"⚠️ No requirements found for selected compliances: {selected_compliance_ids}")
                    return {'success': False, 'error': f'No requirements found for selected compliances: {selected_compliance_ids}'}

                logger.info(f"✅ Loaded {len(rows)} requirement(s) for selected compliances: {selected_compliance_ids}")
                requirements = []
                for r in rows:
                    description = r[2] or ''
                    title = description if description and len(description) > 10 else (r[1] or f"Compliance {r[0]}")
                    compliance_title = r[1] or f"Compliance {r[0]}"  # Short title from ComplianceTitle column
                    requirements.append({
                        'compliance_id': r[0],
                        'title': title,
                        'compliance_title': compliance_title,  # Add short compliance title
                        'description': description,
                        'type': r[3] or 'General',
                        'risk': r[4] or 'Medium',
                        'mandatory': (r[5] or '').lower() == 'mandatory',
                        'subpolicy_id': r[6],
                        'subpolicy_name': r[7],
                        'policy_id': r[8],
                        'policy_name': r[9]
                    })
            except Exception as e:
                logger.error(f"❌ Error loading requirements: {e}")
                return {'success': False, 'error': 'Error loading requirements'}
        else:
            # Fallback to policy/subpolicy requirements (only if no selected_compliance_ids)
            logger.info(f"📋 No selected compliances provided - using all compliances for policy_id={policy_id}, subpolicy_id={subpolicy_id}")
            requirements = _get_policy_requirements(policy_id, subpolicy_id)
            if not requirements:
                return {'success': False, 'error': 'No requirements for policy'}

        # =============================
        # ALWAYS EXTRACT TEXT FROM DOCUMENT
        # =============================
        # Extract text content (and inferred schema if Excel)
        text = extract_text_from_document(full_path, doc_type or 'text/plain')
        inferred_schema = None
        try:
            if '__INFERRED_SCHEMA_START__' in text:
                import re, json as _json
                m = re.search(r'__INFERRED_SCHEMA_START__\n([\s\S]*?)\n__INFERRED_SCHEMA_END__', text)
                if m:
                    inferred_schema = _json.loads(m.group(1))
                    text = text.replace(m.group(0), '')
        except Exception:
            inferred_schema = None

        # Do not send extraction errors or empty content to the AI
        _err_prefixes = ("Error extracting text:", "PDF extraction failed:", "Excel extraction failed:", "Word extraction failed:", "TXT extraction failed:", "Document type ")
        _t = (text or "").strip()
        if not _t or len(_t) < 30:
            logger.warning(f"⚠️ [AUTO] Document produced no usable text (length={len(_t or '')})")
            return {'success': False, 'error': 'Document could not be read or has no extractable text. Check file format and that it is not corrupted.'}
        if any(_t.startswith(p) for p in _err_prefixes):
            logger.warning(f"⚠️ [AUTO] Extraction failed: {_t[:120]}")
            return {'success': False, 'error': 'Document text extraction failed. File may be wrong type or corrupted.'}

        # =============================
        # CHECK FOR COMBINED EVIDENCE
        # =============================
        # If both document and database evidence exist for the same compliances,
        # automatically use combined evidence approach (one judgment per requirement)
        compliance_ids_to_check = [r['compliance_id'] for r in requirements]
        
        with connection.cursor() as cursor:
            # Check if database evidence exists for any of these compliances
            placeholders = ",".join(["%s"] * len(compliance_ids_to_check))
            cursor.execute(
                f"""
                SELECT COUNT(DISTINCT d.document_id) as db_evidence_count
                FROM ai_audit_data d
                WHERE d.audit_id = %s
                  AND (d.external_source = 'database_record' OR d.document_type = 'db_record')
                  AND (
                      (d.policy_id = %s AND d.subpolicy_id = %s)
                      OR EXISTS (
                          SELECT 1 FROM compliance c
                          WHERE c.ComplianceId IN ({placeholders})
                            AND c.SubPolicyId = d.subpolicy_id
                      )
                  )
                """,
                [int(audit_id), policy_id, subpolicy_id] + compliance_ids_to_check
            )
            db_evidence_result = cursor.fetchone()
            db_evidence_count = db_evidence_result[0] if db_evidence_result else 0
        
        # If database evidence exists, use combined approach
        if db_evidence_count > 0:
            logger.info(f"🔍 [AUTO] ✅ Both document and database evidence found ({db_evidence_count} DB records).")
            logger.info(f"🔍 [AUTO] 📋 Using COMBINED EVIDENCE approach - will analyze both together in ONE check.")
            
            # Call internal combined evidence function
            try:
                result = _check_compliance_with_combined_evidence_internal(
                    audit_id=audit_id,
                    compliance_ids=compliance_ids_to_check,
                    user_id=user_id,
                    primary_document_id=int(document_id)  # Pass document_id so status can be updated even if download fails
                )
                
                # Clean up temporary file before returning
                if temp_file_created and 'full_path' in locals():
                    try:
                        if os.path.exists(full_path):
                            os.remove(full_path)
                    except Exception:
                        pass
                
                if result.get('success'):
                    # Combined evidence processed all compliances - return success
                    return {
                        'success': True,
                        'document_id': int(document_id),
                        'audit_id': audit_id,
                        'status': 'COMBINED_EVIDENCE_USED',
                        'message': f'Used combined evidence approach with {db_evidence_count} database records',
                        'results': result.get('results', [])
                    }
                else:
                    # If combined evidence failed, fall through to document-only
                    logger.warning(f"⚠️ Combined evidence check failed: {result.get('error')}, falling back to document-only")
                    # Fall through to document-only check
            except Exception as combined_err:
                logger.warning(f"⚠️ Combined evidence check failed, falling back to document-only: {combined_err}")
                # Fall through to document-only check

        # =============================
        # DOCUMENT-ONLY EVIDENCE PATH
        # =============================
        logger.info(f"🔍 [AUTO] Only document evidence found. Using document-only compliance check.")

        # Deterministic signals from schema/sample
        signals = _compute_basic_signals(inferred_schema)

        # AI scoring
        analyses = _ai_score_requirements_with_openai(text, requirements, schema=inferred_schema, audit_id=audit_id, document_id=document_id)
        status_label, confidence = _determine_status(requirements, analyses)

        # Blend signals
        if not signals.get('core_fields_ok'):
            confidence = max(0.0, confidence - 0.2)

        # Persist results
        try:
            import json
            from datetime import datetime
            with connection.cursor() as cursor:
                cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                framework_result = cursor.fetchone()
                if not framework_result or framework_result[0] is None:
                    return {'success': False, 'error': 'Audit not found or has no FrameworkId'}
                framework_id = framework_result[0]
                
                if external_source in ['evidence_attachment', 's3'] and external_id:
                    where_clause = "WHERE audit_id = %s AND external_id = %s"
                    where_params = [int(audit_id), external_id]
                elif doc_path:
                    where_clause = "WHERE audit_id = %s AND document_path = %s"
                    where_params = [int(audit_id), doc_path]
                else:
                    where_clause = "WHERE document_id = %s AND audit_id = %s"
                    where_params = [int(document_id), int(audit_id)]

                update_sql = f"""
                    UPDATE ai_audit_data 
                    SET ai_processing_status = 'completed',
                        compliance_status = %s,
                        confidence_score = %s,
                        compliance_analyses = %s,
                        processing_completed_at = NOW(),
                        FrameworkId = %s
                    {where_clause}
                """

                cursor.execute(
                    update_sql,
                    [
                        status_label,
                        float(confidence),
                        json.dumps({
                            "compliance_status": status_label,
                            "confidence_score": float(confidence),
                            "compliance_analyses": analyses,
                            "processed_at": datetime.now().isoformat(),
                        }),
                        framework_id,
                        *where_params,
                    ],
                )
                rows_updated = cursor.rowcount
                logger.info(f"✅ Updated {rows_updated} record(s) for document {document_id}")
                
                # Check if all documents for this audit are completed, and update audit status accordingly
                try:
                    # Check if this is an AI audit
                    cursor.execute("SELECT AuditType, Status FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                    audit_info = cursor.fetchone()
                    
                    if audit_info and audit_info[0] == 'A':  # AI Audit
                        # Count total documents and completed documents for this audit
                        cursor.execute("""
                            SELECT 
                                COUNT(*) as total,
                                SUM(CASE WHEN ai_processing_status = 'completed' THEN 1 ELSE 0 END) as completed,
                                SUM(CASE WHEN ai_processing_status = 'failed' THEN 1 ELSE 0 END) as failed
                            FROM ai_audit_data 
                            WHERE audit_id = %s
                        """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        
                        doc_stats = cursor.fetchone()
                        total_docs = doc_stats[0] if doc_stats else 0
                        completed_docs = doc_stats[1] if doc_stats else 0
                        failed_docs = doc_stats[2] if doc_stats else 0
                        
                        logger.info(f"📊 Audit {audit_id} document status: total={total_docs}, completed={completed_docs}, failed={failed_docs}")
                        
                        # Update audit status if all documents are processed (completed or failed)
                        if total_docs > 0 and (completed_docs + failed_docs) == total_docs:
                            # All documents processed - mark AI audit as fully completed
                            from ...models import Audit
                            try:
                                # Get tenant_id from audit if not available
                                if 'tenant_id' not in locals() or tenant_id is None:
                                    cursor.execute("SELECT TenantId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                                    tenant_row = cursor.fetchone()
                                    tenant_id = tenant_row[0] if tenant_row and tenant_row[0] else None
                                
                                if tenant_id:
                                    audit_obj = Audit.objects.get(AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id=tenant_id)
                                else:
                                    audit_obj = Audit.objects.get(AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id)
                                current_status = audit_obj.Status
                                logger.info(f"🔍 Current audit status: '{current_status}'")
                                
                                if audit_obj.Status != 'Completed':
                                    # Get auditor user_id for creating audit_version (extract integer ID from User object if needed)
                                    auditor_id = audit_obj.Auditor if hasattr(audit_obj, 'Auditor') and audit_obj.Auditor else None
                                    if not auditor_id:
                                        cursor.execute("SELECT Auditor FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                                        auditor_row = cursor.fetchone()
                                        auditor_id = auditor_row[0] if auditor_row and auditor_row[0] else None
                                    
                                    # Extract integer ID if auditor_id is a User object
                                    if auditor_id:
                                        if hasattr(auditor_id, 'UserId'):
                                            auditor_id = auditor_id.UserId
                                        elif hasattr(auditor_id, 'id'):
                                            auditor_id = auditor_id.id
                                        elif hasattr(auditor_id, 'pk'):
                                            auditor_id = auditor_id.pk
                                        elif not isinstance(auditor_id, int):
                                            try:
                                                auditor_id = int(auditor_id)
                                            except (ValueError, TypeError):
                                                auditor_id = None
                                    
                                    audit_obj.Status = 'Under review'
                                    if not getattr(audit_obj, "ReviewStartDate", None):
                                        audit_obj.ReviewStartDate = timezone.now()
                                    audit_obj.save()
                                    logger.info(f"✅ Updated AI audit {audit_id} status from '{current_status}' to 'Under review' (all {total_docs} documents processed)")
                                    
                                    # Automatically create audit_version with all AI-generated findings
                                    try:
                                        from .audit_views import create_audit_version
                                        if auditor_id:
                                            version_result = create_audit_version(audit_id, auditor_id)
                                            logger.info(f"✅ Automatically created audit_version for AI audit {audit_id} (document handling): {version_result}")
                                        else:
                                            logger.warning(f"⚠️ Could not create audit_version for audit {audit_id}: No auditor_id found")
                                    except Exception as version_err:
                                        logger.error(f"❌ Failed to create audit_version automatically (document handling): {version_err}")
                                        import traceback
                                        logger.error(f"Traceback: {traceback.format_exc()}")
                                    
                                    # Create incidents for non-compliant and partially compliant findings
                                    try:
                                        from .reviewing import create_incidents_for_findings
                                        logger.info(f"🔍 Creating incidents for non-compliant findings in AI audit {audit_id}")
                                        # tenant_id is already retrieved above (line 4775-4778)
                                        create_incidents_for_findings(int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id)
                                        logger.info(f"✅ Successfully created incidents for AI audit {audit_id}")
                                    except Exception as incident_err:
                                        logger.error(f"❌ Failed to create incidents for AI audit {audit_id}: {incident_err}")
                                        import traceback
                                        logger.error(f"Traceback: {traceback.format_exc()}")
                                else:
                                    logger.info(f"ℹ️ Audit {audit_id} already in status '{current_status}', skipping status update")
                            except Exception as audit_update_err:
                                logger.error(f"❌ Failed to update audit status: {audit_update_err}")
                                import traceback
                                logger.error(f"Traceback: {traceback.format_exc()}")
                                raise  # Re-raise to be caught by outer except
                except Exception as status_check_err:
                    logger.error(f"⚠️ Could not check/update audit completion status: {status_check_err}")
                    import traceback
                    logger.error(f"Traceback: {traceback.format_exc()}")
                
                # Save to lastchecklistitemverified
                try:
                    save_ai_compliance_to_checklist(
                        audit_id=audit_id,
                        document_id=document_id,
                        analyses=analyses,
                        user_id=user_id,
                        framework_id=framework_id,
                        policy_id=policy_id,
                        subpolicy_id=subpolicy_id
                    )
                except Exception as e:
                    logger.warning(f"⚠️ Could not save to lastchecklistitemverified: {e}")
        except Exception as e:
            logger.warning(f"ℹ️ Could not persist compliance results: {e}")
            return {'success': False, 'error': f'Failed to persist results: {str(e)}'}

        # Clean up temporary file
        if temp_file_created and 'full_path' in locals():
            try:
                if os.path.exists(full_path):
                    os.remove(full_path)
            except Exception:
                pass

        return {
            'success': True,
            'document_id': int(document_id),
            'audit_id': audit_id,
            'status': status_label,
            'confidence': round(confidence, 2),
            'analyses': analyses
        }
    except Exception as e:
        logger.error(f"❌ Error in internal compliance check: {e}")
        import traceback
        traceback.print_exc()
        return {'success': False, 'error': str(e)}


def _compute_basic_signals(inferred_schema: dict) -> dict:
    """Infer simple, deterministic signals from inferred Excel schema/sample rows.
    Signals are generic and resilient to header variations.
    """
    try:
        if not inferred_schema or 'sheets' not in inferred_schema:
            return {'core_fields_ok': False, 'fields': {}, 'coverage': {}}

        # Synonym map
        synonyms = {
            'user_id': ['employee id', 'employee_id', 'user id', 'user_id', 'userid', 'username'],
            'time_in': ['login time', 'time in', 'signin', 'sign_in', 'time_in', 'access time'],
            'time_out': ['logoff time', 'logout time', 'time out', 'signout', 'sign_out', 'time_out'],
            'action': ['action', 'event', 'activity', 'operation'],
            'mfa': ['mfa', 'two factor', '2fa', 'multifactor', 'otp']
        }

        def normalize(h):
            return (h or '').strip().lower()

        found_fields = {}
        coverage = {}

        for sheet in inferred_schema.get('sheets', [])[:1]:  # first sheet is usually primary
            headers = [normalize(h) for h in (sheet.get('headers') or [])]
            sample = sheet.get('sample_rows') or []
            # Map headers by synonym
            for key, syns in synonyms.items():
                idx = next((i for i, h in enumerate(headers) if any(s in h for s in syns)), None)
                if idx is not None:
                    found_fields[key] = headers[idx]
                    if sample:
                        non_empty = 0
                        total = len(sample)
                        for r in sample:
                            val = r.get(headers[idx])
                            if val not in (None, ''):
                                non_empty += 1
                        coverage[key] = round(non_empty / max(total, 1), 2)
                else:
                    coverage[key] = 0.0

        core_fields = ['user_id', 'time_in']
        core_ok = all(found_fields.get(f) for f in core_fields)
        return {'core_fields_ok': core_ok, 'fields': found_fields, 'coverage': coverage}
    except Exception:
        return {'core_fields_ok': False, 'fields': {}, 'coverage': {}}


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def check_document_compliance(request, audit_id, document_id):
    """Run compliance check for a single mapped document using OpenAI.
    MULTI-TENANCY: Only checks compliance for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        # Check authentication using JWT (like other endpoints)
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        logger.info(f"🔍 Compliance check request from user: {user_id}")
        
        # Fetch audit metadata for context-aware analysis
        audit_title = None
        audit_objective = None
        audit_scope = None
        audit_business_unit = None
        audit_type = None
        audit_due_date = None
        
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT Title, Objective, Scope, BusinessUnit, AuditType, DueDate
                FROM audit
                WHERE AuditId = %s AND TenantId = %s
                """,
                [int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id]
            )
            audit_row = cursor.fetchone()
            if audit_row:
                audit_title = audit_row[0] or None
                audit_objective = audit_row[1] or None
                audit_scope = audit_row[2] or None
                audit_business_unit = audit_row[3] or None
                audit_type = audit_row[4] or None
                audit_due_date = audit_row[5] or None
                logger.info(f"📋 Audit context: Title='{audit_title}', Objective='{audit_objective[:100] if audit_objective else None}...', Scope='{audit_scope[:100] if audit_scope else None}...'")
        
        # Lookup document path, mime, and policy mapping from ai_audit_data table (tenant-scoped)
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT d.document_path, COALESCE(d.mime_type, d.document_type) AS doc_type,
                       COALESCE(d.policy_id, a.PolicyId) AS policy_id,
                       COALESCE(d.subpolicy_id, a.SubPolicyId) AS subpolicy_id,
                       d.external_source, d.external_id, d.document_name, d.file_size
                FROM ai_audit_data d
                JOIN audit a ON a.AuditId = d.audit_id AND a.TenantId = %s
                WHERE d.document_id = %s AND d.audit_id = %s
                """,
                [tenant_id, int(document_id), int(audit_id) if str(audit_id).isdigit() else audit_id]
            )
            row = cursor.fetchone()
            if not row:
                logger.warning(f"🔍 Document not found: document_id={document_id}, audit_id={audit_id}, tenant_id={tenant_id}")
                return Response({'success': False, 'error': 'Document not found. Ensure the document is uploaded and linked to this audit.'}, status=status.HTTP_404_NOT_FOUND)
            doc_path, doc_type, policy_id, subpolicy_id, external_source, external_id, document_name, file_size = row
            # Decrypt document_path if stored encrypted (so local fallback can match ai_audit_documents/...)
            try:
                from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
                doc_path = decrypt_any_encrypted_value(doc_path) if doc_path else doc_path
            except Exception:
                pass
            logger.info(f"📋 Document lookup: external_source={external_source}, external_id={external_id} (type: {type(external_id)}), document_name={document_name}, file_size={file_size}")
        
        # Fetch policy and subpolicy names for audit context
        policy_name = None
        subpolicy_name = None
        if policy_id:
            with connection.cursor() as cursor:
                try:
                    cursor.execute("SELECT PolicyName FROM policies WHERE PolicyId = %s", [int(policy_id)])
                    policy_row = cursor.fetchone()
                    if policy_row:
                        policy_name = policy_row[0]
                except Exception as e:
                    logger.warning(f"⚠️ Could not fetch policy name: {e}")
        
        if subpolicy_id:
            with connection.cursor() as cursor:
                try:
                    cursor.execute("SELECT SubPolicyName FROM subpolicies WHERE SubPolicyId = %s", [int(subpolicy_id)])
                    subpolicy_row = cursor.fetchone()
                    if subpolicy_row:
                        subpolicy_name = subpolicy_row[0]
                except Exception as e:
                    logger.warning(f"⚠️ Could not fetch subpolicy name: {e}")

        # Handle file path - check if it's S3 or local
        if external_source in ['s3', 'evidence_attachment'] and external_id:
            # Handle S3 file
            try:
                import json
                s3_metadata = None
                s3_key = None
                file_name_from_db = None
                s3_url_for_fallback = None
                stored_name_for_fallback = None
                
                # For documents from document handling (evidence_attachment), try to look up from file_operations table
                if external_source == 'evidence_attachment':
                    logger.info(f"🔍 Looking up file_operations for external_id: {external_id}")
                    try:
                        with connection.cursor() as cursor:
                            # Try multiple strategies to find the file:
                            # 1. Exact match on s3_key, stored_name, file_name, or original_name
                            # 2. Match where s3_key ends with external_id (in case external_id is just filename)
                            # 3. Match where file_name or original_name matches external_id
                            cursor.execute("""
                                SELECT s3_key, file_name, original_name, s3_url, stored_name
                                FROM file_operations
                                WHERE operation_type = 'upload'
                                  AND status = 'completed'
                                  AND (
                                    s3_key = %s
                                    OR stored_name = %s
                                    OR file_name = %s
                                    OR original_name = %s
                                    OR s3_key LIKE %s
                                    OR s3_key LIKE %s
                                  )
                                ORDER BY 
                                    CASE 
                                        WHEN s3_key = %s THEN 1
                                        WHEN stored_name = %s THEN 2
                                        WHEN file_name = %s THEN 3
                                        WHEN original_name = %s THEN 4
                                        ELSE 5
                                    END
                                LIMIT 1
                            """, [
                                external_id, external_id, external_id, external_id,  # Exact matches
                                f'%/{external_id}', f'%{external_id}',  # Pattern matches (with or without leading slash)
                                external_id, external_id, external_id, external_id  # For ordering
                            ])
                            file_op = cursor.fetchone()
                            
                            if file_op:
                                s3_key_from_db = file_op[0]  # s3_key from file_operations
                                file_name_from_db = file_op[1] or file_op[2]  # file_name or original_name
                                s3_url_from_db = file_op[3]  # s3_url
                                stored_name_from_db = file_op[4] if len(file_op) > 4 else None  # stored_name
                                logger.info(f"✅ Found file_operations record: s3_key={s3_key_from_db}, s3_url={s3_url_from_db}, file_name={file_name_from_db}, stored_name={stored_name_from_db}")
                                
                                # Prioritize extracting s3_key from s3_url as it's more reliable
                                if s3_url_from_db:
                                    if 'amazonaws.com/' in s3_url_from_db:
                                        extracted_key = s3_url_from_db.split('amazonaws.com/')[-1].split('?')[0]
                                        if extracted_key:
                                            s3_key = extracted_key
                                            logger.info(f"✅ Extracted s3_key from s3_url: {s3_key}")
                                    elif 's3.' in s3_url_from_db or 's3-' in s3_url_from_db:
                                        # Handle other S3 URL formats
                                        try:
                                            from urllib.parse import urlparse
                                            parsed = urlparse(s3_url_from_db)
                                            if parsed.path:
                                                s3_key = parsed.path.lstrip('/')
                                                logger.info(f"✅ Extracted s3_key from s3_url (parsed): {s3_key}")
                                        except Exception as parse_err:
                                            logger.warning(f"⚠️ Could not parse s3_url: {parse_err}")
                                
                                # Fallback to s3_key from database if we couldn't extract from URL
                                if not s3_key and s3_key_from_db:
                                    s3_key = s3_key_from_db
                                    logger.info(f"✅ Using s3_key from database: {s3_key}")
                                
                                # Store for potential fallback
                                s3_url_for_fallback = s3_url_from_db if s3_url_from_db else None
                                stored_name_for_fallback = stored_name_from_db
                            else:
                                logger.warning(f"⚠️ File not found in file_operations table for external_id: {external_id}")
                                # Try to use external_id as s3_key if it looks like a valid key
                                if external_id and '/' in external_id:
                                    s3_key = external_id
                                    logger.info(f"⚠️ Using external_id as s3_key: {s3_key}")
                    except Exception as lookup_err:
                        logger.warning(f"⚠️ Error looking up file_operations: {lookup_err}")
                        import traceback
                        traceback.print_exc()
                
                # If we didn't find it in file_operations, try to extract from external_id
                if not s3_key:
                    # external_id might be JSON, a raw S3 key, or a full S3 URL
                    if isinstance(external_id, str):
                        raw_value = external_id.strip()
                        try:
                            s3_metadata = json.loads(raw_value)
                        except Exception:
                            # Not JSON; handle as key or URL
                            if 'amazonaws.com/' in raw_value:
                                s3_key = raw_value.split('amazonaws.com/')[-1].split('?')[0]
                            else:
                                s3_key = raw_value
                    elif isinstance(external_id, dict):
                        # Legacy structure where external_id is a JSON/dict payload
                        s3_metadata = external_id
                    else:
                        # external_id is neither string nor dict (e.g. numeric operation_id) –
                        # we've already tried file_operations lookup above, so nothing more to derive here
                        s3_metadata = None

                    if not s3_key and isinstance(s3_metadata, dict):
                        s3_key = s3_metadata.get('s3_key')
                        if not s3_key and s3_metadata.get('aws_file_link'):
                            aws_link = s3_metadata.get('aws_file_link')
                            if 'amazonaws.com/' in aws_link:
                                s3_key = aws_link.split('amazonaws.com/')[-1].split('?')[0]
                
                # Also check document_path as it might contain the s3_key
                if not s3_key and doc_path:
                    # document_path might be the s3_key itself or contain it
                    if 'amazonaws.com/' in doc_path:
                        s3_key = doc_path.split('amazonaws.com/')[-1].split('?')[0]
                    elif not doc_path.startswith('/') and not doc_path.startswith('ai_audit_documents'):
                        # If it doesn't look like a local path, it might be an s3_key
                        s3_key = doc_path
                
                if not s3_key:
                    return Response({'success': False, 'error': 'S3 key not found in metadata'}, status=status.HTTP_400_BAD_REQUEST)
                
                # Download file from S3 to temporary location
                from ..Global.s3_fucntions import create_direct_mysql_client
                import tempfile
                
                s3_client = create_direct_mysql_client()
                # Use file_name from database if available, otherwise extract from s3_key
                if file_name_from_db:
                    file_name = file_name_from_db
                else:
                    file_name = s3_key.split('/')[-1] if '/' in s3_key else s3_key
                
                # Default filename: use extension from document_name, external_id, or s3_key so type is correct (e.g. .xlsx not .pdf)
                if not file_name or '.' not in file_name:
                    ext = '.pdf'
                    for src in (document_name, (external_id if isinstance(external_id, str) else None), (s3_key.split('/')[-1] if s3_key and '/' in s3_key else s3_key)):
                        if src and '.' in str(src):
                            candidate = ('.' + str(src).rsplit('.', 1)[-1].strip()).lower()
                            if candidate in ('.pdf', '.xlsx', '.xls', '.docx', '.doc', '.txt'):
                                ext = candidate
                                break
                    file_name = f"document_{document_id}{ext}"
                
                logger.info(f"🔍 Downloading from S3: s3_key={s3_key}, file_name={file_name}")
                
                # Validate s3_key before attempting download
                if not s3_key or len(s3_key.strip()) == 0:
                    return Response({'success': False, 'error': 'Invalid S3 key: key is empty'}, status=status.HTTP_400_BAD_REQUEST)
                
                # Create temporary file
                temp_dir = tempfile.gettempdir()
                temp_file_path = os.path.join(temp_dir, f"ai_audit_{document_id}_{file_name}")
                
                # Attempt download with detailed error logging and fallback strategies
                download_success = False
                last_error = None
                # Prefer direct s3_url from file_operations (upload row) so we hit the actual storage, not the microservice
                if s3_url_for_fallback and ('https://' in s3_url_for_fallback or 'http://' in s3_url_for_fallback):
                    try:
                        import requests as req_lib
                        logger.info(f"💡 Trying direct download from stored S3 URL first")
                        url_response = req_lib.get(s3_url_for_fallback, timeout=60, stream=True)
                        if url_response.status_code == 200:
                            with open(temp_file_path, 'wb') as f:
                                for chunk in url_response.iter_content(chunk_size=8192):
                                    f.write(chunk)
                            if os.path.exists(temp_file_path):
                                download_success = True
                                full_path = temp_file_path
                                temp_file_created = True
                                logger.info(f"✅ Direct download from S3 URL successful")
                    except Exception as direct_err:
                        last_error = str(direct_err)
                        logger.warning(f"⚠️ Direct S3 URL download failed: {direct_err}")
                s3_keys_to_try = [s3_key] if s3_key else []
                if stored_name_for_fallback and stored_name_for_fallback not in s3_keys_to_try:
                    s3_keys_to_try.append(stored_name_for_fallback)
                
                # Try microservice only if direct URL did not succeed
                for attempt_key in s3_keys_to_try:
                    try:
                        logger.info(f"🔄 Attempting download with s3_key: {attempt_key}")
                        download_result = s3_client.download(attempt_key, file_name, temp_dir, str(user_id))
                        if download_result.get('success'):
                            download_success = True
                            s3_key = attempt_key  # Update to the working key
                            logger.info(f"✅ Download successful with s3_key: {attempt_key}")
                            break
                        else:
                            last_error = download_result.get('error', 'Unknown error')
                            logger.warning(f"⚠️ Download failed with s3_key {attempt_key}: {last_error}")
                    except Exception as download_ex:
                        last_error = str(download_ex)
                        logger.warning(f"⚠️ Exception with s3_key {attempt_key}: {last_error}")
                        import traceback
                        traceback.print_exc()
                
                if not download_success:
                    error_msg = last_error or 'Unknown error'
                    # Log as WARNING when we may still fall back to local file; ERROR only when we will return 500
                    will_try_local = bool(doc_path and (doc_path.startswith('ai_audit_documents') or doc_path.startswith('/')))
                    if will_try_local:
                        logger.warning(f"⚠️ S3 download failed (will try local file). Last error: {error_msg}")
                        logger.warning(f"⚠️ Tried s3_keys: {s3_keys_to_try}")
                    else:
                        logger.error(f"❌ All S3 download attempts failed. Last error: {error_msg}")
                        logger.error(f"❌ Tried s3_keys: {s3_keys_to_try}")
                        logger.error(f"❌ file_name: {file_name}, external_id: {external_id}")
                    
                    # If we have s3_url, try direct download as last resort
                    if s3_url_for_fallback:
                        logger.info(f"💡 Attempting direct download from S3 URL: {s3_url_for_fallback}")
                        # Try one more time with the s3_url if it's a pre-signed URL
                        if 'https://' in s3_url_for_fallback or 'http://' in s3_url_for_fallback:
                            try:
                                import requests as req_lib
                                logger.info(f"🔄 Attempting direct download from S3 URL...")
                                url_response = req_lib.get(s3_url_for_fallback, timeout=60, stream=True)
                                if url_response.status_code == 200:
                                    # Save the file directly
                                    with open(temp_file_path, 'wb') as f:
                                        for chunk in url_response.iter_content(chunk_size=8192):
                                            f.write(chunk)
                                    if os.path.exists(temp_file_path):
                                        download_success = True
                                        logger.info(f"✅ Direct download from S3 URL successful!")
                                        full_path = temp_file_path
                                        temp_file_created = True
                            except Exception as direct_dl_err:
                                logger.warning(f"⚠️ Direct download from URL also failed: {direct_dl_err}")
                    
                    # If still not successful, try local file when doc_path points to ai_audit_documents
                    if not download_success and doc_path:
                        local_full = doc_path if os.path.isabs(doc_path) else os.path.join(settings.MEDIA_ROOT, doc_path)
                        if (doc_path.startswith('ai_audit_documents') or doc_path.startswith('/')) and os.path.exists(local_full):
                            full_path = local_full
                            temp_file_created = True  # so we don't overwrite full_path in the block below
                            download_success = True
                            logger.warning(f"⚠️ Using local file after S3 failure: {full_path}")
                    
                    # If still not successful, return error
                    if not download_success:
                        return Response({
                            'success': False, 
                            'error': f'Failed to download from S3: {error_msg}',
                            'details': {
                                's3_keys_tried': s3_keys_to_try,
                                'file_name': file_name,
                                'external_id': external_id,
                                'external_source': external_source,
                                's3_url_attempted': s3_url_for_fallback if s3_url_for_fallback else None,
                                'suggestion': 'The file may not exist in S3 with the provided key, or the key format may be incorrect. Please verify the file exists in S3.'
                            }
                        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # If download was successful via microservice (not direct URL), get the file path
                if download_success and not temp_file_created and 'download_result' in locals():
                    temp_file_path = download_result.get('file_path', temp_file_path)
                    if not temp_file_path or not os.path.exists(temp_file_path):
                        return Response({'success': False, 'error': 'Failed to download file from S3'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    
                    full_path = temp_file_path
                    temp_file_created = True
                    logger.info(f"🔍 Downloaded S3 file to: {full_path}")
                
            except Exception as e:
                logger.error(f"❌ Error handling S3 file: {e}")
                import traceback
                traceback.print_exc()
                return Response({'success': False, 'error': f'S3 file handling error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            # Handle local file
            full_path = doc_path if os.path.isabs(doc_path) else os.path.join(settings.MEDIA_ROOT, doc_path)
            if not os.path.exists(full_path):
                return Response({'success': False, 'error': 'File not found on server'}, status=status.HTTP_404_NOT_FOUND)
            temp_file_created = False

        # =============================
        # Load requirements to analyze (needed to check for combined evidence)
        # =============================
        # If frontend sent explicit compliance IDs, build requirements directly from them
        selected_ids = None
        try:
            raw_selected = request.data.get('selected_compliance_ids') or request.data.get('compliance_ids')
        except Exception:
            raw_selected = None

        if raw_selected:
            # Handle JSON string or list
            if isinstance(raw_selected, str):
                try:
                    import json as _json
                    parsed = _json.loads(raw_selected)
                    if isinstance(parsed, list):
                        raw_selected = parsed
                except Exception:
                    pass
            if isinstance(raw_selected, list):
                try:
                    selected_ids = {int(x) for x in raw_selected if str(x).isdigit()}
                except Exception:
                    selected_ids = None

        if selected_ids:
            logger.info(f"🔍 check_document_compliance using explicit selected_compliance_ids: {len(selected_ids)} items")
            # Build requirements list directly from compliance IDs (ignoring policy/subpolicy)
            try:
                with connection.cursor() as cursor:
                    placeholders = ",".join(["%s"] * len(selected_ids))
                    cursor.execute(
                        f"""
                        SELECT c.ComplianceId, c.ComplianceTitle, c.ComplianceItemDescription,
                               c.ComplianceType, c.Criticality, c.MandatoryOptional,
                               sp.SubPolicyId, sp.SubPolicyName
                        FROM compliance c
                        JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                        WHERE c.ComplianceId IN ({placeholders})
                        """,
                        list(selected_ids)
                    )
                    rows = cursor.fetchall()
                if not rows:
                    return Response({'success': False, 'error': 'No requirements found for selected compliances'}, status=status.HTTP_400_BAD_REQUEST)

                requirements = []
                from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
                for r in rows:
                    # Decrypt encrypted fields from database
                    compliance_title = decrypt_any_encrypted_value(r[1]) if r[1] else None
                    description = decrypt_any_encrypted_value(r[2]) if r[2] else ''
                    subpolicy_name = decrypt_any_encrypted_value(r[7]) if r[7] else None
                    
                    title = description if description and len(description) > 10 else (compliance_title or f"Compliance {r[0]}")
                    requirements.append({
                        'compliance_id': r[0],
                        'title': title,
                        'description': description,
                        'type': r[3] or 'General',
                        'risk': r[4] or 'Medium',
                        'mandatory': (r[5] or '').lower() == 'mandatory',
                        'subpolicy_id': r[6],
                        'subpolicy_name': subpolicy_name
                    })
            except Exception as e:
                logger.error(f"❌ Error loading requirements for explicit compliance IDs: {e}")
                return Response({'success': False, 'error': 'Error loading requirements for selected compliances'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            # Default behavior: derive requirements from document's policy/subpolicy mapping
            requirements = _get_policy_requirements(policy_id, subpolicy_id)
            if not requirements:
                return Response({'success': False, 'error': 'No requirements for policy'}, status=status.HTTP_400_BAD_REQUEST)

        # NOTE: Previously we capped to first 10 requirements for performance.
        # This limit has been removed so that all selected compliances are analyzed.

        # =============================
        # ALWAYS EXTRACT TEXT FROM DOCUMENT
        # =============================
        # Extract text content (and inferred schema if Excel)
        text = extract_text_from_document(full_path, doc_type or 'text/plain')
        inferred_schema = None
        try:
            if '__INFERRED_SCHEMA_START__' in text:
                import re, json as _json
                m = re.search(r'__INFERRED_SCHEMA_START__\n([\s\S]*?)\n__INFERRED_SCHEMA_END__', text)
                if m:
                    inferred_schema = _json.loads(m.group(1))
                    # Remove schema marker from text before sending to model
                    text = text.replace(m.group(0), '')
        except Exception:
            inferred_schema = None

        # Do not send extraction errors or empty content to the AI (avoids "NO EVIDENCE FOUND" fallback)
        _extraction_error_prefixes = (
            "Error extracting text:",
            "PDF extraction failed:",
            "Excel extraction failed:",
            "Word extraction failed:",
            "TXT extraction failed:",
            "Document type ",
        )
        text_stripped = (text or "").strip()
        if not text_stripped or len(text_stripped) < 30:
            logger.warning(f"⚠️ Document produced no usable text (length={len(text_stripped or '')})")
            return Response({
                'success': False,
                'error': 'Document could not be read or has no extractable text. Check that the file is a supported format (PDF, Word, Excel, TXT) and not corrupted.',
                'detail': 'extraction_empty_or_failed'
            }, status=status.HTTP_400_BAD_REQUEST)
        if any(text_stripped.startswith(p) for p in _extraction_error_prefixes):
            logger.warning(f"⚠️ Extraction failed; content sent to AI would be error message: {text_stripped[:120]}")
            return Response({
                'success': False,
                'error': 'Document text extraction failed. The file may be the wrong type (e.g. Excel saved as PDF) or corrupted. Try re-uploading in a supported format.',
                'detail': 'extraction_failed'
            }, status=status.HTTP_400_BAD_REQUEST)

        # =============================
        # CHECK FOR COMBINED EVIDENCE
        # =============================
        # IMPORTANT: For manually uploaded documents (external_source='manual'), 
        # DO NOT use combined evidence - only check the uploaded document itself.
        # Combined evidence should only be used for documents from Document Handling
        # that have database evidence for the same compliances.
        use_combined_evidence = False
        
        if external_source != 'manual':
            # If both document and database evidence exist for the same compliances,
            # automatically use combined evidence approach (one judgment per requirement)
            compliance_ids_to_check = [r['compliance_id'] for r in requirements]
            
            with connection.cursor() as cursor:
                # Check if database evidence exists for any of these compliances
                placeholders = ",".join(["%s"] * len(compliance_ids_to_check))
                cursor.execute(
                    f"""
                    SELECT COUNT(DISTINCT d.document_id) as db_evidence_count
                    FROM ai_audit_data d
                    WHERE d.audit_id = %s
                      AND (d.external_source = 'database_record' OR d.document_type = 'db_record')
                      AND (
                          (d.policy_id = %s AND d.subpolicy_id = %s)
                          OR EXISTS (
                              SELECT 1 FROM compliance c
                              WHERE c.ComplianceId IN ({placeholders})
                                AND c.SubPolicyId = d.subpolicy_id
                          )
                      )
                    """,
                    [int(audit_id), policy_id, subpolicy_id] + compliance_ids_to_check
                )
                db_evidence_result = cursor.fetchone()
                db_evidence_count = db_evidence_result[0] if db_evidence_result else 0
            
            # If database evidence exists, use combined approach
            if db_evidence_count > 0:
                use_combined_evidence = True
                logger.info(f"🔍 Both document and database evidence found ({db_evidence_count} DB records). Using combined evidence approach.")
            else:
                logger.info(f"🔍 Only document evidence found. Using document-only compliance check.")
        else:
            logger.info(f"🔍 Manually uploaded document (external_source='manual'). Using document-only compliance check (no combined evidence).")
        
        # If database evidence exists, use combined approach
        if use_combined_evidence:
            
            # Call internal combined evidence function
            from ...rbac.utils import RBACUtils
            user_id = RBACUtils.get_user_id_from_request(request)
            
            # Pass audit context to combined evidence function
            audit_context = {
                'title': audit_title,
                'objective': audit_objective,
                'scope': audit_scope,
                'business_unit': audit_business_unit,
                'audit_type': audit_type,
                'due_date': audit_due_date,
                'policy_id': policy_id,
                'policy_name': policy_name,
                'subpolicy_id': subpolicy_id,
                'subpolicy_name': subpolicy_name
            }
            result = _check_compliance_with_combined_evidence_internal(
                audit_id=audit_id,
                compliance_ids=compliance_ids_to_check,
                user_id=user_id,
                audit_context=audit_context
            )
            
            # Convert dict result to Response
            if result.get('success'):
                return Response(result, status=status.HTTP_200_OK)
            else:
                return Response(result, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # =============================
        # DOCUMENT-ONLY EVIDENCE PATH
        # =============================
        logger.info(f"🔍 Only document evidence found. Using document-only compliance check.")

        # Check if we should use background processing (for large requirement sets)
        USE_BACKGROUND_THRESHOLD = 10  # Use background processing if more than 10 requirements
        if len(requirements) > USE_BACKGROUND_THRESHOLD:
            logger.info(f"🚀 Large requirement set ({len(requirements)} requirements). Using background processing to avoid timeout.")
            
            # Create background job
            from .compliance_job_tracker import ComplianceJobTracker
            from .compliance_background_processor import process_compliance_check_background
            import threading
            
            job_id = ComplianceJobTracker.create_job(
                audit_id=int(audit_id),
                document_id=int(document_id),
                total_requirements=len(requirements),
                tenant_id=tenant_id
            )
            
            # Prepare audit context
            audit_context = {
                'title': audit_title,
                'objective': audit_objective,
                'scope': audit_scope,
                'business_unit': audit_business_unit,
                'audit_type': audit_type,
                'due_date': audit_due_date,
                'policy_id': policy_id,
                'policy_name': policy_name,
                'subpolicy_id': subpolicy_id,
                'subpolicy_name': subpolicy_name
            }
            
            # Get framework_id
            with connection.cursor() as cursor:
                cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id)])
                framework_result = cursor.fetchone()
                framework_id = framework_result[0] if framework_result and framework_result[0] else None
            
            if not framework_id:
                return Response({
                    'success': False,
                    'error': 'Audit not found or has no FrameworkId assigned'
                }, status=status.HTTP_400_BAD_REQUEST)
            
            # Start background thread
            def run_background_job():
                try:
                    process_compliance_check_background(
                        job_id=job_id,
                        audit_id=int(audit_id),
                        document_id=int(document_id),
                        requirements=requirements,
                        document_text=text,
                        inferred_schema=inferred_schema or {},
                        audit_context=audit_context,
                        framework_id=framework_id,
                        tenant_id=tenant_id,
                        user_id=user_id,
                        doc_path=doc_path,
                        document_name=document_name,
                        file_size=file_size,
                        external_source=external_source,
                        external_id=external_id,
                        policy_id=policy_id,
                        subpolicy_id=subpolicy_id
                    )
                except Exception as e:
                    logger.error(f"❌ Background job {job_id} failed: {e}", exc_info=True)
                    ComplianceJobTracker.fail_job(job_id, str(e))
            
            background_thread = threading.Thread(target=run_background_job, daemon=True)
            background_thread.start()
            
            logger.info(f"✅ Started background job {job_id} for {len(requirements)} requirements")
            
            # Return job_id immediately
            return Response({
                'success': True,
                'job_id': job_id,
                'message': f'Compliance check started in background. Processing {len(requirements)} requirements.',
                'status': 'processing',
                'total_requirements': len(requirements),
                'progress_percent': 0
            }, status=status.HTTP_202_ACCEPTED)

        # For small requirement sets, process normally (quick response)
        # Deterministic signals from schema/sample
        signals = _compute_basic_signals(inferred_schema)

        # AI scoring with audit context
        audit_context = {
            'title': audit_title,
            'objective': audit_objective,
            'scope': audit_scope,
            'business_unit': audit_business_unit,
            'audit_type': audit_type,
            'due_date': audit_due_date,
            'policy_id': policy_id,
            'policy_name': policy_name,
            'subpolicy_id': subpolicy_id,
            'subpolicy_name': subpolicy_name
        }
        analyses = _ai_score_requirements_with_openai(text, requirements, schema=inferred_schema, audit_id=audit_id, document_id=document_id, audit_context=audit_context)
        logger.info(f"🔍 Analyses type: {type(analyses)}, length: {len(analyses) if analyses else 0}")
        if analyses:
            logger.info(f"🔍 First analysis item type: {type(analyses[0])}, content: {analyses[0]}")
        status_label, confidence = _determine_status(requirements, analyses)

        # Blend signals: if core fields missing, keep status but reduce confidence
        if not signals.get('core_fields_ok'):
            confidence = max(0.0, confidence - 0.2)

        # Persist results to the new ai_audit_data table
        try:
            import json
            with connection.cursor() as cursor:
                # First, ensure the document exists in ai_audit_data table
                # Get FrameworkId from the audit record
                try:
                    logger.info(f"🔍 Querying FrameworkId for audit {audit_id} in compliance check")
                    cursor.execute("SELECT FrameworkId FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                    framework_result = cursor.fetchone()
                    logger.info(f"🔍 FrameworkId query result: {framework_result}")
                    
                    if framework_result and framework_result[0] is not None:
                        framework_id = framework_result[0]
                        logger.info(f"✅ Found FrameworkId {framework_id} for audit {audit_id}")
                    else:
                        logger.error(f"❌ No FrameworkId found for audit {audit_id}. Audit record may not exist or FrameworkId is NULL.")
                        return JsonResponse({
                            'success': False,
                            'error': f'Audit {audit_id} not found or has no FrameworkId assigned. Please ensure the audit exists and has a framework assigned.'
                        }, status=400)
                except Exception as framework_err:
                    logger.error(f"❌ Error querying FrameworkId for audit {audit_id}: {framework_err}")
                    return JsonResponse({
                        'success': False,
                        'error': f'Database error while retrieving audit framework: {framework_err}'
                    }, status=500)
                
                # UPDATE existing records instead of INSERT to avoid creating duplicates.
                # For documents coming from Document Handling (evidence_attachment),
                # the same physical file can have multiple ai_audit_data rows
                # (one per mapping) that share the same audit_id + external_id.
                # In that case we want to update **all** of them so every mapping
                # is marked as completed and gets the latest analyses.
                if external_source in ['evidence_attachment', 's3'] and external_id:
                    # Documents coming from Document Handling or S3 reuse the same
                    # external_id for all mappings of the physical file.
                    # Convert external_id to string to ensure type matching works
                    external_id_str = str(external_id) if external_id else None
                    # Use external_id as primary match, but also include document_name + file_size as fallback
                    # to ensure we catch all mappings even if external_id doesn't match exactly
                    if document_name and file_size:
                        where_clause = "WHERE audit_id = %s AND external_source = %s AND (external_id = %s OR external_id = %s OR (document_name = %s AND file_size = %s))"
                        where_params = [int(audit_id), external_source, external_id, external_id_str, document_name, file_size]
                        logger.info(
                            f"🔄 Updating all ai_audit_data rows for audit_id={audit_id}, external_source={external_source}, external_id={external_id}/{external_id_str} OR (document_name={document_name}, file_size={file_size})"
                        )
                    else:
                        where_clause = "WHERE audit_id = %s AND external_source = %s AND (external_id = %s OR external_id = %s)"
                        where_params = [int(audit_id), external_source, external_id, external_id_str]
                        logger.info(
                            f"🔄 Updating all ai_audit_data rows for audit_id={audit_id}, external_source={external_source}, external_id={external_id}/{external_id_str}"
                        )
                elif doc_path:
                    # Manually uploaded AI‑audit documents create one ai_audit_data
                    # row per mapping but share the same document_path.
                    # Update ALL such rows so every mapping reflects the latest
                    # compliance check results.
                    where_clause = "WHERE audit_id = %s AND document_path = %s"
                    where_params = [int(audit_id), doc_path]
                    logger.info(
                        f"🔄 Updating all ai_audit_data rows for audit_id={audit_id}, document_path={doc_path}"
                    )
                elif document_name and file_size:
                    # Fallback for manually uploaded documents: match by document_name + file_size
                    # This ensures all mappings for the same physical file are updated
                    # Include 'manual' external_source explicitly to catch manually uploaded documents
                    where_clause = "WHERE audit_id = %s AND document_name = %s AND file_size = %s AND (external_source IS NULL OR external_source = 'manual' OR external_source NOT IN ('evidence_attachment', 's3', 'database_record'))"
                    where_params = [int(audit_id), document_name, file_size]
                    logger.info(
                        f"🔄 Updating all ai_audit_data rows for audit_id={audit_id}, document_name={document_name}, file_size={file_size} (manually uploaded document)"
                    )
                else:
                    # Final fallback: update just this specific document record
                    where_clause = "WHERE document_id = %s AND audit_id = %s"
                    where_params = [int(document_id), int(audit_id)]
                    logger.info(
                        f"🔄 Updating single ai_audit_data row for document_id={document_id}, audit_id={audit_id} (fallback)"
                    )

                update_sql = f"""
                    UPDATE ai_audit_data 
                    SET ai_processing_status = 'completed',
                        compliance_status = %s,
                        confidence_score = %s,
                        compliance_analyses = %s,
                        processing_completed_at = NOW(),
                        FrameworkId = %s
                    {where_clause}
                    """

                cursor.execute(
                    update_sql,
                    [
                        status_label,
                        float(confidence),
                        json.dumps(
                            {
                                "compliance_status": status_label,
                                "confidence_score": float(confidence),
                                "compliance_analyses": analyses,
                                "processed_at": datetime.now().isoformat(),
                            }
                        ),
                        framework_id,
                        *where_params,
                    ],
                )
                rows_updated = cursor.rowcount
                logger.info(f"✅ Updated {rows_updated} record(s) for document {document_id} in ai_audit_data table")
                
                # If no rows were updated, try fallback matching by document_name + file_size for document handling uploads
                if rows_updated == 0 and external_source in ['evidence_attachment', 's3'] and document_name and file_size:
                    logger.warning(f"⚠️ No rows updated with external_id match, trying fallback match by document_name + file_size")
                    cursor.execute("""
                        UPDATE ai_audit_data 
                        SET ai_processing_status = 'completed',
                            compliance_status = %s,
                            confidence_score = %s,
                            compliance_analyses = %s,
                            processing_completed_at = NOW(),
                            FrameworkId = %s
                        WHERE audit_id = %s 
                          AND external_source = %s
                          AND document_name = %s 
                          AND file_size = %s
                    """, [
                        status_label,
                        float(confidence),
                        json.dumps({
                            "compliance_status": status_label,
                            "confidence_score": float(confidence),
                            "compliance_analyses": analyses,
                            "processed_at": datetime.now().isoformat(),
                        }),
                        framework_id,
                        int(audit_id),
                        external_source,
                        document_name,
                        file_size
                    ])
                    rows_updated = cursor.rowcount
                    logger.info(f"✅ Fallback update: Updated {rows_updated} record(s) for document_name={document_name}, file_size={file_size}")
                
                # If still no rows were updated, log a warning
                if rows_updated == 0:
                    logger.warning(f"⚠️ No existing record found for document_id={document_id}, audit_id={audit_id}, external_source={external_source}, external_id={external_id}. Document may need to be uploaded first.")
                        # Save to lastchecklistitemverified for standard compliance tracking
                try:
                    save_ai_compliance_to_checklist(
                        audit_id=audit_id,
                        document_id=document_id,
                        analyses=analyses,  # The compliance analyses array
                        user_id=user_id,
                        framework_id=framework_id,
                        policy_id=policy_id,
                        subpolicy_id=subpolicy_id
                    )
                    logger.info(f"✅ Saved compliance results to lastchecklistitemverified for document {document_id}")
                except Exception as e:
                    logger.warning(f"⚠️ Could not save to lastchecklistitemverified: {e}")
                    # Don't fail the whole request if this fails
                
                # Run SEBI AI Auditor checks automatically (if enabled) - non-blocking
                try:
                    if framework_id:
                        import threading
                        def run_sebi_checks():
                            try:
                                from .sebi_ai_auditor import SEBIAIAuditor
                                auditor = SEBIAIAuditor(framework_id, tenant_id)
                                if auditor.is_sebi_framework:
                                    logger.info(f"🔍 SEBI AI Auditor enabled - running automatic checks for audit {audit_id}")
                                    # Run SEBI checks (non-blocking, results stored for report)
                                    try:
                                        auditor.verify_filing_accuracy(audit_id, document_id)
                                        auditor.check_timeliness_sla(audit_id)
                                        auditor.calculate_risk_score(audit_id)
                                        auditor.detect_patterns(audit_id)
                                        logger.info(f"✅ SEBI AI Auditor checks completed for audit {audit_id}")
                                    except Exception as sebi_err:
                                        logger.warning(f"⚠️ SEBI checks failed (non-critical): {str(sebi_err)}")
                                else:
                                    logger.debug(f"ℹ️ SEBI AI Auditor not enabled for framework {framework_id}")
                            except Exception as e:
                                logger.warning(f"⚠️ SEBI AI Auditor check failed (non-critical): {str(e)}")
                        
                        # Run in background thread (non-blocking)
                        sebi_thread = threading.Thread(target=run_sebi_checks, daemon=True)
                        sebi_thread.start()
                        logger.info(f"🚀 Started SEBI AI Auditor checks in background for audit {audit_id}")
                except Exception as sebi_init_err:
                    logger.warning(f"⚠️ Could not initialize SEBI checks (non-critical): {str(sebi_init_err)}")
                
                # Check if all documents are processed and update audit status to "Under review" for AI audits
                try:
                    cursor.execute("SELECT AuditType, Status FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                    audit_info = cursor.fetchone()
                    
                    if audit_info and audit_info[0] == 'A':  # AI Audit
                        cursor.execute("""
                            SELECT 
                                COUNT(*) as total,
                                SUM(CASE WHEN ai_processing_status = 'completed' THEN 1 ELSE 0 END) as completed,
                                SUM(CASE WHEN ai_processing_status = 'failed' THEN 1 ELSE 0 END) as failed
                            FROM ai_audit_data 
                            WHERE audit_id = %s
                        """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        
                        doc_stats = cursor.fetchone()
                        total_docs = doc_stats[0] if doc_stats else 0
                        completed_docs = doc_stats[1] if doc_stats else 0
                        failed_docs = doc_stats[2] if doc_stats else 0
                        
                        if total_docs > 0 and (completed_docs + failed_docs) == total_docs:
                            from ...models import Audit
                            audit_obj = Audit.objects.get(AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id=tenant_id)
                            if audit_obj.Status != 'Completed':
                                # Get auditor user_id for creating audit_version (extract integer ID from User object if needed)
                                auditor_id = audit_obj.Auditor if hasattr(audit_obj, 'Auditor') and audit_obj.Auditor else None
                                if not auditor_id:
                                    cursor.execute("SELECT Auditor FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                                    auditor_row = cursor.fetchone()
                                    auditor_id = auditor_row[0] if auditor_row and auditor_row[0] else None
                                
                                # Extract integer ID if auditor_id is a User object
                                if auditor_id:
                                    if hasattr(auditor_id, 'UserId'):
                                        auditor_id = auditor_id.UserId
                                    elif hasattr(auditor_id, 'id'):
                                        auditor_id = auditor_id.id
                                    elif hasattr(auditor_id, 'pk'):
                                        auditor_id = auditor_id.pk
                                    elif not isinstance(auditor_id, int):
                                        try:
                                            auditor_id = int(auditor_id)
                                        except (ValueError, TypeError):
                                            auditor_id = None
                                
                                # Set to 'Under review' so reviewer can Start / Accept / Reject (reviewing.py requires Status == 'Under review')
                                audit_obj.Status = 'Under review'
                                if not getattr(audit_obj, "ReviewStartDate", None):
                                    audit_obj.ReviewStartDate = timezone.now()
                                audit_obj.save()
                                logger.info(f"✅ Updated AI audit {audit_id} status to 'Under review' after single document check (all {total_docs} documents processed)")
                                
                                # Automatically create audit_version with all AI-generated findings
                                try:
                                    from .audit_views import create_audit_version
                                    if auditor_id:
                                        version_result = create_audit_version(audit_id, auditor_id)
                                        logger.info(f"✅ Automatically created audit_version for AI audit {audit_id} (single document check): {version_result}")
                                    else:
                                        logger.warning(f"⚠️ Could not create audit_version for audit {audit_id}: No auditor_id found")
                                except Exception as version_err:
                                    logger.error(f"❌ Failed to create audit_version automatically (single document check): {version_err}")
                                    import traceback
                                    logger.error(f"Traceback: {traceback.format_exc()}")
                except Exception as status_check_err:
                    logger.warning(f"⚠️ Could not check/update audit completion status after single document check: {status_check_err}")
                    
        except Exception as e:
            logger.warning(f"ℹ️ Could not persist compliance results for doc {document_id}: {e}")
 

        return Response({
            'success': True,
            'document_id': int(document_id),
            'audit_id': audit_id,
            'policy_id': policy_id,
            'status': status_label,
            'confidence': round(confidence, 2),
            'analyses': analyses,
            'signals': signals
        })
    except Exception as e:
        logger.error(f"❌ Error checking document compliance: {e}")
        return Response({'success': False, 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    finally:
        # Clean up temporary file if it was created from S3
        if 'temp_file_created' in locals() and temp_file_created and 'full_path' in locals():
            try:
                if os.path.exists(full_path):
                    os.remove(full_path)
                    logger.info(f"🗑️ Cleaned up temporary file: {full_path}")
            except Exception as e:
                logger.warning(f"⚠️ Failed to clean up temporary file {full_path}: {e}")
                
        # Additional cleanup for any other temporary files that might have been created
        cleaned_temp_files = cleanup_ai_audit_temp_files(document_id)
        if cleaned_temp_files:
            logger.info(f"🗑️ Additional cleanup - removed {len(cleaned_temp_files)} temporary files")


def _check_compliance_with_combined_evidence_internal(audit_id, compliance_ids, user_id=None, primary_document_id=None, audit_context=None):
    """
    Internal function to check compliance with combined evidence (documents + database).
    This can be called from both API endpoints and background processing.
    
    Args:
        audit_id: The audit ID
        compliance_ids: List of compliance IDs to check
        user_id: Optional user ID
        primary_document_id: Optional document_id that triggered this check (used to update status)
        audit_context: Optional audit context dict (if not provided, will be fetched from database)
    
    Returns: dict with 'success', 'results', etc.
    """
    # Helper function to fix timestamp corruption in s3_key (defined at function scope for reuse)
    def fix_timestamp_corruption(key):
        """Fix corrupted timestamps (extra digits) in s3_key
        
        Handles cases where timestamps have extra digits inserted (e.g., 17665772641288 instead of 1766572641288).
        Timestamps in s3_keys should be exactly 13 digits (milliseconds since epoch).
        """
        if not key:
            return key
        import re
        # Find all 13+ digit sequences (timestamps) - these should be exactly 13 digits
        # Pattern matches sequences of 13 or more consecutive digits
        timestamps = re.findall(r'\d{13,}', key)
        fixed_key = key
        for ts in timestamps:
            if len(ts) > 13:
                # Timestamp should be exactly 13 digits (milliseconds since epoch)
                # Corruption pattern: extra digits get inserted (e.g., 17665772641288 -> should be 1766572641288)
                # Fix by keeping only first 13 digits
                original_ts = ts
                fixed_ts = ts[:13]
                # Use replace with count=1 to replace only the first occurrence of this specific timestamp
                fixed_key = fixed_key.replace(original_ts, fixed_ts, 1)
                logger.info(f"🔧 [COMBINED] Fixed corrupted timestamp: '{original_ts}' (length: {len(original_ts)}) -> '{fixed_ts}' (length: 13)")
        return fixed_key
    try:
        logger.info(f"🔍 Combined evidence check for audit {audit_id}, compliances: {compliance_ids}")
        
        # Get audit framework and metadata
        audit_title = None
        audit_objective = None
        audit_scope = None
        audit_business_unit = None
        audit_type = None
        audit_due_date = None
        
        with connection.cursor() as cursor:
            cursor.execute("SELECT FrameworkId, PolicyId, SubPolicyId, Title, Objective, Scope, BusinessUnit, AuditType, DueDate, TenantId FROM audit WHERE AuditId = %s", [int(audit_id)])
            audit_row = cursor.fetchone()
            if not audit_row:
                return {'success': False, 'error': 'Audit not found'}
            framework_id, policy_id, subpolicy_id = audit_row[0], audit_row[1], audit_row[2]
            tenant_id = audit_row[9] if len(audit_row) > 9 else None
            if len(audit_row) > 3:
                audit_title = audit_row[3] or None
                audit_objective = audit_row[4] or None
                audit_scope = audit_row[5] or None
                audit_business_unit = audit_row[6] or None
                audit_type = audit_row[7] or None
                audit_due_date = audit_row[8] or None
        
        # Fetch policy and subpolicy names if not in audit_context
        policy_name = None
        subpolicy_name = None
        if policy_id:
            with connection.cursor() as cursor:
                try:
                    cursor.execute("SELECT PolicyName FROM policies WHERE PolicyId = %s", [int(policy_id)])
                    policy_row = cursor.fetchone()
                    if policy_row:
                        policy_name = policy_row[0]
                except Exception as e:
                    logger.warning(f"⚠️ Could not fetch policy name: {e}")
        
        if subpolicy_id:
            with connection.cursor() as cursor:
                try:
                    cursor.execute("SELECT SubPolicyName FROM subpolicies WHERE SubPolicyId = %s", [int(subpolicy_id)])
                    subpolicy_row = cursor.fetchone()
                    if subpolicy_row:
                        subpolicy_name = subpolicy_row[0]
                except Exception as e:
                    logger.warning(f"⚠️ Could not fetch subpolicy name: {e}")
        
        # Use provided audit_context or build from fetched data
        if audit_context is None:
            audit_context = {
                'title': audit_title,
                'objective': audit_objective,
                'scope': audit_scope,
                'business_unit': audit_business_unit,
                'audit_type': audit_type,
                'due_date': audit_due_date,
                'policy_id': policy_id,
                'policy_name': policy_name,
                'subpolicy_id': subpolicy_id,
                'subpolicy_name': subpolicy_name
            }
        else:
            # Ensure policy/subpolicy info is in audit_context even if it was provided
            if not audit_context.get('policy_name') and policy_name:
                audit_context['policy_name'] = policy_name
            if not audit_context.get('subpolicy_name') and subpolicy_name:
                audit_context['subpolicy_name'] = subpolicy_name
            if not audit_context.get('policy_id') and policy_id:
                audit_context['policy_id'] = policy_id
            if not audit_context.get('subpolicy_id') and subpolicy_id:
                audit_context['subpolicy_id'] = subpolicy_id
            # Also ensure audit metadata is present
            if not audit_context.get('title') and audit_title:
                audit_context['title'] = audit_title
            if not audit_context.get('objective') and audit_objective:
                audit_context['objective'] = audit_objective
            if not audit_context.get('scope') and audit_scope:
                audit_context['scope'] = audit_scope
        
        # Load compliance requirements
        with connection.cursor() as cursor:
            placeholders = ",".join(["%s"] * len(compliance_ids))
            cursor.execute(
                f"""
                SELECT c.ComplianceId, c.ComplianceTitle, c.ComplianceItemDescription,
                       c.ComplianceType, c.Criticality, c.MandatoryOptional,
                       sp.SubPolicyId, sp.SubPolicyName, p.PolicyId, p.PolicyName
                FROM compliance c
                JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                JOIN policies p ON sp.PolicyId = p.PolicyId
                WHERE c.ComplianceId IN ({placeholders})
                """,
                compliance_ids
            )
            compliance_rows = cursor.fetchall()
        
        if not compliance_rows:
            return {'success': False, 'error': 'No compliance requirements found'}
        
        # Build requirements list
        requirements = []
        from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value
        for r in compliance_rows:
            # Decrypt encrypted fields from database
            compliance_title_raw = r[1] or f"Compliance {r[0]}"
            compliance_title = decrypt_any_encrypted_value(compliance_title_raw)
            description_raw = r[2] or ''
            description = decrypt_any_encrypted_value(description_raw)
            subpolicy_name = decrypt_any_encrypted_value(r[7]) if r[7] else None
            policy_name = decrypt_any_encrypted_value(r[9]) if r[9] else None
            
            title = description if description and len(description) > 10 else compliance_title
            requirements.append({
                'compliance_id': r[0],
                'title': title,
                'compliance_title': compliance_title,  # Add short compliance title
                'description': description,
                'type': r[3] or 'General',
                'risk': r[4] or 'Medium',
                'mandatory': (r[5] or '').lower() == 'mandatory',
                'subpolicy_id': r[6],
                'subpolicy_name': subpolicy_name,
                'policy_id': r[8],
                'policy_name': policy_name
            })
        
        # For each compliance requirement, gather ALL evidence (documents + database)
        all_results = []
        
        # Cache for document text to avoid re-downloading the same document multiple times
        document_text_cache = {}  # key: (external_id or doc_id), value: document_text
        
        # Reuse S3 client across all document downloads to avoid creating multiple MySQL connections
        s3_client_reused = None
        try:
            from ..Global.s3_fucntions import create_direct_mysql_client
            s3_client_reused = create_direct_mysql_client()
            logger.info("✅ [COMBINED] Created reusable S3 client for all document downloads")
        except Exception as s3_err:
            logger.warning(f"⚠️ [COMBINED] Could not create reusable S3 client: {s3_err}")
            s3_client_reused = None
        
        for req in requirements:
            compliance_id = req['compliance_id']
            logger.info(f"🔍 Processing compliance {compliance_id}: {req['title']}")
            
            # Gather document evidence
            document_evidence_texts = []
            document_evidence_details = []
            
            with connection.cursor() as cursor:
                # Get all documents mapped to this compliance (via policy/subpolicy or direct mapping)
                cursor.execute("""
                    SELECT d.document_id, d.document_name, d.document_path, d.external_source, d.external_id,
                           d.mime_type, d.document_type, d.file_size, d.compliance_analyses
                    FROM ai_audit_data d
                    WHERE d.audit_id = %s
                      AND d.external_source != 'database_record'
                      AND d.document_type != 'db_record'
                      AND (
                          (d.policy_id = %s AND d.subpolicy_id = %s)
                          OR EXISTS (
                              SELECT 1 FROM compliance c
                              WHERE c.ComplianceId = %s
                                AND c.SubPolicyId = d.subpolicy_id
                          )
                      )
                    """, [int(audit_id), req['policy_id'], req['subpolicy_id'], compliance_id])
                
                doc_rows = cursor.fetchall()
                
                for doc_row in doc_rows:
                    doc_id, doc_name, doc_path, ext_source, ext_id, mime_type, doc_type, file_size, analyses = doc_row
                    
                    # Check cache first to avoid re-downloading the same document
                    cache_key = str(ext_id) if ext_id else f"doc_{doc_id}"
                    if cache_key in document_text_cache:
                        logger.info(f"📋 [COMBINED] Using cached text for document {doc_id} (cache_key={cache_key[:50]}...)")
                        doc_text = document_text_cache[cache_key]
                    else:
                        # Try to extract text from document
                        doc_text = None
                        # Initialize variables before conditional to avoid UnboundLocalError
                        s3_key = None
                        actual_file_name = None
                        s3_url_for_fallback = None
                        
                        if ext_source in ['s3', 'evidence_attachment'] and ext_id:
                            # For S3 files, try to look up from file_operations
                            # external_id can be: operation_id (numeric), s3_key, stored_name, or file_name
                            
                            try:
                                with connection.cursor() as lookup_cursor:
                                    file_op = None
                                    
                                    # First, try by operation_id if ext_id is numeric
                                    if ext_id and str(ext_id).isdigit():
                                        lookup_cursor.execute("""
                                            SELECT s3_key, file_name, original_name, s3_url, stored_name
                                            FROM file_operations 
                                            WHERE id = %s AND status = 'completed'
                                            LIMIT 1
                                        """, [int(ext_id)])
                                        file_op = lookup_cursor.fetchone()
                                    
                                    # If not found and ext_id is not numeric, try by s3_key, stored_name, or file_name
                                    if not file_op and ext_id:
                                        ext_id_str = str(ext_id).strip()
                                        lookup_cursor.execute("""
                                            SELECT s3_key, file_name, original_name, s3_url, stored_name
                                            FROM file_operations 
                                            WHERE status = 'completed'
                                              AND operation_type = 'upload'
                                              AND (
                                                  s3_key = %s 
                                                  OR stored_name = %s 
                                                  OR file_name = %s
                                                  OR original_name = %s
                                              )
                                            ORDER BY created_at DESC
                                            LIMIT 1
                                        """, [ext_id_str, ext_id_str, ext_id_str, ext_id_str])
                                        file_op = lookup_cursor.fetchone()
                                    
                                    if file_op:
                                        s3_key_from_db = file_op[0]
                                        file_name_from_db = file_op[1]
                                        original_name_from_db = file_op[2] if len(file_op) > 2 else None
                                        s3_url_from_db = file_op[3] if len(file_op) > 3 else None
                                        stored_name_from_db = file_op[4] if len(file_op) > 4 else None
                                        
                                        logger.info(f"🔍 [COMBINED] file_operations data for ext_id {ext_id}: s3_key={s3_key_from_db}, file_name={file_name_from_db}, stored_name={stored_name_from_db}, s3_url={s3_url_from_db[:100] if s3_url_from_db else None}")
                                        
                                        # Helper function to detect corrupted s3_key (e.g., extra digits in timestamps)
                                        # Note: fix_timestamp_corruption is defined at function scope above
                                        def validate_s3_key(key, reference_name=None):
                                            """Check if s3_key looks corrupted by comparing with reference file_name"""
                                            if not key or not reference_name:
                                                return True  # Can't validate, assume OK
                                            # Extract timestamps from both and compare
                                            key_timestamps = re.findall(r'\d{13,}', key)  # Find 13+ digit sequences (timestamps)
                                            ref_timestamps = re.findall(r'\d{13,}', reference_name)
                                            # If we have matching timestamps and they differ, it might be corrupted
                                            if key_timestamps and ref_timestamps:
                                                # Check if any timestamp in key is longer than in reference (likely corruption)
                                                if len(key_timestamps[0]) > len(ref_timestamps[0]):
                                                    return False
                                            return True
                                        
                                        # Try to extract s3_key from s3_url first (most reliable source)
                                        s3_key_from_url = None
                                        if s3_url_from_db and 'amazonaws.com/' in s3_url_from_db:
                                            try:
                                                s3_key_from_url = s3_url_from_db.split('amazonaws.com/')[-1].split('?')[0]
                                                # Fix timestamp corruption immediately after extraction
                                                s3_key_from_url = fix_timestamp_corruption(s3_key_from_url)
                                                logger.info(f"🔍 [COMBINED] Extracted s3_key from s3_url: {s3_key_from_url[:80]}...")
                                            except:
                                                pass
                                        
                                        # Priority: s3_key_from_url (most reliable) > validated stored_name > s3_key_from_db > fallback
                                        s3_key = None
                                        if s3_key_from_url:
                                            s3_key = s3_key_from_url
                                            logger.info(f"🔍 [COMBINED] Using s3_key extracted from s3_url (most reliable): {s3_key[:80]}...")
                                        elif stored_name_from_db:
                                            # Fix timestamp corruption in stored_name first
                                            stored_name_fixed = fix_timestamp_corruption(stored_name_from_db)
                                            # Validate stored_name doesn't look corrupted
                                            ref_name = file_name_from_db or original_name_from_db
                                            if validate_s3_key(stored_name_fixed, ref_name):
                                                s3_key = stored_name_fixed
                                                logger.info(f"🔍 [COMBINED] Using validated stored_name as s3_key: {s3_key[:80]}...")
                                            else:
                                                logger.warning(f"⚠️ [COMBINED] stored_name appears corrupted, trying alternatives: {stored_name_from_db[:80]}...")
                                                # Try s3_key_from_db instead
                                                if s3_key_from_db:
                                                    s3_key = fix_timestamp_corruption(s3_key_from_db)
                                                    logger.info(f"🔍 [COMBINED] Using s3_key_from_db as fallback: {s3_key[:80]}...")
                                                elif s3_url_from_db and 'amazonaws.com/' in s3_url_from_db:
                                                    s3_key = s3_url_from_db.split('amazonaws.com/')[-1].split('?')[0]
                                                    s3_key = fix_timestamp_corruption(s3_key)
                                                    logger.info(f"🔍 [COMBINED] Using s3_key extracted from s3_url as fallback: {s3_key[:80]}...")
                                        elif s3_key_from_db:
                                            s3_key = fix_timestamp_corruption(s3_key_from_db)
                                            logger.info(f"🔍 [COMBINED] Using s3_key_from_db: {s3_key[:80]}...")
                                        
                                        # Store s3_url for potential direct download fallback
                                        s3_url_for_fallback = s3_url_from_db
                                        
                                        # Determine file_name - should match s3_key for download API
                                        if s3_key:
                                            if '/' in s3_key:
                                                # Extract filename from s3_key if it contains path
                                                actual_file_name = s3_key.split('/')[-1]
                                            else:
                                                # If s3_key doesn't have path, use it directly as file_name
                                                actual_file_name = s3_key
                                        else:
                                            actual_file_name = file_name_from_db or original_name_from_db or doc_name
                                        
                                        logger.info(f"🔍 [COMBINED] Final s3_key: {s3_key[:80] if s3_key else 'None'}..., file_name: {actual_file_name[:80] if actual_file_name else 'None'}... for document {doc_id}")
                            except Exception as lookup_err:
                                logger.warning(f"⚠️ Error looking up file_operations for ext_id {ext_id}: {lookup_err}")
                        
                        # If we have s3_key, try to download
                        if s3_key and s3_key.strip():
                            try:
                                import tempfile
                                
                                # Make a clean copy of s3_key to prevent any modification
                                # Store original for validation
                                original_s3_key = s3_key
                                download_s3_key = str(s3_key).strip() if s3_key else None
                                
                                # Validate s3_key hasn't been corrupted (check for common corruption patterns)
                                # Check if timestamp corruption occurred (extra digits in 13+ digit sequences)
                                if download_s3_key and original_s3_key:
                                    import re
                                    orig_timestamps = re.findall(r'\d{13,}', original_s3_key)
                                    download_timestamps = re.findall(r'\d{13,}', download_s3_key)
                                    if orig_timestamps and download_timestamps and len(orig_timestamps) == len(download_timestamps):
                                        # Compare timestamps - if download version is longer, it's corrupted
                                        for i, (orig_ts, dl_ts) in enumerate(zip(orig_timestamps, download_timestamps)):
                                            if len(dl_ts) > len(orig_ts):
                                                logger.warning(f"⚠️ [COMBINED] Detected timestamp corruption in s3_key! Original timestamp: {orig_ts}, Corrupted: {dl_ts}")
                                                logger.warning(f"⚠️ [COMBINED] Using original s3_key to avoid corruption: {original_s3_key[:80]}...")
                                                download_s3_key = original_s3_key
                                                break
                                
                                if download_s3_key != original_s3_key and download_s3_key != original_s3_key.strip():
                                    logger.warning(f"⚠️ [COMBINED] s3_key was modified during copy! Original: {original_s3_key[:80]}, Modified: {download_s3_key[:80] if download_s3_key else None}")
                                    # Use original if validation fails
                                    download_s3_key = original_s3_key
                                
                                # REUSE the S3 client created earlier to avoid creating multiple MySQL connections
                                if not s3_client_reused:
                                    from ..Global.s3_fucntions import create_direct_mysql_client
                                    s3_client_reused = create_direct_mysql_client()
                                    logger.info("✅ [COMBINED] Created S3 client (reused for all downloads)")
                                
                                s3_client = s3_client_reused
                                file_name = actual_file_name or doc_name or f"document_{doc_id}.pdf"
                                # Ensure file_name is a clean string
                                file_name = str(file_name).strip() if file_name else f"document_{doc_id}.pdf"
                                
                                temp_dir = tempfile.gettempdir()
                                temp_file_path = os.path.join(temp_dir, f"ai_audit_combined_{doc_id}_{file_name}")
                                
                                # FINAL VALIDATION: ALWAYS extract s3_key from s3_url if available (most reliable source)
                                # The s3_key stored in database may be corrupted, but s3_url should be correct
                                # This prevents any corruption that might have occurred in the s3_key variable
                                if s3_url_for_fallback and 'amazonaws.com/' in s3_url_for_fallback:
                                    try:
                                        # Extract s3_key from URL (most reliable source - bypasses any database corruption)
                                        final_s3_key_from_url = s3_url_for_fallback.split('amazonaws.com/')[-1].split('?')[0]
                                        if final_s3_key_from_url:
                                            if final_s3_key_from_url != download_s3_key:
                                                logger.info(f"🔄 [COMBINED] Using s3_key from s3_url (bypassing potential DB corruption). Old: {download_s3_key[:60] if download_s3_key else None}..., New: {final_s3_key_from_url[:60]}...")
                                            download_s3_key = final_s3_key_from_url
                                            # CRITICAL: Also update file_name to match the corrected s3_key
                                            # This ensures file_name and s3_key are in sync (both from the reliable s3_url source)
                                            if '/' in download_s3_key:
                                                file_name = download_s3_key.split('/')[-1]
                                            else:
                                                file_name = download_s3_key
                                            logger.info(f"🔄 [COMBINED] Updated file_name to match s3_key from s3_url: {file_name[:60]}...")
                                    except Exception as verify_err:
                                        logger.warning(f"⚠️ [COMBINED] Could not extract s3_key from s3_url: {verify_err}")
                                        # Fallback: try to fix known corruption patterns if s3_url extraction fails
                                        if download_s3_key:
                                            download_s3_key = fix_timestamp_corruption(download_s3_key)
                                elif download_s3_key:
                                    # If no s3_url available, try to fix known corruption patterns
                                    download_s3_key = fix_timestamp_corruption(download_s3_key)
                                    # Also ensure file_name matches the fixed s3_key
                                    if download_s3_key and (not file_name or file_name != download_s3_key):
                                        if '/' in download_s3_key:
                                            file_name = download_s3_key.split('/')[-1]
                                        else:
                                            file_name = download_s3_key
                                
                                logger.info(f"⬇️ [COMBINED] Downloading document {doc_id} from S3: s3_key={download_s3_key[:80] if download_s3_key else None}..., file_name={file_name[:60] if file_name else None}...")
                                logger.info(f"🔍 [COMBINED] s3_key length: {len(download_s3_key) if download_s3_key else 0}, file_name length: {len(file_name)}")
                                download_result = s3_client.download(download_s3_key, file_name, temp_dir, str(user_id) if user_id else 'system')
                                
                                # Get the actual downloaded file path
                                if download_result.get('success'):
                                    downloaded_path = download_result.get('file_path') or temp_file_path
                                    if os.path.exists(downloaded_path):
                                        doc_text = extract_text_from_document(downloaded_path, mime_type or doc_type or 'text/plain')
                                        logger.info(f"✅ [COMBINED] Successfully extracted text from document {doc_id} ({len(doc_text)} chars)")
                                        # Clean up temp file
                                        try:
                                            os.remove(downloaded_path)
                                        except:
                                            pass
                                    else:
                                        logger.warning(f"⚠️ [COMBINED] Downloaded file not found at {downloaded_path}")
                                else:
                                    error_msg = download_result.get('error', 'Unknown error')
                                    logger.warning(f"⚠️ [COMBINED] Download failed for document {doc_id}: {error_msg}")
                                    
                                    # Fallback: Try downloading directly from s3_url if available (for 404 errors)
                                    if ('404' in str(error_msg) or 'Not Found' in str(error_msg)) and s3_url_for_fallback:
                                        try:
                                            logger.info(f"🔄 [COMBINED] Attempting direct download from s3_url for document {doc_id}")
                                            import requests
                                            response = requests.get(s3_url_for_fallback, timeout=60, stream=True)
                                            if response.status_code == 200:
                                                # Save to temp file
                                                temp_file_path_direct = os.path.join(temp_dir, f"ai_audit_direct_{doc_id}_{file_name}")
                                                with open(temp_file_path_direct, 'wb') as f:
                                                    for chunk in response.iter_content(chunk_size=8192):
                                                        f.write(chunk)
                                                
                                                if os.path.exists(temp_file_path_direct):
                                                    doc_text = extract_text_from_document(temp_file_path_direct, mime_type or doc_type or 'text/plain')
                                                    logger.info(f"✅ [COMBINED] Successfully extracted text via direct s3_url download for document {doc_id} ({len(doc_text)} chars)")
                                                    # Clean up temp file
                                                    try:
                                                        os.remove(temp_file_path_direct)
                                                    except:
                                                        pass
                                            else:
                                                logger.warning(f"⚠️ [COMBINED] Direct s3_url download also failed: HTTP {response.status_code}")
                                        except Exception as direct_download_err:
                                            logger.warning(f"⚠️ [COMBINED] Direct s3_url download failed: {direct_download_err}")
                            except Exception as download_err:
                                logger.error(f"❌ [COMBINED] Could not download/extract document {doc_id}: {download_err}", exc_info=True)
                    
                    # Try local file path if S3 download failed or not applicable
                    if not doc_text and doc_path and os.path.exists(doc_path):
                        try:
                            doc_text = extract_text_from_document(doc_path, mime_type or doc_type or 'text/plain')
                        except:
                            pass
                    
                    if doc_text:
                        # Cache the document text for reuse in this compliance check run
                        if cache_key not in document_text_cache:
                            document_text_cache[cache_key] = doc_text
                            logger.info(f"💾 [COMBINED] Cached text for document {doc_id} (cache_key={cache_key[:50]}...), length: {len(doc_text)} chars")
                        
                        document_evidence_texts.append(f"[DOCUMENT: {doc_name}]\n{doc_text}")
                        document_evidence_details.append({
                            'type': 'document',
                            'id': doc_id,
                            'name': doc_name,
                            'source': 'document_evidence'
                        })
            
            # Gather database evidence
            database_evidence_texts = []
            database_evidence_details = []
            
            with connection.cursor() as cursor:
                # Get all database records mapped to this compliance
                cursor.execute("""
                    SELECT d.document_id, d.document_name, d.compliance_analyses, d.external_id
                    FROM ai_audit_data d
                    WHERE d.audit_id = %s
                      AND (d.external_source = 'database_record' OR d.document_type = 'db_record')
                      AND (
                          (d.policy_id = %s AND d.subpolicy_id = %s)
                          OR EXISTS (
                              SELECT 1 FROM compliance c
                              WHERE c.ComplianceId = %s
                                AND c.SubPolicyId = d.subpolicy_id
                          )
                      )
                """, [int(audit_id), req['policy_id'], req['subpolicy_id'], compliance_id])
                
                db_rows = cursor.fetchall()
                
                for db_row in db_rows:
                    db_id, db_name, analyses, external_id = db_row
                    
                    # Extract evidence from compliance_analyses
                    if analyses:
                        try:
                            # Parse JSON if it's a string
                            if isinstance(analyses, str):
                                analyses_json = json.loads(analyses)
                            else:
                                analyses_json = analyses
                            
                            # Try to extract evidence from the structure
                            db_text = None
                            
                            if isinstance(analyses_json, dict):
                                # Try compliance_analyses array first
                                analyses_list = analyses_json.get('compliance_analyses', [])
                                if isinstance(analyses_list, list):
                                    for a in analyses_list:
                                        if isinstance(a, dict) and a.get('compliance_id') == compliance_id:
                                            reason = a.get('relevance_reason', '') or a.get('reason', '')
                                            score = a.get('relevance_score', 0) or a.get('score', 0)
                                            db_text = f"Relevance: {score:.2f}. {reason}" if reason else f"Relevance score: {score:.2f}"
                                            break
                                
                                # Fallback: try direct fields
                                if not db_text:
                                    reason = analyses_json.get('relevance_reason', '') or analyses_json.get('reason', '')
                                    score = analyses_json.get('relevance_score', 0) or analyses_json.get('score', 0)
                                    if reason or score:
                                        db_text = f"Relevance: {score:.2f}. {reason}" if reason else f"Relevance score: {score:.2f}"
                            
                            # If we found evidence text, add it
                            if db_text:
                                database_evidence_texts.append(f"[DATABASE EVIDENCE: {db_name}]\n{db_text}")
                                database_evidence_details.append({
                                    'type': 'database',
                                    'id': db_id,
                                    'name': db_name,
                                    'source': 'database_evidence',
                                    'external_id': external_id
                                })
                            else:
                                # Even if parsing fails, include the record name as evidence
                                database_evidence_texts.append(f"[DATABASE EVIDENCE: {db_name}]\nDatabase record available for this compliance requirement.")
                                database_evidence_details.append({
                                    'type': 'database',
                                    'id': db_id,
                                    'name': db_name,
                                    'source': 'database_evidence',
                                    'external_id': external_id
                                })
                        except Exception as e:
                            logger.warning(f"⚠️ Error parsing database evidence analyses for {db_name}: {e}")
                            # Fallback: include the record anyway
                            database_evidence_texts.append(f"[DATABASE EVIDENCE: {db_name}]\nDatabase record available for this compliance requirement.")
                            database_evidence_details.append({
                                'type': 'database',
                                'id': db_id,
                                'name': db_name,
                                'source': 'database_evidence',
                                'external_id': external_id
                            })
            
            # Combine all evidence
            combined_evidence = []
            if document_evidence_texts:
                combined_evidence.append("=== DOCUMENT EVIDENCE (Intent/Design) ===\n")
                combined_evidence.extend(document_evidence_texts)
                combined_evidence.append("\n")
                logger.info(f"📄 [COMBINED] Found {len(document_evidence_texts)} document(s) for compliance {compliance_id}")
            
            if database_evidence_texts:
                combined_evidence.append("=== DATABASE EVIDENCE (Operational/Factual) ===\n")
                combined_evidence.extend(database_evidence_texts)
                combined_evidence.append("\n")
                logger.info(f"💾 [COMBINED] Found {len(database_evidence_texts)} database record(s) for compliance {compliance_id}")
            
            if not combined_evidence:
                logger.warning(f"⚠️ No evidence found for compliance {compliance_id}")
                no_evidence_result = {
                    'compliance_id': compliance_id,
                    'status': 'NO_EVIDENCE',
                    'error': 'No evidence available'
                }
                all_results.append(no_evidence_result)
                # Save NO_EVIDENCE to lastchecklistitemverified so results appear in UI
                try:
                    no_evidence_analysis = [{
                        'compliance_id': compliance_id,
                        'compliance_status': 'NON_COMPLIANT',
                        'status': 'NO_EVIDENCE',
                        'reason': 'No evidence available',
                        'compliance_score': 0.0
                    }]
                    save_ai_compliance_to_checklist(
                        audit_id=audit_id,
                        document_id=None,
                        analyses=no_evidence_analysis,
                        user_id=user_id,
                        framework_id=framework_id,
                        policy_id=req.get('policy_id'),
                        subpolicy_id=req.get('subpolicy_id')
                    )
                except Exception as save_err:
                    logger.warning(f"⚠️ Could not save NO_EVIDENCE to checklist: {save_err}")
                continue
            
            combined_text = "\n".join(combined_evidence)
            
            # Run AI analysis on combined evidence
            logger.info(f"🤖 [COMBINED] Analyzing compliance {compliance_id} with {len(document_evidence_texts)} document(s) + {len(database_evidence_texts)} database record(s) = ONE combined check")
            
            # Create a single-requirement batch for analysis
            single_req_batch = [req]
            analyses = _ai_score_requirements_with_openai(
                combined_text,
                single_req_batch,
                schema=None,
                audit_id=audit_id,
                document_id=None,  # No single document_id since we're combining evidence
                audit_context=audit_context
            )
            
            if not analyses or len(analyses) == 0:
                logger.warning(f"⚠️ No analysis returned for compliance {compliance_id}")
                analysis_failed_result = {
                    'compliance_id': compliance_id,
                    'status': 'ANALYSIS_FAILED',
                    'error': 'AI analysis failed'
                }
                all_results.append(analysis_failed_result)
                # Save ANALYSIS_FAILED to lastchecklistitemverified so results appear in UI
                try:
                    failed_analysis = [{
                        'compliance_id': compliance_id,
                        'compliance_status': 'NON_COMPLIANT',
                        'status': 'ANALYSIS_FAILED',
                        'reason': 'AI analysis failed',
                        'compliance_score': 0.0
                    }]
                    save_ai_compliance_to_checklist(
                        audit_id=audit_id,
                        document_id=None,
                        analyses=failed_analysis,
                        user_id=user_id,
                        framework_id=framework_id,
                        policy_id=req.get('policy_id'),
                        subpolicy_id=req.get('subpolicy_id')
                    )
                except Exception as save_err:
                    logger.warning(f"⚠️ Could not save ANALYSIS_FAILED to checklist: {save_err}")
                continue
            
            analysis = analyses[0] if analyses else {}
            status_label, confidence = _determine_status([req], analyses)
            
            # Save to checklist (once per compliance, combining all evidence)
            try:
                # Get the primary policy/subpolicy for this compliance
                primary_policy_id = req.get('policy_id')
                primary_subpolicy_id = req.get('subpolicy_id')
                
                # Create a combined analysis result
                combined_analysis = {
                    **analysis,
                    'evidence_sources': {
                        'documents': document_evidence_details,
                        'database': database_evidence_details
                    },
                    'combined_evidence_count': len(document_evidence_texts) + len(database_evidence_texts)
                }
                
                # Save to lastchecklistitemverified
                save_ai_compliance_to_checklist(
                    audit_id=audit_id,
                    document_id=None,  # No single document since we combined evidence
                    analyses=[combined_analysis],
                    user_id=user_id,
                    framework_id=framework_id,
                    policy_id=primary_policy_id,
                    subpolicy_id=primary_subpolicy_id
                )
                
                # Update ALL documents that were used in the combined evidence check
                # This ensures all documents show the same combined analysis result
                # Each document will have the Details button showing the combined result
                documents_to_update = []
                if document_evidence_details and len(document_evidence_details) > 0:
                    # Get all document IDs that were part of the combined check
                    documents_to_update = [doc_detail['id'] for doc_detail in document_evidence_details]
                elif primary_document_id:
                    # If document download failed but we have the primary_document_id, use it
                    documents_to_update = [primary_document_id]
                    logger.info(f"📝 Using primary_document_id {primary_document_id} to update status (document download failed)")
                
                if documents_to_update:
                    try:
                        from datetime import datetime
                        # Create a unique group ID for this combined check (using first document ID + compliance ID + timestamp)
                        combined_check_group_id = f"combined_{compliance_id}_{documents_to_update[0]}_{int(datetime.now().timestamp())}"
                        
                        combined_analysis_json = json.dumps({
                            "compliance_status": status_label,
                            "confidence_score": float(confidence),
                            "compliance_analyses": [combined_analysis],
                            "processed_at": datetime.now().isoformat(),
                            "combined_evidence": True,
                            "combined_check_group_id": combined_check_group_id,  # Unique ID for this combined check group
                            "combined_with_document_ids": documents_to_update,  # List of all document IDs in this combined check
                            "evidence_sources": {
                                "documents": document_evidence_details,
                                "database": database_evidence_details
                            }
                        })
                        
                        with connection.cursor() as cursor:
                            # Update ALL documents that were part of the combined check
                            # This ensures all documents show the same combined analysis result
                            # AND they all have the same combined_check_group_id so frontend can group them
                            # For manually uploaded documents, also update all mappings that share the same document_path or document_name+file_size
                            
                            # First, get document_path and document_name+file_size for the documents being updated
                            # so we can also update other mappings for the same physical file
                            placeholders = ','.join(['%s'] * len(documents_to_update))
                            cursor.execute(f"""
                                SELECT DISTINCT document_path, document_name, file_size
                                FROM ai_audit_data
                                WHERE document_id IN ({placeholders})
                                  AND audit_id = %s
                            """, [*documents_to_update, int(audit_id)])
                            
                            file_identifiers = cursor.fetchall()
                            
                            # Build WHERE clause to match:
                            # 1. The specific document IDs
                            # 2. All records with the same document_path (for manually uploaded)
                            # 3. All records with the same document_name + file_size (fallback for manually uploaded)
                            where_conditions = [f"document_id IN ({placeholders})"]
                            where_params = list(documents_to_update)
                            
                            # Collect unique document_paths and document_name+file_size pairs
                            doc_paths_to_match = set()
                            doc_name_size_pairs = set()
                            
                            for doc_path, doc_name, doc_size in file_identifiers:
                                if doc_path:
                                    doc_paths_to_match.add(doc_path)
                                elif doc_name and doc_size:
                                    doc_name_size_pairs.add((doc_name, doc_size))
                            
                            # Add conditions for document_path matches
                            if doc_paths_to_match:
                                path_placeholders = ','.join(['%s'] * len(doc_paths_to_match))
                                where_conditions.append(f"document_path IN ({path_placeholders})")
                                where_params.extend(list(doc_paths_to_match))
                            
                            # Add conditions for document_name + file_size matches (for manually uploaded without document_path)
                            if doc_name_size_pairs:
                                name_size_conditions = []
                                for doc_name, doc_size in doc_name_size_pairs:
                                    name_size_conditions.append("(document_name = %s AND file_size = %s)")
                                    where_params.extend([doc_name, doc_size])
                                if name_size_conditions:
                                    where_conditions.append("(" + " OR ".join(name_size_conditions) + " AND (external_source IS NULL OR external_source NOT IN ('evidence_attachment', 's3', 'database_record')))")
                            
                            # Combine all conditions with OR
                            where_clause = "(" + " OR ".join(where_conditions) + ")"
                            
                            cursor.execute(f"""
                                UPDATE ai_audit_data 
                                SET ai_processing_status = 'completed',
                                    compliance_status = %s,
                                    confidence_score = %s,
                                    compliance_analyses = %s,
                                    processing_completed_at = NOW()
                                WHERE {where_clause}
                                  AND audit_id = %s
                                  AND (external_source != 'database_record' AND document_type != 'db_record')
                            """, [
                                status_label,
                                float(confidence),
                                combined_analysis_json,
                                *where_params,
                                int(audit_id)
                            ])
                            updated_count = cursor.rowcount
                            logger.info(f"✅ Updated {updated_count} document(s) with combined evidence result (group_id={combined_check_group_id}): {documents_to_update}")
                    except Exception as update_err:
                        logger.warning(f"⚠️ Could not update documents ai_audit_data: {update_err}")
                
                # Mark database evidence items as part of combined check (so frontend hides their Details button)
                # Documents will show Details button with combined result, but database records won't
                if database_evidence_details and len(database_evidence_details) > 0:
                    try:
                        with connection.cursor() as cursor:
                            # Use first document ID as reference (all documents now have the same combined result)
                            primary_doc_id_for_ref = documents_to_update[0] if documents_to_update else None
                            
                            for db_evidence in database_evidence_details:
                                db_evidence_id = db_evidence.get('id')
                                if db_evidence_id:
                                    # Update the database evidence item to mark it as part of combined check
                                    cursor.execute("""
                                        UPDATE ai_audit_data 
                                        SET compliance_analyses = JSON_SET(
                                            COALESCE(compliance_analyses, '{}'),
                                            '$.part_of_combined_check', true,
                                            '$.combined_with_document_id', %s,
                                            '$.combined_compliance_id', %s
                                        )
                                        WHERE document_id = %s 
                                          AND audit_id = %s
                                          AND (external_source = 'database_record' OR document_type = 'db_record')
                                    """, [
                                        primary_doc_id_for_ref,
                                        compliance_id,
                                        db_evidence_id,
                                        int(audit_id)
                                    ])
                            logger.info(f"✅ Marked {len(database_evidence_details)} database evidence items as part of combined check")
                    except Exception as db_mark_err:
                        logger.warning(f"⚠️ Could not mark database evidence items: {db_mark_err}")
                
                logger.info(f"✅ Saved combined evidence result for compliance {compliance_id}")
                
                all_results.append({
                    'compliance_id': compliance_id,
                    'status': status_label,
                    'confidence': round(confidence, 2),
                    'analysis': combined_analysis,
                    'evidence_count': {
                        'documents': len(document_evidence_texts),
                        'database': len(database_evidence_texts),
                        'total': len(document_evidence_texts) + len(database_evidence_texts)
                    }
                })
            except Exception as e:
                logger.error(f"❌ Error saving combined evidence result: {e}")
                all_results.append({
                    'compliance_id': compliance_id,
                    'status': 'SAVE_FAILED',
                    'error': str(e)
                })
        
        # Run SEBI AI Auditor checks automatically (if enabled) - non-blocking
        try:
            if framework_id and tenant_id:
                import threading
                # Capture variables for closure
                sebi_framework_id = framework_id
                sebi_tenant_id = tenant_id
                sebi_audit_id = audit_id
                
                def run_sebi_checks():
                    try:
                        from .sebi_ai_auditor import SEBIAIAuditor
                        auditor = SEBIAIAuditor(sebi_framework_id, sebi_tenant_id)
                        if auditor.is_sebi_framework:
                            logger.info(f"🔍 SEBI AI Auditor enabled - running automatic checks for audit {sebi_audit_id} (combined evidence)")
                            # Run SEBI checks (non-blocking, results stored for report)
                            try:
                                auditor.verify_filing_accuracy(sebi_audit_id)
                                auditor.check_timeliness_sla(sebi_audit_id)
                                auditor.calculate_risk_score(sebi_audit_id)
                                auditor.detect_patterns(sebi_audit_id)
                                logger.info(f"✅ SEBI AI Auditor checks completed for audit {sebi_audit_id} (combined evidence)")
                            except Exception as sebi_err:
                                logger.warning(f"⚠️ SEBI checks failed (non-critical): {str(sebi_err)}")
                        else:
                            logger.debug(f"ℹ️ SEBI AI Auditor not enabled for framework {sebi_framework_id}")
                    except Exception as e:
                        logger.warning(f"⚠️ SEBI AI Auditor check failed (non-critical): {str(e)}")
                
                # Run in background thread (non-blocking)
                sebi_thread = threading.Thread(target=run_sebi_checks, daemon=True)
                sebi_thread.start()
                logger.info(f"🚀 Started SEBI AI Auditor checks in background for audit {sebi_audit_id} (combined evidence)")
        except Exception as sebi_init_err:
            logger.warning(f"⚠️ Could not initialize SEBI checks (non-critical): {str(sebi_init_err)}")
        
        return {
            'success': True,
            'audit_id': audit_id,
            'results': all_results,
            'message': f'Processed {len(all_results)} compliance requirements with combined evidence'
        }
        
    except Exception as e:
        logger.error(f"❌ Error in combined evidence check: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def check_compliance_with_combined_evidence(request, audit_id):
    """
    Check compliance by combining ALL evidence (documents + database records) for each requirement.
    This implements the correct audit approach: one judgment per compliance requirement based on all available evidence.
    MULTI-TENANCY: Only checks compliance for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        # Get compliance IDs to check (can be single ID or list)
        compliance_ids = request.data.get('compliance_ids') or request.data.get('compliance_id')
        if isinstance(compliance_ids, (int, str)):
            compliance_ids = [int(compliance_ids)]
        elif isinstance(compliance_ids, str):
            try:
                compliance_ids = json.loads(compliance_ids)
            except:
                compliance_ids = [int(compliance_ids)]
        
        if not compliance_ids:
            return Response({
                'success': False,
                'error': 'compliance_ids required'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Call internal function
        result = _check_compliance_with_combined_evidence_internal(audit_id, compliance_ids, user_id)
        
        # Convert dict result to Response
        if result.get('success'):
            return Response(result, status=status.HTTP_200_OK)
        else:
            return Response(result, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
    except Exception as e:
        logger.error(f"❌ Error in combined evidence check: {e}")
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def check_all_documents_compliance(request, audit_id):
    """
    Run compliance check for all mapped documents in an audit.
    Automatically uses combined evidence approach when both document and database evidence exist.
    MULTI-TENANCY: Only checks compliance for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        # Check authentication using JWT (like other endpoints)
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        logger.info(f"🔍 Bulk compliance check request from user: {user_id}")
        
        # Get all unique compliance IDs from all documents and database evidence
        with connection.cursor() as cursor:
            # Get compliances from documents (simpler approach: get compliances from subpolicies)
            cursor.execute(
                """
                SELECT DISTINCT c.ComplianceId
                FROM ai_audit_data d
                JOIN subpolicies sp ON sp.SubPolicyId = d.subpolicy_id AND sp.PolicyId = d.policy_id
                JOIN compliance c ON c.SubPolicyId = sp.SubPolicyId
                WHERE d.audit_id = %s
                  AND d.external_source != 'database_record'
                  AND d.document_type != 'db_record'
                  AND d.subpolicy_id IS NOT NULL
                  AND d.policy_id IS NOT NULL
                """,
                [int(audit_id) if str(audit_id).isdigit() else audit_id]
            )
            doc_compliance_ids = {r[0] for r in cursor.fetchall()}
            
            # Get compliances from database evidence
            cursor.execute(
                """
                SELECT DISTINCT c.ComplianceId
                FROM ai_audit_data d
                JOIN subpolicies sp ON sp.SubPolicyId = d.subpolicy_id AND sp.PolicyId = d.policy_id
                JOIN compliance c ON c.SubPolicyId = sp.SubPolicyId
                WHERE d.audit_id = %s
                  AND (d.external_source = 'database_record' OR d.document_type = 'db_record')
                  AND d.subpolicy_id IS NOT NULL
                  AND d.policy_id IS NOT NULL
                """,
                [int(audit_id) if str(audit_id).isdigit() else audit_id]
            )
            db_compliance_ids = {r[0] for r in cursor.fetchall()}
            
            # Combine all compliance IDs
            all_compliance_ids = list(doc_compliance_ids | db_compliance_ids)
        
        if not all_compliance_ids:
            return Response({
                'success': False,
                'error': 'No compliance requirements found for this audit'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Check if both document and database evidence exist
        has_doc_evidence = len(doc_compliance_ids) > 0
        has_db_evidence = len(db_compliance_ids) > 0
        
        if has_doc_evidence and has_db_evidence:
            # Use combined evidence approach (one check per compliance, combining all evidence)
            logger.info(f"🔍 Both document and database evidence found. Using combined evidence approach for {len(all_compliance_ids)} compliance requirements.")
            
            # Call internal combined evidence function
            result = _check_compliance_with_combined_evidence_internal(
                audit_id=audit_id,
                compliance_ids=all_compliance_ids,
                user_id=user_id
            )

            # If combined check succeeded, try to move AI audit into "Under review"
            if result.get('success'):
                try:
                    status_update = check_and_update_ai_audit_status(audit_id)
                    logger.info(f"🔁 Bulk check: status update result for audit {audit_id}: {status_update}")
                except Exception as status_err:
                    logger.error(f"⚠️ Bulk check: could not update audit status after combined evidence check: {status_err}")
                return Response(result, status=status.HTTP_200_OK)
            else:
                return Response(result, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        else:
            # Fallback to document-by-document check (original behavior)
            logger.info(f"🔍 Only {'document' if has_doc_evidence else 'database'} evidence found. Using individual document checks.")
            
            # Get all documents
            with connection.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT DISTINCT d.document_id
                    FROM ai_audit_data d
                    WHERE d.audit_id = %s
                      AND d.external_source != 'database_record'
                      AND d.document_type != 'db_record'
                    """,
                    [int(audit_id) if str(audit_id).isdigit() else audit_id]
                )
                doc_ids = [r[0] for r in cursor.fetchall()]

            results = []
            for doc_id in doc_ids:
                try:
                    r = check_document_compliance(request, audit_id, doc_id)
                    # If called internally, r may be a DRF Response already
                    res_data = r.data if hasattr(r, 'data') else r
                    if isinstance(res_data, dict) and res_data.get('success'):
                        results.append(res_data)
                except Exception as e:
                    logger.warning(f"⚠️ Error checking document {doc_id}: {e}")

            # Aggregate simple rollup
            summary = { 'compliant': 0, 'partially_compliant': 0, 'non_compliant': 0 }
            for r in results:
                s = r.get('status')
                if s in summary:
                    summary[s] += 1
            # After individual document checks, also try to move AI audit into "Under review"
            try:
                status_update = check_and_update_ai_audit_status(audit_id)
                logger.info(f"🔁 Bulk doc-by-doc check: status update result for audit {audit_id}: {status_update}")
            except Exception as status_err:
                logger.error(f"⚠️ Bulk doc-by-doc check: could not update audit status after checks: {status_err}")

            return Response({
                'success': True,
                'audit_id': audit_id,
                'summary': summary,
                'results': results
            })
    except Exception as e:
        logger.error(f"❌ Error checking all documents: {e}")
        return Response({'success': False, 'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def _get_audit_document_view_url(audit_id_int, doc_id_int, request):
    """Resolve view URL for an AI audit document. Returns (view_url, error) or (None, error_message)."""
    from ...rbac.utils import RBACUtils
    from ...routes.Global.s3_fucntions import RenderS3Client
    from ...models import FileOperations

    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT id, external_source, external_id, document_name, document_path
            FROM ai_audit_data
            WHERE audit_id = %s AND (id = %s OR document_id = %s)
            LIMIT 1
            """,
            [audit_id_int, doc_id_int, doc_id_int],
        )
        row = cursor.fetchone()

    if not row:
        # Fallback: treat document_id as file_operations id (e.g. schedule document_details preview)
        try:
            file_op = FileOperations.objects.get(id=doc_id_int, operation_type='upload')
        except FileOperations.DoesNotExist:
            return None, 'Document not found'
        bucket = (file_op.s3_bucket or '').strip()
        key = (file_op.s3_key or '').strip()
        file_name = file_op.original_name or file_op.file_name or 'document'
        if not bucket or not key:
            # Do not return raw s3_url - private buckets fail in the browser. Use serve URL so backend can presign or stream.
            return f'/api/ai-audit/{audit_id_int}/documents/{doc_id_int}/serve/', None
        try:
            client = RenderS3Client()
            return client.presign_get(bucket=bucket, key=key, file_name=file_name, expires_in=900, disposition='inline'), None
        except Exception as e:
            logger.warning("Presign failed for file_op %s: %s", doc_id_int, e)
            return f'/api/ai-audit/{audit_id_int}/documents/{doc_id_int}/serve/', None

    aad_id, external_source, external_id, document_name, document_path = row
    if external_source == 'evidence_attachment' and external_id:
        try:
            file_op = FileOperations.objects.get(id=int(external_id), operation_type='upload')
        except (FileOperations.DoesNotExist, ValueError):
            return None, 'Document file not found'
        bucket = (file_op.s3_bucket or '').strip()
        key = (file_op.s3_key or '').strip()
        file_name = file_op.original_name or file_op.file_name or document_name or 'document'
        if not bucket or not key:
            # Do not return raw s3_url - private buckets fail in the browser. Use serve URL so backend can presign or stream.
            return f'/api/ai-audit/{audit_id_int}/documents/{doc_id_int}/serve/', None
        try:
            client = RenderS3Client()
            return client.presign_get(bucket=bucket, key=key, file_name=file_name, expires_in=900, disposition='inline'), None
        except Exception as e:
            logger.warning("Presign failed for ai_audit_data %s (external_id=%s): %s", aad_id, external_id, e)
            return f'/api/ai-audit/{audit_id_int}/documents/{doc_id_int}/serve/', None

    if document_path:
        return f'/api/ai-audit/{audit_id_int}/documents/{doc_id_int}/serve/', None
    return None, 'Document type cannot be viewed here'


@csrf_exempt
@api_view(['GET'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@tenant_filter
def get_audit_document_view_url(request, audit_id, document_id):
    """Return a read-only view URL for an AI audit document. Opens in browser (inline, not download)."""
    from ...rbac.utils import RBACUtils

    user_id = RBACUtils.get_user_id_from_request(request)
    if not user_id:
        return Response({'success': False, 'error': 'Authentication required'}, status=status.HTTP_401_UNAUTHORIZED)

    audit_id_int = int(audit_id) if str(audit_id).isdigit() else None
    doc_id_int = int(document_id) if str(document_id).isdigit() else None
    if not audit_id_int or doc_id_int is None:
        return Response({'success': False, 'error': 'Invalid audit or document id'}, status=status.HTTP_400_BAD_REQUEST)

    view_url, err = _get_audit_document_view_url(audit_id_int, doc_id_int, request)
    if err:
        return Response({'success': False, 'error': err}, status=status.HTTP_404_NOT_FOUND if err == 'Document not found' else status.HTTP_400_BAD_REQUEST)
    return Response({'success': True, 'viewUrl': view_url}, status=status.HTTP_200_OK)


@csrf_exempt
@api_view(['GET'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@tenant_filter
def serve_audit_document(request, audit_id, document_id):
    """Serve document file for read-only viewing (inline). Tries local file first, then redirects to S3 presigned URL."""
    from ...rbac.utils import RBACUtils
    from ...models import FileOperations
    from ...routes.Global.s3_fucntions import RenderS3Client
    from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value

    user_id = RBACUtils.get_user_id_from_request(request)
    if not user_id:
        return HttpResponse('Unauthorized', status=401)

    audit_id_int = int(audit_id) if str(audit_id).isdigit() else None
    doc_id_int = int(document_id) if str(document_id).isdigit() else None
    if not audit_id_int or doc_id_int is None:
        return HttpResponse('Bad request', status=400)

    with connection.cursor() as cursor:
        cursor.execute(
            """
            SELECT id, external_source, external_id, document_name, document_path
            FROM ai_audit_data
            WHERE audit_id = %s AND (id = %s OR document_id = %s)
            LIMIT 1
            """,
            [audit_id_int, doc_id_int, doc_id_int],
        )
        row = cursor.fetchone()

    local_path = None
    file_name = 'document'
    file_op_for_s3 = None  # use to try S3 presign when no local file

    if row:
        aad_id, external_source, external_id, document_name, document_path = row
        file_name = (document_name or 'document').strip() or 'document'
        if document_path:
            try:
                document_path = decrypt_any_encrypted_value(document_path) if document_path else document_path
            except Exception:
                pass
            if document_path:
                full = os.path.join(settings.MEDIA_ROOT, document_path)
                if os.path.isfile(full):
                    local_path = full
        if not local_path and external_source == 'evidence_attachment' and external_id:
            try:
                file_op = FileOperations.objects.get(id=int(external_id), operation_type='upload')
                file_name = file_op.original_name or file_op.file_name or file_name
                stored = (file_op.s3_key or file_op.stored_name or '').strip()
                if stored:
                    full = os.path.join(settings.MEDIA_ROOT, stored)
                    if os.path.isfile(full):
                        local_path = full
                if not local_path and (file_op.s3_bucket or file_op.s3_key):
                    file_op_for_s3 = file_op
            except (FileOperations.DoesNotExist, ValueError):
                pass
    else:
        try:
            file_op = FileOperations.objects.get(id=doc_id_int, operation_type='upload')
            file_name = file_op.original_name or file_op.file_name or 'document'
            stored = (file_op.s3_key or file_op.stored_name or '').strip()
            if stored:
                full = os.path.join(settings.MEDIA_ROOT, stored)
                if os.path.isfile(full):
                    local_path = full
            if not local_path and (file_op.s3_bucket or file_op.s3_key):
                file_op_for_s3 = file_op
        except FileOperations.DoesNotExist:
            pass

    if not local_path or not os.path.isfile(local_path):
        # No local file: try redirect to S3 presigned URL so the document can still be viewed
        if file_op_for_s3:
            bucket = (file_op_for_s3.s3_bucket or '').strip()
            key = (file_op_for_s3.s3_key or '').strip()
            file_name_for_presign = file_op_for_s3.original_name or file_op_for_s3.file_name or file_name
            if not key and (file_op_for_s3.s3_url or '').strip():
                # Parse key from direct S3 URL (e.g. https://bucket.s3.region.amazonaws.com/key)
                from urllib.parse import urlparse
                try:
                    parsed = urlparse((file_op_for_s3.s3_url or '').strip())
                    if parsed.path:
                        key = parsed.path.lstrip('/')
                except Exception as parse_err:
                    logger.warning("serve_audit_document: could not parse s3_url for doc %s: %s", doc_id_int, parse_err)
            if key:
                try:
                    client = RenderS3Client()
                    presigned = client.presign_get(
                        bucket=bucket, key=key,
                        file_name=file_name_for_presign,
                        expires_in=900, disposition='inline',
                    )
                    if presigned:
                        return HttpResponseRedirect(presigned)
                except Exception as e:
                    logger.warning("serve_audit_document: presign failed for doc %s: %s", doc_id_int, e)
            # Do not redirect to raw s3_url - it fails for private buckets
        return HttpResponse('File not found', status=404)

    ext = os.path.splitext(file_name)[1].lower()
    mime = 'application/octet-stream'
    if ext in ('.pdf',):
        mime = 'application/pdf'
    elif ext in ('.doc', '.docx',):
        mime = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' if ext == '.docx' else 'application/msword'
    elif ext in ('.xls', '.xlsx',):
        mime = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet' if ext == '.xlsx' else 'application/vnd.ms-excel'
    elif ext in ('.png',):
        mime = 'image/png'
    elif ext in ('.jpg', '.jpeg',):
        mime = 'image/jpeg'
    elif ext in ('.txt',):
        mime = 'text/plain'
    elif ext in ('.csv',):
        mime = 'text/csv'
    elif ext in ('.json',):
        mime = 'application/json'
    elif ext in ('.xbrl', '.xml'):
        mime = 'application/xml'

    safe_name = os.path.basename(str(file_name).strip()) or 'document'
    response = FileResponse(open(local_path, 'rb'), as_attachment=False, filename=safe_name)
    response['Content-Type'] = mime
    response['Content-Disposition'] = f'inline; filename="{safe_name}"'
    return response


@csrf_exempt
@api_view(['DELETE'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([AuditConductPermission, AuditReviewPermission])
@audit_conduct_required
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def delete_audit_document_api(request, audit_id, document_id):
    """Delete an audit document
    MULTI-TENANCY: Only deletes documents for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        logger.info(f"🗑️ Deleting document {document_id} from audit {audit_id}")
        
        # Check authentication using JWT
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        logger.info(f"🗑️ Delete request from user: {user_id}")
        
        # Check audit status first – do not allow delete while running/completed
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT Status FROM audit WHERE AuditId = %s",
                [int(audit_id) if str(audit_id).isdigit() else audit_id],
            )
            status_row = cursor.fetchone()
        audit_status = status_row[0] if status_row else None
        if audit_status in ('Work In Progress', 'Under review', 'Completed'):
            return Response(
                {
                    'success': False,
                    'error': f'Cannot delete evidence while audit status is "{audit_status}".',
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        # Check if document exists and belongs to the audit
        # IMPORTANT: ai_audit_data table uses 'id' as primary key, not 'document_id'
        # The 'document_id' column is often 0 or NULL, so we use 'id' instead
        # Also, a file group may have multiple records (one per mapping), so we need to find all related records
        with connection.cursor() as cursor:
            # Try to find by 'id' first (primary key)
            cursor.execute(
                """
                SELECT id, document_name, document_path, document_id, external_id, file_size
                FROM ai_audit_data 
                WHERE id = %s AND audit_id = %s
                """,
                [int(document_id), int(audit_id) if str(audit_id).isdigit() else audit_id]
            )
            row = cursor.fetchone()
            
            # If not found by id, try by document_id (for backward compatibility)
            if not row:
                cursor.execute(
                    """
                    SELECT id, document_name, document_path, document_id, external_id, file_size
                    FROM ai_audit_data 
                    WHERE document_id = %s AND audit_id = %s
                    LIMIT 1
                    """,
                    [int(document_id), int(audit_id) if str(audit_id).isdigit() else audit_id]
                )
                row = cursor.fetchone()
        
        if not row:
            # Debug: Check if document exists at all (without audit_id filter)
            with connection.cursor() as debug_cursor:
                debug_cursor.execute(
                    """
                    SELECT id, document_id, audit_id, document_name 
                    FROM ai_audit_data 
                    WHERE id = %s OR document_id = %s
                    LIMIT 5
                    """,
                    [int(document_id), int(document_id)]
                )
                debug_rows = debug_cursor.fetchall()
                if debug_rows:
                    logger.warning(f"🗑️ Document {document_id} exists but not in audit {audit_id}. Found in audits: {[r[2] for r in debug_rows]}")
                else:
                    logger.warning(f"🗑️ Document {document_id} does not exist in ai_audit_data table at all")
            
            return Response({
                'success': False,
                'error': 'Document not found or does not belong to this audit'
            }, status=status.HTTP_404_NOT_FOUND)
        
        doc_id, doc_name, doc_path, doc_document_id, external_id, file_size = row
        logger.info(f"🗑️ Found document: id={doc_id}, document_id={doc_document_id}, name={doc_name}, external_id={external_id}")
        
        # Find ALL records for this file group (same external_id + file_size, or same document_path)
        # This handles the case where one file has multiple mappings (multiple database records)
        related_record_ids = [doc_id]  # Start with the found record
        
        if external_id or doc_path:
            with connection.cursor() as related_cursor:
                # Find all records with same external_id + file_size (same physical file)
                if external_id:
                    related_cursor.execute(
                        """
                        SELECT id FROM ai_audit_data 
                        WHERE audit_id = %s 
                          AND external_id = %s 
                          AND file_size = %s
                          AND id != %s
                        """,
                        [int(audit_id) if str(audit_id).isdigit() else audit_id, external_id, file_size, doc_id]
                    )
                    related_ids = [r[0] for r in related_cursor.fetchall()]
                    related_record_ids.extend(related_ids)
                    logger.info(f"🗑️ Found {len(related_ids)} additional records with same external_id={external_id} and file_size={file_size}")
                
                # Also find records with same document_path (if different from external_id)
                if doc_path and not external_id:
                    related_cursor.execute(
                        """
                        SELECT id FROM ai_audit_data 
                        WHERE audit_id = %s 
                          AND document_path = %s
                          AND id != %s
                        """,
                        [int(audit_id) if str(audit_id).isdigit() else audit_id, doc_path, doc_id]
                    )
                    related_ids = [r[0] for r in related_cursor.fetchall()]
                    related_record_ids.extend(related_ids)
                    logger.info(f"🗑️ Found {len(related_ids)} additional records with same document_path={doc_path}")
        
        # Remove duplicates
        related_record_ids = list(set(related_record_ids))
        logger.info(f"🗑️ Will delete {len(related_record_ids)} record(s) total: {related_record_ids}")
        
        # Delete ALL related database records FIRST (all mappings for this file group)
        # This ensures that when deleting a file group, all its mappings are removed
        # Do this BEFORE file cleanup to ensure DB consistency even if file cleanup fails
        deleted_count = 0
        try:
            with connection.cursor() as cursor:
                for record_id in related_record_ids:
                    try:
                        cursor.execute(
                            """
                            DELETE FROM ai_audit_data 
                            WHERE id = %s AND audit_id = %s
                            """,
                            [record_id, int(audit_id) if str(audit_id).isdigit() else audit_id]
                        )
                        deleted_count += cursor.rowcount
                        logger.info(f"🗑️ Deleted record id={record_id}")
                    except Exception as delete_err:
                        logger.error(f"❌ Error deleting record {record_id}: {delete_err}")
                        # Continue with other records even if one fails
                        continue
        except Exception as db_err:
            logger.error(f"❌ Database error during delete: {db_err}")
            return Response({
                'success': False,
                'error': f'Database error during delete: {str(db_err)}'
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Delete the physical file (if present) - do this after DB deletion
        try:
            if doc_path:
                full_path = doc_path if os.path.isabs(doc_path) else os.path.join(settings.MEDIA_ROOT, doc_path)
                if os.path.exists(full_path):
                    os.remove(full_path)
                    logger.info(f"🗑️ Deleted file: {full_path}")
                else:
                    logger.warning(f"🗑️ File not found on disk: {full_path}")
        except Exception as e:
            logger.warning(f"🗑️ Could not delete file {doc_path}: {e}")
            # Don't fail the entire operation if file deletion fails
        
        # Clean up any temporary files that might have been created during AI processing
        # Do this last and don't block on it - it's cleanup, not critical
        try:
            cleaned_temp_files = cleanup_ai_audit_temp_files(document_id, doc_id)
            if cleaned_temp_files:
                logger.info(f"🗑️ Cleaned up {len(cleaned_temp_files)} temporary files for document {document_id}")
        except Exception as cleanup_err:
            logger.warning(f"🗑️ Could not clean up temp files: {cleanup_err}")
            # Don't fail the operation if cleanup fails
        
        if deleted_count == 0:
            return Response({
                'success': False,
                'error': 'Document could not be deleted'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        logger.info(f"✅ Document {doc_name} deleted successfully")
        
        return Response({
            'success': True,
            'message': f'Document {doc_name} deleted successfully',
            'document_id': document_id,
            'audit_id': audit_id
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        logger.error(f"❌ Error deleting document: {e}")
        return Response({
            'success': False,
            'error': f'Failed to delete document: {str(e)}'
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
@api_view(['DELETE'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([AuditConductPermission, AuditReviewPermission])
@audit_conduct_required
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def delete_all_audit_documents_api(request, audit_id):
    """
    Delete ALL ai_audit_data rows for a given audit.
    Used by the UI 'Delete All' to quickly clear evidence/documents.
    MULTI-TENANCY: Only deletes documents for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        logger.info(f"🗑️ Bulk delete requested for all documents/evidence of audit {audit_id}")

        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)

        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)

        converted_audit_id = int(audit_id) if str(audit_id).isdigit() else audit_id

        with connection.cursor() as cursor:
            cursor.execute(
                """
                DELETE FROM ai_audit_data
                WHERE audit_id = %s
                """,
                [converted_audit_id]
            )
            deleted_count = cursor.rowcount

        logger.info(f"✅ Bulk delete completed for audit {audit_id}: {deleted_count} row(s) removed from ai_audit_data")

        return Response({
            'success': True,
            'message': f'Deleted {deleted_count} evidence/document record(s) for audit {audit_id}',
            'deleted_count': deleted_count
        }, status=status.HTTP_200_OK)
    except Exception as e:
        logger.error(f"❌ Error in bulk delete for audit {audit_id}: {e}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        return {}


@csrf_exempt
@api_view(['GET'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant
@tenant_filter
def get_ai_annual_consolidation_report(request, audit_id):
    """
    Annual consolidation / issue-summary style view for a single AI audit.

    For the given audit, aggregates ai_audit_data rows across the selected year
    (default = year of today) and returns counts of compliant / partially
    compliant / non-compliant items per compliance (or policy/subpolicy when
    compliance_id is not available).

    Additionally, it highlights **unresolved items from the last cycle**:
    - If the same control/policy/subpolicy was NON_COMPLIANT in the
      immediately previous year (year - 1) for this tenant, and it is still
      NON_COMPLIANT in the selected year, it is flagged as
      `is_unresolved_from_previous_cycles = True`.
    """
    tenant_id = get_tenant_id_from_request(request)

    try:
        year_param = request.GET.get('year')
        if year_param and str(year_param).isdigit():
            year = int(year_param)
        else:
            year = datetime.now().year

        converted_audit_id = int(audit_id) if str(audit_id).isdigit() else audit_id

        with connection.cursor() as cursor:
            # Scope to tenant's audit to avoid cross-tenant leakage
            cursor.execute(
                """
                SELECT TenantId
                FROM audit
                WHERE AuditId = %s
                """,
                [converted_audit_id],
            )
            row = cursor.fetchone()
            audit_tenant_id = row[0] if row else None
            if tenant_id and audit_tenant_id and str(tenant_id) != str(audit_tenant_id):
                return Response(
                    {
                        'success': False,
                        'error': 'Access denied for this audit in current tenant',
                    },
                    status=status.HTTP_403_FORBIDDEN,
                )

            # Some deployments don't have compliance_id on ai_audit_data yet – detect safely
            has_compliance_id = _check_ai_audit_data_has_compliance_id()

            # 1) Aggregate for the selected year
            if has_compliance_id:
                cursor.execute(
                    """
                    SELECT
                        COALESCE(CAST(compliance_id AS CHAR), '')     AS compliance_id,
                        COALESCE(CAST(policy_id AS CHAR), '')         AS policy_id,
                        COALESCE(CAST(subpolicy_id AS CHAR), '')      AS subpolicy_id,
                        compliance_status,
                        COUNT(*) AS count_items
                    FROM ai_audit_data
                    WHERE audit_id = %s
                      AND YEAR(uploaded_date) = %s
                    GROUP BY
                        COALESCE(CAST(compliance_id AS CHAR), ''),
                        COALESCE(CAST(policy_id AS CHAR), ''),
                        COALESCE(CAST(subpolicy_id AS CHAR), ''),
                        compliance_status
                    """,
                    [converted_audit_id, year],
                )
            else:
                # Fallback: group only by policy/subpolicy when compliance_id column is missing
                cursor.execute(
                    """
                    SELECT
                        ''                                         AS compliance_id,
                        COALESCE(CAST(policy_id AS CHAR), '')     AS policy_id,
                        COALESCE(CAST(subpolicy_id AS CHAR), '')  AS subpolicy_id,
                        compliance_status,
                        COUNT(*) AS count_items
                    FROM ai_audit_data
                    WHERE audit_id = %s
                      AND YEAR(uploaded_date) = %s
                    GROUP BY
                        COALESCE(CAST(policy_id AS CHAR), ''),
                        COALESCE(CAST(subpolicy_id AS CHAR), ''),
                        compliance_status
                    """,
                    [converted_audit_id, year],
                )

            rows = cursor.fetchall()

            # 2) Find non-compliant items from the **immediately previous year**
            #    for the same control (policy / subpolicy / compliance) across
            #    any audit. We keep per-control counts so the UI can show how
            #    many non-compliant hits occurred last year.
            unresolved_map = {}
            if year > 1900:
                previous_year = year - 1
                if has_compliance_id:
                    cursor.execute(
                        """
                        SELECT
                            COALESCE(CAST(compliance_id AS CHAR), '')     AS compliance_id,
                            COALESCE(CAST(policy_id AS CHAR), '')         AS policy_id,
                            COALESCE(CAST(subpolicy_id AS CHAR), '')      AS subpolicy_id,
                            COUNT(*) AS cnt
                        FROM ai_audit_data d
                        WHERE YEAR(d.uploaded_date) = %s
                          AND UPPER(TRIM(d.compliance_status)) = 'NON_COMPLIANT'
                        GROUP BY
                            COALESCE(CAST(compliance_id AS CHAR), ''),
                            COALESCE(CAST(policy_id AS CHAR), ''),
                            COALESCE(CAST(subpolicy_id AS CHAR), '')
                        """,
                        [previous_year],
                    )
                else:
                    cursor.execute(
                        """
                        SELECT
                            ''                                         AS compliance_id,
                            COALESCE(CAST(policy_id AS CHAR), '')       AS policy_id,
                            COALESCE(CAST(subpolicy_id AS CHAR), '')    AS subpolicy_id,
                            COUNT(*) AS cnt
                        FROM ai_audit_data d
                        WHERE YEAR(d.uploaded_date) = %s
                          AND UPPER(TRIM(d.compliance_status)) = 'NON_COMPLIANT'
                        GROUP BY
                            COALESCE(CAST(policy_id AS CHAR), ''),
                            COALESCE(CAST(subpolicy_id AS CHAR), '')
                        """,
                        [previous_year],
                    )
                for prev_compliance_id, prev_policy_id, prev_subpolicy_id, cnt in cursor.fetchall():
                    key = (
                        prev_compliance_id or '',
                        prev_policy_id or '',
                        prev_subpolicy_id or '',
                    )
                    unresolved_map[key] = unresolved_map.get(key, 0) + int(cnt or 0)

        # Post-process into a friendlier structure
        aggregates = {}
        for compliance_id, policy_id, subpolicy_id, status_label, cnt in rows:
            key = (compliance_id or '', policy_id or '', subpolicy_id or '')
            bucket = aggregates.setdefault(
                key,
                {
                    'compliance_id': compliance_id or None,
                    'policy_id': policy_id or None,
                    'subpolicy_id': subpolicy_id or None,
                    'compliant': 0,
                    'partially_compliant': 0,
                    'non_compliant': 0,
                    'unknown': 0,
                },
            )
            normalized = (status_label or '').strip().upper()
            if normalized == 'COMPLIANT':
                bucket['compliant'] += cnt
            elif normalized == 'PARTIALLY_COMPLIANT':
                bucket['partially_compliant'] += cnt
            elif normalized == 'NON_COMPLIANT':
                bucket['non_compliant'] += cnt
            else:
                bucket['unknown'] += cnt

        # Look up human-readable names for all distinct IDs
        policy_ids = set()
        subpolicy_ids = set()
        compliance_ids = set()
        for b in aggregates.values():
            if b['policy_id']:
                policy_ids.add(int(b['policy_id'])) if str(b['policy_id']).isdigit() else None
            if b['subpolicy_id']:
                subpolicy_ids.add(int(b['subpolicy_id'])) if str(b['subpolicy_id']).isdigit() else None
            if b['compliance_id']:
                compliance_ids.add(int(b['compliance_id'])) if str(b['compliance_id']).isdigit() else None

        policy_names = {}
        subpolicy_names = {}
        compliance_names = {}
        try:
            with connection.cursor() as cursor:
                if policy_ids:
                    placeholders = ",".join(["%s"] * len(policy_ids))
                    cursor.execute(
                        f"SELECT PolicyId, PolicyName FROM policies WHERE PolicyId IN ({placeholders})",
                        list(policy_ids),
                    )
                    for pid, pname in cursor.fetchall():
                        policy_names[str(pid)] = pname
                if subpolicy_ids:
                    placeholders = ",".join(["%s"] * len(subpolicy_ids))
                    cursor.execute(
                        f"SELECT SubPolicyId, SubPolicyName FROM subpolicies WHERE SubPolicyId IN ({placeholders})",
                        list(subpolicy_ids),
                    )
                    for sid, sname in cursor.fetchall():
                        subpolicy_names[str(sid)] = sname
                if compliance_ids:
                    placeholders = ",".join(["%s"] * len(compliance_ids))
                    cursor.execute(
                        f"SELECT ComplianceId, ComplianceTitle FROM compliances WHERE ComplianceId IN ({placeholders})",
                        list(compliance_ids),
                    )
                    for cid, cname in cursor.fetchall():
                        compliance_names[str(cid)] = cname
        except Exception:
            # Name lookup is best-effort; continue even if it fails
            pass

        items = []
        total_non_compliant = 0
        total_items = 0
        total_unresolved_from_previous_cycles = 0

        for bucket in aggregates.values():
            # Skip completely unlinked rows (no compliance/policy/subpolicy)
            if not bucket['compliance_id'] and not bucket['policy_id'] and not bucket['subpolicy_id']:
                continue
            bucket_total = (
                bucket['compliant']
                + bucket['partially_compliant']
                + bucket['non_compliant']
                + bucket['unknown']
            )
            total_items += bucket_total
            total_non_compliant += bucket['non_compliant']

            key_for_unresolved = (
                bucket['compliance_id'] if bucket['compliance_id'] is not None else '',
                bucket['policy_id'] if bucket['policy_id'] is not None else '',
                bucket['subpolicy_id'] if bucket['subpolicy_id'] is not None else '',
            )
            previous_year_non_compliant = unresolved_map.get(key_for_unresolved, 0)
            is_unresolved_from_previous = (
                bucket['non_compliant'] > 0 and previous_year_non_compliant > 0
            )
            if is_unresolved_from_previous:
                total_unresolved_from_previous_cycles += 1

            # Keep existing repeat-finding logic (multiple non-compliant hits in the same year)
            is_repeat = bucket['non_compliant'] > 1
            policy_id_str = str(bucket['policy_id']) if bucket['policy_id'] is not None else None
            subpolicy_id_str = str(bucket['subpolicy_id']) if bucket['subpolicy_id'] is not None else None
            compliance_id_str = str(bucket['compliance_id']) if bucket['compliance_id'] is not None else None
            items.append(
                {
                    'compliance_id': bucket['compliance_id'],
                    'policy_id': bucket['policy_id'],
                    'subpolicy_id': bucket['subpolicy_id'],
                    'policy_name': policy_names.get(policy_id_str),
                    'subpolicy_name': subpolicy_names.get(subpolicy_id_str),
                    'compliance_name': compliance_names.get(compliance_id_str),
                    'compliant': bucket['compliant'],
                    'partially_compliant': bucket['partially_compliant'],
                    'non_compliant': bucket['non_compliant'],
                    'unknown': bucket['unknown'],
                    'is_repeat_finding': is_repeat,
                    'is_unresolved_from_previous_cycles': is_unresolved_from_previous,
                    'previous_year_non_compliant': previous_year_non_compliant,
                }
            )

        response_payload = {
            'success': True,
            'audit_id': converted_audit_id,
            'year': year,
            'summary': {
                'total_items': total_items,
                'total_non_compliant': total_non_compliant,
                'total_repeat_findings': sum(1 for i in items if i['is_repeat_finding']),
                'total_unresolved_from_previous_cycles': total_unresolved_from_previous_cycles,
            },
            'items': items,
        }

        return Response(response_payload, status=status.HTTP_200_OK)

    except Exception as e:
        logger.error(f"❌ Error in get_ai_annual_consolidation_report for audit {audit_id}: {e}")
        return Response(
            {'success': False, 'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR,
        )


# AI processing functions using OpenAI


@csrf_exempt
@api_view(['GET'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def download_audit_report(request, audit_id):
    """Generate a comprehensive AI audit report and return it as a download.
    MULTI-TENANCY: Only generates reports for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        logger.info(f"📊 Generating comprehensive AI audit report for audit {audit_id}")
        
        # Check authentication using JWT
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        # Get audit information
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT a.AuditId, a.AuditType, a.Status, a.AssignedDate, a.DueDate, a.CompletionDate,
                       p.PolicyName, sp.SubPolicyName, f.FrameworkName, a.FrameworkId,
                       u1.UserName as AuditorName,
                       u2.UserName as AssigneeName,
                       u3.UserName as ReviewerName
                FROM audit a
                LEFT JOIN policies p ON a.PolicyId = p.PolicyId
                LEFT JOIN subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                LEFT JOIN frameworks f ON a.FrameworkId = f.FrameworkId
                LEFT JOIN users u1 ON a.Auditor = u1.UserId
                LEFT JOIN users u2 ON a.Assignee = u2.UserId
                LEFT JOIN users u3 ON a.Reviewer = u3.UserId
                WHERE a.AuditId = %s
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
            
            audit_row = cursor.fetchone()
            framework_id = audit_row[9] if audit_row else None
            
            # Check if this is actually an AI audit by looking for AI audit data
            cursor.execute("""
                SELECT COUNT(*) FROM ai_audit_data WHERE audit_id = %s
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
            
            ai_data_count = cursor.fetchone()[0]
            is_actual_ai_audit = ai_data_count > 0
            
        if not audit_row:
            return Response({
                'success': False,
                'error': f'Audit {audit_id} not found'
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Optional filter: limit report to specific document_ids passed from UI, or to schedule scope
        selected_document_ids = None
        schedule_compliance_ids = None  # When set, only include these compliance IDs in the report (schedule scope)
        try:
            raw_ids = request.GET.get('document_ids') or request.query_params.get('document_ids')  # type: ignore[attr-defined]
        except Exception:
            raw_ids = None
        
        if raw_ids:
            try:
                id_list = [
                    int(x) for x in str(raw_ids).split(',')
                    if str(x).strip().isdigit()
                ]
                if id_list:
                    selected_document_ids = id_list
            except Exception:
                selected_document_ids = None
        
        # When schedule_id is provided, limit report to that schedule's documents and compliances (same as Compliance Results)
        schedule_id_param = (request.GET.get('schedule_id') or getattr(request, 'query_params', None) and request.query_params.get('schedule_id')) or ''
        schedule_id_param = str(schedule_id_param).strip() if schedule_id_param else ''
        if schedule_id_param and schedule_id_param.isdigit():
            converted_audit_id = int(audit_id) if str(audit_id).isdigit() else audit_id
            try:
                from ...models import AIAuditSchedule
                schedule = AIAuditSchedule.objects.get(id=int(schedule_id_param), audit_id=converted_audit_id)
            except Exception:
                schedule = None
            if schedule:
                from .ai_audit_schedule_api import _get_document_ids_for_schedule
                schedule_audit_data_ids = _get_document_ids_for_schedule(schedule)
                if schedule_audit_data_ids:
                    with connection.cursor() as cur:
                        placeholders = ','.join(['%s'] * len(schedule_audit_data_ids))
                        cur.execute(
                            f"SELECT DISTINCT document_id FROM ai_audit_data WHERE id IN ({placeholders}) AND document_id IS NOT NULL",
                            list(schedule_audit_data_ids)
                        )
                        selected_document_ids = [r[0] for r in cur.fetchall()]
                else:
                    selected_document_ids = []
                sel = getattr(schedule, 'selected_compliance_ids', None)
                if sel:
                    if isinstance(sel, list):
                        schedule_compliance_ids = set(int(x) for x in sel if x is not None and str(x).strip().isdigit())
                    elif isinstance(sel, str):
                        try:
                            parsed = json.loads(sel)
                            schedule_compliance_ids = set(int(x) for x in (parsed if isinstance(parsed, list) else []) if x is not None and str(x).strip().isdigit())
                        except Exception:
                            pass
                if schedule_compliance_ids is not None and not schedule_compliance_ids:
                    schedule_compliance_ids = None
                logger.info(f"📊 Report limited to schedule {schedule_id_param}: {len(selected_document_ids or [])} doc(s), {len(schedule_compliance_ids or [])} compliance(s)")
        
        # Get all (or filtered) AI audit documents and their processing results
        with connection.cursor() as cursor:
            converted_audit_id = int(audit_id) if str(audit_id).isdigit() else audit_id
            
            if selected_document_ids:
                placeholders = ",".join(["%s"] * len(selected_document_ids))
                cursor.execute(
                    f"""
                    SELECT document_id, document_name, document_type, file_size, mime_type,
                           ai_processing_status, compliance_status, confidence_score,
                           compliance_analyses, processing_completed_at, uploaded_date,
                           policy_id, subpolicy_id, external_source, external_id, document_path
                    FROM ai_audit_data 
                    WHERE audit_id = %s
                      AND document_id IN ({placeholders})
                    """,
                    [converted_audit_id, *selected_document_ids]
                )
            else:
                cursor.execute(
                    """
                    SELECT document_id, document_name, document_type, file_size, mime_type,
                           ai_processing_status, compliance_status, confidence_score,
                           compliance_analyses, processing_completed_at, uploaded_date,
                           policy_id, subpolicy_id, external_source, external_id, document_path
                    FROM ai_audit_data 
                    WHERE audit_id = %s 
                    """,
                    [converted_audit_id]
                )
            
            documents = cursor.fetchall()
            columns = [col[0] for col in cursor.description]
            
            # Sort in Python to avoid MySQL sort memory issues
            # Sort by uploaded_date DESC (newest first)
            uploaded_date_idx = columns.index('uploaded_date') if 'uploaded_date' in columns else None
            if uploaded_date_idx is not None:
                documents = sorted(documents, key=lambda x: x[uploaded_date_idx] if x[uploaded_date_idx] else None, reverse=True)
        
        # Process documents data - GROUP BY PHYSICAL FILES
        # Group mappings by physical file to count unique files, not mapping records
        # Use document_path or external_id to identify the same physical file, fallback to (name, size)
        file_groups = {}  # Key: unique file identifier, Value: list of mapping records
        from grc.utils.auto_decrypt_helper import decrypt_any_encrypted_value as _decrypt_doc
        from grc.utils.auto_decrypt_helper import decrypt_all_encrypted_in_dict as _decrypt_dict
        for doc_row in documents:
            doc_dict = dict(zip(columns, doc_row))
            # Decrypt document_name and document_path (may be stored encrypted from FileOperations)
            raw_name = doc_dict.get('document_name')
            raw_path = doc_dict.get('document_path')
            if raw_name:
                doc_dict['document_name'] = _decrypt_doc(raw_name) or raw_name
            if raw_path:
                doc_dict['document_path'] = _decrypt_doc(raw_path) or raw_path
            # Try to use document_path or external_id to identify the same physical file
            document_path = doc_dict.get('document_path', '')
            external_id = doc_dict.get('external_id', '')
            
            # Create a unique key for the physical file
            if document_path:
                # Use document_path as primary identifier (same path = same file)
                file_key = ('path', document_path)
            elif external_id:
                # Use external_id (S3 key or external identifier) as identifier
                file_key = ('external', str(external_id))
            else:
                # Fallback to (document_name, file_size) if no path or external_id
                file_key = ('name_size', doc_dict.get('document_name', ''), doc_dict.get('file_size', 0))
            
            if file_key not in file_groups:
                file_groups[file_key] = []
            file_groups[file_key].append(doc_dict)
        
        # Count unique physical files
        total_physical_files = len(file_groups)
        completed_physical_files = 0
        failed_physical_files = 0
        
        processed_documents = []
        total_mapping_records = len(documents)  # Total mapping records (for reference)
        completed_mapping_records = 0
        failed_mapping_records = 0
        # Mapping-level summary (one count per mapping record)
        compliance_summary = {'compliant': 0, 'partially_compliant': 0, 'non_compliant': 0, 'unknown': 0}
        # Requirement-level summary (one count per individual requirement / analysis item)
        # IMPORTANT: Track (document_id, compliance_id) pairs to avoid counting duplicates
        requirement_summary = {'compliant': 0, 'partially_compliant': 0, 'non_compliant': 0, 'unknown': 0}
        counted_requirements = set()  # Track (document_id, compliance_id) pairs already counted
        
        # Collect unique policy/sub-policy IDs from uploaded documents (multi-select approach)
        selected_policy_ids = set()
        selected_subpolicy_ids = set()
        # Collect unique compliance IDs seen in analyses (requirement-level)
        selected_compliance_ids = set()
        policy_names_map = {}
        subpolicy_names_map = {}
        
        # Determine actual audit status based on AI processing results
        actual_audit_status = audit_row[2]  # Default to original status
        actual_completion_date = audit_row[5]  # Default to original completion date
        
        # Process each physical file group
        for file_key, mapping_records in file_groups.items():
            # Extract document_name and file_size from the first record
            first_record = mapping_records[0]
            document_name = first_record.get('document_name', '')
            file_size = first_record.get('file_size', 0)
            
            # Check if this physical file has any completed mappings
            file_has_completed = any(
                record.get('ai_processing_status') == 'completed' 
                for record in mapping_records
            )
            file_has_failed = any(
                record.get('ai_processing_status') == 'failed' 
                for record in mapping_records
            )
            
            if file_has_completed:
                completed_physical_files += 1
            elif file_has_failed:
                failed_physical_files += 1
            
            # Process each mapping record for this file
            file_mappings = []
            # Track requirements for this physical file to avoid duplicates across mappings
            # Use a string representation of the file identifier for reliable hashing
            file_identifier = None
            if first_record.get('document_path'):
                file_identifier = f"path:{first_record.get('document_path')}"
            elif first_record.get('external_id'):
                file_identifier = f"external:{first_record.get('external_id')}"
            else:
                file_identifier = f"name_size:{first_record.get('document_name', '')}:{first_record.get('file_size', 0)}"
            
            file_requirements_counted = set()
            for doc_dict in mapping_records:
                # Parse compliance analyses if available
                compliance_analyses = None
                analyses_list = None
                if doc_dict.get('compliance_analyses'):
                    try:
                        try:
                            compliance_analyses = json.loads(doc_dict['compliance_analyses'])
                            if compliance_analyses:
                                compliance_analyses = _decrypt_dict(compliance_analyses)
                        except Exception as e:
                            logger.error(f"Failed to parse compliance analyses JSON: {e}")
                            compliance_analyses = None
                    except Exception:
                        compliance_analyses = None
                
                # Build a flat list of requirement-level analyses (if present)
                if compliance_analyses:
                    if isinstance(compliance_analyses, dict) and 'compliance_analyses' in compliance_analyses:
                        analyses_list = compliance_analyses.get('compliance_analyses') or []
                    elif isinstance(compliance_analyses, list):
                        analyses_list = compliance_analyses
                    else:
                        analyses_list = None
                
                # Count compliance status (only for completed mappings; when schedule scope, only if mapping has schedule's compliances)
                processing_status = doc_dict.get('ai_processing_status', 'pending')
                if processing_status == 'completed':
                    has_schedule_compliance = True
                    if schedule_compliance_ids is not None and analyses_list and isinstance(analyses_list, list):
                        try:
                            has_schedule_compliance = any(
                                isinstance(a, dict) and a.get('compliance_id') is not None
                                and int(a.get('compliance_id')) in schedule_compliance_ids
                                for a in analyses_list
                            )
                        except (TypeError, ValueError):
                            has_schedule_compliance = False
                    if schedule_compliance_ids is None or has_schedule_compliance:
                        status = doc_dict.get('compliance_status', 'unknown')
                        if status in compliance_summary:
                            compliance_summary[status] += 1
                        completed_mapping_records += 1
                elif processing_status == 'failed':
                    failed_mapping_records += 1

                # Requirement-level compliance counting (each individual requirement)
                # IMPORTANT: Only count each compliance_id once per physical file (file_key)
                # to avoid counting the same requirement multiple times when file has multiple mappings
                if analyses_list and isinstance(analyses_list, list):
                    for analysis in analyses_list:
                        if not isinstance(analysis, dict):
                            continue
                        comp_id = analysis.get('compliance_id')
                        if schedule_compliance_ids is not None and comp_id is not None:
                            try:
                                if int(comp_id) not in schedule_compliance_ids:
                                    continue
                            except (TypeError, ValueError):
                                continue
                        # Track compliance IDs for audit-level summary
                        try:
                            if comp_id is not None:
                                selected_compliance_ids.add(int(comp_id))
                        except (TypeError, ValueError):
                            pass
                        
                        # Create unique key using file_identifier and compliance_id to avoid duplicates
                        requirement_key = None
                        if comp_id is not None and file_identifier:
                            try:
                                comp_id_int = int(comp_id)
                                # Use file_identifier (string) + compliance_id for reliable hashing
                                requirement_key = f"{file_identifier}:{comp_id_int}"
                            except (TypeError, ValueError):
                                pass
                        
                        # Only count if we haven't seen this (file_identifier, compliance_id) combination before
                        if requirement_key and requirement_key in file_requirements_counted:
                            continue  # Skip - already counted this requirement for this physical file
                        
                        # Skip requirements without compliance_id (can't deduplicate them reliably, would cause duplicates)
                        if not requirement_key:
                            continue
                        
                        raw_status = (
                            str(analysis.get('status') or analysis.get('compliance_status') or '')
                            .strip()
                            .lower()
                        )
                        if not raw_status:
                            key = 'unknown'
                        elif 'non_compliant' in raw_status or raw_status == 'non-compliant':
                            key = 'non_compliant'
                        elif 'partially' in raw_status:
                            key = 'partially_compliant'
                        elif 'compliant' in raw_status:
                            # If it says compliant but not partially/non_compliant, treat as compliant
                            key = 'compliant'
                        else:
                            key = 'unknown'
                        if key in requirement_summary:
                            requirement_summary[key] += 1
                            # Mark this requirement as counted for this physical file
                            file_requirements_counted.add(requirement_key)
                            counted_requirements.add(requirement_key)
                
                # Collect policy/sub-policy IDs for multi-select display
                if doc_dict.get('policy_id'):
                    selected_policy_ids.add(doc_dict['policy_id'])
                if doc_dict.get('subpolicy_id'):
                    selected_subpolicy_ids.add(doc_dict['subpolicy_id'])
                
                # When schedule scope: only include analyses for schedule's compliances in the report
                report_analyses = compliance_analyses
                if schedule_compliance_ids is not None and compliance_analyses and analyses_list:
                    if isinstance(compliance_analyses, list):
                        try:
                            report_analyses = [a for a in compliance_analyses if isinstance(a, dict) and a.get('compliance_id') is not None and int(a.get('compliance_id')) in schedule_compliance_ids]
                        except (TypeError, ValueError):
                            report_analyses = compliance_analyses
                    elif isinstance(compliance_analyses, dict) and 'compliance_analyses' in compliance_analyses:
                        try:
                            sub = [a for a in (compliance_analyses.get('compliance_analyses') or []) if isinstance(a, dict) and a.get('compliance_id') is not None and int(a.get('compliance_id')) in schedule_compliance_ids]
                            report_analyses = dict(compliance_analyses)
                            report_analyses['compliance_analyses'] = sub
                        except (TypeError, ValueError):
                            report_analyses = compliance_analyses
                # Store mapping record
                file_mappings.append({
                    'document_id': doc_dict.get('document_id'),
                    'policy_id': doc_dict.get('policy_id'),
                    'subpolicy_id': doc_dict.get('subpolicy_id'),
                    'ai_processing_status': processing_status,
                    'compliance_status': doc_dict.get('compliance_status'),
                    'confidence_score': doc_dict.get('confidence_score'),
                    'compliance_analyses': report_analyses,
                    'processing_completed_at': doc_dict.get('processing_completed_at'),
                })
            
            # Create one entry per physical file (with all its mappings)
            # Use the first record for file-level metadata
            first_record = mapping_records[0]
            processed_documents.append({
                'document_name': document_name,
                'document_type': first_record.get('document_type'),
                'file_size': file_size,
                'mime_type': first_record.get('mime_type'),
                'uploaded_date': first_record.get('uploaded_date'),
                'ai_processing_status': 'completed' if file_has_completed else ('failed' if file_has_failed else 'pending'),
                'mappings': file_mappings,  # All mappings for this physical file
                'mapping_count': len(file_mappings)
            })
        
        # Fetch policy and sub-policy names for selected IDs (decrypt if stored encrypted)
        if selected_policy_ids:
            policy_ids_list = sorted(selected_policy_ids)
            placeholders = ','.join(['%s'] * len(policy_ids_list))
            with connection.cursor() as cursor:
                cursor.execute(
                    f"""
                    SELECT PolicyId, PolicyName FROM policies 
                    WHERE PolicyId IN ({placeholders})
                    """,
                    policy_ids_list,
                )
                for row in cursor.fetchall():
                    policy_names_map[row[0]] = _decrypt_doc(row[1]) if row[1] else row[1]
        
        if selected_subpolicy_ids:
            subpolicy_ids_list = sorted(selected_subpolicy_ids)
            placeholders = ','.join(['%s'] * len(subpolicy_ids_list))
            with connection.cursor() as cursor:
                cursor.execute(
                    f"""
                    SELECT SubPolicyId, SubPolicyName FROM subpolicies 
                    WHERE SubPolicyId IN ({placeholders})
                    """,
                    subpolicy_ids_list,
                )
                for row in cursor.fetchall():
                    subpolicy_names_map[row[0]] = _decrypt_doc(row[1]) if row[1] else row[1]

        # Fetch compliance titles for selected compliance IDs (decrypt if stored encrypted)
        compliance_names_map = {}
        if selected_compliance_ids:
            compliance_ids_list = sorted(selected_compliance_ids)
            placeholders = ','.join(['%s'] * len(compliance_ids_list))
            with connection.cursor() as cursor:
                cursor.execute(
                    f"""
                    SELECT ComplianceId, ComplianceTitle
                    FROM compliance
                    WHERE ComplianceId IN ({placeholders})
                    """,
                    compliance_ids_list,
                )
                for row in cursor.fetchall():
                    cid, title = row[0], row[1] or f"Compliance {row[0]}"
                    compliance_names_map[cid] = _decrypt_doc(title) if title else title
        
        # Build selected policies/sub-policies/compliances display strings
        # 1) Prefer explicit multi-select (selected_policy_ids / selected_subpolicy_ids)
        selected_policies_display = ', '.join(
            [policy_names_map.get(pid, f'Policy {pid}') for pid in sorted(selected_policy_ids)]
        ) if selected_policy_ids else ''
        selected_subpolicies_display = ', '.join(
            [subpolicy_names_map.get(sid, f'Sub-policy {sid}') for sid in sorted(selected_subpolicy_ids)]
        ) if selected_subpolicy_ids else ''

        # 2) If no explicit policy/subpolicy selected but we have compliances in scope,
        #    derive the policy / sub-policy names from those compliances so the report
        #    does not show "Policy: Not Specified" when the scope is actually clear.
        if (not selected_policy_ids or not selected_subpolicy_ids) and selected_compliance_ids:
            try:
                with connection.cursor() as cursor:
                    placeholders = ','.join(['%s'] * len(selected_compliance_ids))
                    cursor.execute(
                        f"""
                        SELECT DISTINCT c.ComplianceId,
                                        c.SubPolicyId,
                                        sp.SubPolicyName,
                                        sp.PolicyId,
                                        p.PolicyName
                        FROM compliance c
                        JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                        JOIN policies p ON sp.PolicyId = p.PolicyId
                        WHERE c.ComplianceId IN ({placeholders})
                        """,
                        list(selected_compliance_ids)
                    )
                    derived_policy_names: dict[int, str] = {}
                    derived_subpolicy_names: dict[int, str] = {}
                    for cid, subpol_id, subpol_name, pol_id, pol_name in cursor.fetchall():
                        if pol_id and pol_id not in derived_policy_names:
                            derived_policy_names[pol_id] = _decrypt_doc(pol_name) if pol_name else pol_name
                        if subpol_id and subpol_id not in derived_subpolicy_names:
                            derived_subpolicy_names[subpol_id] = _decrypt_doc(subpol_name) if subpol_name else subpol_name

                # Only override when we actually derived something
                if not selected_policies_display and derived_policy_names:
                    selected_policies_display = ', '.join(
                        [derived_policy_names[pid] or f'Policy {pid}' for pid in sorted(derived_policy_names.keys())]
                    )
                if not selected_subpolicies_display and derived_subpolicy_names:
                    selected_subpolicies_display = ', '.join(
                        [derived_subpolicy_names[sid] or f'Sub-policy {sid}' for sid in sorted(derived_subpolicy_names.keys())]
                    )
            except Exception as e:
                logger.warning(f"⚠️ Could not derive policy/sub-policy names from compliances for AI report: {e}")

        # 3) Final fallback to audit row (legacy single policy/sub-policy) or 'Not Specified'
        if not selected_policies_display:
            selected_policies_display = audit_row[6] or 'Not Specified'
        if not selected_subpolicies_display:
            selected_subpolicies_display = audit_row[7] or 'Not Specified'
        if selected_compliance_ids:
            sorted_cids = sorted(selected_compliance_ids)
            compliance_titles = [compliance_names_map.get(cid, f"Compliance {cid}") for cid in sorted_cids]
            if len(compliance_titles) <= 5:
                selected_compliances_display = f"{len(sorted_cids)} compliance(s): " + ", ".join(compliance_titles)
            else:
                selected_compliances_display = f"{len(sorted_cids)} compliance(s): " + ", ".join(compliance_titles[:5]) + ", ..."
        else:
            selected_compliances_display = "None selected"
        
        # Calculate actual audit status based on AI processing results (using physical files)
        if is_actual_ai_audit and total_physical_files > 0:
            completion_percentage = (completed_physical_files / total_physical_files) * 100
            
            if completion_percentage == 100:
                actual_audit_status = 'Completed'
                # Set completion date to the latest processing completion date
                latest_completion = None
                for doc in processed_documents:
                    if doc.get('processing_completed_at'):
                        if latest_completion is None or doc['processing_completed_at'] > latest_completion:
                            latest_completion = doc['processing_completed_at']
                if latest_completion:
                    actual_completion_date = latest_completion
            elif completion_percentage > 0:
                actual_audit_status = 'Work In Progress'
        
        # Check if SEBI AI Auditor is enabled and get SEBI insights
        # IMPORTANT: Run SEBI checks BEFORE generating report so insights are included
        sebi_insights = None
        sebi_enabled = False
        try:
            from .sebi_ai_auditor import SEBIAIAuditor
            if framework_id:
                auditor = SEBIAIAuditor(framework_id, tenant_id)
                if auditor.is_sebi_framework:
                    sebi_enabled = True
                    logger.info(f"🔍 SEBI AI Auditor enabled - running checks for audit {audit_id} before report generation...")
                    # Get SEBI insights - RUN SYNCHRONOUSLY to ensure they're ready before report generation
                    logger.info(f"📊 Calculating filing accuracy for audit {audit_id}...")
                    filing_accuracy = auditor.verify_filing_accuracy(audit_id)
                    
                    logger.info(f"⏱️ Checking timeliness SLA for audit {audit_id}...")
                    timeliness_sla = auditor.check_timeliness_sla(audit_id)
                    
                    logger.info(f"🎯 Calculating risk score for audit {audit_id}...")
                    risk_score = auditor.calculate_risk_score(audit_id)
                    
                    logger.info(f"🔍 Detecting patterns for audit {audit_id}...")
                    patterns = auditor.detect_patterns(audit_id)
                    
                    sebi_insights = {
                        'filing_accuracy': filing_accuracy,
                        'timeliness_sla': timeliness_sla,
                        'risk_score': risk_score,
                        'patterns': patterns
                    }
                    logger.info(f"✅ SEBI AI Auditor checks completed - insights ready for report generation")
        except Exception as e:
            logger.warning(f"⚠️ SEBI insights not available (non-critical): {str(e)}")
            # Continue with report generation even if SEBI checks fail
        
        # Generate comprehensive report data
        report_data = {
            'report_metadata': {
                'generated_at': datetime.now().isoformat(),
                'generated_by_user_id': user_id,
                'report_type': 'AI Audit Comprehensive Report',
                'audit_id': audit_id
            },
            'audit_information': {
                'audit_id': audit_row[0],
                'audit_type': 'AI Audit' if (audit_row[1] == 'A' or is_actual_ai_audit) else ('Internal' if audit_row[1] == 'I' else ('External' if audit_row[1] == 'E' else ('Self-Audit' if audit_row[1] == 'S' else audit_row[1]))),
                'status': actual_audit_status,
                'assigned_date': audit_row[3].isoformat() if audit_row[3] else None,
                'due_date': audit_row[4].isoformat() if audit_row[4] else None,
                'completion_date': actual_completion_date.isoformat() if actual_completion_date else None,
                'policy_name': selected_policies_display,  # Show selected policies from multi-select
                'subpolicy_name': selected_subpolicies_display,  # Show selected sub-policies from multi-select
                'framework_name': audit_row[8],  # Framework is fixed per audit
                'framework_id': framework_id,
                'selected_policy_count': len(selected_policy_ids),
                'selected_subpolicy_count': len(selected_subpolicy_ids),
                'selected_compliance_count': len(selected_compliance_ids),
                'compliance_display': selected_compliances_display,
                'auditor_name': 'AI System (Automated)' if (audit_row[1] == 'A' or is_actual_ai_audit) else (audit_row[9] or 'Not Assigned'),
                'assignee_name': audit_row[10] or 'Not Assigned',
                'reviewer_name': audit_row[11] or 'Not Assigned',
                'is_ai_audit_detected': is_actual_ai_audit,
                'ai_data_records': ai_data_count,
                'audit_assignment_issue': '⚠️ WARNING: This audit was created as Internal but has AI processing data. It should be treated as an AI Audit.' if (audit_row[1] == 'I' and is_actual_ai_audit) else None
            },
            'processing_summary': {
                'total_documents': total_physical_files,  # Count physical files, not mapping records
                'completed_documents': completed_physical_files,
                'failed_documents': failed_physical_files,
                'pending_documents': total_physical_files - completed_physical_files - failed_physical_files,
                'completion_percentage': round((completed_physical_files / total_physical_files * 100) if total_physical_files > 0 else 0, 2),
                'total_mapping_records': total_mapping_records,  # For reference
                'completed_mapping_records': completed_mapping_records
            },
            'compliance_summary': compliance_summary,
            'requirement_compliance_summary': requirement_summary,
            'ai_processing_details': {
                'processing_method': 'OpenAI AI/ML Analysis',
                'analysis_engine': 'gpt-4o-mini',
                'compliance_checking': 'Structured Compliance Analysis',
                'text_extraction': 'Multi-format Document Processing',
                'metadata_extraction': 'AI-powered Document Analysis'
            },
            'documents_processed': processed_documents,
            'key_findings': {
                # Calculate overall compliance rate using requirement-level data for more accuracy
                # Count fully compliant as 100%, partially compliant as 50%, non-compliant as 0%
                'overall_compliance_rate': (
                    round(
                        ((requirement_summary.get('compliant', 0) * 1.0 +
                          requirement_summary.get('partially_compliant', 0) * 0.5 +
                          requirement_summary.get('non_compliant', 0) * 0.0 +
                          requirement_summary.get('unknown', 0) * 0.0) /
                         max(requirement_summary.get('compliant', 0) + 
                             requirement_summary.get('partially_compliant', 0) + 
                             requirement_summary.get('non_compliant', 0) + 
                             requirement_summary.get('unknown', 0), 1)) * 100,
                        2
                    )
                    if (requirement_summary.get('compliant', 0) + 
                        requirement_summary.get('partially_compliant', 0) + 
                        requirement_summary.get('non_compliant', 0) + 
                        requirement_summary.get('unknown', 0)) > 0
                    else round((compliance_summary['compliant'] / completed_mapping_records * 100) if completed_mapping_records > 0 else 0, 2)
                ),
                'average_confidence_score': round(sum(
                    mapping.get('confidence_score') or 0 
                    for doc in processed_documents 
                    for mapping in doc.get('mappings', [])
                    if mapping.get('ai_processing_status') == 'completed'
                ) / completed_mapping_records if completed_mapping_records > 0 else 0, 2),
                'documents_requiring_attention': compliance_summary['non_compliant'] + compliance_summary['partially_compliant'],
                'processing_success_rate': round((completed_physical_files / total_physical_files * 100) if total_physical_files > 0 else 0, 2),
                'audit_assignment_status': 'Complete' if (audit_row[1] == 'A' and audit_row[11]) or (audit_row[1] != 'A' and audit_row[9] and audit_row[10] and audit_row[11]) else 'Incomplete - Missing Reviewer Assignment'
            },
            'recommendations': _generate_audit_recommendations(compliance_summary, processed_documents, audit_row),
            'technical_details': {
                'ai_model_used': 'llama3.2:3b',
                'processing_timestamp': datetime.now().isoformat(),
                'report_format': 'JSON',
                'data_sources': ['ai_audit_data', 'audit', 'policies', 'subpolicies', 'frameworks', 'grc_users']
            },
            'compliance_names_map': compliance_names_map,  # Add compliance names mapping for report generation
            'sebi_insights': sebi_insights if sebi_enabled else None,  # SEBI AI Auditor insights
            'sebi_enabled': sebi_enabled  # Flag indicating if SEBI AI Auditor is enabled
        }

        # Compute hash for the full report payload and attach to metadata
        report_hash = _compute_ai_audit_report_hash(report_data)
        report_data['report_metadata']['hash'] = report_hash

        # Persist digital sign-off when the caller is the assigned reviewer
        try:
            from ...models import Audit
            audit_obj = Audit.objects.filter(AuditId=audit_id).select_related('Reviewer').first()
            if audit_obj and audit_obj.Reviewer_id == user_id:
                audit_obj.AIReportSignoffHash = report_hash
                audit_obj.AIReportSignedBy = audit_obj.Reviewer
                audit_obj.AIReportSignedAt = timezone.now()
                audit_obj.save(update_fields=['AIReportSignoffHash', 'AIReportSignedBy', 'AIReportSignedAt'])
                # Surface sign-off details back into the report payload for rendering
                report_data['digital_signoff'] = {
                    'hash': report_hash,
                    'signed_by_id': audit_obj.Reviewer_id,
                    'signed_by_name': getattr(audit_obj.Reviewer, 'UserName', None) or report_data['audit_information'].get('reviewer_name'),
                    'signed_at': audit_obj.AIReportSignedAt.isoformat() if audit_obj.AIReportSignedAt else None,
                }
        except Exception as e:
            logger.warning("Could not persist AI audit report sign-off for audit %s: %s", audit_id, e)
        
        # Generate formatted PDF document (from HTML report)
        logger.info(f"📄 Starting PDF document generation for audit {audit_id}")
        try:
            pdf_content = _generate_pdf_document(report_data)
            logger.info(f"✅ PDF document generation completed for audit {audit_id}")
        except Exception as pdf_gen_error:
            logger.error(f"❌ Error in PDF document generation: {pdf_gen_error}")
            import traceback
            logger.error(f"❌ Traceback: {traceback.format_exc()}")
            raise
        
        # Create response with PDF document
        response = HttpResponse(
            pdf_content,
            content_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="AI_Audit_Report_{audit_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf"'
            }
        )
        
        logger.info(f"✅ Generated comprehensive AI audit report (PDF) for audit {audit_id}")
        return response
        
    except Exception as e:
        logger.error(f"❌ Error generating AI audit report: {e}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=500)


def _format_datetime(dt_value):
    """Safely format datetime values that could be datetime objects or strings"""
    if not dt_value:
        return 'N/A'
    
    if isinstance(dt_value, datetime):
        return dt_value.strftime('%Y-%m-%d %H:%M:%S')
    elif isinstance(dt_value, str):
        # If it's already a string, try to parse and reformat
        try:
            # Handle ISO format strings
            if 'T' in dt_value:
                parsed_dt = datetime.fromisoformat(dt_value.replace('Z', '+00:00'))
                return parsed_dt.strftime('%Y-%m-%d %H:%M:%S')
            else:
                # If it's already in a good format, return as is
                return dt_value
        except (ValueError, TypeError):
            return str(dt_value)
    else:
        return str(dt_value)


def _get_report_logo_path():
    """Resolve path to RiskaVaire.png for report header. Returns None if not found."""
    candidates = [
        os.path.join(settings.BASE_DIR, '..', 'grc_frontend', 'src', 'assets', 'RiskaVaire.png'),
        os.path.join(settings.BASE_DIR, 'grc_frontend', 'src', 'assets', 'RiskaVaire.png'),
    ]
    for path in candidates:
        abs_path = os.path.abspath(path)
        if os.path.isfile(abs_path):
            return abs_path
    return None


def _get_report_logo_data_url():
    """Return data URL for report logo (for HTML embedding) or empty string if not found."""
    path = _get_report_logo_path()
    if not path:
        return ''
    try:
        with open(path, 'rb') as f:
            data = base64.b64encode(f.read()).decode('ascii')
        ext = os.path.splitext(path)[1].lower()
        mime = 'image/png' if ext == '.png' else ('image/jpeg' if ext in ('.jpg', '.jpeg') else 'image/png')
        return f'data:{mime};base64,{data}'
    except Exception:
        return ''


def _escape_pdf_text(text):
    """Escape special XML chars for reportlab Paragraph."""
    if text is None:
        return ''
    s = str(text)
    s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;').replace('"', '&quot;').replace("'", '&#39;')
    return s


def _generate_pdf_document(report_data):
    """Generate a PDF document that matches the Word layout - plain document style, no boxes."""
    import io
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError as e:
        logger.error(f"reportlab not installed: {e}")
        raise ValueError("PDF generation requires reportlab. Install with: pip install reportlab")

    metadata = report_data['report_metadata']
    audit_info = report_data['audit_information']
    processing = report_data['processing_summary']
    compliance = report_data['compliance_summary']
    requirement_compliance = report_data.get('requirement_compliance_summary')
    documents = report_data['documents_processed']
    findings = report_data['key_findings']
    recommendations = report_data['recommendations']
    compliance_names_map = report_data.get('compliance_names_map', {})
    digital_signoff = report_data.get('digital_signoff') or {}

    buffer = io.BytesIO()
    logo_path = _get_report_logo_path()
    header_reserved_height = 0.9 * inch  # space reserved for header logo on every page
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=inch + (header_reserved_height if logo_path else 0),
        bottomMargin=inch,
        leftMargin=inch,
        rightMargin=inch
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(name='H1', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=14, spaceAfter=6)
    h2 = ParagraphStyle(name='H2', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, spaceAfter=4)
    h3 = ParagraphStyle(name='H3', parent=styles['Heading3'], fontName='Helvetica-Bold', fontSize=11, spaceAfter=3)
    normal = ParagraphStyle(name='Normal', parent=styles['Normal'], fontSize=10, spaceAfter=4)
    bullet = ParagraphStyle(name='Bullet', parent=styles['Normal'], fontSize=10, leftIndent=20, spaceAfter=2)
    bold_label = ParagraphStyle(name='BoldLabel', fontName='Helvetica-Bold', fontSize=10, spaceAfter=2)
    content_next = ParagraphStyle(name='ContentNext', fontSize=10, leftIndent=15, spaceAfter=6)

    story = []

    def _draw_logo_header(canvas, _doc):
        """Draw the logo in the header area on every page."""
        if not logo_path:
            return
        try:
            if not os.path.isfile(logo_path):
                return
            from reportlab.lib.utils import ImageReader

            img = ImageReader(logo_path)
            iw, ih = img.getSize()
            if not iw or not ih:
                return

            target_w = 1.5 * inch
            scale = target_w / float(iw)
            target_h = float(ih) * scale

            page_w, page_h = A4
            x = (page_w - target_w) / 2.0
            # Place logo near top, inside reserved header area
            y_top = page_h - 0.55 * inch
            y = y_top - target_h
            canvas.drawImage(img, x, y, width=target_w, height=target_h, mask='auto')
        except Exception as e:
            logger.warning(f"Could not draw logo on PDF header: {e}")

    # Title
    story.append(Paragraph(_escape_pdf_text('AI Audit Comprehensive Report'), ParagraphStyle(name='Title', alignment=TA_CENTER, fontName='Helvetica-Bold', fontSize=16, spaceAfter=6)))
    story.append(Paragraph(_escape_pdf_text(f'Generated on: {_format_datetime(metadata["generated_at"]).replace(" ", " at ")}'), ParagraphStyle(name='Sub', alignment=TA_CENTER, fontSize=9, spaceAfter=2)))
    story.append(Paragraph(_escape_pdf_text(f'Audit ID: {audit_info["audit_id"]}'), ParagraphStyle(name='Sub2', alignment=TA_CENTER, fontSize=9, spaceAfter=12)))
    story.append(Spacer(1, 12))

    if audit_info.get('audit_assignment_issue'):
        story.append(Paragraph(_escape_pdf_text(f'WARNING: {audit_info["audit_assignment_issue"]}'), ParagraphStyle(name='Warn', fontName='Helvetica-Bold', fontSize=10, textColor='red', spaceAfter=12)))

    # Audit Information
    story.append(Paragraph(_escape_pdf_text('Audit Information'), h1))
    story.append(Paragraph(_escape_pdf_text('Audit Details'), h2))
    story.append(Paragraph(_escape_pdf_text(f'• Type: {audit_info["audit_type"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Status: {audit_info["status"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Framework: {audit_info["framework_name"]}'), bullet))
    policy_count = audit_info.get('selected_policy_count', 0)
    story.append(Paragraph(_escape_pdf_text(f'• Policy: {audit_info["policy_name"]}' if policy_count <= 1 else f'• Selected Policies ({policy_count}): {audit_info["policy_name"]}'), bullet))
    subpolicy_count = audit_info.get('selected_subpolicy_count', 0)
    story.append(Paragraph(_escape_pdf_text(f'• Sub-policy: {audit_info.get("subpolicy_name", "")}' if subpolicy_count <= 1 else f'• Selected Sub-policies ({subpolicy_count}): {audit_info["subpolicy_name"]}'), bullet))
    comp_count = audit_info.get("selected_compliance_count", 0)
    story.append(Paragraph(_escape_pdf_text(f'• Compliances ({comp_count}): {audit_info.get("compliance_display", "")}' if comp_count > 0 else '• Compliances: None selected'), bullet))

    story.append(Paragraph(_escape_pdf_text('Assignment'), h2))
    story.append(Paragraph(_escape_pdf_text(f'• Auditor: {audit_info["auditor_name"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Assignee: {audit_info["assignee_name"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Reviewer: {audit_info["reviewer_name"]}'), bullet))
    story.append(Spacer(1, 12))

    # Processing Summary
    story.append(Paragraph(_escape_pdf_text('Processing Summary'), h1))
    story.append(Paragraph(_escape_pdf_text(f'• Total Documents: {processing["total_documents"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Completed: {processing["completed_documents"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Pending: {processing["pending_documents"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Completion Rate: {processing["completion_percentage"]}%'), bullet))

    if requirement_compliance:
        story.append(Paragraph(_escape_pdf_text('Requirement-Level Compliance Summary'), h2))
        story.append(Paragraph(_escape_pdf_text(f'• Compliant: {requirement_compliance.get("compliant", 0)}'), bullet))
        story.append(Paragraph(_escape_pdf_text(f'• Partially Compliant: {requirement_compliance.get("partially_compliant", 0)}'), bullet))
        story.append(Paragraph(_escape_pdf_text(f'• Non-Compliant: {requirement_compliance.get("non_compliant", 0)}'), bullet))
        story.append(Paragraph(_escape_pdf_text(f'• Unknown: {requirement_compliance.get("unknown", 0)}'), bullet))
    story.append(Spacer(1, 12))

    # Documents Processed
    story.append(Paragraph(_escape_pdf_text('Documents Processed (Completed Only)'), h1))
    completed_documents = [d for d in documents if d.get('ai_processing_status') == 'completed']
    if not completed_documents:
        story.append(Paragraph(_escape_pdf_text('No documents have been completed yet.'), normal))
    else:
        story.append(Paragraph(_escape_pdf_text(f'Showing {len(completed_documents)} of {processing["total_documents"]} total physical file(s).'), normal))
        story.append(Spacer(1, 8))

        for doc_item in completed_documents:
            story.append(Paragraph(_escape_pdf_text(doc_item['document_name']), h2))
            story.append(Paragraph(_escape_pdf_text(f'• Type: {doc_item["document_type"]} | Size: {doc_item["file_size"]:,} bytes'), bullet))
            story.append(Paragraph(_escape_pdf_text(f'• Processing: {doc_item["ai_processing_status"]}'), bullet))
            story.append(Paragraph(_escape_pdf_text(f'• Uploaded: {_format_datetime(doc_item["uploaded_date"])}'), bullet))
            story.append(Paragraph(_escape_pdf_text(f'• Mappings: {doc_item.get("mapping_count", 0)} mapping(s)'), bullet))
            story.append(Spacer(1, 6))

            mappings = doc_item.get('mappings', [])
            completed_mappings = [m for m in mappings if m.get('ai_processing_status') == 'completed']
            all_analyses_map = {}
            for mapping in completed_mappings:
                compliance_analyses = mapping.get('compliance_analyses')
                if compliance_analyses:
                    analyses_list = compliance_analyses.get('compliance_analyses', []) if isinstance(compliance_analyses, dict) else (compliance_analyses if isinstance(compliance_analyses, list) else [])
                    for analysis in analyses_list:
                        if isinstance(analysis, dict) and analysis.get('compliance_id') and analysis.get('compliance_id') not in all_analyses_map:
                            all_analyses_map[analysis['compliance_id']] = analysis

            unique_analyses = list(all_analyses_map.values())[:150]
            if unique_analyses:
                story.append(Paragraph(_escape_pdf_text('Compliance Analysis Results'), h3))
                story.append(Paragraph(_escape_pdf_text(f'Showing {len(unique_analyses)} unique compliance analyses.'), normal))
                story.append(Spacer(1, 4))

                for i, analysis in enumerate(unique_analyses, 1):
                    if not isinstance(analysis, dict):
                        continue
                    compliance_id = analysis.get('compliance_id')
                    compliance_name = None
                    if compliance_id:
                        try:
                            compliance_name = compliance_names_map.get(int(compliance_id))
                        except (ValueError, TypeError):
                            compliance_name = compliance_names_map.get(compliance_id)
                    if not compliance_name:
                        compliance_name = analysis.get('requirement_title', f'Compliance Requirement {i}')
                    story.append(Paragraph(_escape_pdf_text(compliance_name), h3))

                    req_desc = analysis.get('requirement_description') or analysis.get('requirement_title', '')
                    if req_desc and req_desc != compliance_name and len(req_desc) > 20:
                        story.append(Paragraph(_escape_pdf_text('What This Requirement Means:'), bold_label))
                        story.append(Paragraph(_escape_pdf_text(req_desc), content_next))

                    score = analysis.get('compliance_score') or analysis.get('relevance')
                    if score is not None:
                        story.append(Paragraph(_escape_pdf_text('Compliance Score:'), bold_label))
                        story.append(Paragraph(_escape_pdf_text(str(round(float(score), 2))), content_next))
                    status_text = (analysis.get('status', '') or '').replace('_', ' ').title() or 'Unknown'
                    story.append(Paragraph(_escape_pdf_text('Status:'), bold_label))
                    story.append(Paragraph(_escape_pdf_text(status_text), content_next))

                    if analysis.get('evidence'):
                        ev = analysis['evidence']
                        ev_text = ', '.join(str(x) for x in (ev if isinstance(ev, list) else [ev])) if ev else ''
                        if ev_text:
                            story.append(Paragraph(_escape_pdf_text('Evidence Found:'), bold_label))
                            story.append(Paragraph(_escape_pdf_text(ev_text), content_next))
                    if analysis.get('missing'):
                        miss = analysis['missing']
                        miss_items = miss if isinstance(miss, list) else [miss] if miss else []
                        if miss_items:
                            story.append(Paragraph(_escape_pdf_text('Gaps:'), bold_label))
                            _gap_headings = ('NO EVIDENCE FOUND:', 'REQUIREMENT NEEDS:', 'DOCUMENT IS ABOUT:', 'DATA REQUIRED FOR THIS AUDIT:', 'EXPECTED DOCUMENT TYPE:', 'WHAT IS NEEDED:', 'POLICY CONTEXT:')
                            full_text = ' '.join(str(x).strip() for x in miss_items if x)
                            parts = re.split(r'(' + '|'.join(re.escape(h) for h in _gap_headings) + r')', full_text, flags=re.IGNORECASE)
                            for i in range(1, len(parts), 2):
                                sub_head = parts[i].strip() if i < len(parts) else ''
                                sub_content = (parts[i + 1].strip() if i + 1 < len(parts) else '').strip()
                                if sub_head:
                                    story.append(Paragraph(_escape_pdf_text(sub_head), bold_label))
                                if sub_content:
                                    story.append(Paragraph(_escape_pdf_text(sub_content), content_next))
                            if len(parts) == 1 and full_text:
                                story.append(Paragraph(_escape_pdf_text(full_text), content_next))

                    comp_score = float(analysis.get('compliance_score') or analysis.get('relevance', 0))
                    if comp_score < 0.4:
                        rec = "CRITICAL: Immediate action required"
                    elif comp_score < 0.6:
                        rec = "HIGH PRIORITY: Significant improvements needed"
                    elif comp_score < 0.8:
                        rec = "MEDIUM PRIORITY: Minor improvements recommended"
                    else:
                        rec = "LOW PRIORITY: Document adequately addresses this requirement"
                    story.append(Paragraph(_escape_pdf_text('Recommendations:'), bold_label))
                    story.append(Paragraph(_escape_pdf_text(rec), content_next))
                    story.append(Spacer(1, 6))
            story.append(Spacer(1, 8))

    # Key Findings
    story.append(Paragraph(_escape_pdf_text('Key Findings'), h1))
    story.append(Paragraph(_escape_pdf_text(f'• Overall Compliance Rate: {findings["overall_compliance_rate"]}%'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Average Confidence Score: {findings["average_confidence_score"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Documents Requiring Attention: {findings["documents_requiring_attention"]}'), bullet))
    story.append(Paragraph(_escape_pdf_text(f'• Processing Success Rate: {findings["processing_success_rate"]}%'), bullet))
    story.append(Spacer(1, 12))

    # Recommendations
    story.append(Paragraph(_escape_pdf_text('Recommendations'), h1))
    for rec in recommendations:
        story.append(Paragraph(_escape_pdf_text(f'• {rec}'), bullet))
    story.append(Spacer(1, 12))

    # SEBI section
    sebi_insights = report_data.get('sebi_insights')
    sebi_enabled = report_data.get('sebi_enabled', False)
    if sebi_enabled and sebi_insights:
        story.append(Paragraph(_escape_pdf_text('SEBI AI Auditor Insights'), h1))
        filing = sebi_insights.get('filing_accuracy', {})
        if filing:
            story.append(Paragraph(_escape_pdf_text('Filing Accuracy Analysis'), h2))
            story.append(Paragraph(_escape_pdf_text(f'• Overall Accuracy: {filing.get("accuracy_score", "N/A")}%'), bullet))
            story.append(Paragraph(_escape_pdf_text(f'• Total Filings: {filing.get("total_filings", 0)}'), bullet))
        timeliness = sebi_insights.get('timeliness_sla', {})
        if timeliness:
            story.append(Paragraph(_escape_pdf_text('Timeliness & SLA Compliance'), h2))
            story.append(Paragraph(_escape_pdf_text(f'• SLA Compliance Rate: {timeliness.get("sla_compliance_rate", "N/A")}%'), bullet))
        risk = sebi_insights.get('risk_score', {})
        if risk:
            story.append(Paragraph(_escape_pdf_text('Risk Assessment'), h2))
            story.append(Paragraph(_escape_pdf_text(f'• Risk Score: {risk.get("risk_score", risk.get("overall_risk_score", "N/A"))}/100'), bullet))
        story.append(Spacer(1, 12))

    # Digital sign-off section (Reviewer acting as signatory)
    if metadata.get('hash') or digital_signoff:
        story.append(Paragraph(_escape_pdf_text('Digital Sign-Off'), h1))
        signer_name = digital_signoff.get('signed_by_name') or audit_info.get('reviewer_name') or 'Reviewer'
        signed_at = digital_signoff.get('signed_at') or metadata.get('generated_at')
        story.append(Paragraph(_escape_pdf_text(f'• Signed by: {signer_name} (Reviewer)'), bullet))
        if signed_at:
            story.append(Paragraph(_escape_pdf_text(f'• Signed at: {_format_datetime(signed_at)}'), bullet))
        story.append(Paragraph(_escape_pdf_text(f'• Report Content Hash (SHA-256): {metadata.get("hash", digital_signoff.get("hash", "N/A"))}'), bullet))
        story.append(Spacer(1, 12))

    # Footer
    story.append(Paragraph(_escape_pdf_text('This report was generated automatically by the AI Audit System'), ParagraphStyle(name='Foot', alignment=TA_CENTER, fontSize=8, spaceAfter=2)))
    footer_line = f'Report ID: {metadata["audit_id"]} | Generated by User: {metadata["generated_by_user_id"]}'
    if metadata.get('hash'):
        footer_line += f' | Hash: {metadata["hash"][:16]}...'
    story.append(Paragraph(_escape_pdf_text(footer_line), ParagraphStyle(name='Foot2', alignment=TA_CENTER, fontSize=8)))

    doc.build(story, onFirstPage=_draw_logo_header, onLaterPages=_draw_logo_header)
    return buffer.getvalue()


def _generate_html_report(report_data):
    """Generate a formatted HTML report from the audit data"""
    
    # Extract key data for easier access
    metadata = report_data['report_metadata']
    audit_info = report_data['audit_information']
    processing = report_data['processing_summary']
    compliance = report_data['compliance_summary']
    requirement_compliance = report_data.get('requirement_compliance_summary', {'compliant': 0, 'partially_compliant': 0, 'non_compliant': 0, 'unknown': 0})
    documents = report_data['documents_processed']
    findings = report_data['key_findings']
    recommendations = report_data['recommendations']
    technical = report_data['technical_details']
    logo_data_url = _get_report_logo_data_url()
    
    # Generate HTML content
    html_content = f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>AI Audit Report - Audit {audit_info['audit_id']}</title>
        <style>
            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                line-height: 1.6;
                margin: 0;
                padding: 20px;
                background-color: #f5f5f5;
            }}
            .container {{
                max-width: 1200px;
                margin: 0 auto;
                background: white;
                padding: 30px;
                border-radius: 10px;
                box-shadow: 0 0 20px rgba(0,0,0,0.1);
            }}
            .header {{
                text-align: center;
                border-bottom: 3px solid #2c3e50;
                padding-bottom: 20px;
                margin-bottom: 30px;
            }}
            .header h1 {{
                color: #2c3e50;
                margin: 0;
                font-size: 2.5em;
            }}
            .header p {{
                color: #7f8c8d;
                margin: 10px 0 0 0;
                font-size: 1.1em;
            }}
            .section {{
                margin-bottom: 30px;
                padding: 20px;
                border-left: 4px solid #3498db;
                background-color: #f8f9fa;
            }}
            .section h2 {{
                color: #2c3e50;
                margin-top: 0;
                font-size: 1.8em;
                border-bottom: 2px solid #ecf0f1;
                padding-bottom: 10px;
            }}
            .section h3 {{
                color: #34495e;
                margin-top: 25px;
                font-size: 1.4em;
            }}
            .info-grid {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
                gap: 20px;
                margin: 20px 0;
            }}
            .info-card {{
                background: white;
                padding: 15px;
                border-radius: 8px;
                border: 1px solid #e1e8ed;
            }}
            .info-card h4 {{
                color: #2c3e50;
                margin: 0 0 10px 0;
                font-size: 1.1em;
            }}
            .info-card p {{
                margin: 5px 0;
                color: #555;
            }}
            .status-badge {{
                display: inline-block;
                padding: 5px 15px;
                border-radius: 20px;
                font-weight: bold;
                font-size: 0.9em;
            }}
            .status-compliant {{ background-color: #d4edda; color: #155724; }}
            .status-non-compliant {{ background-color: #f8d7da; color: #721c24; }}
            .status-pending {{ background-color: #fff3cd; color: #856404; }}
            .status-ai {{ background-color: #d1ecf1; color: #0c5460; }}
            .document-item {{
                background: white;
                padding: 15px;
                margin: 10px 0;
                border-radius: 8px;
                border-left: 4px solid #3498db;
            }}
            .document-header {{
                display: flex;
                justify-content: space-between;
                align-items: center;
                margin-bottom: 10px;
            }}
            .document-name {{
                font-weight: bold;
                color: #2c3e50;
            }}
            .confidence-score {{
                font-weight: bold;
                padding: 5px 10px;
                border-radius: 15px;
                background-color: #ecf0f1;
            }}
            .recommendations {{
                background: #fff3cd;
                border-left: 4px solid #ffc107;
                padding: 15px;
                margin: 10px 0;
            }}
            .recommendations ul {{
                margin: 10px 0;
                padding-left: 20px;
            }}
            .recommendations li {{
                margin: 8px 0;
                color: #856404;
            }}
            .warning {{
                background: #f8d7da;
                border-left: 4px solid #dc3545;
                padding: 15px;
                margin: 10px 0;
                color: #721c24;
            }}
            .summary-stats {{
                display: grid;
                grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
                gap: 15px;
                margin: 20px 0;
            }}
            .stat-card {{
                background: white;
                padding: 20px;
                text-align: center;
                border-radius: 8px;
                border: 1px solid #e1e8ed;
            }}
            .stat-number {{
                font-size: 2em;
                font-weight: bold;
                color: #2c3e50;
            }}
            .stat-label {{
                color: #7f8c8d;
                margin-top: 5px;
            }}
            .footer {{
                text-align: center;
                margin-top: 40px;
                padding-top: 20px;
                border-top: 2px solid #ecf0f1;
                color: #7f8c8d;
            }}
            .report-logo {{
                text-align: center;
                margin-bottom: 20px;
            }}
            .report-logo img {{
                max-width: 200px;
                max-height: 60px;
                object-fit: contain;
            }}
            @media print {{
                .report-logo {{
                    position: fixed;
                    top: 0;
                    left: 0;
                    right: 0;
                    text-align: center;
                    z-index: 9999;
                    margin-bottom: 0;
                    padding: 8px 0;
                }}
                body {{ padding-top: 50px; }}
            }}
        </style>
    </head>
    <body>
        {f'<div class="report-logo"><img src="{logo_data_url}" alt="Logo" /></div>' if logo_data_url else ''}
        <div class="container">
            <div class="header">
                <h1>AI Audit Comprehensive Report</h1>
                <p>Generated on {_format_datetime(metadata['generated_at']).replace(' ', ' at ')} | Audit ID: {audit_info['audit_id']}</p>
            </div>
            
            {f'<div class="warning"><strong>Warning:</strong> {audit_info["audit_assignment_issue"]}</div>' if audit_info.get('audit_assignment_issue') else ''}
            
            <div class="section">
                <h2>Audit Information</h2>
                <div class="info-grid">
                    <div class="info-card">
                        <h4>Audit Details</h4>
                        <p><strong>Type:</strong> <span class="status-badge status-ai">{audit_info['audit_type']}</span></p>
                        <p><strong>Status:</strong> {audit_info['status']}</p>
                        <p><strong>Framework:</strong> {audit_info['framework_name']}</p>
                        <p><strong>Policy:</strong> {audit_info['policy_name']}</p>
                        <p><strong>Sub-policy:</strong> {audit_info['subpolicy_name']}</p>
                        <p><strong>Compliances:</strong> {audit_info['compliance_display']}</p>
                    </div>
                    <div class="info-card">
                        <h4>Assignment</h4>
                        <p><strong>Auditor:</strong> {audit_info['auditor_name']}</p>
                        <p><strong>Assignee:</strong> {audit_info['assignee_name']}</p>
                        <p><strong>Reviewer:</strong> {audit_info['reviewer_name']}</p>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h2>Processing Summary</h2>
                <div class="summary-stats">
                    <div class="stat-card">
                        <div class="stat-number">{processing['total_documents']}</div>
                        <div class="stat-label">Total Documents</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-number">{processing['completed_documents']}</div>
                        <div class="stat-label">Completed</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-number">{processing['pending_documents']}</div>
                        <div class="stat-label">Pending</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-number">{processing['completion_percentage']}%</div>
                        <div class="stat-label">Completion Rate</div>
                    </div>
                </div>
            </div>

            <div class="section">
                <h2>Requirement-Level Compliance Summary</h2>
                <div class="summary-stats">
                    <div class="stat-card">
                        <div class="stat-number" style="color: #28a745;">{requirement_compliance['compliant']}</div>
                        <div class="stat-label">Compliant Requirements</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-number" style="color: #ffc107;">{requirement_compliance['partially_compliant']}</div>
                        <div class="stat-label">Partially Compliant Requirements</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-number" style="color: #dc3545;">{requirement_compliance['non_compliant']}</div>
                        <div class="stat-label">Non-Compliant Requirements</div>
                    </div>
                    <div class="stat-card">
                        <div class="stat-number" style="color: #6c757d;">{requirement_compliance['unknown']}</div>
                        <div class="stat-label">Unknown Requirements</div>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h2>Documents Processed</h2>
                {''.join([f'''
                <div class="document-item">
                    <div class="document-header">
                        <div class="document-name">{doc['document_name']}</div>
                    </div>
                    <p><strong>Type:</strong> {doc['document_type']} | <strong>Size:</strong> {doc['file_size']:,} bytes</p>
                    <p><strong>Processing:</strong> {doc['ai_processing_status']}</p>
                    <p><strong>Uploaded:</strong> {_format_datetime(doc['uploaded_date'])}</p>
                </div>
                ''' for doc in documents])}
            </div>
            
            <div class="section">
                <h2>Key Findings</h2>
                <div class="info-grid">
                    <div class="info-card">
                        <h4>Compliance Rate</h4>
                        <p><strong>{findings['overall_compliance_rate']}%</strong> overall compliance</p>
                    </div>
                    <div class="info-card">
                        <h4>Confidence Score</h4>
                        <p><strong>{findings['average_confidence_score']}</strong> average confidence</p>
                    </div>
                    <div class="info-card">
                        <h4>Attention Required</h4>
                        <p><strong>{findings['documents_requiring_attention']}</strong> documents need attention</p>
                    </div>
                    <div class="info-card">
                        <h4>Processing Success</h4>
                        <p><strong>{findings['processing_success_rate']}%</strong> success rate</p>
                    </div>
                </div>
            </div>
            
            <div class="section">
                <h2>Recommendations</h2>
                <div class="recommendations">
                    <ul>
                        {''.join([f'<li>{rec}</li>' for rec in recommendations])}
                    </ul>
                </div>
            </div>
            
            
            <div class="footer">
                <p>This report was generated automatically by the AI Audit System</p>
                <p>Report ID: {metadata['audit_id']} | Generated by User: {metadata['generated_by_user_id']}</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content


def _make_heading_bold(heading):
    """Helper function to make a heading bold"""
    for run in heading.runs:
        run.bold = True
    return heading


def _add_bold_label_paragraph(doc, text):
    """Helper function to add a paragraph with bold label"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    return para


def _generate_word_document(report_data):
    """Generate a proper Word document (.docx) from the audit data"""
    
    # Extract key data for easier access
    metadata = report_data['report_metadata']
    audit_info = report_data['audit_information']
    processing = report_data['processing_summary']
    compliance = report_data['compliance_summary']
    requirement_compliance = report_data.get('requirement_compliance_summary')
    documents = report_data['documents_processed']
    findings = report_data['key_findings']
    recommendations = report_data['recommendations']
    technical = report_data['technical_details']
    compliance_names_map = report_data.get('compliance_names_map', {})  # Get compliance names mapping
    
    # Create a new Word document
    doc = Document()
    
    # Set document margins and add logo to header (top center on every page)
    logo_path = _get_report_logo_path()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        if logo_path:
            try:
                header = section.header
                header_para = header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = header_para.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
            except Exception as e:
                logger.warning(f"Could not add logo to Word report header: {e}")
    
    # Add title
    title = doc.add_heading('AI Audit Comprehensive Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _make_heading_bold(title)
    
    # Add subtitle
    subtitle = doc.add_paragraph(f'Generated on: {_format_datetime(metadata["generated_at"]).replace(" ", " at ")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Inches(0.15)
    
    audit_id_para = doc.add_paragraph(f'Audit ID: {audit_info["audit_id"]}')
    audit_id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audit_id_run = audit_id_para.runs[0]
    audit_id_run.font.size = Inches(0.15)
    
    doc.add_paragraph()  # Add space
    
    # Add warning if exists
    if audit_info.get('audit_assignment_issue'):
        warning = doc.add_paragraph(f'WARNING: {audit_info["audit_assignment_issue"]}')
        warning_run = warning.runs[0]
        warning_run.font.color.rgb = None  # Red color
        warning_run.bold = True
        doc.add_paragraph()
    
    # Audit Information Section
    audit_info_heading = doc.add_heading('Audit Information', level=1)
    _make_heading_bold(audit_info_heading)
    
    # Audit Details
    audit_details_heading = doc.add_heading('Audit Details', level=2)
    _make_heading_bold(audit_details_heading)
    doc.add_paragraph(f'• Type: {audit_info["audit_type"]}')
    doc.add_paragraph(f'• Status: {audit_info["status"]}')
    doc.add_paragraph(f'• Framework: {audit_info["framework_name"]}')
    # Show selected policies (multi-select approach)
    policy_count = audit_info.get('selected_policy_count', 0)
    if policy_count > 1:
        doc.add_paragraph(f'• Selected Policies ({policy_count}): {audit_info["policy_name"]}')
    else:
        doc.add_paragraph(f'• Policy: {audit_info["policy_name"]}')
    # Show selected sub-policies (multi-select approach)
    subpolicy_count = audit_info.get('selected_subpolicy_count', 0)
    if subpolicy_count > 1:
        doc.add_paragraph(f'• Selected Sub-policies ({subpolicy_count}): {audit_info["subpolicy_name"]}')
    elif audit_info.get("subpolicy_name"):
        doc.add_paragraph(f'• Sub-policy: {audit_info["subpolicy_name"]}')

    # Show selected compliances and count
    comp_count = audit_info.get("selected_compliance_count", 0)
    if comp_count > 0:
        doc.add_paragraph(f'• Compliances ({comp_count}): {audit_info.get("compliance_display", "")}')
    else:
        doc.add_paragraph('• Compliances: None selected')
    
    # Assignment
    assignment_heading = doc.add_heading('Assignment', level=2)
    _make_heading_bold(assignment_heading)
    doc.add_paragraph(f'• Auditor: {audit_info["auditor_name"]}')
    doc.add_paragraph(f'• Assignee: {audit_info["assignee_name"]}')
    doc.add_paragraph(f'• Reviewer: {audit_info["reviewer_name"]}')
    
    # Processing Summary Section
    processing_heading = doc.add_heading('Processing Summary', level=1)
    _make_heading_bold(processing_heading)
    doc.add_paragraph(f'• Total Documents: {processing["total_documents"]}')
    doc.add_paragraph(f'• Completed: {processing["completed_documents"]}')
    doc.add_paragraph(f'• Pending: {processing["pending_documents"]}')
    doc.add_paragraph(f'• Completion Rate: {processing["completion_percentage"]}%')
    
    # Requirement-Level Compliance Summary (optional)
    if requirement_compliance:
        req_heading = doc.add_heading('Requirement-Level Compliance Summary', level=2)
        _make_heading_bold(req_heading)
        doc.add_paragraph(f'• Compliant Requirements: {requirement_compliance.get("compliant", 0)}')
        doc.add_paragraph(f'• Partially Compliant Requirements: {requirement_compliance.get("partially_compliant", 0)}')
        doc.add_paragraph(f'• Non-Compliant Requirements: {requirement_compliance.get("non_compliant", 0)}')
        doc.add_paragraph(f'• Unknown Status Requirements: {requirement_compliance.get("unknown", 0)}')
    
    # Documents Processed Section (Only Completed)
    documents_heading = doc.add_heading('Documents Processed (Completed Only)', level=1)
    _make_heading_bold(documents_heading)
    
    # Filter to only show completed physical files (files with at least one completed mapping)
    completed_documents = [d for d in documents if d.get('ai_processing_status') == 'completed']
    
    if not completed_documents:
        doc.add_paragraph('No documents have been completed yet. Please run compliance checks to see results here.')
    else:
        # Use the total_documents from processing_summary which counts unique physical files
        total_files = processing['total_documents']
        doc.add_paragraph(f'Showing {len(completed_documents)} of {total_files} total physical file(s) (completed only).')
        doc.add_paragraph()  # Add space
    
    for doc_item in completed_documents:
        doc_name_heading = doc.add_heading(doc_item['document_name'], level=2)
        _make_heading_bold(doc_name_heading)
        doc.add_paragraph(f'• Type: {doc_item["document_type"]} | Size: {doc_item["file_size"]:,} bytes')
        doc.add_paragraph(f'• Processing: {doc_item["ai_processing_status"]}')
        doc.add_paragraph(f'• Uploaded: {_format_datetime(doc_item["uploaded_date"])}')
        doc.add_paragraph(f'• Mappings: {doc_item.get("mapping_count", 0)} mapping(s) for this file')
        doc.add_paragraph()  # Add space
        
        # Process compliance analyses for this physical file
        # Deduplicate analyses across all mappings - each compliance should appear only once per file
        mappings = doc_item.get('mappings', [])
        completed_mappings = [m for m in mappings if m.get('ai_processing_status') == 'completed']
        total_completed_mappings = len(completed_mappings)
        
        # Collect all unique compliance analyses across all mappings (deduplicate by compliance_id)
        all_analyses_map = {}  # Key: compliance_id, Value: analysis dict
        for mapping in completed_mappings:
            compliance_analyses = mapping.get('compliance_analyses')
            if compliance_analyses:
                # Handle nested structure
                if isinstance(compliance_analyses, dict) and 'compliance_analyses' in compliance_analyses:
                    analyses_list = compliance_analyses['compliance_analyses']
                elif isinstance(compliance_analyses, list):
                    analyses_list = compliance_analyses
                else:
                    analyses_list = []
                
                # Deduplicate by compliance_id - keep first occurrence
                for analysis in analyses_list:
                    if isinstance(analysis, dict):
                        compliance_id = analysis.get('compliance_id')
                        if compliance_id and compliance_id not in all_analyses_map:
                            all_analyses_map[compliance_id] = analysis
        
        # Convert to list and limit to 150 analyses
        unique_analyses = list(all_analyses_map.values())[:150]
        
        if unique_analyses:
            doc.add_paragraph()
            detailed_analysis_heading = doc.add_heading('Compliance Analysis Results', level=3)
            _make_heading_bold(detailed_analysis_heading)
            doc.add_paragraph(f'Showing {len(unique_analyses)} unique compliance analyses (deduplicated across {total_completed_mappings} mapping(s))')
            doc.add_paragraph()
            
            for i, analysis in enumerate(unique_analyses, 1):
                if isinstance(analysis, dict):
                    # Get the actual compliance name from compliance_names_map if available
                    compliance_id = analysis.get('compliance_id')
                    compliance_name = None
                    if compliance_id:
                        # Try both int and string versions of compliance_id
                        try:
                            compliance_id_int = int(compliance_id) if compliance_id else None
                            if compliance_id_int and compliance_id_int in compliance_names_map:
                                compliance_name = compliance_names_map[compliance_id_int]
                            elif compliance_id in compliance_names_map:
                                compliance_name = compliance_names_map[compliance_id]
                        except (ValueError, TypeError):
                            pass
                    
                    # Use compliance name as heading, or fallback to requirement_title
                    if not compliance_name:
                        compliance_name = analysis.get('requirement_title', f'Compliance Requirement {i}')
                    
                    # Compliance Requirement heading
                    compliance_heading = doc.add_heading(compliance_name, level=3)
                    _make_heading_bold(compliance_heading)
                    
                    # Show full requirement description (what needs to be complied with)
                    requirement_title = analysis.get('requirement_title', '')
                    requirement_description = analysis.get('requirement_description', '')
                    
                    # Use description if available, otherwise use title
                    full_requirement_text = requirement_description if requirement_description else requirement_title
                    
                    # Only show if it's different from the heading and not empty
                    if full_requirement_text and full_requirement_text != compliance_name and len(full_requirement_text) > 20:
                        req_para = doc.add_paragraph()
                        req_run = req_para.add_run('What This Requirement Means: ')
                        req_run.bold = True
                        req_para.add_run(full_requirement_text)
                    
                    # Compliance Score (bullet point)
                    compliance_score_value = analysis.get('compliance_score') or analysis.get('relevance')
                    if compliance_score_value is not None:
                        score_para = doc.add_paragraph(style='List Bullet')
                        score_run = score_para.add_run('Compliance Score: ')
                        score_run.bold = True
                        score_para.add_run(str(round(float(compliance_score_value), 2)))
                    
                    # Status (bullet point)
                    status_text = analysis.get('status', '').replace('_', ' ').title() or 'Unknown'
                    status_para = doc.add_paragraph(style='List Bullet')
                    status_run = status_para.add_run('Status: ')
                    status_run.bold = True
                    status_para.add_run(status_text)
                    
                    # Evidence found (bullet points)
                    if analysis.get('evidence'):
                        evidence = analysis['evidence']
                        if isinstance(evidence, list) and evidence:
                            evid_header = doc.add_paragraph(style='List Bullet')
                            evid_run = evid_header.add_run('Evidence Found:')
                            evid_run.bold = True
                            for evid_item in evidence:
                                if evid_item:
                                    doc.add_paragraph(str(evid_item), style='List Bullet 2')
                        elif evidence:
                            evid_header = doc.add_paragraph(style='List Bullet')
                            evid_run = evid_header.add_run('Evidence Found: ')
                            evid_run.bold = True
                            evid_header.add_run(str(evidence))
                    
                    # Gaps/Missing elements (bullet points)
                    if analysis.get('missing'):
                        missing_elements = analysis['missing']
                        if isinstance(missing_elements, list) and missing_elements:
                            gap_header = doc.add_paragraph(style='List Bullet')
                            gap_run = gap_header.add_run('Gaps:')
                            gap_run.bold = True
                            for gap_item in missing_elements:
                                if gap_item:
                                    doc.add_paragraph(str(gap_item), style='List Bullet 2')
                        elif missing_elements:
                            gap_header = doc.add_paragraph(style='List Bullet')
                            gap_run = gap_header.add_run('Gaps: ')
                            gap_run.bold = True
                            gap_header.add_run(str(missing_elements))
                    
                    # Recommendations (bullet points)
                    compliance_score = float(analysis.get('compliance_score') or analysis.get('relevance', 0))
                    rec_header = doc.add_paragraph(style='List Bullet')
                    rec_run = rec_header.add_run('Recommendations:')
                    rec_run.bold = True
                    
                    recommendations = []
                    if compliance_score < 0.4:
                        recommendations.append("CRITICAL: Immediate action required - document does not adequately address this requirement")
                    elif compliance_score < 0.6:
                        recommendations.append("HIGH PRIORITY: Significant improvements needed to meet compliance standards")
                    elif compliance_score < 0.8:
                        recommendations.append("MEDIUM PRIORITY: Minor improvements recommended for better compliance")
                    else:
                        recommendations.append("LOW PRIORITY: Document adequately addresses this requirement")
                    
                    missing_elements = analysis.get('missing', [])
                    if missing_elements:
                        if isinstance(missing_elements, list):
                            recommendations.append(f"Address missing elements: {', '.join(str(m) for m in missing_elements[:3])}")
                        else:
                            recommendations.append(f"Address missing elements: {missing_elements}")
                    
                    for rec in recommendations:
                        doc.add_paragraph(rec, style='List Bullet 2')
                    
                    doc.add_paragraph()  # Add space between analyses
            
            total_unique_analyses = len(all_analyses_map)
            if total_unique_analyses > 150:
                doc.add_paragraph(f'Note: Showing first 150 of {total_unique_analyses} unique compliance analyses (showing all would cause timeout)')
        else:
            doc.add_paragraph('• Detailed Analysis: No detailed analysis available')
        
        doc.add_paragraph()  # Add space between files
    
    # Key Findings Section
    findings_heading = doc.add_heading('Key Findings', level=1)
    _make_heading_bold(findings_heading)
    doc.add_paragraph(f'• Overall Compliance Rate: {findings["overall_compliance_rate"]}%')
    doc.add_paragraph(f'• Average Confidence Score: {findings["average_confidence_score"]}')
    doc.add_paragraph(f'• Documents Requiring Attention: {findings["documents_requiring_attention"]}')
    doc.add_paragraph(f'• Processing Success Rate: {findings["processing_success_rate"]}%')
    
    # Recommendations Section
    recommendations_heading = doc.add_heading('Recommendations', level=1)
    _make_heading_bold(recommendations_heading)
    for rec in recommendations:
        doc.add_paragraph(f'• {rec}')
    
    # SEBI AI Auditor Insights Section (if enabled)
    sebi_insights = report_data.get('sebi_insights')
    sebi_enabled = report_data.get('sebi_enabled', False)
    if sebi_enabled and sebi_insights:
        sebi_heading = doc.add_heading('SEBI AI Auditor Insights', level=1)
        _make_heading_bold(sebi_heading)
        
        # Filing Accuracy
        filing_accuracy = sebi_insights.get('filing_accuracy', {})
        if filing_accuracy:
            filing_heading = doc.add_heading('Filing Accuracy Analysis', level=2)
            _make_heading_bold(filing_heading)
            doc.add_paragraph(f'• Overall Accuracy: {filing_accuracy.get("accuracy_score", "N/A")}%')
            doc.add_paragraph(f'• Total Filings Analyzed: {filing_accuracy.get("total_filings", 0)}')
            doc.add_paragraph(f'• Accurate Filings: {filing_accuracy.get("accurate_filings", 0)}')
            doc.add_paragraph(f'• Issues Found: {filing_accuracy.get("issues_count", 0)}')
            if filing_accuracy.get('issues'):
                issues_para = doc.add_paragraph('• Key Issues: ')
                issues_para.add_run(', '.join(filing_accuracy['issues'][:5]))
        
        # Timeliness SLA
        timeliness = sebi_insights.get('timeliness_sla', {})
        if timeliness:
            timeliness_heading = doc.add_heading('Timeliness & SLA Compliance', level=2)
            _make_heading_bold(timeliness_heading)
            doc.add_paragraph(f'• SLA Compliance Rate: {timeliness.get("sla_compliance_rate", "N/A")}%')
            doc.add_paragraph(f'• On-Time Filings: {timeliness.get("on_time_filings", 0)}')
            doc.add_paragraph(f'• Late Filings: {timeliness.get("late_filings", 0)}')
            doc.add_paragraph(f'• Average Delay: {timeliness.get("average_delay_days", "N/A")} days')
            if timeliness.get('violations'):
                violations_para = doc.add_paragraph('• SLA Violations: ')
                violations_para.add_run(', '.join(timeliness['violations'][:5]))
        
        # Risk Score
        risk_score = sebi_insights.get('risk_score', {})
        if risk_score:
            risk_heading = doc.add_heading('Risk Assessment', level=2)
            _make_heading_bold(risk_heading)
            risk_score_value = risk_score.get("risk_score", risk_score.get("overall_risk_score", 0))
            if isinstance(risk_score_value, (int, float)) and risk_score_value > 0:
                doc.add_paragraph(f'• Overall Risk Score: {risk_score_value}/100')
            else:
                doc.add_paragraph(f'• Overall Risk Score: N/A/100')
            doc.add_paragraph(f'• Risk Level: {risk_score.get("risk_level", "N/A")}')
            doc.add_paragraph(f'• High-Risk Areas: {risk_score.get("high_risk_areas", risk_score.get("high_risk_count", 0))}')
            key_factors = risk_score.get('key_risk_factors', risk_score.get('risk_factors', []))
            if key_factors:
                factors_para = doc.add_paragraph('• Key Risk Factors: ')
                factors_list = key_factors if isinstance(key_factors, list) else []
                factors_para.add_run(', '.join(str(f) for f in factors_list[:5]))
        
        # Patterns
        patterns = sebi_insights.get('patterns', {})
        if patterns:
            patterns_heading = doc.add_heading('Pattern Analysis', level=2)
            _make_heading_bold(patterns_heading)
            patterns_count = patterns.get("patterns_detected", patterns.get("patterns_count", 0))
            doc.add_paragraph(f'• Patterns Detected: {patterns_count}')
            pattern_types = patterns.get('pattern_types', [])
            if pattern_types:
                types_para = doc.add_paragraph('• Pattern Types: ')
                types_list = pattern_types if isinstance(pattern_types, list) else []
                types_para.add_run(', '.join(str(t) for t in types_list[:5]))
            elif patterns.get('detected_patterns'):
                for pattern in patterns['detected_patterns'][:5]:
                    pattern_para = doc.add_paragraph(f'• {pattern.get("pattern_type", "Unknown")}: ')
                    pattern_para.add_run(pattern.get("description", "No description"))
        
        doc.add_paragraph()  # Add space before footer
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph('This report was generated automatically by the AI Audit System')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Inches(0.12)
    
    footer2 = doc.add_paragraph(f'Report ID: {metadata["audit_id"]} | Generated by User: {metadata["generated_by_user_id"]}')
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2_run = footer2.runs[0]
    footer2_run.font.size = Inches(0.12)
    
    # Save document to bytes
    import io
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer.getvalue()


def _generate_audit_recommendations(compliance_summary, documents, audit_row=None):
    """Generate recommendations based on audit results"""
    recommendations = []
    
    total_docs = sum(compliance_summary.values())
    if total_docs == 0:
        return ["No documents processed for analysis"]
    
    # Audit assignment recommendations
    if audit_row:
        is_ai_audit = audit_row[1] == 'A'
        if is_ai_audit:
            # For AI audits, only reviewer is required
            if not audit_row[11]:
                recommendations.append("ADMINISTRATIVE: AI audit assignment is incomplete. Please assign a reviewer through the Assign Audit page.")
        else:
            # For regular audits, auditor, assignee, and reviewer are required
            if not audit_row[9] or not audit_row[10] or not audit_row[11]:
                recommendations.append("ADMINISTRATIVE: Regular audit assignment is incomplete. Please assign auditor, assignee, and reviewer through the Assign Audit page.")
    
    # Compliance-based recommendations
    non_compliant_rate = compliance_summary['non_compliant'] / total_docs
    if non_compliant_rate > 0.3:
        recommendations.append("HIGH PRIORITY: Significant number of non-compliant documents detected. Immediate remediation required.")
    
    partially_compliant_rate = compliance_summary['partially_compliant'] / total_docs
    if partially_compliant_rate > 0.4:
        recommendations.append("MEDIUM PRIORITY: Many documents show partial compliance. Review and strengthen controls.")
    
    # Processing-based recommendations
    failed_docs = sum(1 for doc in documents if doc['ai_processing_status'] == 'failed')
    if failed_docs > 0:
        recommendations.append(f"TECHNICAL: {failed_docs} document(s) failed AI processing. Review document formats and retry processing.")
    
    # Confidence-based recommendations
    # Some document entries may not have a top-level confidence_score (e.g. when grouped by file),
    # so we use .get(...) and default to 0 to avoid KeyError.
    low_confidence_docs = sum(
        1
        for doc in documents
        if (doc.get('confidence_score') or 0) < 0.6
    )
    if low_confidence_docs > 0:
        recommendations.append(f"QUALITY: {low_confidence_docs} document(s) have low confidence scores. Consider manual review or document improvement.")
    
    # General recommendations
    if compliance_summary['compliant'] / total_docs > 0.8:
        recommendations.append("POSITIVE: High compliance rate achieved. Continue monitoring and maintain current controls.")
    else:
        recommendations.append("IMPROVEMENT NEEDED: Overall compliance rate below 80%. Implement additional controls and training.")
    
    return recommendations


@csrf_exempt
@api_view(['POST'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def test_structured_compliance_api(request, audit_id):
    """Test the structured compliance checking with a sample document
    MULTI-TENANCY: Only tests for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        # TODO: Implement structured_compliance_checker module
        # from .structured_compliance_checker import check_document_structured_compliance
        
        # Sample audit document for testing
        sample_document = """
        AUDIT REPORT - DATA SECURITY COMPLIANCE
        
        Executive Summary:
        This audit was conducted to assess our data security controls and ensure compliance 
        with our Data Protection Policy and Access Control Standards.
        
        Key Findings:
        1. Data Encryption: All sensitive data is encrypted using AES-256 encryption
        2. Access Controls: User access is properly managed with role-based permissions
        3. Audit Logging: Comprehensive audit logs are maintained for all data access
        4. Data Retention: Data retention policies are properly implemented
        
        Technical Details:
        - Database encryption: AES-256 enabled
        - Access control: Multi-factor authentication implemented
        - Audit trail: All access attempts logged with timestamps
        - Data classification: Sensitive data properly labeled
        
        Compliance Status:
        - Data Protection Policy: COMPLIANT
        - Access Control Standards: COMPLIANT
        - Audit Requirements: COMPLIANT
        """
        
        # Test structured compliance checking
        # TODO: Implement structured_compliance_checker module
        # result = check_document_structured_compliance(
        #     audit_id=audit_id,
        #     document_id="test_doc_001",
        #     document_text=sample_document,
        #     document_name="Sample Audit Report"
        # )
        result = {'success': False, 'error': 'structured_compliance_checker module not implemented'}
        
        return Response({
            'success': True,
            'message': 'Structured compliance checking test completed',
            'audit_id': audit_id,
            'test_result': result
        })
        
    except Exception as e:
        logger.error(f"❌ Error testing structured compliance: {e}")
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
@api_view(['GET'])
@authentication_classes([CsrfExemptSessionAuthentication, BasicAuthentication])
@permission_classes([IsAuthenticated])
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def get_relevant_documents_for_audit(request, audit_id):
    """
    Get documents from file_operations that are relevant to this audit.
    Filters by framework and optionally by selected policies/subpolicies/compliances.
    MULTI-TENANCY: Only returns documents for audits in user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        logger.info(f"📋 Getting relevant documents for audit {audit_id}")
        
        # Check authentication
        from ...rbac.utils import RBACUtils
        user_id = RBACUtils.get_user_id_from_request(request)
        if not user_id:
            return Response({
                'success': False,
                'error': 'Authentication required'
            }, status=status.HTTP_401_UNAUTHORIZED)
        
        # Get audit details to get framework_id
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT AuditId, FrameworkId, PolicyId, SubPolicyId
                FROM audit
                WHERE AuditId = %s AND TenantId = %s
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id, tenant_id])
            audit_row = cursor.fetchone()
            
            if not audit_row:
                return Response({
                    'success': False,
                    'error': 'Audit not found'
                }, status=status.HTTP_404_NOT_FOUND)
            
            framework_id = audit_row[1]
        
        if not framework_id:
            logger.warning(f"⚠️ Audit {audit_id} has no FrameworkId assigned")
            return Response({
                'success': True,
                'documents': [],
                'count': 0,
                'audit_id': audit_id,
                'framework_id': None,
                'message': 'Audit has no framework assigned. Please assign a framework to see relevant documents.'
            })
        
        # Get optional filter parameters
        policy_ids = request.GET.get('policy_ids', '').split(',') if request.GET.get('policy_ids') else []
        policy_ids = [int(p) for p in policy_ids if p.strip().isdigit()]
        
        subpolicy_ids = request.GET.get('subpolicy_ids', '').split(',') if request.GET.get('subpolicy_ids') else []
        subpolicy_ids = [int(s) for s in subpolicy_ids if s.strip().isdigit()]
        
        compliance_ids = request.GET.get('compliance_ids', '').split(',') if request.GET.get('compliance_ids') else []
        compliance_ids = [int(c) for c in compliance_ids if c.strip().isdigit()]
        
        # First, check if there are any documents in file_operations for this framework
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT COUNT(*) as total_files, 
                       COUNT(CASE WHEN status = 'completed' THEN 1 END) as completed_files
                FROM file_operations 
                WHERE FrameworkId = %s
            """, [framework_id])
            file_stats = cursor.fetchone()
            total_files = file_stats[0] if file_stats else 0
            completed_files = file_stats[1] if file_stats else 0
            
            logger.info(f"📊 File operations stats for framework {framework_id}: {total_files} total, {completed_files} completed")
            
            # Check if any documents have been analyzed for this audit
            cursor.execute("""
                SELECT COUNT(*) as analyzed_count
                FROM document_audit_relevance 
                WHERE audit_id = %s
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
            analyzed_row = cursor.fetchone()
            analyzed_count = analyzed_row[0] if analyzed_row else 0
            
            logger.info(f"📊 Documents analyzed for audit {audit_id}: {analyzed_count}")
        
        # Build query to get relevant documents
        query = """
            SELECT 
                fo.id as file_operation_id,
                fo.file_name,
                fo.original_name,
                COALESCE(dar.s3_url, fo.s3_url) as s3_url,
                fo.s3_key,
                fo.file_type,
                fo.file_size,
                fo.summary,
                fo.metadata,
                fo.created_at,
                fo.module,
                dar.relevance_score,
                dar.relevance_reason,
                dar.matched_policies,
                dar.matched_subpolicies,
                dar.matched_compliances
            FROM file_operations fo
            INNER JOIN document_audit_relevance dar ON fo.id = dar.file_operation_id
            WHERE dar.audit_id = %s
              AND dar.relevance_score >= 0.6
              AND fo.status = 'completed'
              AND fo.FrameworkId = %s
        """
        
        params = [int(audit_id) if str(audit_id).isdigit() else audit_id, framework_id]
        
        # Add policy filter if provided
        if policy_ids:
            # Check if matched_policies JSON contains any of the policy_ids
            policy_conditions = ' OR '.join(['JSON_CONTAINS(dar.matched_policies, %s)'] * len(policy_ids))
            query += f" AND ({policy_conditions})"
            params.extend([json.dumps([pid]) for pid in policy_ids])
        
        # Add subpolicy filter if provided
        if subpolicy_ids:
            subpolicy_conditions = ' OR '.join(['JSON_CONTAINS(dar.matched_subpolicies, %s)'] * len(subpolicy_ids))
            query += f" AND ({subpolicy_conditions})"
            params.extend([json.dumps([sid]) for sid in subpolicy_ids])
        
        # Add compliance filter if provided
        if compliance_ids:
            compliance_conditions = ' OR '.join(['JSON_CONTAINS(dar.matched_compliances, %s)'] * len(compliance_ids))
            query += f" AND ({compliance_conditions})"
            params.extend([json.dumps([cid]) for cid in compliance_ids])
        
        query += " ORDER BY dar.relevance_score DESC, fo.created_at DESC"
        
        with connection.cursor() as cursor:
            cursor.execute(query, params)
            columns = [col[0] for col in cursor.description]
            documents = [dict(zip(columns, row)) for row in cursor.fetchall()]
        
        # Get policy/subpolicy/compliance names for matched IDs
        with connection.cursor() as cursor:
            # Get all unique policy IDs
            all_policy_ids = set()
            all_subpolicy_ids = set()
            all_compliance_ids = set()
            for doc in documents:
                matched_policies = json.loads(doc.get('matched_policies', '[]')) if doc.get('matched_policies') else []
                matched_subpolicies = json.loads(doc.get('matched_subpolicies', '[]')) if doc.get('matched_subpolicies') else []
                matched_compliances = json.loads(doc.get('matched_compliances', '[]')) if doc.get('matched_compliances') else []
                all_policy_ids.update(matched_policies)
                all_subpolicy_ids.update(matched_subpolicies)
                all_compliance_ids.update(matched_compliances)
            
            # Fetch policy names
            policy_names = {}
            if all_policy_ids:
                policy_ids_str = ','.join(map(str, all_policy_ids))
                cursor.execute(f"""
                    SELECT PolicyId, PolicyName 
                    FROM policies 
                    WHERE PolicyId IN ({policy_ids_str})
                """)
                for row in cursor.fetchall():
                    policy_names[row[0]] = row[1]
            
            # Fetch subpolicy names
            subpolicy_names = {}
            if all_subpolicy_ids:
                subpolicy_ids_str = ','.join(map(str, all_subpolicy_ids))
                cursor.execute(f"""
                    SELECT SubPolicyId, SubPolicyName 
                    FROM subpolicies 
                    WHERE SubPolicyId IN ({subpolicy_ids_str})
                """)
                for row in cursor.fetchall():
                    subpolicy_names[row[0]] = row[1]
            
            # Fetch compliance names with their parent subpolicy and policy IDs
            compliance_names = {}
            compliance_subpolicy_map = {}  # Map compliance_id -> subpolicy_id
            compliance_policy_map = {}    # Map compliance_id -> policy_id
            if all_compliance_ids:
                compliance_ids_str = ','.join(map(str, all_compliance_ids))
                cursor.execute(f"""
                    SELECT c.ComplianceId, c.ComplianceTitle, c.SubPolicyId, sp.PolicyId
                    FROM compliance c
                    LEFT JOIN subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                    WHERE c.ComplianceId IN ({compliance_ids_str})
                """)
                for row in cursor.fetchall():
                    compliance_id = row[0]
                    compliance_names[compliance_id] = row[1]
                    if row[2]:  # SubPolicyId
                        compliance_subpolicy_map[compliance_id] = row[2]
                    if row[3]:  # PolicyId
                        compliance_policy_map[compliance_id] = row[3]
        
        # Format response
        formatted_documents = []
        for doc in documents:
            # Parse JSON fields
            matched_policies = json.loads(doc.get('matched_policies', '[]')) if doc.get('matched_policies') else []
            matched_subpolicies = json.loads(doc.get('matched_subpolicies', '[]')) if doc.get('matched_subpolicies') else []
            matched_compliances = json.loads(doc.get('matched_compliances', '[]')) if doc.get('matched_compliances') else []
            metadata = json.loads(doc.get('metadata', '{}')) if doc.get('metadata') else {}
            
            # Create arrays with IDs and names
            matched_policies_with_names = [
                {'id': pid, 'name': policy_names.get(pid, f'Policy {pid}')}
                for pid in matched_policies
            ]
            matched_subpolicies_with_names = [
                {'id': sid, 'name': subpolicy_names.get(sid, f'Subpolicy {sid}')}
                for sid in matched_subpolicies
            ]
            matched_compliances_with_names = [
                {
                    'id': cid, 
                    'name': compliance_names.get(cid, f'Compliance {cid}'),
                    'subpolicy_id': compliance_subpolicy_map.get(cid),
                    'policy_id': compliance_policy_map.get(cid)
                }
                for cid in matched_compliances
            ]
            
            formatted_documents.append({
                'file_operation_id': doc.get('file_operation_id'),
                'file_name': doc.get('file_name'),
                'original_name': doc.get('original_name'),
                's3_url': doc.get('s3_url'),
                's3_key': doc.get('s3_key'),
                'file_type': doc.get('file_type'),
                'file_size': doc.get('file_size'),
                'summary': doc.get('summary'),
                'metadata': metadata,
                'created_at': doc.get('created_at').isoformat() if doc.get('created_at') else None,
                'module': doc.get('module'),
                'relevance_score': float(doc.get('relevance_score', 0)),
                'relevance_reason': doc.get('relevance_reason'),
                'matched_policies': matched_policies,
                'matched_policies_with_names': matched_policies_with_names,
                'matched_subpolicies': matched_subpolicies,
                'matched_subpolicies_with_names': matched_subpolicies_with_names,
                'matched_compliances': matched_compliances,
                'matched_compliances_with_names': matched_compliances_with_names
            })
        
        logger.info(f"✅ Found {len(formatted_documents)} relevant documents for audit {audit_id}")
        
        # Add helpful message if no documents found
        message = None
        if len(formatted_documents) == 0:
            if total_files == 0:
                message = 'No documents found in Document Handling for this framework. Upload documents in Document Handling to see AI-suggested documents here.'
            elif completed_files == 0:
                message = f'Found {total_files} document(s) in Document Handling, but none are completed yet. Please wait for uploads to complete.'
            elif analyzed_count == 0:
                message = f'Found {completed_files} completed document(s) in Document Handling, but they have not been analyzed for this audit yet. The analysis runs automatically in the background - please check back in a few moments.'
            else:
                message = f'Found {completed_files} completed document(s) in Document Handling, but none have a relevance score >= 60% for this audit. Documents need to be highly relevant to appear here.'
        
        return Response({
            'success': True,
            'documents': formatted_documents,
            'count': len(formatted_documents),
            'audit_id': audit_id,
            'framework_id': framework_id,
            'message': message,
            'stats': {
                'total_files_in_framework': total_files,
                'completed_files_in_framework': completed_files,
                'documents_analyzed_for_audit': analyzed_count
            }
        })
        
    except Exception as e:
        logger.error(f"❌ Error getting relevant documents: {e}")
        import traceback
        traceback.print_exc()
        return Response({
            'success': False,
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


def check_and_update_ai_audit_status(audit_id):
    """
    Helper function to check if all documents for an AI audit are processed,
    and update the audit status to 'Under review' if all are done.
    
    Returns: dict with 'updated': bool, 'status': str, 'message': str
    """
    try:
        with connection.cursor() as cursor:
            # Check if this is an AI audit
            cursor.execute("SELECT AuditType, Status FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
            audit_info = cursor.fetchone()
            
            if not audit_info:
                return {'updated': False, 'status': None, 'message': f'Audit {audit_id} not found'}
            
            audit_type, current_status = audit_info
            
            if audit_type != 'A':  # Not an AI audit
                return {'updated': False, 'status': current_status, 'message': f'Audit {audit_id} is not an AI audit (type: {audit_type})'}
            
            # Count total documents and completed documents for this audit
            cursor.execute("""
                SELECT 
                    COUNT(*) as total,
                    SUM(CASE WHEN ai_processing_status = 'completed' THEN 1 ELSE 0 END) as completed,
                    SUM(CASE WHEN ai_processing_status = 'failed' THEN 1 ELSE 0 END) as failed
                FROM ai_audit_data 
                WHERE audit_id = %s
            """, [int(audit_id) if str(audit_id).isdigit() else audit_id])
            
            doc_stats = cursor.fetchone()
            total_docs = (doc_stats[0] if doc_stats and doc_stats[0] is not None else 0) or 0
            completed_docs = (doc_stats[1] if doc_stats and doc_stats[1] is not None else 0) or 0
            failed_docs = (doc_stats[2] if doc_stats and doc_stats[2] is not None else 0) or 0
            
            logger.info(f"📊 Manual status check for audit {audit_id}: total={total_docs}, completed={completed_docs}, failed={failed_docs}, current_status='{current_status}'")
            
            # Update audit status if all documents are processed (completed or failed)
            if total_docs > 0 and (completed_docs + failed_docs) == total_docs:
                # All documents processed - update audit status to "Under review" if not already
                if current_status not in ['Under review', 'Completed']:
                    from ...models import Audit
                    audit_obj = Audit.objects.get(AuditId=int(audit_id) if str(audit_id).isdigit() else audit_id)
                    
                    # Get auditor user_id for creating audit_version (extract integer ID from User object if needed)
                    auditor_id = audit_obj.Auditor if hasattr(audit_obj, 'Auditor') and audit_obj.Auditor else None
                    if not auditor_id:
                        # Fallback: try to get from audit table directly
                        cursor.execute("SELECT Auditor FROM audit WHERE AuditId = %s", [int(audit_id) if str(audit_id).isdigit() else audit_id])
                        auditor_row = cursor.fetchone()
                        auditor_id = auditor_row[0] if auditor_row and auditor_row[0] else None
                    
                    # Extract integer ID if auditor_id is a User object
                    if auditor_id:
                        if hasattr(auditor_id, 'UserId'):
                            auditor_id = auditor_id.UserId
                        elif hasattr(auditor_id, 'id'):
                            auditor_id = auditor_id.id
                        elif hasattr(auditor_id, 'pk'):
                            auditor_id = auditor_id.pk
                        elif not isinstance(auditor_id, int):
                            try:
                                auditor_id = int(auditor_id)
                            except (ValueError, TypeError):
                                auditor_id = None
                    
                    # Update audit status to Under review so reviewer can Start / Accept / Reject
                    audit_obj.Status = 'Under review'
                    if not getattr(audit_obj, "ReviewStartDate", None):
                        audit_obj.ReviewStartDate = timezone.now()
                    audit_obj.save()
                    logger.info(f"✅ Updated AI audit {audit_id} status from '{current_status}' to 'Under review'")
                    
                    # Automatically create audit_version with all AI-generated findings
                    try:
                        from .audit_views import create_audit_version
                        if auditor_id:
                            version_result = create_audit_version(audit_id, auditor_id)
                            logger.info(f"✅ Automatically created audit_version for AI audit {audit_id}: {version_result}")
                        else:
                            logger.warning(f"⚠️ Could not create audit_version for audit {audit_id}: No auditor_id found")
                    except Exception as version_err:
                        logger.error(f"❌ Failed to create audit_version automatically: {version_err}")
                        import traceback
                        logger.error(f"Traceback: {traceback.format_exc()}")
                        # Don't fail the whole operation if version creation fails
                    
                    return {
                        'updated': True,
                        'status': 'Under review',
                        'message': f'Status updated from "{current_status}" to "Under review" (all {total_docs} documents processed). Audit version created automatically.'
                    }
                else:
                    return {
                        'updated': False,
                        'status': current_status,
                        'message': f'All documents processed, but status is already "{current_status}"'
                    }
            else:
                completed = completed_docs if completed_docs is not None else 0
                failed = failed_docs if failed_docs is not None else 0
                remaining = total_docs - (completed + failed)
                return {
                    'updated': False,
                    'status': current_status,
                    'message': f'Not all documents processed yet: {remaining} remaining (completed: {completed}, failed: {failed}, total: {total_docs})'
                }
    except Exception as e:
        logger.error(f"❌ Error checking/updating audit status: {e}")
        import traceback
        logger.error(f"Traceback: {traceback.format_exc()}")
        return {'updated': False, 'status': None, 'message': f'Error: {str(e)}'}

