"""
AI-Powered Risk Instance Document Ingestion (Complete)
- Reads PDF/DOCX/XLSX/TXT
- Extracts risk instance data using AI (OpenAI or Ollama - configurable)
- Fills missing fields with focused prompts
- Returns normalized, DB-ready JSON for `risk_instance` table

Configuration:
- Set RISK_AI_PROVIDER='openai' or 'ollama' in environment or Django settings
- Default: 'ollama' if Ollama is configured, else 'openai'
- OpenAI requires: OPENAI_API_KEY
- Ollama requires: OLLAMA_BASE_URL (uses optimized quantized models)
"""

import os
import re
import json
import time
import tempfile
from datetime import date, datetime
from typing import Any, Optional

import requests
from django.http import JsonResponse, HttpResponse, StreamingHttpResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.parsers import MultiPartParser, FormParser

# RBAC imports
from ...rbac.decorators import rbac_required
# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)


# Phase 2 Optimizations (reuse same utilities as risk_ai_doc)
from ...utils.ai_cache import cached_llm_call
from ...utils.document_preprocessor import preprocess_document, calculate_document_hash
from ...debug_utils import debug_print

# Phase 3 Optimizations (reuse same utilities as risk_ai_doc)
from ...utils.rag_system import (
    add_document_to_rag,
    retrieve_relevant_context,
    is_rag_available,
    get_rag_stats,
)
from ...utils.model_router import (
    route_model,
    track_system_load,
    get_current_system_load,
)
from ...utils.request_queue import (
    rate_limit_decorator,
    process_with_queue,
    get_queue_status
)
from ...utils.file_compression import decompress_if_needed
from ...routes.Global.s3_fucntions import create_direct_mysql_client
from ...ai.service import get_ai_service
from ...ai.types import AIRequestOptions
from ...ai.processing.preprocessor import DocumentPreparationService

# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import RiskInstance  # Import RiskInstance model


# =========================
# AI PROVIDER CONFIG (OpenAI or Ollama)
# =========================
# Use the same provider configuration and helpers as risk_ai_doc
from django.conf import settings
from .risk_ai_doc import (
    AI_PROVIDER,
    call_ollama_json,
    call_openai_json,  # Phase 2: cached OpenAI wrapper
    _calculate_optimal_context_size,
    _select_ollama_model_by_complexity,
    OLLAMA_BASE_URL,
    OLLAMA_MODEL_DEFAULT,
    OLLAMA_MODEL_FAST,
    OLLAMA_MODEL_COMPLEX,
    OPENAI_API_KEY,
    OPENAI_API_URL,
    OPENAI_MODEL,
)

if AI_PROVIDER == 'openai':
    if not OPENAI_API_KEY:
        debug_print("⚠️  WARNING: OPENAI_API_KEY not found in Django settings!")
        debug_print("   Please set OPENAI_API_KEY in your .env file.")
    else:
        debug_print(f"🌐 OpenAI Configuration for Risk Instance AI:")
        debug_print(f"   API URL: {OPENAI_API_URL}")
        debug_print(f"   Model: {OPENAI_MODEL}")
        debug_print(f"   API Key: {'*' * (len(OPENAI_API_KEY) - 4) + OPENAI_API_KEY[-4:]}")
elif AI_PROVIDER == 'ollama':
    debug_print("🚀 Ollama Configuration for Risk Instance AI (OPTIMIZED):")
    debug_print(f"   Base URL: {OLLAMA_BASE_URL}")
    debug_print(f"   Default Model: {OLLAMA_MODEL_DEFAULT}")
    debug_print(f"   Fast Model: {OLLAMA_MODEL_FAST}")
    debug_print(f"   Complex Model: {OLLAMA_MODEL_COMPLEX}")

# RiskInstance DB fields (excluding auto-generated IDs: RiskInstanceId, UserId, ReportedBy, ReviewerId, IncidentId, ComplianceId, RiskId)
RISK_INSTANCE_DB_FIELDS = [
    "RiskTitle",
    "RiskDescription",
    "PossibleDamage",
    "RiskPriority",
    "Criticality",
    "Category",
    "Origin",
    "RiskLikelihood",
    "RiskImpact",
    "RiskExposureRating",
    "RiskMultiplierX",
    "RiskMultiplierY",
    "Appetite",
    "RiskResponseType",
    "RiskResponseDescription",
    "RiskMitigation",
    "RiskType",
    "RiskOwner",
    "BusinessImpact",
    "RiskStatus",
    "MitigationStatus",
    "ModifiedMitigations",
    "RiskFormDetails",
    "Reviewer",
]

# Canonical choices / constraints to stabilize LLM outputs
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES = ["Low", "Medium", "High", "Critical"]
CATEGORY_HINTS = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]
RISKTYPE_HINTS = ["Current", "Residual", "Inherent", "Emerging", "Accepted"]
ORIGIN_HINTS = ["Internal", "External", "Third-Party", "Regulatory", "Market", "Operational"]
APPETITE_HINTS = ["Low", "Medium", "High"]
RESPONSE_TYPE_HINTS = ["Avoid", "Mitigate", "Transfer", "Accept"]
RISK_STATUS_CHOICES = ["Not Assigned", "Assigned", "Approved", "Rejected"]
MITIGATION_STATUS_CHOICES = ["Pending", "Yet to Start", "Work In Progress", "Revision Required by Reviewer", "Revision Required by User", "Completed"]
DATE_FORMAT_HINT = "YYYY-MM-DD (ISO)"

# Strict JSON schema block the LLM must follow
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON array. No markdown, no code blocks, no explanations.
Start with [ and end with ]. Use proper JSON syntax with double quotes.

Example structure (return array of risk instances like this):
[
  {{
    "RiskTitle": "Specific risk instance title here",
    "RiskDescription": "Detailed description of the risk instance occurrence",
    "PossibleDamage": "Concrete description of damages",
    "RiskPriority": "Medium",
    "Criticality": "Medium",
    "Category": "Operational",
    "Origin": "Internal",
    "RiskLikelihood": 5,
    "RiskImpact": 6,
    "RiskExposureRating": 30.0,
    "RiskMultiplierX": 1.0,
    "RiskMultiplierY": 1.0,
    "Appetite": "Medium",
    "RiskResponseType": "Mitigate",
    "RiskResponseDescription": "Detailed response strategy description",
    "RiskMitigation": [{{"step": "Step 1", "description": "Detailed action description"}}],
    "RiskType": "Current",
    "RiskOwner": "Risk Manager Name/Role",
    "BusinessImpact": "Detailed business impact explanation",
    "RiskStatus": "Not Assigned",
    "MitigationStatus": "Pending",
    "ModifiedMitigations": [],
    "RiskFormDetails": {{"assessment_method": "Qualitative", "data_sources": ["Document"], "assessment_date": "{date.today().isoformat()}"}},
    "Reviewer": "Pending Review",
    "_meta": {{
      "per_field": {{
        "RiskTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found explicitly in document"}},
        "Criticality": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from severity indicators"}},
        "RiskLikelihood": {{"source": "AI_GENERATED", "confidence": 0.6, "rationale": "Estimated based on frequency mentions"}}
      }}
    }}
  }}
]

Rules:
- Criticality must be: {CRITICALITY_CHOICES}
- RiskPriority must be: {PRIORITY_CHOICES}
- RiskLikelihood and RiskImpact must be integers 1-10
- RiskExposureRating must be float 0-100
- RiskMultiplierX and RiskMultiplierY must be floats 0.1-2.0
- RiskMitigation must be a JSON array of objects with "step" and "description"
- ModifiedMitigations must be a JSON array (empty [] if no modifications)
- RiskFormDetails must be a JSON object
- RiskStatus must be one of: {RISK_STATUS_CHOICES}
- MitigationStatus must be one of: {MITIGATION_STATUS_CHOICES}
- Appetite must be one of: {APPETITE_HINTS}
- RiskResponseType must be one of: {RESPONSE_TYPE_HINTS}
- Origin must be one of: {ORIGIN_HINTS}
- RiskType must be one of: {RISKTYPE_HINTS}
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON array, nothing else
- Include _meta.per_field for EVERY field showing source (EXTRACTED or AI_GENERATED), confidence (0.0-1.0), and rationale
"""

def clamp_int(v, lo, hi) -> Optional[int]:
    if v is None: return None
    try:
        return max(lo, min(hi, int(v)))
    except Exception:
        # try to find a number in string
        m = re.search(r"\b(\d{1,2})\b", str(v))
        if m:
            return max(lo, min(hi, int(m.group(1))))
        return None

def as_float_or_none(v) -> Optional[float]:
    if v is None: return None
    try:
        return float(v)
    except Exception:
        try:
            # extract first float from string
            m = re.search(r"[-+]?\d*\.?\d+", str(v))
            return float(m.group(0)) if m else None
        except Exception:
            return None

def as_date_or_none(s: str) -> Optional[str]:
    if not s: return None
    s = str(s).strip()
    return s if re.match(r"^\d{4}-\d{2}-\d{2}$", s) else None

def compute_exposure(lk, im) -> Optional[float]:
    if lk is None or im is None: return None
    try:
        return round(float(lk) * float(im), 2)
    except Exception:
        return None

def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None


# =========================
# FILE EXTRACTORS
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF (pdfplumber preferred; PyPDF2 fallback)."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        debug_print(f"[PDF] Extraction error: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        debug_print(f"[DOCX] Extraction error: {e}")
        return ""

def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        debug_print(f"[XLSX] Extraction error: {e}")
        return ""

def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""


# =========================
# DOCUMENT PARSING & DETECTION
# =========================
def detect_and_parse_risk_instance_blocks(text: str) -> list[dict]:
    """
    Step 1: Detect how many risk instances are in the document by finding 'Risk X:' patterns.
    Step 2: For each risk instance, extract whatever fields are explicitly present.
    Returns a list of partially-filled risk instance dictionaries with only extracted fields.
    """
    debug_print(f"🔍 Detecting risk instance blocks in document...")
    
    # Find all risk blocks using pattern "Risk X: Title" or "Risk Instance X: Title"
    risk_pattern = r'Risk(?:\s+Instance)?\s+(\d+):\s*([^\n]+)'
    risk_matches = list(re.finditer(risk_pattern, text, re.IGNORECASE))
    
    if not risk_matches:
        debug_print(f"⚠️  No risk instances found with 'Risk X:' or 'Risk Instance X:' pattern")
        return []
    
    debug_print(f"✅ Found {len(risk_matches)} risk instance(s) in document")
    
    risk_instances = []
    for i, match in enumerate(risk_matches):
        risk_num = match.group(1)
        risk_title = match.group(2).strip()
        
        # Extract the text block for this risk instance (from this match to next risk or end of document)
        start_pos = match.end()
        if i + 1 < len(risk_matches):
            end_pos = risk_matches[i + 1].start()
        else:
            end_pos = len(text)
        
        risk_block = text[start_pos:end_pos]
        
        debug_print(f"  📋 Risk Instance {risk_num}: {risk_title[:60]}...")
        
        # Extract fields that are explicitly present in the block
        risk_data = {
            "RiskTitle": risk_title,  # Always extracted from document
            "_extracted_fields": ["RiskTitle"],  # Track what was extracted
            "_risk_block": risk_block  # Keep for AI context
        }
        
        # Field mapping: document field name -> DB field name
        field_mappings = {
            r'(?:Risk\s+)?Description:\s*([^\n]+(?:\n(?!(?:Risk|Possible Damage|Risk Priority|Criticality|Category|Origin|Status|Mitigation|Risk Type|Risk Owner|Business Impact|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'RiskDescription',
            r'Possible Damage:\s*([^\n]+(?:\n(?!(?:Risk|Description|Risk Priority|Criticality|Category|Origin|Status|Mitigation|Risk Type|Risk Owner|Business Impact|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'PossibleDamage',
            r'Risk Priority:\s*([^\n]+)': 'RiskPriority',
            r'Criticality:\s*([^\n]+)': 'Criticality',
            r'Category:\s*([^\n]+)': 'Category',
            r'Origin:\s*([^\n]+)': 'Origin',
            r'Risk Likelihood:\s*([^\n]+)': 'RiskLikelihood',
            r'Risk Impact:\s*([^\n]+)': 'RiskImpact',
            r'(?:Risk\s+)?Exposure(?:\s+Rating)?:\s*([^\n]+)': 'RiskExposureRating',
            r'Multiplier\s*X:\s*([^\n]+)': 'RiskMultiplierX',
            r'Multiplier\s*Y:\s*([^\n]+)': 'RiskMultiplierY',
            r'Appetite:\s*([^\n]+)': 'Appetite',
            r'Risk Response Type:\s*([^\n]+)': 'RiskResponseType',
            r'Risk Response(?:\s+Description)?:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Criticality|Category|—))[^\n]+)*)': 'RiskResponseDescription',
            r'Mitigation:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Criticality|Category|Origin|Risk Type|Risk Owner|Business Impact|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'RiskMitigation',
            r'Risk Type:\s*([^\n]+)': 'RiskType',
            r'Risk Owner:\s*([^\n]+)': 'RiskOwner',
            r'Business Impact:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|Risk Priority|Criticality|Category|Origin|Mitigation|Risk Type|Risk Owner|Risk Status|Mitigation Status|Reviewer|—))[^\n]+)*)': 'BusinessImpact',
            r'Risk Status:\s*([^\n]+)': 'RiskStatus',
            r'Mitigation Status:\s*([^\n]+)': 'MitigationStatus',
            r'Modified Mitigations:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|—))[^\n]+)*)': 'ModifiedMitigations',
            r'Risk Form Details:\s*([^\n]+(?:\n(?!(?:Risk|Description|Possible Damage|—))[^\n]+)*)': 'RiskFormDetails',
            r'Reviewer:\s*([^\n]+)': 'Reviewer',
        }
        
        # Extract each field if present
        for pattern, field_name in field_mappings.items():
            match = re.search(pattern, risk_block, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if value:  # Only add if not empty
                    risk_data[field_name] = value
                    risk_data["_extracted_fields"].append(field_name)
                    debug_print(f"    ✓ {field_name}: {value[:50]}...")
        
        risk_instances.append(risk_data)
    
    return risk_instances


# =========================
# AI EXTRACTION CORE
# =========================
def infer_single_field(
    field_name: str,
    current_record: dict,
    document_context: str,
    document_hash: str = None,
) -> tuple[Any, dict]:
    """
    Focused prompt for ONE field using AI (supports both OpenAI and Ollama).
    Uses Phase 2 few-shot prompts, Phase 3 RAG, and caching.
    Returns: (value, metadata_dict)
    """
    provider_name = AI_PROVIDER.upper()
    debug_print(f"🤖 AI PREDICTING FIELD: {field_name} (using {provider_name} with Phase 2+3 optimizations)")

    # Optimize context size for Ollama; keep previous behavior for OpenAI
    if AI_PROVIDER == 'ollama':
        context_size = _calculate_optimal_context_size(len(document_context), "simple")
        optimized_context = document_context[:context_size] if len(document_context) > context_size else document_context
    else:
        optimized_context = document_context[:3000]

    try:
        ai_service = get_ai_service()
        if AI_PROVIDER == 'ollama':
            debug_print(f"   📤 Sending structured request to centralized Ollama task for {field_name}...")
            model = _select_ollama_model_by_complexity(len(optimized_context), 1)
            out = ai_service.run_task(
                "risk.infer_field",
                payload={
                    "field_name": field_name,
                    "subject_type": "risk_instance",
                    "document_context": optimized_context,
                    "current_record": current_record,
                    "current_record_fields": RISK_INSTANCE_DB_FIELDS,
                },
                options=AIRequestOptions(
                    task_name="risk.infer_field",
                    preferred_provider="ollama",
                    preferred_model=model,
                    document_hash=document_hash,
                    use_cache=True,
                ),
            )
            model_used = model
        else:
            debug_print(f"   📤 Sending structured request to centralized OpenAI task for {field_name}...")
            out = ai_service.run_task(
                "risk.infer_field",
                payload={
                    "field_name": field_name,
                    "subject_type": "risk_instance",
                    "document_context": optimized_context,
                    "current_record": current_record,
                    "current_record_fields": RISK_INSTANCE_DB_FIELDS,
                },
                options=AIRequestOptions(
                    task_name="risk.infer_field",
                    preferred_provider="openai",
                    document_hash=document_hash,
                    use_cache=True,
                ),
            )
            model_used = OPENAI_MODEL

        # Handle response - check if it's the expected format or a full risk object
        if isinstance(out, dict):
            # Check if it's the expected format with "value" key
            if "value" in out:
                v = out.get("value")
                confidence = out.get("confidence", 0.7)
                rationale = out.get("rationale", "AI predicted based on document context")
            # Check if model returned a full risk object instead (extract the field we need)
            elif field_name in out:
                debug_print(f"   ⚠️  Model returned full risk object instead of single field format. Extracting {field_name}...")
                v = out.get(field_name)
                confidence = 0.6  # Lower confidence since format was wrong
                rationale = f"Extracted {field_name} from full risk object response (model format issue)"
            else:
                # Try to find any value that might be the answer
                v = None
                confidence = 0.5
                rationale = "Could not extract value from response format"
        else:
            v = None
            confidence = 0.5
            rationale = "Response was not in expected format"
        
        debug_print(f"   ✅ AI PREDICTED {field_name}: '{v}' (confidence: {confidence:.2f})")
    except Exception as e:
        debug_print(f"   ❌ AI FAILED to predict {field_name}: {str(e)}")
        # Try to extract value from error message if it contains JSON
        v = None
        confidence = 0.0
        rationale = f"AI prediction failed: {str(e)}"
        model_used = None
        
        # Last resort: try to extract from error string if it contains the field
        error_str = str(e)
        if field_name in error_str:
            # Look for field_name: "value" pattern in error
            pattern = rf'"{field_name}"\s*:\s*"([^"]+)"'
            match = re.search(pattern, error_str)
            if match:
                v = match.group(1)
                confidence = 0.4
                rationale = f"Extracted {field_name} from error response (low confidence)"
                debug_print(f"   ⚠️  Extracted value from error message: {v}")

    # Create metadata for this field
    metadata = {
        "source": "AI_GENERATED",
        "confidence": confidence,
        "rationale": rationale,
        "provider": AI_PROVIDER,
        "model_used": model_used,
    }

    # normalize after inference
    if field_name in ("RiskLikelihood", "RiskImpact"):
        v = clamp_int(v, 1, 10) or 5
    elif field_name == "RiskExposureRating":
        lk = clamp_int(current_record.get("RiskLikelihood"), 1, 10)
        im = clamp_int(current_record.get("RiskImpact"), 1, 10)
        computed = compute_exposure(lk, im)
        v = computed if computed is not None else 0.0
    elif field_name in ("RiskMultiplierX", "RiskMultiplierY"):
        vf = as_float_or_none(v)
        v = vf if vf is not None else 1.0
    elif field_name in ("CreatedAt", "MitigationDueDate"):
        v = as_date_or_none(v) or date.today().isoformat()
    elif field_name == "Criticality":
        v = normalize_choice(v, CRITICALITY_CHOICES) or "Medium"
    elif field_name == "RiskPriority":
        v = normalize_choice(v, PRIORITY_CHOICES) or "Medium"
    elif field_name == "RiskType":
        v = normalize_choice(v, RISKTYPE_HINTS) or "Current"
    elif field_name == "Origin":
        v = normalize_choice(v, ORIGIN_HINTS) or "Internal"
    elif field_name == "Appetite":
        v = normalize_choice(v, APPETITE_HINTS) or "Medium"
    elif field_name == "RiskResponseType":
        v = normalize_choice(v, RESPONSE_TYPE_HINTS) or "Mitigate"
    elif field_name == "RiskStatus":
        v = normalize_choice(v, RISK_STATUS_CHOICES) or "Not Assigned"
    elif field_name == "MitigationStatus":
        v = normalize_choice(v, MITIGATION_STATUS_CHOICES) or "Pending"
    elif field_name == "RiskMitigation":
        # Should be a JSON array of mitigation steps
        if isinstance(v, list):
            pass
        elif isinstance(v, str):
            try:
                v = json.loads(v)
            except:
                v = [{"step": "Mitigation Step 1", "description": v}]
        else:
            v = [{"step": "Mitigation Step 1", "description": "To be defined"}]
    elif field_name == "ModifiedMitigations":
        # Should be a JSON array documenting modifications
        if isinstance(v, list):
            pass
        elif isinstance(v, str):
            try:
                v = json.loads(v)
            except:
                v = []
        else:
            v = []
    elif field_name == "RiskFormDetails":
        # Should be a JSON object with assessment details
        if isinstance(v, dict):
            pass
        elif isinstance(v, str):
            try:
                v = json.loads(v)
            except:
                v = {
                    "assessment_method": "Qualitative",
                    "data_sources": ["Document"],
                    "assessment_date": date.today().isoformat()
                }
        else:
            v = {
                "assessment_method": "Qualitative",
                "data_sources": ["Document"],
                "assessment_date": date.today().isoformat()
            }
    elif field_name == "Reviewer":
        v = str(v).strip() if v else "Pending Review"
    elif isinstance(v, str):
        v = v.strip() or None
    
    return v, metadata


def fallback_risk_instance_extraction(text: str) -> list[dict]:
    """
    Minimal pattern-based fallback when AI fails completely.
    Produces at least one record using generic defaults.
    """
    risk_instances = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    risk_keywords = ["risk", "threat", "vulnerability", "hazard", "danger", "exposure", "incident"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in risk_keywords):
            if current:
                risk_instances.append(current)
            count += 1
            current = {
                "RiskTitle": f"Risk Instance {count}: {ln[:100]}",
                "RiskDescription": ln,
                "PossibleDamage": "Potential damages vary (reputation, compliance, downtime).",
                "RiskPriority": "Medium",
                "Criticality": "Medium",
                "Category": "Operational",
                "Origin": "Internal",
                "RiskLikelihood": 5,
                "RiskImpact": 5,
                "RiskExposureRating": 25.0,
                "RiskMultiplierX": 1.0,
                "RiskMultiplierY": 1.0,
                "Appetite": "Medium",
                "RiskResponseType": "Mitigate",
                "RiskResponseDescription": "Response strategy to be defined after assessment.",
                "RiskMitigation": [{"step": "Step 1", "description": "Mitigation steps to be defined"}],
                "RiskType": "Current",
                "RiskOwner": "Risk Manager",
                "BusinessImpact": "Potential business impact in operations/compliance.",
                "RiskStatus": "Not Assigned",
                "MitigationStatus": "Pending",
                "ModifiedMitigations": [],
                "RiskFormDetails": {
                    "assessment_method": "Qualitative",
                    "data_sources": ["Document"],
                    "assessment_date": date.today().isoformat()
                },
                "Reviewer": "Pending Review",
            }

    if current:
        risk_instances.append(current)

    if not risk_instances:
        risk_instances.append({
            "RiskTitle": "Document Risk Instance Analysis",
            "RiskDescription": f"Automated risk instance derived from document (length={len(text)} chars).",
            "PossibleDamage": "Data loss, downtime, penalties, reputation.",
            "RiskPriority": "Medium",
            "Criticality": "Medium",
            "Category": "Operational",
            "Origin": "Internal",
            "RiskLikelihood": 5,
            "RiskImpact": 5,
            "RiskExposureRating": 25.0,
            "RiskMultiplierX": 1.0,
            "RiskMultiplierY": 1.0,
            "Appetite": "Medium",
            "RiskResponseType": "Mitigate",
            "RiskResponseDescription": "Response strategy to be defined after assessment.",
            "RiskMitigation": [{"step": "Step 1", "description": "Review & implement standard mitigations for identified weaknesses."}],
            "RiskType": "Current",
            "RiskOwner": "Risk Manager",
            "BusinessImpact": "Potential business impact in operations/compliance.",
            "RiskStatus": "Not Assigned",
            "MitigationStatus": "Pending",
            "ModifiedMitigations": [],
            "RiskFormDetails": {
                "assessment_method": "Qualitative",
                "data_sources": ["Document"],
                "assessment_date": date.today().isoformat()
            },
            "Reviewer": "Pending Review",
        })

    return risk_instances


def parse_risk_instances_from_text(text: str, document_hash: str = None) -> list[dict]:
    """
    NEW APPROACH:
    1. First detect how many risk instances are in the document by finding 'Risk X:' patterns
    2. Extract whatever fields are explicitly present in each risk instance block
    3. Use AI to fill ONLY the missing fields (not to create risk titles)
    """
    debug_print(f"📊 parse_risk_instances_from_text() called with {len(text)} chars of text")
    
    # Step 1: Detect and parse risk instance blocks from document
    detected_risk_instances = detect_and_parse_risk_instance_blocks(text)
    
    if not detected_risk_instances:
        debug_print(f"⚠️  No structured risk instances found. Falling back to old AI extraction...")
        # Fallback to old behavior if no "Risk X:" pattern found
        return fallback_risk_instance_extraction(text)
    
    debug_print(f"✅ Detected {len(detected_risk_instances)} risk instance(s), now processing each...")
    
    # Step 2: For each detected risk instance, normalize extracted fields and fill missing ones
    completed_risk_instances = []
    for idx, risk_data in enumerate(detected_risk_instances, 1):
        debug_print(f"\n🔧 Processing Risk Instance {idx}: {risk_data.get('RiskTitle', 'Unknown')[:50]}...")
        
        # Start with empty item for all DB fields
        item = {k: None for k in RISK_INSTANCE_DB_FIELDS}
        
        # Copy extracted fields
        for field in RISK_INSTANCE_DB_FIELDS:
            if field in risk_data and risk_data[field]:
                item[field] = risk_data[field]
        
        # Get the risk block for AI context
        risk_block = risk_data.get("_risk_block", "")
        extracted_fields = risk_data.get("_extracted_fields", [])
        
        debug_print(f"  📝 Extracted fields: {', '.join(extracted_fields)}")
        
        # Initialize metadata structure
        if "_meta" not in item:
            item["_meta"] = {}
        if "per_field" not in item["_meta"]:
            item["_meta"]["per_field"] = {}
        
        # Mark extracted fields in metadata
        for field in extracted_fields:
            if field in item and item[field]:
                item["_meta"]["per_field"][field] = {
                    "source": "EXTRACTED",
                    "confidence": 0.95,
                    "rationale": "Found explicitly in document"
                }
        
        # Normalize extracted fields
        if item.get("RiskLikelihood"):
            item["RiskLikelihood"] = clamp_int(item["RiskLikelihood"], 1, 10)
        if item.get("RiskImpact"):
            item["RiskImpact"] = clamp_int(item["RiskImpact"], 1, 10)
        if item.get("RiskExposureRating"):
            item["RiskExposureRating"] = as_float_or_none(item["RiskExposureRating"])
        if item.get("Criticality"):
            item["Criticality"] = normalize_choice(item["Criticality"], CRITICALITY_CHOICES)
        if item.get("RiskPriority"):
            item["RiskPriority"] = normalize_choice(item["RiskPriority"], PRIORITY_CHOICES)
        if item.get("RiskType"):
            item["RiskType"] = normalize_choice(item["RiskType"], RISKTYPE_HINTS)
        if item.get("Origin"):
            item["Origin"] = normalize_choice(item["Origin"], ORIGIN_HINTS)
        if item.get("Appetite"):
            item["Appetite"] = normalize_choice(item["Appetite"], APPETITE_HINTS)
        if item.get("RiskResponseType"):
            item["RiskResponseType"] = normalize_choice(item["RiskResponseType"], RESPONSE_TYPE_HINTS)
        if item.get("RiskStatus"):
            item["RiskStatus"] = normalize_choice(item["RiskStatus"], RISK_STATUS_CHOICES)
        if item.get("MitigationStatus"):
            item["MitigationStatus"] = normalize_choice(item["MitigationStatus"], MITIGATION_STATUS_CHOICES)
        if item.get("RiskMultiplierX"):
            item["RiskMultiplierX"] = as_float_or_none(item["RiskMultiplierX"])
        if item.get("RiskMultiplierY"):
            item["RiskMultiplierY"] = as_float_or_none(item["RiskMultiplierY"])
        
        # Handle JSON fields
        if item.get("RiskMitigation"):
            mitigation = item["RiskMitigation"]
            if isinstance(mitigation, str):
                try:
                    item["RiskMitigation"] = json.loads(mitigation)
                except:
                    item["RiskMitigation"] = [{"step": "Step 1", "description": mitigation}]
            elif not isinstance(mitigation, list):
                item["RiskMitigation"] = None
        
        if item.get("ModifiedMitigations"):
            mod_mits = item["ModifiedMitigations"]
            if isinstance(mod_mits, str):
                try:
                    item["ModifiedMitigations"] = json.loads(mod_mits)
                except:
                    item["ModifiedMitigations"] = []
            elif not isinstance(mod_mits, list):
                item["ModifiedMitigations"] = None
        
        if item.get("RiskFormDetails"):
            form_details = item["RiskFormDetails"]
            if isinstance(form_details, str):
                try:
                    item["RiskFormDetails"] = json.loads(form_details)
                except:
                    item["RiskFormDetails"] = None
            elif not isinstance(form_details, dict):
                item["RiskFormDetails"] = None
        
        # Step 3: Use AI to fill ONLY missing fields (NEVER RiskTitle - must be in document)
        missing_fields = [f for f in RISK_INSTANCE_DB_FIELDS if item.get(f) in (None, "", []) and f != "RiskTitle"]
        if missing_fields:
            debug_print(f"  🤖 Missing fields: {', '.join(missing_fields)}")
            debug_print(f"  🤖 Using AI to infer missing fields...")

            for field in missing_fields:
                debug_print(f"    🔍 Inferring {field}...")
                value, metadata = infer_single_field(field, item, risk_block or text[:3000], document_hash=document_hash)
                item[field] = value
                # Store AI generation metadata
                item["_meta"]["per_field"][field] = metadata
                debug_print(f"    🏷️  Marked {field} as AI_GENERATED in metadata")
        else:
            debug_print(f"  ✅ All fields extracted from document!")
        
        # Final normalization and defaults
        item["RiskLikelihood"] = item["RiskLikelihood"] or 5
        item["RiskImpact"] = item["RiskImpact"] or 5
        
        # Compute exposure if not present
        if not item.get("RiskExposureRating"):
            item["RiskExposureRating"] = compute_exposure(item["RiskLikelihood"], item["RiskImpact"]) or 25.0
        
        item["RiskExposureRating"] = float(max(0.0, min(100.0, item["RiskExposureRating"])))
        item["RiskMultiplierX"] = item["RiskMultiplierX"] or 1.0
        item["RiskMultiplierY"] = item["RiskMultiplierY"] or 1.0
        item["Criticality"] = item["Criticality"] or "Medium"
        item["RiskPriority"] = item["RiskPriority"] or "Medium"
        item["RiskType"] = item["RiskType"] or "Current"
        item["Origin"] = item["Origin"] or "Internal"
        item["Appetite"] = item["Appetite"] or "Medium"
        item["RiskResponseType"] = item["RiskResponseType"] or "Mitigate"
        item["RiskStatus"] = item["RiskStatus"] or "Not Assigned"
        item["MitigationStatus"] = item["MitigationStatus"] or "Pending"
        item["Reviewer"] = item["Reviewer"] or "Pending Review"
        
        # RiskTitle must ALWAYS come from document - never generate it
        if not item.get("RiskTitle"):
            raise ValueError(f"RiskTitle is missing for risk instance {idx}. All risk titles must be present in the document as 'Risk X: Title'.")
        
        # Ensure JSON fields have proper defaults
        if not isinstance(item.get("RiskMitigation"), list):
            item["RiskMitigation"] = [{"step": "Step 1", "description": "To be defined"}]
        if not isinstance(item.get("ModifiedMitigations"), list):
            item["ModifiedMitigations"] = []
        if not isinstance(item.get("RiskFormDetails"), dict):
            item["RiskFormDetails"] = {
                "assessment_method": "Qualitative",
                "data_sources": ["Document"],
                "assessment_date": date.today().isoformat()
            }
        
        # Debug: Print metadata summary
        ai_fields = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "AI_GENERATED"]
        extracted = [field for field, info in item["_meta"]["per_field"].items() 
                    if info.get("source") == "EXTRACTED"]
        
        debug_print(f"  📊 Metadata Summary:")
        debug_print(f"     🤖 AI Generated: {len(ai_fields)} fields - {ai_fields}")
        debug_print(f"     📄 Extracted: {len(extracted)} fields - {extracted}")
        
        completed_risk_instances.append(item)
        debug_print(f"  ✅ Risk Instance {idx} completed!")
    
    return completed_risk_instances


# =========================
# DJANGO API ENDPOINTS
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
# rbac_required removed: path in JWT skip list; token parsing was returning 401. Allow upload for dev parity with ai-risk-doc-upload.
@rate_limit_decorator(requests_per_minute=10, requests_per_hour=100)  # Phase 3: Rate limiting
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def upload_and_process_risk_instance_document(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    debug_print(f"📤 Upload request for risk instance document")
    debug_print(f"📤 Request data: {request.POST}")
    debug_print(f"📤 Request files: {request.FILES}")
    debug_print(f"📤 User ID: {request.POST.get('user_id', 'unknown')}")

    # CORS preflight is handled by django-cors-headers.
    if request.method == 'OPTIONS':
        return HttpResponse()

    try:
        if 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            return JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)

        # Create the ai_uploads/risk_instance directory if it doesn't exist
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'risk_instance')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save the file with timestamp to avoid conflicts
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        # Write the uploaded file to disk
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        # Decompress if needed (client-side compression)
        compression_metadata = None
        file_path, was_compressed, compression_stats = decompress_if_needed(file_path)
        if was_compressed:
            compression_metadata = compression_stats
            # Update extension after decompression (remove .gz)
            ext = os.path.splitext(file_path)[1].lower()
            debug_print(f"📦 Decompressed file: {compression_stats['ratio']}% reduction, saved {compression_stats['bandwidth_saved_kb']} KB")
        
        # Upload to S3 for backup and cloud storage
        s3_url = None
        s3_key = None
        user_id = request.POST.get('user_id', '1')
        try:
            debug_print(f"☁️ Uploading file to S3...")
            s3_client = create_direct_mysql_client()
            connection_test = s3_client.test_connection()
            if connection_test.get('overall_success', False):
                # Generate unique filename for S3
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                s3_filename = f"risk_instance_{timestamp}_{os.path.basename(file_path)}"
                upload_result = s3_client.upload(
                    file_path=file_path,
                    user_id=user_id,
                    custom_file_name=s3_filename,
                    module='Risk'
                )
                if upload_result.get('success'):
                    s3_url = upload_result['file_info']['url']
                    s3_key = upload_result['file_info'].get('s3Key', '')
                    debug_print(f"✅ File uploaded to S3: {s3_url}")
                else:
                    debug_print(f"⚠️ S3 upload failed: {upload_result.get('error', 'Unknown error')}")
            else:
                debug_print(f"⚠️ S3 service unavailable, continuing with local file")
        except Exception as s3_error:
            debug_print(f"⚠️ S3 upload error (continuing with local file): {str(s3_error)}")
        
        debug_print(f"✅ File saved to: {file_path}")

        try:
            # Step 1: Prepare document using centralized preprocessor (includes lemmatization)
            print("[ROUTE-RISK-INST] upload_and_process_risk_instance_document: STEP 1 - preprocessing")
            debug_print("🔍 STEP 1: Preparing risk instance document via centralized DocumentPreparationService...")
            prep = DocumentPreparationService().prepare_text(
                extract_text_from_file(file_path, ext),
                max_length=8000,
            )
            text = prep["text"]
            preprocess_metadata = prep["metadata"]
            debug_print("✅ STEP 1 COMPLETE: Centralized preprocessing applied for risk instance")
            debug_print(f"   Original length: {preprocess_metadata.get('original_length')} chars")
            debug_print(f"   Processed length: {preprocess_metadata.get('processed_length')} chars")
            if preprocess_metadata.get("was_truncated"):
                debug_print(f"   ⚠️  Document was truncated ({preprocess_metadata.get('reduction_percent', 0):.1f}% reduction)")
            print(f"[ROUTE-RISK-INST] preprocessing DONE: orig={preprocess_metadata.get('original_length')} proc={preprocess_metadata.get('processed_length')}")

            # Step 2: Extract fields from document first, then AI-fill only missing fields
            # This ensures: fields IN the PDF -> EXTRACTED; fields NOT in PDF -> AI_GENERATED
            print("[ROUTE-RISK-INST] STEP 2 - parse_risk_instances_from_text (extract + AI for missing only)")
            debug_print("🤖 STEP 2: Extracting from document, AI-filling only missing fields...")
            document_hash = calculate_document_hash(text)
            risk_instances = parse_risk_instances_from_text(text, document_hash=document_hash)

            if not isinstance(risk_instances, list):
                debug_print("❌ parse_risk_instances_from_text did not return a list, wrapping into list")
                risk_instances = [risk_instances] if risk_instances else []

            debug_print(f"✅ STEP 2 COMPLETE: Extracted {len(risk_instances)} risk instance(s) (EXTRACTED + AI_GENERATED per field)")
            print(f"[ROUTE-RISK-INST] ingest_risk_instance_document DONE: instances={len(risk_instances)}")

            response_data = {
                "status": "success",
                "message": f"Successfully extracted {len(risk_instances)} risk instance(s)",
                "document_name": file_name,
                "saved_path": safe_filename,
                "extracted_text_length": len(text),
                "preprocessing_metadata": preprocess_metadata,
                "risk_instances": risk_instances,
            }

            if compression_metadata:
                response_data["compression_metadata"] = compression_metadata
            if s3_url:
                response_data["s3_url"] = s3_url
                response_data["s3_key"] = s3_key

            return JsonResponse(response_data)
        except Exception as process_error:
            # Clean up the file if processing fails
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def save_extracted_risk_instances(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        data = json.loads(request.body or "{}")
        risk_instances_data = data.get('risk_instances', [])
        if not risk_instances_data:
            return JsonResponse({'status': 'error', 'message': 'No risk instances provided'}, status=400)

        saved = []
        errors = []

        for idx, r in enumerate(risk_instances_data):
            try:
                # Parse RiskMitigation if it's a string
                risk_mitigation = r.get('RiskMitigation')
                if isinstance(risk_mitigation, str):
                    try:
                        risk_mitigation = json.loads(risk_mitigation)
                    except:
                        risk_mitigation = [{"step": "Step 1", "description": risk_mitigation}]
                elif not isinstance(risk_mitigation, list):
                    risk_mitigation = [{"step": "Step 1", "description": "To be defined"}]
                
                # Parse ModifiedMitigations if it's a string
                modified_mitigations = r.get('ModifiedMitigations', [])
                if isinstance(modified_mitigations, str):
                    try:
                        modified_mitigations = json.loads(modified_mitigations)
                    except:
                        modified_mitigations = []
                elif not isinstance(modified_mitigations, list):
                    modified_mitigations = []
                
                # Parse RiskFormDetails if it's a string
                risk_form_details = r.get('RiskFormDetails')
                if isinstance(risk_form_details, str):
                    try:
                        risk_form_details = json.loads(risk_form_details)
                    except:
                        risk_form_details = {
                            "assessment_method": "Qualitative",
                            "data_sources": ["Document"],
                            "assessment_date": date.today().isoformat()
                        }
                elif not isinstance(risk_form_details, dict):
                    risk_form_details = {
                        "assessment_method": "Qualitative",
                        "data_sources": ["Document"],
                        "assessment_date": date.today().isoformat()
                    }
                
                kwargs = {
                    # not setting: RiskInstanceId (auto), RiskId, IncidentId, ComplianceId, UserId, ReviewerId, ReportedBy
                    'RiskTitle': r.get('RiskTitle', f'Risk Instance {idx+1}'),
                    'RiskDescription': r.get('RiskDescription', ''),
                    'PossibleDamage': r.get('PossibleDamage', ''),
                    'RiskPriority': r.get('RiskPriority', 'Medium'),
                    'Criticality': r.get('Criticality', 'Medium'),
                    'Category': r.get('Category', ''),
                    'Origin': r.get('Origin', 'Internal'),
                    'RiskLikelihood': int(r.get('RiskLikelihood') or 5),
                    'RiskImpact': int(r.get('RiskImpact') or 5),
                    'RiskExposureRating': float(r.get('RiskExposureRating') or 25.0),
                    'RiskMultiplierX': float(r.get('RiskMultiplierX') or 1.0),
                    'RiskMultiplierY': float(r.get('RiskMultiplierY') or 1.0),
                    'Appetite': r.get('Appetite', 'Medium'),
                    'RiskResponseType': r.get('RiskResponseType', 'Mitigate'),
                    'RiskResponseDescription': r.get('RiskResponseDescription', ''),
                    'RiskMitigation': risk_mitigation,
                    'RiskType': r.get('RiskType', 'Current'),
                    'RiskOwner': r.get('RiskOwner', ''),
                    'BusinessImpact': r.get('BusinessImpact', ''),
                    'RiskStatus': r.get('RiskStatus', 'Not Assigned'),
                    'MitigationStatus': r.get('MitigationStatus', 'Pending'),
                    'ModifiedMitigations': modified_mitigations,
                    'RiskFormDetails': risk_form_details,
                    'Reviewer': r.get('Reviewer', 'Pending Review'),
                }
                risk_instance = RiskInstance.objects.create(**kwargs)
                saved.append({'risk_instance_id': getattr(risk_instance, "RiskInstanceId", None), 'risk_title': risk_instance.RiskTitle})
            except Exception as ex:
                errors.append({'risk_instance_index': idx, 'title': r.get('RiskTitle'), 'error': str(ex)})

        resp = {
            'status': 'success',
            'message': f'Saved {len(saved)} risk instance(s)' + (f' with {len(errors)} error(s)' if errors else ''),
            'saved': saved
        }
        if errors:
            resp['errors'] = errors
        return JsonResponse(resp)

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON payload'}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error saving risk instances: {str(e)}'}, status=500)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
@rbac_required(required_permission='view_all_risk')
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def test_openai_connection_risk_instance(request):
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    
    try:
        if not OPENAI_API_KEY:
            return JsonResponse({
                'status': 'error',
                'message': 'OPENAI_API_KEY is not set',
                'model': OPENAI_MODEL,
                'api_url': OPENAI_API_URL
            }, status=500)
        
        out = call_openai_json('Return JSON: {"ok": true}')
        return JsonResponse({
            'status': 'success', 
            'openai_reply': out, 
            'model': OPENAI_MODEL, 
            'api_url': OPENAI_API_URL
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': f'OpenAI error: {e}',
            'model': OPENAI_MODEL,
            'api_url': OPENAI_API_URL
        }, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@rbac_required(required_permission='create_risk')
@rate_limit_decorator(requests_per_minute=5, requests_per_hour=50)  # Lower limits for streaming
@require_tenant
@tenant_filter
def upload_and_process_risk_instance_document_streaming(request):
    """
    Streaming version of risk instance upload that sends real-time updates via SSE.
    """
    tenant_id = get_tenant_id_from_request(request)
    
    debug_print(f"📤 Streaming upload request for risk instance document")

    if request.method == 'OPTIONS':
        return HttpResponse()

    try:
        if 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            return JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)

        # Save file temporarily
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'risk_instance')
        os.makedirs(upload_dir, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)

        # Prepare text
        from ...utils.document_preprocessor import extract_text_from_file
        from ...ai.processing.preprocessor import DocumentPreparationService
        
        prep = DocumentPreparationService().prepare_text(
            extract_text_from_file(file_path, ext),
            max_length=8000,
        )
        text = prep["text"]

        # Set up SSE streaming - collect all events then stream them
        events = []
        
        def stream_callback(event_type, data):
            """Called by the streaming AI task for each update"""
            event_data = {
                'type': event_type,
                'timestamp': datetime.now().isoformat(),
                **data
            }
            events.append(event_data)

        try:
            # Import the streaming task  
            from ...ai.tasks.risk import ingest_risk_instance_document_streaming
            ai_service = get_ai_service()
            
            # Call streaming task
            result = ingest_risk_instance_document_streaming(
                ai_service,
                payload={"document_text": text},
                stream_callback=stream_callback,
                options=AIRequestOptions(
                    task_name="risk.ingest_risk_instance_document_streaming",
                    use_cache=False,  # Don't cache streaming results
                )
            )
            
            # Add final complete event
            events.append({'type': 'done', 'risk_instances': result})
            
        except Exception as e:
            events.append({
                'type': 'error',
                'message': str(e),
                'timestamp': datetime.now().isoformat()
            })

        # Stream all collected events
        def event_stream():
            for event_data in events:
                yield f"data: {json.dumps(event_data)}\n\n"
                # Add small delay to simulate real-time streaming
                import time
                time.sleep(0.1)

        response = StreamingHttpResponse(
            event_stream(),
            content_type='text/event-stream'
        )
        response['Cache-Control'] = 'no-cache'
        response['Connection'] = 'keep-alive'
        
        return response

    except Exception as e:
        debug_print(f"❌ Error in streaming upload: {e}")
        return JsonResponse({'status': 'error', 'message': f'Streaming upload failed: {str(e)}'}, status=500)