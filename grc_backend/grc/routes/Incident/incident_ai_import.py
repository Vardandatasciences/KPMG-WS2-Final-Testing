"""
AI-Powered Incident Document Ingestion
- Reads PDF/DOCX/XLSX/TXT
- Extracts incident data using OpenAI GPT models
- Fills missing fields with focused prompts
- Returns normalized, DB-ready JSON for `incidents` table
"""

import os
import re
import json
import time
from datetime import date, datetime
from typing import Any, Optional

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from rest_framework.decorators import api_view, permission_classes, parser_classes
from rest_framework.permissions import AllowAny
from rest_framework.parsers import MultiPartParser, FormParser
import requests

# Centralized AI imports
from ...ai.service import get_ai_service
from ...ai.processing.preprocessor import DocumentPreparationService
from ...utils.file_compression import decompress_if_needed
from ...routes.Global.s3_fucntions import create_direct_mysql_client
from ...debug_utils import debug_print

# --- Optional parsers (install as needed) ---
try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import pandas as pd
except Exception:
    pd = None

# --- Your models ---
from grc.models import Incident

# MULTI-TENANCY: Import tenant utilities for data isolation
from ...tenant_utils import (
    require_tenant, tenant_filter, get_tenant_id_from_request,
    validate_tenant_access, get_tenant_aware_queryset
)


# Initialize centralized AI service
ai_service = get_ai_service()

debug_print("\n🤖 Incident Import AI - Using Centralized AI Module")
debug_print(f"   AI Service initialized: {ai_service is not None}")
debug_print(f"   Centralized preprocessing, lemmatization, and optimization: enabled")

# Fields we want to extract from incidents table (excluding ID and date fields)
INCIDENT_DB_FIELDS = [
    "IncidentTitle",
    "Description",
    "Mitigation",
    "Origin",
    "Comments",
    "RiskCategory",
    "IncidentCategory",
    "RiskPriority",
    "Attachments",
    "Status",
    "RepeatedNot",
    "CostOfIncident",
    "ReopenedNot",
    "RejectionSource",
    "AffectedBusinessUnit",
    "SystemsAssetsInvolved",
    "GeographicLocation",
    "Criticality",
    "InitialImpactAssessment",
    "InternalContacts",
    "ExternalPartiesInvolved",
    "RegulatoryBodies",
    "RelevantPoliciesProceduresViolated",
    "ControlFailures",
    "LessonsLearned",
    "IncidentClassification",
    "PossibleDamage",
    "IncidentFormDetails",
]

# Canonical choices / constraints to stabilize LLM outputs
CRITICALITY_CHOICES = ["Low", "Medium", "High", "Critical"]
PRIORITY_CHOICES = ["Low", "Medium", "High", "Critical"]
STATUS_CHOICES = ["New", "In Progress", "Under Investigation", "Resolved", "Closed", "Escalated", "Risk Mitigated"]
ORIGIN_CHOICES = ["AUDIT_FINDING", "MANUAL", "AUTOMATED", "EXTERNAL_REPORT", "INTERNAL_DETECTION"]
REJECTION_SOURCE_CHOICES = ["INCIDENT", "RISK"]
INCIDENT_CATEGORY_HINTS = [
    "Security Breach", "Data Loss", "System Outage", "Compliance Violation", 
    "Operational Failure", "Third-Party Issue", "Human Error", "Natural Disaster",
    "Cyber Attack", "Privacy Incident", "Safety Incident", "Financial Loss"
]
RISK_CATEGORY_HINTS = [
    "Operational", "Financial", "Strategic", "Compliance", "Technical",
    "Reputational", "Information Security", "Process Risk", "Third-Party",
    "Regulatory", "Governance"
]

# Legacy field prompts - now managed by centralized AI tasks

# Strict JSON schema block
STRICT_SCHEMA_BLOCK = f"""
CRITICAL: Return ONLY a valid JSON array. No markdown, no code blocks, no explanations.
Start with [ and end with ]. Use proper JSON syntax with double quotes.

Example structure (return array of incidents like this):
[
  {{
    "IncidentTitle": "Brief incident title here",
    "Description": "Detailed description",
    "Mitigation": [{{"step": "Action taken", "status": "Completed", "responsible": "Team", "deadline": "2024-01-01"}}],
    "Origin": "MANUAL",
    "Comments": "Additional context",
    "RiskCategory": "Operational",
    "IncidentCategory": "Security Breach",
    "RiskPriority": "High",
    "Attachments": "",
    "Status": "In Progress",
    "RepeatedNot": false,
    "CostOfIncident": "$50,000",
    "ReopenedNot": false,
    "RejectionSource": null,
    "AffectedBusinessUnit": "IT Department",
    "SystemsAssetsInvolved": "Production Server, Database",
    "GeographicLocation": "New York Office",
    "Criticality": "High",
    "InitialImpactAssessment": "Initial assessment details",
    "InternalContacts": "John Smith (IT Manager)",
    "ExternalPartiesInvolved": "",
    "RegulatoryBodies": "",
    "RelevantPoliciesProceduresViolated": "Security Policy Section 4.2",
    "ControlFailures": "Firewall misconfiguration",
    "LessonsLearned": "Key insights from incident",
    "IncidentClassification": "P1: Critical",
    "PossibleDamage": "Service outage, revenue loss",
    "IncidentFormDetails": {{"reported_by": "Security Team", "detection_method": "Automated", "response_time_minutes": 30}},
    "_meta": {{
      "per_field": {{
        "IncidentTitle": {{"source": "EXTRACTED", "confidence": 0.9, "rationale": "Found in document"}},
        "Description": {{"source": "AI_GENERATED", "confidence": 0.7, "rationale": "Inferred from context"}}
      }}
    }}
  }}
]

Rules:
- Criticality must be: Low, Medium, High, or Critical
- RiskPriority must be: Low, Medium, High, or Critical
- Status must be one of: {STATUS_CHOICES}
- Origin must be one of: {ORIGIN_CHOICES}
- RepeatedNot and ReopenedNot must be boolean (true/false)
- Mitigation must be JSON array
- IncidentFormDetails must be JSON object
- NO trailing commas
- NO comments in JSON
- Return ONLY the JSON array, nothing else
"""


# =========================
# UTILITIES / VALIDATORS
# =========================
def _json_from_llm_text(text: str) -> Any:
    """
    Extract the first valid JSON array/object from the LLM response.
    Handles malformed JSON, trailing commas, and extracts nested values.
    Reuses the improved version from risk_ai_doc.py
    """
    # Import the improved version from risk_ai_doc
    from ..Risk.risk_ai_doc import _json_from_llm_text as improved_json_parser
    return improved_json_parser(text)


# Note: call_ollama_json and call_openai_json are now imported from risk_ai_doc.py
# They include Phase 2 caching and Phase 3 optimizations


def normalize_choice(val: str, choices: list[str]) -> Optional[str]:
    if not val: return None
    v = str(val).strip()
    for c in choices:
        if v.lower() == c.lower():
            return c
    for c in choices:
        if v.lower().startswith(c.lower()[0:3]):
            return c
    return None


def as_boolean(v) -> bool:
    """Convert various inputs to boolean."""
    if isinstance(v, bool):
        return v
    if isinstance(v, str):
        return v.lower() in ('true', 'yes', '1', 't', 'y')
    return bool(v)


# =========================
# FILE EXTRACTORS
# =========================
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF."""
    try:
        if pdfplumber:
            parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    if t.strip():
                        parts.append(t)
            return "\n".join(parts)
        elif PyPDF2:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for p in reader.pages:
                    text += (p.extract_text() or "") + "\n"
                return text
        return ""
    except Exception as e:
        debug_print(f"[PDF] Extraction error: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        if not docx:
            raise RuntimeError("python-docx not installed")
        d = docx.Document(file_path)
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for t in d.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except Exception as e:
        debug_print(f"[DOCX] Extraction error: {e}")
        return ""


def extract_text_from_excel(file_path: str) -> str:
    try:
        if not pd:
            raise RuntimeError("pandas/openpyxl not installed")
        df_sheets = pd.read_excel(file_path, sheet_name=None)
        out = []
        for name, df in df_sheets.items():
            out.append(f"=== Sheet: {name} ===")
            out.append(df.to_string(index=False))
        return "\n".join(out)
    except Exception as e:
        debug_print(f"[XLSX] Extraction error: {e}")
        return ""


def extract_text_from_file(file_path: str, file_extension: str) -> str:
    ext = file_extension.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    if ext in [".xlsx", ".xls"]:
        return extract_text_from_excel(file_path)
    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    return ""


# =========================
# AI EXTRACTION CORE
# =========================
def infer_single_field(field_name: str, current_record: dict, document_context: str,
                       document_hash: str = None) -> tuple:
    """
    Focused prompt for ONE field using centralized AI service.
    Returns: (value, metadata_dict) for per_field tracking.
    """
    print(f"[INFER-FIELD] Calling AI for field: {field_name}")
    metadata = {"source": "AI_GENERATED", "confidence": 0.8, "rationale": "Inferred from document by AI field-by-field extraction."}

    try:
        result = ai_service.run_task(
            "incident.infer_field",
            payload={
                "field_name": field_name,
                "document_text": document_context[:3000],
                "current_record": current_record
            },
            metadata={"document_hash": document_hash}
        )
        print(f"[INFER-FIELD] AI returned for {field_name}: {repr(result)[:80]}")
        return (result, metadata)
    except Exception as e:
        print(f"[INFER-FIELD] AI FAILED for {field_name}: {e}")
        if field_name in ("RepeatedNot", "ReopenedNot"):
            return (False, {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default (AI failed)"})
        if field_name == "Criticality":
            return ("Medium", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default criticality"})
        if field_name == "RiskPriority":
            return ("Medium", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default priority"})
        if field_name == "Status":
            return ("New", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default status"})
        if field_name == "Origin":
            return ("MANUAL", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default origin"})
        if field_name == "IncidentCategory":
            return ("Operational Failure", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default category"})
        if field_name == "RiskCategory":
            return ("Operational", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default risk category"})
        if field_name in ("Mitigation", "IncidentFormDetails"):
            return ([] if field_name == "Mitigation" else {}, {"source": "DEFAULT", "confidence": 0.5, "rationale": "Empty default"})
        return (None, {"source": "DEFAULT", "confidence": 0.3, "rationale": "AI failed, no default"})


def fallback_incident_extraction(text: str) -> list[dict]:
    """Minimal pattern-based fallback when AI fails completely."""
    incidents = []
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    incident_keywords = ["incident", "breach", "outage", "failure", "violation", "attack"]
    current = None
    count = 0

    for i, ln in enumerate(lines):
        if any(k in ln.lower() for k in incident_keywords):
            if current:
                incidents.append(current)
            count += 1
            current = {
                "IncidentTitle": f"Incident {count}: {ln[:100]}",
                "Description": ln,
                "Mitigation": [],
                "Origin": "MANUAL",
                "Comments": "",
                "RiskCategory": "Operational",
                "IncidentCategory": "Operational Failure",
                "RiskPriority": "Medium",
                "Attachments": "",
                "Status": "New",
                "RepeatedNot": False,
                "CostOfIncident": "Not assessed",
                "ReopenedNot": False,
                "RejectionSource": None,
                "AffectedBusinessUnit": "Unknown",
                "SystemsAssetsInvolved": "Unknown",
                "GeographicLocation": "Unknown",
                "Criticality": "Medium",
                "InitialImpactAssessment": "Requires investigation",
                "InternalContacts": "",
                "ExternalPartiesInvolved": "",
                "RegulatoryBodies": "",
                "RelevantPoliciesProceduresViolated": "",
                "ControlFailures": "",
                "LessonsLearned": "",
                "IncidentClassification": "Standard",
                "PossibleDamage": "To be assessed",
                "IncidentFormDetails": {},
            }

    if current:
        incidents.append(current)

    if not incidents:
        incidents.append({
            "IncidentTitle": "Document Incident Analysis",
            "Description": f"Automated incident derived from document (length={len(text)} chars).",
            "Mitigation": [],
            "Origin": "MANUAL",
            "Comments": "Extracted from uploaded document",
            "RiskCategory": "Operational",
            "IncidentCategory": "Operational Failure",
            "RiskPriority": "Medium",
            "Attachments": "",
            "Status": "New",
            "RepeatedNot": False,
            "CostOfIncident": "Not assessed",
            "ReopenedNot": False,
            "RejectionSource": None,
            "AffectedBusinessUnit": "To be determined",
            "SystemsAssetsInvolved": "To be determined",
            "GeographicLocation": "To be determined",
            "Criticality": "Medium",
            "InitialImpactAssessment": "Requires detailed assessment",
            "InternalContacts": "",
            "ExternalPartiesInvolved": "",
            "RegulatoryBodies": "",
            "RelevantPoliciesProceduresViolated": "",
            "ControlFailures": "",
            "LessonsLearned": "",
            "IncidentClassification": "Standard",
            "PossibleDamage": "To be assessed through investigation",
            "IncidentFormDetails": {},
        })

    return incidents


def detect_incident_blocks(text: str) -> tuple[int, list[str]]:
    """
    Preprocessing: detect how many incidents and split document.
    Strategy 1: Split by 'Incident N:' pattern (e.g. "Incident 1: Title", "Incident 2: Title").
    Strategy 2 (fallback): If repeated main-field markers (e.g. "1. Description:") appear,
       treat each as a new incident - split so title/description repetition = new incident.
    Returns (count, list of text segments). If no pattern found, returns (1, [full_text]).
    """
    if not text or not text.strip():
        return (0, [])

    # Strategy 1: "Incident 1:", "Incident 2:", ...
    pattern_incident_n = re.compile(r"Incident\s*\d+\s*:", re.IGNORECASE)
    matches = list(pattern_incident_n.finditer(text))
    if matches:
        segments = []
        for i, m in enumerate(matches):
            start = m.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            segment = text[start:end].strip()
            if segment:
                segments.append(segment)
        count = len(segments)
        print(f"[ROUTE-INCIDENT] PREPROCESS: Detected {count} incident block(s) via 'Incident N:' pattern")
        for i, seg in enumerate(segments, 1):
            first_line = seg.split("\n")[0][:80] if seg else ""
            print(f"[ROUTE-INCIDENT]   Block {i}: {first_line}...")
        return (count, segments)

    # Strategy 2: Repeated "1. Description:" or line-start "Description:" / "Title" = new incident
    # Match "1. Description:" or "Description:" at line start (main field repeated per incident)
    pattern_desc = re.compile(r"^(?:\d+\.\s*)?(?:Description|Title)\s*:\s*", re.IGNORECASE | re.MULTILINE)
    desc_matches = list(pattern_desc.finditer(text))
    if len(desc_matches) >= 2:
        segments = []
        for i in range(len(desc_matches)):
            start = desc_matches[i].start()
            end = desc_matches[i + 1].start() if i + 1 < len(desc_matches) else len(text)
            segment = text[start:end].strip()
            if segment:
                segments.append(segment)
        count = len(segments)
        print(f"[ROUTE-INCIDENT] PREPROCESS: Detected {count} incident block(s) via repeated 'Description/Title' pattern")
        for i, seg in enumerate(segments, 1):
            first_line = seg.split("\n")[0][:80] if seg else ""
            print(f"[ROUTE-INCIDENT]   Block {i}: {first_line}...")
        return (count, segments)

    print(f"[ROUTE-INCIDENT] PREPROCESS: No multi-incident pattern found -> treating as 1 incident")
    return (1, [text.strip()])


def parse_incidents_from_text(text: str, document_hash: str = None) -> list[dict]:
    """
    Extract ALL incident fields using centralized AI service.
    Step 1 (preprocess): Detect incident count and split by "Incident N:" if present.
    Step 2: Extract incidents (one doc per block when multiple blocks, else full doc).
    Step 3: Fill missing fields per incident.
    """
    print(f"[ROUTE-INCIDENT] parse_incidents_from_text: text_len={len(text)}")
    debug_print(f"📊 parse_incidents_from_text() called with {len(text)} chars of text")
    
    try:
        # Step 1: Preprocessing - identify how many incidents and split document
        incident_count, segments = detect_incident_blocks(text)
        if incident_count == 0:
            return []
        
        # Step 2: Extract incidents (one ingestion per segment when multiple)
        incidents = []
        if incident_count >= 2:
            for idx, segment in enumerate(segments, 1):
                print(f"[ROUTE-INCIDENT] Extracting incident {idx}/{incident_count} from block ({len(segment)} chars)")
                block_incidents = ai_service.run_task(
                    "incident.ingest_document",
                    payload={"document_text": segment, "single_incident_block": True},
                    metadata={"document_hash": document_hash, "block_index": idx}
                )
                if isinstance(block_incidents, list):
                    incidents.extend(block_incidents)
                elif isinstance(block_incidents, dict):
                    incidents.append(block_incidents)
            print(f"[ROUTE-INCIDENT] incident.ingest_document (per-block) DONE: {len(incidents)} incident(s)")
        else:
            debug_print(f"🚀 Calling centralized AI service to extract incidents (single doc)...")
            incidents = ai_service.run_task(
                "incident.ingest_document",
                payload={"document_text": text},
                metadata={"document_hash": document_hash}
            )
            if not isinstance(incidents, list):
                raise ValueError("Incidents must be a JSON array")
            print(f"[ROUTE-INCIDENT] incident.ingest_document DONE: incidents={len(incidents)} (will fill missing fields for each)")
        
        debug_print(f"✅ Centralized AI service returned {len(incidents)} incident(s)")
        
        # The centralized service handles field normalization and metadata
        # Just ensure backward compatibility with any missing fields
        cleaned = []
        for idx, inc in enumerate(incidents, 1):
            print(f"\n[ROUTE-INCIDENT] ========== Incident {idx}/{len(incidents)} ==========")
            debug_print(f"📋 Processing incident {idx}/{len(incidents)}")

            # Ensure all expected fields and metadata
            item = {}
            for field in INCIDENT_DB_FIELDS:
                item[field] = inc.get(field, "")
            item["_meta"] = inc.get("_meta", {})
            item["_meta"].setdefault("per_field", {})

            # Print: fields IN document (from ingest) vs fields AI NEEDS to generate
            empty_vals = (None, "", [], {})
            from_doc = [f for f in INCIDENT_DB_FIELDS if item.get(f) not in empty_vals]
            missing_fields = [f for f in INCIDENT_DB_FIELDS if item.get(f) in empty_vals]
            print(f"[ROUTE-INCIDENT] FROM DOCUMENT ({len(from_doc)}): {', '.join(from_doc)}")
            print(f"[ROUTE-INCIDENT] AI MUST GENERATE ({len(missing_fields)}): {', '.join(missing_fields)}")

            # Phase 2: Fill ALL missing fields with AI for this incident
            if missing_fields:
                # Use this incident's block as context when we split by "Incident N:"
                if incident_count >= 2 and idx <= len(segments):
                    doc_ctx = segments[idx - 1][:8000]
                else:
                    doc_ctx = text[:8000] if text else ""
                for i, field in enumerate(missing_fields, 1):
                    print(f"[ROUTE-INCIDENT] Generating {i}/{len(missing_fields)}: {field} ...")
                    try:
                        value, meta = infer_single_field(field, item, doc_ctx, document_hash=document_hash)
                        item[field] = value
                        item["_meta"]["per_field"][field] = meta
                        val_preview = str(value)[:60] + "..." if value is not None and len(str(value)) > 60 else str(value)
                        print(f"[ROUTE-INCIDENT]   -> {field} = {val_preview} (source={meta.get('source', '?')})")
                    except Exception as ex:
                        print(f"[ROUTE-INCIDENT]   -> {field} EXCEPTION: {ex}")
                        fallback_val, fallback_meta = (None, {"source": "DEFAULT", "confidence": 0.3, "rationale": str(ex)})
                        if field in ("RepeatedNot", "ReopenedNot"):
                            fallback_val, fallback_meta = (False, {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default"})
                        elif field in ("Criticality", "RiskPriority", "Status"):
                            fallback_val = "Medium" if field != "Status" else "New"
                            fallback_meta = {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default"}
                        elif field == "Origin":
                            fallback_val, fallback_meta = ("MANUAL", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default"})
                        elif field == "IncidentCategory":
                            fallback_val, fallback_meta = ("Operational Failure", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default"})
                        elif field == "RiskCategory":
                            fallback_val, fallback_meta = ("Operational", {"source": "DEFAULT", "confidence": 0.5, "rationale": "Default"})
                        elif field in ("Mitigation", "IncidentFormDetails"):
                            fallback_val = [] if field == "Mitigation" else {}
                            fallback_meta = {"source": "DEFAULT", "confidence": 0.5, "rationale": "Empty"}
                        item[field] = fallback_val
                        item["_meta"]["per_field"][field] = fallback_meta

            # Mark fields from initial ingest as EXTRACTED (from document), rest as AI_GENERATED
            for f in from_doc:
                if f in item["_meta"]["per_field"]:
                    item["_meta"]["per_field"][f]["source"] = "EXTRACTED"
                    item["_meta"]["per_field"][f]["rationale"] = "Extracted from document by initial AI ingestion"

            # Deduplicate per_field: keep only canonical INCIDENT_DB_FIELDS keys
            ALIAS_TO_CANONICAL = {"Possible Damage": "PossibleDamage", "Risk Priority": "RiskPriority", "Title": "IncidentTitle"}
            canonical_per_field = {}
            pf = item["_meta"]["per_field"]
            for f in INCIDENT_DB_FIELDS:
                if f in pf:
                    canonical_per_field[f] = pf[f]
                else:
                    for alias, canon in ALIAS_TO_CANONICAL.items():
                        if canon == f and alias in pf:
                            canonical_per_field[f] = pf[alias]
                            break
            item["_meta"]["per_field"] = canonical_per_field

            extracted_count = sum(1 for info in item["_meta"]["per_field"].values() if info.get("source") == "EXTRACTED")
            ai_count = sum(1 for info in item["_meta"]["per_field"].values() if info.get("source") == "AI_GENERATED")
            default_count = sum(1 for info in item["_meta"]["per_field"].values() if info.get("source") == "DEFAULT")
            print(f"[ROUTE-INCIDENT] Incident {idx} DONE: EXTRACTED={extracted_count}, AI_GENERATED={ai_count}, DEFAULT={default_count}")
            cleaned.append(item)
        
        return cleaned

    except Exception as e:
        debug_print(f"Centralized AI extraction failed, using fallback extractor: {e}")
        return fallback_incident_extraction(text)


# =========================
# DJANGO API ENDPOINTS
# =========================
@api_view(['POST'])
@permission_classes([AllowAny])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def upload_and_process_incident_document(request):
    """
    Upload and process incident document
    MULTI-TENANCY: Extracted incidents will be automatically assigned to user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    """
    Upload a document and process it to extract incident data.
    Phase 2: Uses document preprocessing, few-shot prompts, and caching.
    Phase 3: Uses RAG context retrieval, model routing, request queuing, and system load tracking.
    """
    debug_print(f"📤 Upload request for incident document")

    try:
        if 'file' not in request.FILES:
            return JsonResponse({'status': 'error', 'message': 'No file uploaded'}, status=400)

        uploaded_file = request.FILES['file']
        file_name = uploaded_file.name
        ext = os.path.splitext(file_name)[1].lower()

        allowed = ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.txt']
        if ext not in allowed:
            return JsonResponse({'status': 'error', 'message': f'Invalid file type. Allowed: {", ".join(allowed)}'}, status=400)

        # Create upload directory
        from django.conf import settings
        upload_dir = os.path.join(settings.MEDIA_ROOT, 'ai_uploads', 'incident')
        os.makedirs(upload_dir, exist_ok=True)
        
        # Save file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_filename = f"{timestamp}_{file_name}"
        file_path = os.path.join(upload_dir, safe_filename)
        
        with open(file_path, 'wb') as f:
            for chunk in uploaded_file.chunks():
                f.write(chunk)
        
        # Decompress if needed (client-side compression)
        compression_metadata = None
        file_path, was_compressed, compression_stats = decompress_if_needed(file_path)
        if was_compressed:
            compression_metadata = compression_stats
            # Update extension after decompression (remove .gz)
            ext = os.path.splitext(file_path)[1].lower()
            debug_print(f"📦 Decompressed file: {compression_stats['ratio']}% reduction, saved {compression_stats['bandwidth_saved_kb']} KB")
        
        # Upload to S3 for backup and cloud storage
        s3_url = None
        s3_key = None
        user_id = request.POST.get('user_id', '1')
        try:
            debug_print(f"☁️ Uploading file to S3...")
            s3_client = create_direct_mysql_client()
            connection_test = s3_client.test_connection()
            if connection_test.get('overall_success', False):
                # Generate unique filename for S3
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                s3_filename = f"incident_ai_{timestamp}_{os.path.basename(file_path)}"
                upload_result = s3_client.upload(
                    file_path=file_path,
                    user_id=user_id,
                    custom_file_name=s3_filename,
                    module='Incident'
                )
                if upload_result.get('success'):
                    s3_url = upload_result['file_info']['url']
                    s3_key = upload_result['file_info'].get('s3Key', '')
                    debug_print(f"✅ File uploaded to S3: {s3_url}")
                else:
                    debug_print(f"⚠️ S3 upload failed: {upload_result.get('error', 'Unknown error')}")
            else:
                debug_print(f"⚠️ S3 service unavailable, continuing with local file")
        except Exception as s3_error:
            debug_print(f"⚠️ S3 upload error (continuing with local file): {str(s3_error)}")
        
        debug_print(f"✅ File saved to: {file_path}")

        try:
            # Step 1: Extract text
            debug_print(f"🔍 STEP 1: Starting text extraction from {ext} file...")
            raw_text = extract_text_from_file(file_path, ext)
            
            if not raw_text or len(raw_text.strip()) < 50:
                return JsonResponse({'status': 'error', 'message': 'Could not extract meaningful text from document'}, status=400)

            debug_print(f"✅ STEP 1A COMPLETE: Extracted {len(raw_text)} characters from document")
            
            # Step 1B: Preprocess document using centralized service
            print("[ROUTE-INCIDENT] upload_and_process_incident_document: STEP 1B - preprocessing")
            debug_print(f"🔍 STEP 1B: Preprocessing document with centralized service...")
            prep_service = DocumentPreparationService()
            prep_result = prep_service.prepare_text(
                raw_text,
                max_length=8000,
            )
            text = prep_result.get("text", "")
            preprocess_metadata = prep_result.get("metadata", {})
            debug_print(f"✅ STEP 1B COMPLETE: Centralized preprocessing complete")
            debug_print(f"   Original length: {preprocess_metadata['original_length']} chars")
            debug_print(f"   Processed length: {preprocess_metadata['processed_length']} chars")
            debug_print(f"   Lemmatization applied: {preprocess_metadata.get('lemmatization_applied', False)}")
            if preprocess_metadata['was_truncated']:
                debug_print(f"   ⚠️  Document was truncated ({preprocess_metadata['reduction_percent']:.1f}% reduction)")
            
            # Document hash is included in preprocessing metadata
            document_hash = preprocess_metadata.get('document_hash')
            debug_print(f"📝 Document hash: {document_hash[:16] if document_hash else 'None'}... (for caching)")
            print(f"[ROUTE-INCIDENT] preprocessing DONE: orig={preprocess_metadata.get('original_length')} proc={preprocess_metadata.get('processed_length')}")
            
            # Step 2: Verify AI service is available
            debug_print(f"🔍 STEP 2: Checking centralized AI service...")
            if not ai_service:
                debug_print(f"❌ ERROR: Centralized AI service is not available")
                return JsonResponse({
                    'status': 'error', 
                    'message': 'Centralized AI service is not available. Please check AI module configuration.'
                }, status=503)
            
            debug_print(f"✅ STEP 2 COMPLETE: Centralized AI service is ready")
            
            # Step 3: Process with centralized AI service
            start_time = time.time()
            
            print("[ROUTE-INCIDENT] STEP 3 - calling parse_incidents_from_text (incident.ingest_document)")
            debug_print(f"🤖 STEP 3: Calling centralized AI service to extract incidents...")
            
            # Process document using centralized AI service
            incidents = parse_incidents_from_text(text, document_hash=document_hash)
            
            # Track processing time
            processing_time = time.time() - start_time
            
            debug_print(f"✅ STEP 3 COMPLETE: AI extracted {len(incidents)} incident(s) from document")
            print(f"[ROUTE-INCIDENT] upload_and_process DONE: incidents={len(incidents)}")
            for idx, incident in enumerate(incidents, 1):
                debug_print(f"  Incident {idx}: {incident.get('IncidentTitle', 'Untitled')[:50]}...")

            # Include processing metadata
            ai_metadata = {
                "centralized_ai": True,
                "processing_time": processing_time,
                "lemmatization_applied": preprocess_metadata.get('lemmatization_applied', False)
            }
            
            response_data = {
                'status': 'success',
                'message': f'Successfully extracted {len(incidents)} incident(s)',
                'document_name': file_name,
                'saved_path': safe_filename,
                'extracted_text_length': len(text),
                'preprocessing_metadata': preprocess_metadata,
                'ai_metadata': ai_metadata,  # Centralized AI stats
                'incidents': incidents
            }
            
            # Include compression metadata if file was compressed
            if compression_metadata:
                response_data['compression_metadata'] = compression_metadata
            
            # Include S3 info if uploaded successfully
            if s3_url:
                response_data['s3_url'] = s3_url
                response_data['s3_key'] = s3_key
            
            return JsonResponse(response_data)
        except Exception as process_error:
            if os.path.exists(file_path):
                os.unlink(file_path)
            raise process_error

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error processing document: {str(e)}'}, status=500)


@api_view(['POST'])
@permission_classes([AllowAny])
@csrf_exempt
@require_tenant  # MULTI-TENANCY: Ensure tenant is present
@tenant_filter   # MULTI-TENANCY: Add tenant_id to request
def save_extracted_incidents(request):
    """
    Save extracted/reviewed incidents to the database.
    MULTI-TENANCY: Incidents will be automatically assigned to user's tenant
    """
    # MULTI-TENANCY: Extract tenant_id from request
    tenant_id = get_tenant_id_from_request(request)
    debug_print(f"💾 Save incidents request received")
    
    try:
        # Try to get data from request.body first
        try:
            data = json.loads(request.body or "{}")
        except Exception as body_error:
            debug_print(f"⚠️  Error reading request.body: {body_error}")
            # Fallback to request.data if available
            if hasattr(request, 'data') and request.data:
                data = request.data
                debug_print(f"✅ Using request.data as fallback")
            else:
                raise body_error
        
        incidents_data = data.get('incidents', [])
        user_id = data.get('user_id', '1')
        
        debug_print(f"💾 Processing {len(incidents_data)} incident(s) for user {user_id}")
        
        if not incidents_data:
            return JsonResponse({'status': 'error', 'message': 'No incidents provided'}, status=400)

        saved = []
        errors = []

        for idx, inc in enumerate(incidents_data):
            try:
                # Parse JSON fields
                mitigation = inc.get('Mitigation')
                if isinstance(mitigation, str):
                    try:
                        mitigation = json.loads(mitigation)
                    except:
                        mitigation = []
                
                form_details = inc.get('IncidentFormDetails')
                if isinstance(form_details, str):
                    try:
                        form_details = json.loads(form_details)
                    except:
                        form_details = {}
                
                # Set required fields with defaults, skip IDs and foreign keys
                kwargs = {
                    'IncidentTitle': inc.get('IncidentTitle', f'Incident {idx+1}'),
                    'Description': inc.get('Description', ''),
                    'Mitigation': mitigation,
                    # Date and Time are required in the model
                    'Date': date.today(),
                    'Time': datetime.now().time(),
                    # Skip UserId and other IDs - they can be null
                    'Origin': inc.get('Origin', 'MANUAL'),
                    'Comments': inc.get('Comments', ''),
                    'RiskCategory': inc.get('RiskCategory', ''),
                    'IncidentCategory': inc.get('IncidentCategory', ''),
                    'RiskPriority': inc.get('RiskPriority', 'Medium'),
                    'Attachments': inc.get('Attachments', ''),
                    'Status': 'Open',  # Always set status to 'Open' for new incidents
                    'RepeatedNot': inc.get('RepeatedNot', False),
                    'CostOfIncident': inc.get('CostOfIncident', ''),
                    'ReopenedNot': inc.get('ReopenedNot', False),
                    'RejectionSource': inc.get('RejectionSource'),
                    'AffectedBusinessUnit': inc.get('AffectedBusinessUnit', ''),
                    'SystemsAssetsInvolved': inc.get('SystemsAssetsInvolved', ''),
                    'GeographicLocation': inc.get('GeographicLocation', ''),
                    'Criticality': inc.get('Criticality', 'Medium'),
                    'InitialImpactAssessment': inc.get('InitialImpactAssessment', ''),
                    'InternalContacts': inc.get('InternalContacts', ''),
                    'ExternalPartiesInvolved': inc.get('ExternalPartiesInvolved', ''),
                    'RegulatoryBodies': inc.get('RegulatoryBodies', ''),
                    'RelevantPoliciesProceduresViolated': inc.get('RelevantPoliciesProceduresViolated', ''),
                    'ControlFailures': inc.get('ControlFailures', ''),
                    'LessonsLearned': inc.get('LessonsLearned', ''),
                    'IncidentClassification': inc.get('IncidentClassification', ''),
                    'PossibleDamage': inc.get('PossibleDamage', ''),
                    'IncidentFormDetails': form_details,
                }
                incident = Incident.objects.create(**kwargs)
                debug_print(f"✅ Saved incident {idx+1}: {incident.IncidentTitle} (ID: {incident.IncidentId})")
                saved.append({'incident_id': incident.IncidentId, 'incident_title': incident.IncidentTitle})
            except Exception as ex:
                debug_print(f"❌ Error saving incident {idx+1}: {str(ex)}")
                errors.append({'incident_index': idx, 'title': inc.get('IncidentTitle'), 'error': str(ex)})

        debug_print(f"💾 Save complete: {len(saved)} saved, {len(errors)} errors")
        
        resp = {
            'status': 'success',
            'message': f'Saved {len(saved)} incident(s)' + (f' with {len(errors)} error(s)' if errors else ''),
            'saved': saved
        }
        if errors:
            resp['errors'] = errors
        return JsonResponse(resp)

    except json.JSONDecodeError:
        return JsonResponse({'status': 'error', 'message': 'Invalid JSON payload'}, status=400)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'status': 'error', 'message': f'Error saving incidents: {str(e)}'}, status=500)


@api_view(['GET'])
@permission_classes([AllowAny])
def test_openai_connection_incident(request):
    """Quick check that OpenAI API responds for incident module."""
    try:
        test_prompt = 'Return a JSON object with a single key "status" set to "ok" and a key "message" with value "OpenAI connection successful".'
        out = call_openai_json(test_prompt)
        return JsonResponse({
            'status': 'success', 
            'openai_reply': out, 
            'model': OPENAI_MODEL,
            'module': 'incident',
            'api_provider': 'OpenAI',
            'message': 'OpenAI API is working correctly for incident module'
        })
    except Exception as e:
        return JsonResponse({
            'status': 'error', 
            'message': f'OpenAI API error: {str(e)}', 
            'model': OPENAI_MODEL,
            'module': 'incident',
            'api_provider': 'OpenAI',
            'suggestion': 'Check OPENAI_API_KEY environment variable'
        }, status=500)