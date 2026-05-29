#!/usr/bin/env python
"""Create a Word document with 2-column tables and proper flowchart diagrams."""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading = parse_xml(r'<w:shd {} w:fill="{}" w:val="clear"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading)

def add_colored_cell_text(cell, text, color_hex):
    """Add colored text to a cell."""
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    set_cell_shading(cell, color_hex)
    if color_hex in ['4472C4', 'D32F2F', '388E3C']:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def draw_box2(ax, x, y, w, h, text, color, fontsize=9):
    rect = FancyBboxPatch((x - w/2, y - h/2), w, h,
                          boxstyle="round,pad=0.08", 
                          facecolor=color, edgecolor='black', linewidth=1.5)
    ax.add_patch(rect)
    ax.text(x, y, text, ha='center', va='center',
            fontsize=fontsize, color='white', weight='bold',
            linespacing=1.3)

def draw_diamond2(ax, x, y, text, color, fontsize=8):
    diamond = patches.Polygon([
        (x, y + 0.5), (x + 1.8, y), (x, y - 0.5), (x - 1.8, y)
    ], closed=True, facecolor=color, edgecolor='black', linewidth=1.5)
    ax.add_patch(diamond)
    ax.text(x, y, text, ha='center', va='center',
            fontsize=fontsize, color='white', weight='bold')

def draw_arrow2(ax, x1, y1, x2, y2):
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
               arrowprops=dict(arrowstyle='->', color='black', lw=2))

def create_workflow_as_is():
    """Create Current (As-Is) workflow flowchart."""
    fig, ax = plt.subplots(1, 1, figsize=(8, 16))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 16)
    ax.axis('off')
    
    cx = 4.0
    y = 15.0
    
    steps = [
        ("User selects framework\nfrom dropdown", '#4472C4', 'box'),
        ("Click 'Check the updates'", '#388E3C', 'box'),
        ("AI searches internet for\nlatest amendment PDF", '#FF9800', 'box'),
        ("Document found?", '#9C27B0', 'decision'),
        ("Download PDF to server\nMEDIA_ROOT/change_management/", '#388E3C', 'box'),
        ("Click 'Start Analysis' button", '#2196F3', 'box'),
        ("AI reads PDF & extracts\npolicies, sub-policies,\ncompliances (5-10 min)", '#FF9800', 'box'),
        ("Show extracted rules\non LEFT panel", '#4CAF50', 'box'),
        ("Click 'Match Compliances'", '#2196F3', 'box'),
        ("AI compares new rules vs\nexisting database rules\n(semantic understanding)", '#FF9800', 'box'),
        ("Show results on RIGHT panel\nMatched vs Not Following", '#4CAF50', 'box'),
        ("What type?", '#9C27B0', 'decision'),
    ]
    
    for i, (text, color, stype) in enumerate(steps):
        if stype == 'box':
            h = 0.9 if '\n' in text else 0.7
            draw_box2(ax, cx, y, 4.5, h, text, color, 9)
            prev_h = h
        else:
            draw_diamond2(ax, cx, y, text, color, 8)
            prev_h = 1.0
            # YES branch
            ax.annotate('', xy=(cx + 2.5, y), xytext=(cx + 1.8, y),
                       arrowprops=dict(arrowstyle='->', color='green', lw=2))
            ax.text(cx + 2.2, y + 0.3, 'YES', fontsize=8, color='green', weight='bold')
            # NO branch
            ax.annotate('', xy=(cx - 2.5, y), xytext=(cx - 1.8, y),
                       arrowprops=dict(arrowstyle='->', color='red', lw=2))
            ax.text(cx - 2.2, y + 0.3, 'NO', fontsize=8, color='red', weight='bold')
            if i == 3:  # Document found?
                ax.text(cx - 3.5, y, 'Show error\n"No document found"', fontsize=8, color='red', ha='center',
                       bbox=dict(boxstyle='round', facecolor='#FFEBEE', edgecolor='red'))
        
        y -= (prev_h + 0.6)
    
    # Branch labels at bottom
    ax.annotate('', xy=(cx + 2.5, y + 0.5), xytext=(cx + 1.8, y + 1.1),
               arrowprops=dict(arrowstyle='->', color='#2196F3', lw=2))
    ax.text(cx + 2.2, y + 0.9, 'NEW', fontsize=8, color='#2196F3', weight='bold')
    draw_box2(ax, cx + 3.5, y + 0.5, 2.5, 0.6, "Click + to add", '#2196F3', 8)
    
    ax.annotate('', xy=(cx - 2.5, y + 0.5), xytext=(cx - 1.8, y + 1.1),
               arrowprops=dict(arrowstyle='->', color='#757575', lw=2))
    ax.text(cx - 2.2, y + 0.9, 'MATCH', fontsize=8, color='#757575', weight='bold')
    draw_box2(ax, cx - 3.5, y + 0.5, 2.5, 0.6, "Show checkmark\n'Already have'", '#757575', 8)
    ax.annotate('', xy=(cx - 3.5, y - 0.2), xytext=(cx - 3.5, y + 0.2),
               arrowprops=dict(arrowstyle='->', color='#757575', lw=1.5))
    draw_box2(ax, cx - 3.5, y - 0.6, 2.5, 0.5, "END", '#BDBDBD', 8)
    
    plt.tight_layout()
    plt.savefig(r'e:\KPMG-WS2-Final-Testing\flowchart_as_is.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return r'e:\KPMG-WS2-Final-Testing\flowchart_as_is.png'

def create_workflow_to_be():
    """Create Proposed (To-Be) workflow flowchart."""
    fig, ax = plt.subplots(1, 1, figsize=(10, 18))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 18)
    ax.axis('off')
    
    cx = 5.0
    y = 17.0
    
    steps = [
        ("User selects framework\nfrom dropdown", '#4472C4', 'box'),
        ("Click 'Check the updates'", '#388E3C', 'box'),
        ("AI searches internet for\nlatest amendment PDF", '#FF9800', 'box'),
        ("Document found?", '#9C27B0', 'decision'),
        ("Download PDF to server", '#388E3C', 'box'),
        ("Click 'Start Analysis' button", '#2196F3', 'box'),
        ("AI reads PDF & extracts\npolicies, sub-policies,\ncompliances (5-10 min)", '#FF9800', 'box'),
        ("Show extracted rules\non LEFT panel", '#4CAF50', 'box'),
        ("Click 'Match Compliances'", '#2196F3', 'box'),
        ("AI compares new rules vs\nexisting database rules\n(semantic understanding)", '#FF9800', 'box'),
        ("Show results on RIGHT panel\nMatched vs Not Following", '#4CAF50', 'box'),
        ("What type?", '#9C27B0', 'decision'),
    ]
    
    for i, (text, color, stype) in enumerate(steps):
        if stype == 'box':
            h = 0.9 if '\n' in text else 0.7
            draw_box2(ax, cx, y, 4.5, h, text, color, 9)
            prev_h = h
        else:
            draw_diamond2(ax, cx, y, text, color, 8)
            prev_h = 1.0
            ax.annotate('', xy=(cx + 2.5, y), xytext=(cx + 1.8, y),
                       arrowprops=dict(arrowstyle='->', color='green', lw=2))
            ax.text(cx + 2.2, y + 0.3, 'YES', fontsize=8, color='green', weight='bold')
            ax.annotate('', xy=(cx - 2.5, y), xytext=(cx - 1.8, y),
                       arrowprops=dict(arrowstyle='->', color='red', lw=2))
            ax.text(cx - 2.2, y + 0.3, 'NO', fontsize=8, color='red', weight='bold')
            if i == 3:
                ax.text(cx - 3.5, y, 'Show error\n"No document found"', fontsize=8, color='red', ha='center',
                       bbox=dict(boxstyle='round', facecolor='#FFEBEE', edgecolor='red'))
        
        y -= (prev_h + 0.6)
    
    # NEW branch (right side)
    ax.annotate('', xy=(cx + 3.0, y + 0.5), xytext=(cx + 1.8, y + 1.1),
               arrowprops=dict(arrowstyle='->', color='#2196F3', lw=2))
    ax.text(cx + 2.5, y + 0.9, 'NEW', fontsize=8, color='#2196F3', weight='bold')
    draw_box2(ax, cx + 4.0, y + 0.5, 2.5, 0.6, "Click + to add", '#2196F3', 8)
    ax.annotate('', xy=(cx + 4.0, y - 0.1), xytext=(cx + 4.0, y + 0.2),
               arrowprops=dict(arrowstyle='->', color='#2196F3', lw=1.5))
    draw_box2(ax, cx + 4.0, y - 0.5, 2.8, 0.7, "Fill form +\npick reviewer", '#2196F3', 8)
    ax.annotate('', xy=(cx + 4.0, y - 1.0), xytext=(cx + 4.0, y - 0.8),
               arrowprops=dict(arrowstyle='->', color='#2196F3', lw=1.5))
    draw_box2(ax, cx + 4.0, y - 1.4, 2.8, 0.5, "Send for Approval", '#388E3C', 8)
    ax.annotate('', xy=(cx + 4.0, y - 1.8), xytext=(cx + 4.0, y - 1.6),
               arrowprops=dict(arrowstyle='->', color='#388E3C', lw=1.5))
    draw_box2(ax, cx + 4.0, y - 2.2, 2.8, 0.5, "Reviewer Approves", '#4CAF50', 8)
    ax.annotate('', xy=(cx + 4.0, y - 2.6), xytext=(cx + 4.0, y - 2.4),
               arrowprops=dict(arrowstyle='->', color='#4CAF50', lw=1.5))
    draw_box2(ax, cx + 4.0, y - 3.0, 2.8, 0.5, "Compliance Active", '#2E7D32', 8)
    
    # MATCHED branch with CHANGE DETECTION (left side) - KEY DIFFERENCE
    ax.annotate('', xy=(cx - 3.0, y + 0.5), xytext=(cx - 1.8, y + 1.1),
               arrowprops=dict(arrowstyle='->', color='#FF9800', lw=2))
    ax.text(cx - 2.5, y + 0.9, 'MATCH', fontsize=8, color='#FF9800', weight='bold')
    draw_box2(ax, cx - 4.0, y + 0.5, 2.5, 0.6, "Did rule\nchange?", '#FF9800', 8)
    
    # YES - Rule changed
    ax.annotate('', xy=(cx - 4.0, y - 0.2), xytext=(cx - 4.0, y + 0.2),
               arrowprops=dict(arrowstyle='->', color='#D32F2F', lw=1.5))
    ax.text(cx - 3.6, y - 0.1, 'YES', fontsize=8, color='#D32F2F', weight='bold')
    draw_box2(ax, cx - 4.0, y - 0.7, 3.0, 0.9, "Show UPDATE badge\nSide-by-side diff\nOLD vs NEW text", '#D32F2F', 8)
    ax.annotate('', xy=(cx - 4.0, y - 1.3), xytext=(cx - 4.0, y - 1.15),
               arrowprops=dict(arrowstyle='->', color='#D32F2F', lw=1.5))
    draw_box2(ax, cx - 4.0, y - 1.7, 2.8, 0.7, "Click Update\npick reviewer", '#D32F2F', 8)
    ax.annotate('', xy=(cx - 4.0, y - 2.1), xytext=(cx - 4.0, y - 2.0),
               arrowprops=dict(arrowstyle='->', color='#D32F2F', lw=1.5))
    draw_box2(ax, cx - 4.0, y - 2.5, 2.8, 0.5, "Send for Approval", '#388E3C', 8)
    ax.annotate('', xy=(cx - 4.0, y - 2.9), xytext=(cx - 4.0, y - 2.7),
               arrowprops=dict(arrowstyle='->', color='#388E3C', lw=1.5))
    draw_box2(ax, cx - 4.0, y - 3.3, 2.8, 0.5, "Reviewer Approves", '#4CAF50', 8)
    ax.annotate('', xy=(cx - 4.0, y - 3.7), xytext=(cx - 4.0, y - 3.5),
               arrowprops=dict(arrowstyle='->', color='#4CAF50', lw=1.5))
    draw_box2(ax, cx - 4.0, y - 4.1, 2.8, 0.5, "Existing Compliance\nUpdated in DB", '#2E7D32', 8)
    
    # Notify Risk & Audit
    ax.annotate('', xy=(cx - 4.0, y - 4.6), xytext=(cx - 4.0, y - 4.35),
               arrowprops=dict(arrowstyle='->', color='#1976D2', lw=1.5))
    draw_box2(ax, cx - 4.0, y - 5.0, 2.8, 0.7, "Check linked\nRisk entries", '#1976D2', 8)
    ax.annotate('', xy=(cx - 4.0, y - 5.4), xytext=(cx - 4.0, y - 5.35),
               arrowprops=dict(arrowstyle='->', color='#1976D2', lw=1.5))
    draw_box2(ax, cx - 4.0, y - 5.8, 2.8, 0.5, "Notify Risk Owners", '#1976D2', 8)
    
    # NO - Rule unchanged
    ax.annotate('', xy=(cx - 5.5, y + 0.2), xytext=(cx - 5.5, y + 0.5),
               arrowprops=dict(arrowstyle='->', color='#757575', lw=1.5))
    ax.text(cx - 5.2, y + 0.4, 'NO', fontsize=8, color='#757575', weight='bold')
    draw_box2(ax, cx - 6.5, y + 0.2, 2.0, 0.5, "No change\nEND", '#757575', 8)
    
    plt.tight_layout()
    plt.savefig(r'e:\KPMG-WS2-Final-Testing\flowchart_to_be.png', dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return r'e:\KPMG-WS2-Final-Testing\flowchart_to_be.png'

def build_docx():
    """Build the final Word document."""
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # ===== TITLE PAGE =====
    title = doc.add_heading('Framework Comparison', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Implementation Plan & Gap Analysis')
    run.font.size = Pt(16)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Date: 28 May 2026\nPrepared for: Change Management Team')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_page_break()
    
    # ===== CONTENTS =====
    doc.add_heading('Contents', level=1)
    for item in ['1. Existing Functions', '2. What Is Missing', '3. Proposed New Functions', 
                 '4. Current Workflow (As-Is)', '5. Proposed Workflow (To-Be)', 
                 '6. Implementation Priority', '7. Files to Modify']:
        doc.add_paragraph(item, style='List Number')
    doc.add_page_break()
    
    # ===== SECTION 1: EXISTING FUNCTIONS (2 columns only) =====
    doc.add_heading('1. Existing Functions', level=1)
    doc.add_paragraph('These functions are already built and working in the system.')
    doc.add_paragraph()
    
    existing_functions = [
        ("check_framework_updates", "Searches the internet for the latest official amendment PDF for the selected framework and downloads it"),
        ("query_perplexity_api", "Calls the Perplexity AI search engine to find official amendment document links on regulator websites"),
        ("start_amendment_analysis", "Triggers the AI to read the downloaded PDF and extract all rules in a background thread so the user does not wait"),
        ("process_downloaded_amendment", "Main pipeline: extracts sections from PDF, then policies, then sub-policies, then compliance records using AI"),
        ("get_framework_origin_data", "Fetches your current framework data from the database (policies, sub-policies, compliances you already have)"),
        ("get_framework_target_data", "Fetches the AI-extracted rules from the amendment PDF stored in Framework.Amendment JSON"),
        ("match_amendments_compliances", "Compares every new rule from the amendment against your existing database rules using AI semantic matching"),
        ("find_control_matches", "Finds the best matching existing compliance for a single new control from the amendment"),
        ("batch_match_controls", "Matches multiple controls at once against your existing database for bulk comparison"),
        ("add_compliance_from_amendment", "Creates a brand-new compliance, sub-policy, and policy in the database and sends it for reviewer approval"),
        ("get_framework_comparison_summary", "Returns summary counts: how many new controls, modified controls, and deprecated controls exist"),
        ("get_migration_gap_analysis", "Returns a detailed breakdown of all gaps with priority levels (High, Medium, Low) and required actions"),
        ("get_amendment_document_info", "Returns the current document processing status so the frontend can poll for completion"),
        ("calculate_hybrid_similarity", "Calculates a similarity score using text matching, keyword overlap, and AI embeddings combined"),
        ("match_all_amendments_compliances", "Bulk matching of all amendment compliances against database compliances with AI fallback"),
        ("_perform_ai_gap_analysis", "Runs AI-powered gap analysis comparing the original framework version against the amended version"),
        ("_assess_ai_compliance_impact", "Assesses how framework changes will impact your organization's overall compliance posture using AI"),
        ("cancel_amendment_analysis", "Allows the user to cancel a long-running background analysis before it completes"),
    ]
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    add_colored_cell_text(table.rows[0].cells[0], 'Function', '4472C4')
    add_colored_cell_text(table.rows[0].cells[1], 'What It Does', '4472C4')
    
    for func, desc in existing_functions:
        row = table.add_row()
        row.cells[0].text = func
        row.cells[1].text = desc
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.5)
    
    doc.add_page_break()
    
    # ===== SECTION 2: WHAT IS MISSING (2 columns only) =====
    doc.add_heading('2. What Is Missing', level=1)
    doc.add_paragraph('These are the critical gaps in the current implementation.')
    doc.add_paragraph()
    
    missing = [
        ("update_existing_rules", "When an amendment changes an existing rule, the system shows 'You already have this' and offers NO way to update the existing compliance record. The old rule stays outdated in the database."),
        ("risk_module_connection", "When a compliance changes, the system does NOT show which Risk entries in the Risk module are linked to it. Risk and Framework Comparison are completely disconnected."),
        ("audit_module_connection", "When a policy or compliance changes, the system does NOT flag which audits cover that area and may need re-auditing. Audit and Framework Comparison are completely disconnected."),
        ("automatic_periodic_checking", "A human must click 'Check the updates' every time. The system does NOT automatically search for new amendments on a weekly or monthly schedule."),
        ("policy_level_impact_summary", "There is no high-level dashboard showing '3 policies, 7 sub-policies, 12 compliances affected.' Only rule-by-rule matching exists."),
    ]
    
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    add_colored_cell_text(table2.rows[0].cells[0], 'Missing Feature', 'D32F2F')
    add_colored_cell_text(table2.rows[0].cells[1], 'What It Should Do', 'D32F2F')
    
    for feature, desc in missing:
        row = table2.add_row()
        row.cells[0].text = feature
        row.cells[1].text = desc
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    table2.columns[0].width = Inches(2.5)
    table2.columns[1].width = Inches(4.5)
    
    doc.add_page_break()
    
    # ===== SECTION 3: PROPOSED NEW FUNCTIONS (2 columns only) =====
    doc.add_heading('3. Proposed New Functions', level=1)
    doc.add_paragraph('These functions need to be implemented to close the gaps.')
    doc.add_paragraph()
    
    new_functions = [
        ("detect_control_changes", "Compares the amendment version of a rule against the database version and detects if wording, criticality, or requirements changed. Returns a change flag and a side-by-side diff."),
        ("get_compliance_change_detail", "Returns the old values vs new values for every field that changed in a matched compliance, so the user can see exactly what is different."),
        ("create_compliance_update_request", "Creates an approval workflow record for updating an EXISTING compliance. Similar to adding a new compliance, but modifies an existing record instead of creating one."),
        ("apply_compliance_update", "Applies approved changes to the actual Compliance record in the database, saves the old version to history, and triggers risk/audit notifications."),
        ("get_pending_compliance_updates", "Lists all compliance update requests that are currently pending reviewer approval for a given framework."),
        ("reject_compliance_update", "Rejects an update request with a reason, keeping the existing compliance unchanged and notifying the requester."),
        ("get_affected_risks", "Finds all Risk entries in the Risk module that are linked to a specific compliance or policy that changed in the amendment."),
        ("get_risk_impact_summary", "Returns a count of how many risks are affected by the amendment, broken down by impact level (High, Medium, Low)."),
        ("notify_risk_owners", "Sends an email or in-app notification to risk owners when a compliance linked to their risk entry has changed."),
        ("get_affected_audits", "Finds all audits in the Audit module that cover a specific policy, sub-policy, or compliance that changed in the amendment."),
        ("get_audit_impact_summary", "Returns a count of how many audits are impacted and how many require re-auditing due to framework changes."),
        ("flag_audit_for_reaudit", "Marks an audit record as needing re-audit because the underlying compliance or policy has changed."),
        ("schedule_framework_update_check", "A scheduled background task (cron job) that automatically checks ALL active frameworks for new amendments on a weekly basis without human intervention."),
        ("send_update_notification", "Sends an email notification to configured administrators when a new amendment is automatically found and downloaded."),
        ("get_policy_impact_summary", "Returns a high-level dashboard summary showing total policies, sub-policies, compliances, risks, and audits affected by the amendment."),
    ]
    
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    add_colored_cell_text(table3.rows[0].cells[0], 'New Function', '388E3C')
    add_colored_cell_text(table3.rows[0].cells[1], 'What It Does', '388E3C')
    
    for func, desc in new_functions:
        row = table3.add_row()
        row.cells[0].text = func
        row.cells[1].text = desc
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    table3.columns[0].width = Inches(2.5)
    table3.columns[1].width = Inches(4.5)
    
    doc.add_page_break()
    
    # ===== SECTION 4: CURRENT WORKFLOW DIAGRAM =====
    doc.add_heading('4. Current Workflow (As-Is)', level=1)
    doc.add_paragraph('This is how the system works today.')
    doc.add_paragraph()
    
    img_as_is = create_workflow_as_is()
    if img_as_is and os.path.exists(img_as_is):
        doc.add_picture(img_as_is, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ===== SECTION 5: PROPOSED WORKFLOW DIAGRAM =====
    doc.add_heading('5. Proposed Workflow (To-Be)', level=1)
    doc.add_paragraph('This is how the system SHOULD work after implementation.')
    doc.add_paragraph()
    
    img_to_be = create_workflow_to_be()
    if img_to_be and os.path.exists(img_to_be):
        doc.add_picture(img_to_be, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ===== SECTION 6: IMPLEMENTATION PRIORITY (2 columns only) =====
    doc.add_heading('6. Implementation Priority', level=1)
    doc.add_paragraph()
    
    priorities = [
        ("Phase 1: Update Existing Rules", "CRITICAL", "When an amendment changes an existing rule, detect the change, show a diff, and allow updating through approval workflow. This is the #1 gap."),
        ("Phase 2: Risk Module Integration", "HIGH", "Show which Risk entries are affected when a compliance changes. Notify risk owners automatically."),
        ("Phase 3: Audit Module Integration", "HIGH", "Show which Audits are impacted. Flag audits that need re-auditing due to framework changes."),
        ("Phase 4: Automatic Checking", "MEDIUM", "Schedule weekly automatic checks so no human needs to click 'Check the updates' manually."),
        ("Phase 5: Impact Summary Dashboard", "MEDIUM", "Show a high-level summary card: X policies, Y sub-policies, Z compliances, N risks affected."),
    ]
    
    table4 = doc.add_table(rows=1, cols=2)
    table4.style = 'Table Grid'
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    add_colored_cell_text(table4.rows[0].cells[0], 'Phase', '4472C4')
    add_colored_cell_text(table4.rows[0].cells[1], 'What It Does', '4472C4')
    
    for phase, priority, desc in priorities:
        row = table4.add_row()
        row.cells[0].text = f"{phase}\n(Priority: {priority})"
        row.cells[1].text = desc
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
        
        if 'CRITICAL' in priority:
            set_cell_shading(row.cells[0], 'D32F2F')
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        elif 'HIGH' in priority:
            set_cell_shading(row.cells[0], 'FF9800')
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.bold = True
        else:
            set_cell_shading(row.cells[0], 'FFC107')
    
    table4.columns[0].width = Inches(2.5)
    table4.columns[1].width = Inches(4.5)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Estimated Timeline: 4-6 weeks with 1-2 developers')
    run.font.bold = True
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ===== SECTION 7: FILES TO MODIFY (2 columns only) =====
    doc.add_heading('7. Files to Modify', level=1)
    doc.add_paragraph()
    
    files = [
        ("framework_comparison.py", "Add new API endpoints for update requests, risk/audit lookups, impact summaries, and automatic checking"),
        ("similarity_matcher.py", "Add change detection logic to flag when a matched compliance has material changes vs the amendment"),
        ("amendment_processor.py", "Add hooks to store amendment source references on extracted compliances"),
        ("framework_update_checker.py", "Add automatic scheduling logic and email notification functions"),
        ("FrameworkComparisonUpdated.vue", "Add new UI: update modals, diff viewers, risk/audit expandable sections, impact summary card"),
        ("frameworkComparisonService.js", "Add new API service methods for update requests, risk/audit lookups, and auto-check"),
        ("models.py (Compliance)", "Add ComplianceHistory JSON field and AmendmentSource field"),
        ("New: compliance_update_request model", "Create ComplianceUpdateRequest table for approval workflow on existing compliance updates"),
        ("New: risk_integration.py", "Helper module to query Risk module and return linked risks for a compliance"),
        ("New: audit_integration.py", "Helper module to query Audit module and return linked audits for a compliance"),
    ]
    
    table5 = doc.add_table(rows=1, cols=2)
    table5.style = 'Table Grid'
    table5.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    add_colored_cell_text(table5.rows[0].cells[0], 'File', '4472C4')
    add_colored_cell_text(table5.rows[0].cells[1], 'What to Change', '4472C4')
    
    for fname, change in files:
        row = table5.add_row()
        row.cells[0].text = fname
        row.cells[1].text = change
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    
    table5.columns[0].width = Inches(2.5)
    table5.columns[1].width = Inches(4.5)
    
    output_path = r'e:\KPMG-WS2-Final-Testing\Framework_Comparison_Plan.docx'
    doc.save(output_path)
    print(f"Word document created successfully at: {output_path}")
    return output_path

if __name__ == '__main__':
    build_docx()
